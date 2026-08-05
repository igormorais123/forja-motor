"""FORJA N2 - F6/F8: renderiza peça em markdown sobre o template oficial + QA visual.

Evolução do piloto M4 com os achados corrigidos:
  - tabelas markdown viram tabelas Word (achado M4: assinaturas/quadros se perdiam);
  - bloco de assinaturas detectado e centralizado sem recuo;
  - negrito inline **texto** preservado;
  - QA completo: PDF via Word COM, render de todas as páginas, grep de placeholders,
    contact sheet; metadados institucionais.

Uso: python forja_render_docx.py <peca.md> <saida_dir> [<titulo>]
"""

import json
import re
import shutil
import sys
from pathlib import Path

FORJA = Path(__file__).resolve().parent
RAIZ = FORJA.parent
sys.path.insert(0, str(RAIZ / "_FERRAMENTAS"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

from word_visual_pipeline import docx_para_pdf, render_paginas
from forja_docx_layout import normalize_medina_body
from forja_n3_common import atomic_write_json, sha256_file
from forja_visual_review import build_pending_review
from forja_metadata import sanitize_final_artifacts

TEMPLATE = RAIZ / "_FERRAMENTAS" / "TEMPLATE_MEDINA_OSORIO_PETICAO.docx"
if not TEMPLATE.exists():
    raise SystemExit(f"TEMPLATE do escritório não encontrado em {TEMPLATE} — sem ele nenhuma peça pode ser gerada (timbre é arte do template).")
PLACEHOLDER_RE = re.compile(r"\[(?!VERIFICAR)[^\]\n]{1,60}\]")


def limpar_corpo(doc):
    body = doc.element.body
    for el in list(body):
        if el.tag.endswith("}p") or el.tag.endswith("}tbl"):
            body.remove(el)


def add_runs_com_negrito(p, texto, base_bold=False):
    # Lição 41(a): além de **negrito**, converter *itálico* (antes vazava asterisco literal no DOCX).
    # O itálico exige conteúdo sem espaço nas bordas ("*ênfase*"), para não capturar multiplicação ("3 * 4").
    partes = re.split(r"(\*\*[^*]+\*\*|\*\S(?:[^*\n]*\S)?\*)", texto)
    for parte in partes:
        if not parte:
            continue
        bold = base_bold
        italic = False
        if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
            parte = parte[2:-2]
            bold = True
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            parte = parte[1:-1]
            italic = True
        run = p.add_run(parte)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic


def eh_assinatura(linha, contexto_pos_deferimento):
    t = linha.strip().strip("*")
    if not contexto_pos_deferimento or not t:
        return False
    return (t == t.upper() and 2 <= len(t.split()) <= 6) or re.match(r"^OAB[/ ]", t, re.I) or t.startswith("Brasília")


def _tipo_produto(texto, titulo):
    abertura = (titulo + "\n" + texto[:1800]).upper()
    if re.search(r"\b(ESTUDO|DIAGN[ÓO]STICO|RELAT[ÓO]RIO|PARECER)\b", abertura):
        return "estudo"
    if re.search(r"^\s*(E-?MAIL|MENSAGEM DE ENTREGA)\b", titulo.upper()):
        return "email"
    return "peca"


def render(md_path, out_dir, titulo="Peça FORJA", tipo=None, *, case_dir=None,
           ledger_path=None, base_dir=None):
    md_path, out_dir = Path(md_path), Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    texto = md_path.read_text(encoding="utf-8")

    # F7 fail-closed (review adversarial 09/07/2026): gates e métricas rodam ANTES
    # de qualquer artefato. Se o verificador/métricas/persistência do F7 falharem,
    # a exceção sobe e NENHUM DOCX/PDF é gerado — nada entregável sem F7.
    sys.path.insert(0, str(Path(__file__).parent))
    from forja_verificador import verificar as gates_forja
    from forja_lastro import material_economico
    from forja_metricas_f7 import metricas_f7
    from datetime import datetime, timezone
    tipo = tipo or _tipo_produto(texto, titulo)
    ledger = None
    if ledger_path:
        ledger = json.loads(Path(ledger_path).read_text(encoding="utf-8"))
    # A rota de render não pode escapar do gate por não conhecer uma pasta de
    # caso. Material econômico sem contexto documental deve reprovar em L9;
    # produto sem marcador monetário segue a incidência estreita calibrada.
    viol = gates_forja(
        texto, tipo, ledger=ledger, base_dir=base_dir, case_dir=case_dir,
        exigir_economico=material_economico(texto),
    )
    gate = {"total": len(viol), "p0": sum(1 for x in viol if x["sev"] == "P0"),
            "p1": sum(1 for x in viol if x["sev"] == "P1"),
            "violacoes": viol}
    import hashlib
    f7 = {"arquivo": str(md_path), "tipo": "peca",
          # Lastro anti-fraude (régua 10/07/2026): o F10 recomputa este hash do fonte;
          # F7 escrito à mão ou desatualizado não fecha demanda.
          "mdSha256": hashlib.sha256(texto.encode("utf-8")).hexdigest(),
          **gate, **metricas_f7(texto),
          "geradoEm": datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")}
    from forja_metadata import retry_transient_io
    retry_transient_io(lambda: (out_dir / "F7_VERIFICADOR_FORJA.json").write_text(
        json.dumps(f7, ensure_ascii=False, indent=2), encoding="utf-8"))

    # O relatório é persistido para orientar a correção, mas nenhum DOCX/PDF nasce
    # com P0. Antes, o comentário dizia fail-closed e o código ainda seguia o render.
    if gate["p0"]:
        amostra = "; ".join(
            f"{item['gate']}: {item['problema']}" for item in viol if item["sev"] == "P0"
        )[:900]
        raise RuntimeError(f"F7 REPROVADO — {gate['p0']} P0: {amostra}")

    destino = out_dir / (md_path.stem + ".docx")
    shutil.copy2(TEMPLATE, destino)
    doc = Document(str(destino))
    limpar_corpo(doc)

    linhas = texto.splitlines()
    i, pos_deferimento, primeiro = 0, False, True
    while i < len(linhas):
        linha = linhas[i].rstrip()

        # tabela markdown
        if linha.strip().startswith("|") and i + 1 < len(linhas) and re.match(r"^\s*\|[\s:|-]+\|\s*$", linhas[i + 1]):
            celulas_header = [c.strip() for c in linha.strip().strip("|").split("|")]
            corpo_tab = []
            j = i + 2
            while j < len(linhas) and linhas[j].strip().startswith("|"):
                corpo_tab.append([c.strip() for c in linhas[j].strip().strip("|").split("|")])
                j += 1
            tabela = doc.add_table(rows=1 + len(corpo_tab), cols=len(celulas_header))
            tabela.style = "Table Grid"
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, txt in enumerate(celulas_header):
                par = tabela.rows[0].cells[c].paragraphs[0]
                add_runs_com_negrito(par, txt.replace("**", ""), base_bold=True)
            for r, row in enumerate(corpo_tab, start=1):
                for c, txt in enumerate(row[:len(celulas_header)]):
                    add_runs_com_negrito(tabela.rows[r].cells[c].paragraphs[0], txt)
            for row in tabela.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        par.paragraph_format.line_spacing = 1.15
                        for run in par.runs:
                            run.font.size = Pt(10.5)
            doc.add_paragraph()
            i = j
            continue

        if not linha.strip():
            i += 1
            continue

        # separador horizontal de markdown (---, ***, ___): nunca vai para o Word como texto
        if re.match(r"^\s*([-*_])\1{2,}\s*$", linha):
            i += 1
            continue

        if re.search(r"pede deferimento", linha, re.I):
            pos_deferimento = True

        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.line_spacing = 1.5

        m_header = re.match(r"^(#{1,3})\s+(.*)$", linha)
        conteudo = m_header.group(2) if m_header else linha
        conteudo = conteudo.strip()

        if m_header:
            add_runs_com_negrito(p, conteudo.replace("**", ""), base_bold=True)
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER if primeiro else WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Cm(0)
            fmt.space_before = Pt(12)
            primeiro = False
        elif eh_assinatura(conteudo, pos_deferimento):
            add_runs_com_negrito(p, conteudo, base_bold=(conteudo.replace("**", "") == conteudo.replace("**", "").upper()))
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = Cm(0)
        elif conteudo == conteudo.upper() and len(conteudo) > 25 and primeiro:
            # endereçamento inicial em caixa alta
            add_runs_com_negrito(p, conteudo, base_bold=True)
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Cm(0)
            primeiro = False
        else:
            add_runs_com_negrito(p, conteudo)
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Cm(2.0)
            primeiro = False
        i += 1

    doc.core_properties.author = "Medina Osório Advogados"
    doc.core_properties.last_modified_by = "Medina Osório Advogados"
    doc.core_properties.title = titulo
    pre_layout = out_dir / f"{md_path.stem}.pre_layout_raw.docx"
    doc.save(str(pre_layout))

    # O template histórico trazia um fólio VML largo, capaz de invadir o corpo.
    # A saída promovida nasce de uma cópia intermediária e passa pela
    # normalização hash-bound; conteúdo e controle de alterações têm de ficar
    # idênticos. Só então o PDF pode ser gerado.
    layout_audit = normalize_medina_body(pre_layout, destino)
    atomic_write_json(out_dir / "DOCX_LAYOUT_AUDIT.json", layout_audit)
    if not layout_audit["approved"]:
        codes = ", ".join(item["code"] for item in layout_audit["findings"][:8])
        raise RuntimeError(f"DIAGRAMAÇÃO WORD REPROVADA — {codes}")

    pdf = out_dir / (md_path.stem + ".pdf")
    docx_para_pdf(str(destino), str(pdf))
    sanitize_final_artifacts(destino, pdf)
    paginas_dir = out_dir / "paginas"
    if paginas_dir.exists():
        shutil.rmtree(paginas_dir)
    render_paginas(str(pdf), str(paginas_dir), dpi=110)
    arquivos = sorted(paginas_dir.glob("*.png"))
    generator_run_id = f"render-{sha256_file(pdf)[:16]}"
    rendered_pages = [
        {
            "page": number,
            "imagePath": str(path),
            "imageSha256": sha256_file(path),
        }
        for number, path in enumerate(arquivos, 1)
    ]
    pending_review_path = out_dir / "VISUAL_REVIEW_PENDING.json"
    build_pending_review(
        pending_review_path,
        pdf=pdf,
        docx=destino,
        rendered_pages=rendered_pages,
        generator_run_id=generator_run_id,
    )

    texto_final = "\n".join(par.text for par in Document(str(destino)).paragraphs)
    for tab in Document(str(destino)).tables:
        for row in tab.rows:
            for cell in row.cells:
                texto_final += "\n" + cell.text
    placeholders = sorted({m.group(0) for m in PLACEHOLDER_RE.finditer(texto_final)})
    verificar = sorted({m.group(0) for m in re.finditer(r"\[VERIFICAR[^\]]*\]", texto_final)})

    # Fail-fast de marcação vazada (review 09/07/2026): asterisco literal no DOCX
    # final significa markdown não convertido — mutação silenciosa é pior que erro.
    if "*" in texto_final:
        trecho = texto_final[max(0, texto_final.index("*") - 60):texto_final.index("*") + 60]
        raise RuntimeError(f"ASTERISCO LITERAL no DOCX renderizado (markdown não convertido): ...{trecho}...")

    from PIL import Image
    thumbs = []
    for a in arquivos:
        im = Image.open(a)
        im.thumbnail((350, 500))
        thumbs.append(im)
    if thumbs:
        cols = 5
        rows = (len(thumbs) + cols - 1) // cols
        w = max(t.width for t in thumbs) + 8
        h = max(t.height for t in thumbs) + 8
        sheet = Image.new("RGB", (cols * w, rows * h), "white")
        for k, t in enumerate(thumbs):
            sheet.paste(t, ((k % cols) * w + 4, (k // cols) * h + 4))
        from forja_metadata import retry_transient_io
        retry_transient_io(lambda: sheet.save(out_dir / "contact_sheet.png"))

    resumo = {"docx": str(destino), "pdf": str(pdf), "paginas": len(arquivos), "tipoProduto": tipo,
              "placeholdersProibidos": placeholders, "verificarDeliberados": verificar,
              "gatesForjaVerificador": gate,
              "layoutAudit": str(out_dir / "DOCX_LAYOUT_AUDIT.json"),
              "visualReviewStatus": "pending",
              "visualReviewTemplate": str(pending_review_path),
              "deliveryReady": False,
              "contactSheet": str(out_dir / "contact_sheet.png")}
    print(json.dumps(resumo, ensure_ascii=False, indent=2))
    return resumo


if __name__ == "__main__":
    args = sys.argv[1:]
    if len(args) < 2:
        raise SystemExit("uso: python forja_render_docx.py <peca.md> <saida_dir> [titulo] [--case-dir D] [--ledger F] [--base-dir D]")
    kwargs = {}
    for flag, key in (("--case-dir", "case_dir"), ("--ledger", "ledger_path"), ("--base-dir", "base_dir")):
        if flag in args:
            pos = args.index(flag)
            if pos + 1 >= len(args):
                raise SystemExit(f"{flag} exige um valor")
            kwargs[key] = args[pos + 1]
    titulo = args[2] if len(args) > 2 and not args[2].startswith("--") else "Peça FORJA"
    render(args[0], args[1], titulo, **kwargs)
