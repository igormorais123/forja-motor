from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import threading
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document

from forja_document_compare import (
    LAYERS,
    LAYER_CAUSE,
    classify_change,
    compare_documents,
    extract_document,
)
from forja_learning import LAYERS as LEARNING_LAYERS, LAYER_CAUSES as LEARNING_LAYER_CAUSES
from forja_learning_registry import active_rules, register_promoted_rule, suite_learning_findings
from forja_n3_common import ForjaN3Error, RevisionConflict, read_json, sha256_file
from forja_n4_common import ARTIFACT_SPECS, build_envelope, write_artifact
from forja_post_protocol import (
    _case_for_demand,
    _sender_allowed,
    _select_return_parts,
    backfill_baseline_from_gmail,
    content_key,
    evidence_key,
    ingest_return,
    promote_learning,
    rebuild_comparison,
    resolve_ai_baseline,
)
from forja_post_protocol_contracts import (
    LAYERS as CONTRACT_LAYERS,
    LAYER_CAUSES as CONTRACT_LAYER_CAUSES,
    validate_document_comparison,
    validate_learning_candidate,
)
from forja_state_machine import derive_state, initialize_case, record_event


ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent
OFFICE_SCRIPTS = REPO / "gestao_escritorio" / "scripts"
sys.path.insert(0, str(OFFICE_SCRIPTS))
from sync_forja_gestao import _post_protocol_summary  # noqa: E402


def write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


class PostProtocolContractTests(unittest.TestCase):
    def test_catalog_uses_exact_new_schema_names(self) -> None:
        catalog = json.loads((ROOT / "n4_schemas" / "ARTIFACT_CATALOG.json").read_text(encoding="utf-8"))
        expected = {
            "F10_POST_PROTOCOL_RETURN.json": "post_protocol_return.schema.json",
            "F10_PROTOCOL_EVIDENCE.json": "protocol_evidence.schema.json",
            "F10_POST_PROTOCOL_BASELINE_BACKFILL.json": "post_protocol_baseline_backfill.schema.json",
            "F10_POST_PROTOCOL_DOCUMENT_COMPARISON.json": "document_comparison.schema.json",
            "F10_LEARNING_CANDIDATE.json": "learning_candidate.schema.json",
        }
        for artifact, schema in expected.items():
            self.assertEqual(schema, ARTIFACT_SPECS[artifact]["schema"])
            self.assertEqual(schema, catalog["artifacts"][artifact]["schema"])
            self.assertTrue((ROOT / "n4_schemas" / schema).is_file())

    def test_layer_to_cause_is_total_and_default_deny(self) -> None:
        self.assertEqual(LAYERS, set(LAYER_CAUSE))
        layer, cause, impact, confidence, reasons = classify_change("", "\u0000")
        self.assertIn(layer, LAYERS)
        self.assertIn(cause, set(LAYER_CAUSE.values()))
        self.assertGreaterEqual(confidence, 0.0)
        self.assertEqual(CONTRACT_LAYERS, LEARNING_LAYERS)
        self.assertEqual(CONTRACT_LAYER_CAUSES, LEARNING_LAYER_CAUSES)

    def test_tracked_comparison_rejects_raw_legal_text(self) -> None:
        findings = validate_document_comparison({
            "baseline": {"artifactId": "a", "sha256": "a" * 64, "path": "x"},
            "humanArtifact": {"artifactId": "b", "sha256": "b" * 64, "path": "y"},
            "privateComparisonHash": "c" * 64,
            "changes": [{"before": "texto jurídico"}],
        })
        self.assertIn("PP-CONTRACT-RAW-TEXT", {item["code"] for item in findings})

    def test_idempotency_keys_separate_content_from_evidence(self) -> None:
        digest = "a" * 64
        self.assertEqual(content_key("case-a", digest), content_key("case-a", digest))
        self.assertNotEqual(
            evidence_key("me", "thread-a", "message-a", "attachment-a"),
            evidence_key("me", "thread-b", "message-b", "attachment-b"),
        )


class PostProtocolStateTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.case = Path(self.temp.name) / "case-post"
        self.case.mkdir()
        initialize_case(self.case, demand_id="demand-post")

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_post_protocol_does_not_reopen_fulfilled_delivery(self) -> None:
        record_event(
            self.case,
            "case_fulfilled",
            expected_revision=1,
            idempotency_key="fulfilled",
        )
        record_event(
            self.case,
            "post_protocol_candidate_detected",
            expected_revision=2,
            idempotency_key="post-candidate",
            payload={"contentKey": "x" * 64, "openReasonCodes": []},
        )
        state = derive_state(self.case)
        self.assertEqual("fulfilled_by_forja_f10", state["lifecycleStatus"])
        self.assertEqual("candidate_detected", state["postProtocol"]["status"])

    def test_concurrent_post_protocol_revision_conflict(self) -> None:
        outcomes: list[str] = []
        barrier = threading.Barrier(2)

        def worker(index: int) -> None:
            barrier.wait()
            try:
                record_event(
                    self.case,
                    "post_protocol_candidate_detected",
                    expected_revision=1,
                    idempotency_key=f"candidate-{index}",
                    payload={"contentKey": f"{index}" * 64, "openReasonCodes": []},
                )
                outcomes.append("ok")
            except RevisionConflict:
                outcomes.append("conflict")

        threads = [threading.Thread(target=worker, args=(index,)) for index in (1, 2)]
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join()
        self.assertEqual(["conflict", "ok"], sorted(outcomes))

    def test_panel_projection_rejects_free_prose(self) -> None:
        summary = _post_protocol_summary({
            "postProtocol": {
                "status": "review_pending",
                "protocolStatus": "protocol_claimed",
                "contentKey": "a" * 64,
                "openReasonCodes": ["PP-03", "texto livre", "N4-X"],
                "legalText": "conteúdo da peça não pode sair",
                "emailBody": "trecho de e-mail",
                "learningCandidateIds": ["one"],
            }
        })
        serialized = json.dumps(summary, ensure_ascii=False)
        self.assertNotIn("conteúdo da peça", serialized)
        self.assertNotIn("trecho de e-mail", serialized)
        self.assertEqual(["PP-03"], summary["reasonCodes"])

    def test_reason_codes_are_not_erased_by_later_success_events(self) -> None:
        record_event(
            self.case,
            "post_protocol_claimed",
            expected_revision=1,
            idempotency_key="claimed",
            payload={
                "contentKey": "a" * 64,
                "protocolEvidenceId": "protocol-a",
                "protocolStatus": "protocol_claimed",
                "openReasonCodes": ["PP-03"],
                "reasonSource": "protocol",
            },
        )
        record_event(
            self.case,
            "post_protocol_review_pending",
            expected_revision=2,
            idempotency_key="review",
            payload={"contentKey": "a" * 64, "openReasonCodes": [], "reasonSource": "review"},
        )
        self.assertEqual(["PP-03"], derive_state(self.case)["postProtocol"]["openReasonCodes"])


class PostProtocolPipelineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.case = Path(self.temp.name) / "case-pipeline"
        self.case.mkdir()
        initialize_case(self.case, demand_id="demand-pipeline")
        self.baseline = Path(self.temp.name) / "baseline.docx"
        write_docx(
            self.baseline,
            [
                "EXCELENTÍSSIMO SENHOR DESEMBARGADOR",
                "Processo 0000001-11.2026.8.07.0001",
                "A parte requer o provimento do recurso.",
            ],
        )
        selected_hash = sha256_file(self.baseline)
        delivery = build_envelope(
            self.case,
            "F10_DELIVERY_INTEGRITY.json",
            {
                "packageArtifactId": "package-one",
                "selectedArtifactId": "petition-one",
                "packageHash": selected_hash,
                "selectedHash": selected_hash,
                "selectedPath": str(self.baseline),
                "preSendMatch": True,
                "postDeliveryVerification": {
                    "mode": "channel_hash",
                    "deliveredHash": selected_hash,
                    "deliveryEvidenceId": "email-sent-one",
                    "status": "confirmed",
                },
                "deliveredAt": "2026-07-01T10:00:00-03:00",
            },
            source_hashes=[selected_hash],
            producer_run_id="delivery-producer",
            reviewer_run_id="delivery-reviewer",
            status="approved",
        )
        write_artifact(self.case, "F10_DELIVERY_INTEGRITY.json", delivery)
        self.human = Path(self.temp.name) / "human.docx"
        write_docx(
            self.human,
            [
                "EXCELENTÍSSIMO SENHOR DESEMBARGADOR",
                "Processo 0000001-11.2026.8.07.0001",
                "A parte requer o provimento integral do recurso.",
            ],
        )

    def tearDown(self) -> None:
        self.temp.cleanup()

    def ingest(self, *, message: str = "message-one", explicit_links=None, declaration="") -> dict:
        return ingest_return(
            self.case,
            self.human,
            account_id="gmail-me",
            thread_id="thread-one",
            message_id=message,
            attachment_id=f"attachment-{message}",
            received_at="2026-07-02T10:00:00-03:00",
            original_name="Petição final humana.docx",
            piece_name="Agravo interno",
            process_id="0000001-11.2026.8.07.0001",
            declaration_text=declaration,
            explicit_evidence_links=explicit_links,
        )

    def test_capture_diff_report_and_case_limited_learning(self) -> None:
        result = self.ingest()
        self.assertEqual("review_pending", result["status"])
        self.assertEqual("human_final_received", result["protocolStatus"])
        folder = Path(result["folder"])
        self.assertTrue(folder.name.startswith("VERSÃO HUMANA FINAL"))
        self.assertTrue((folder / "MUDANÇAS_IA_VS_PEÇA_PROTOCOLADA.md").is_file())
        self.assertTrue((folder / "COMPARAÇÃO_PRIVADA_IA_VS_HUMANO.json").is_file())
        comparison = json.loads(
            (self.case / "n4_artifacts" / "F10_POST_PROTOCOL_DOCUMENT_COMPARISON.json").read_text(encoding="utf-8")
        )
        self.assertEqual(sha256_file(self.baseline), comparison["baseline"]["sha256"])
        for change in comparison["changes"]:
            self.assertNotIn("before", change)
            self.assertNotIn("after", change)
            self.assertEqual("case", change["scopeCeiling"])
        learning = json.loads(
            (self.case / "n4_artifacts" / "F10_LEARNING_CANDIDATE.json").read_text(encoding="utf-8")
        )
        self.assertTrue(learning["candidates"])
        self.assertTrue(all(item["scope"] == "case" for item in learning["candidates"]))
        self.assertTrue(all(item["status"] == "observed" for item in learning["candidates"]))

    def test_protocol_claim_does_not_use_protocolled_name(self) -> None:
        result = self.ingest(declaration="Confirmo que a peça foi protocolada.")
        self.assertEqual("protocol_claimed", result["protocolStatus"])
        self.assertTrue(Path(result["folder"]).name.startswith("VERSÃO HUMANA FINAL"))

    def test_verified_file_link_allows_protocolled_name(self) -> None:
        result = self.ingest(explicit_links=[{
            "evidenceId": "receipt-linked",
            "kind": "filing_receipt",
            "strength": "verified_file_link",
            "evidencePath": str(self.human),
            "sha256": sha256_file(self.human),
            "externalProtocolId": "PROTOCOLO-2026-000001",
        }])
        self.assertEqual("protocol_verified", result["protocolStatus"])
        self.assertTrue(Path(result["folder"]).name.startswith("PEÇA PROTOCOLADA"))

    def test_incomplete_explicit_protocol_link_is_rejected(self) -> None:
        with self.assertRaises(ForjaN3Error):
            self.ingest(explicit_links=[{
                "evidenceId": "receipt-linked",
                "kind": "filing_receipt",
                "strength": "verified_file_link",
            }])

    def test_electronic_signature_alone_is_not_protocol_verification(self) -> None:
        signed = Path(self.temp.name) / "signed.docx"
        write_docx(signed, ["Documento assinado eletronicamente."])
        result = ingest_return(
            self.case,
            self.human,
            account_id="gmail-me",
            thread_id="thread-signed",
            message_id="message-signed",
            attachment_id="attachment-signed",
            received_at="2026-07-02T10:00:00-03:00",
            evidence_paths=[signed],
        )
        self.assertEqual("human_final_received", result["protocolStatus"])

    def test_resent_attachment_does_not_duplicate_capture_or_diff(self) -> None:
        first = self.ingest(message="message-one")
        event_count = len(list((self.case / "events").glob("*.json")))
        second = self.ingest(message="message-two")
        self.assertTrue(first["created"])
        self.assertFalse(second["created"])
        self.assertEqual("duplicate_content", second["status"])
        folders = list(self.case.glob("VERSÃO HUMANA FINAL — *"))
        self.assertEqual(1, len(folders))
        self.assertEqual(event_count + 1, len(list((self.case / "events").glob("*.json"))))

    def test_baseline_hash_drift_blocks_diff_but_keeps_capture(self) -> None:
        write_docx(self.baseline, ["arquivo alterado depois da entrega"])
        result = self.ingest()
        self.assertEqual("captured_baseline_unresolved", result["status"])
        self.assertIn("PP-BASELINE-HASH", result["reasonCodes"])
        self.assertTrue(Path(result["folder"]).is_dir())
        self.assertFalse(
            (self.case / "n4_artifacts" / "F10_POST_PROTOCOL_DOCUMENT_COMPARISON.json").exists()
        )

    def test_interrupted_capture_resumes_instead_of_becoming_duplicate(self) -> None:
        from forja_post_protocol import _write_artifact_checked as real_write

        calls = {"count": 0}

        def fail_once(*args, **kwargs):
            calls["count"] += 1
            if calls["count"] == 1:
                raise RuntimeError("injected")
            return real_write(*args, **kwargs)

        with patch("forja_post_protocol._write_artifact_checked", side_effect=fail_once):
            with self.assertRaises(RuntimeError):
                self.ingest(message="interrupted")
        resumed = self.ingest(message="interrupted")
        self.assertEqual("review_pending", resumed["status"])
        self.assertFalse(resumed["created"])
        index = read_json(self.case / "private" / "post_protocol" / "INDEX.json", {})
        entry = next(iter(index["contents"].values()))
        self.assertEqual("review_pending", entry["state"])

    def test_low_confidence_pdf_is_persisted_as_blocked(self) -> None:
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF indisponível")
        blank = Path(self.temp.name) / "blank.pdf"
        document = fitz.open()
        document.new_page()
        document.save(blank)
        document.close()
        self.human = blank
        result = self.ingest(message="blank-pdf")
        self.assertEqual("blocked", result["status"])
        self.assertIn("PP-OCR-LOW-CONFIDENCE", result["reasonCodes"])
        self.assertEqual("blocked", derive_state(self.case)["postProtocol"]["status"])

    def test_naive_delivery_timestamp_is_normalized(self) -> None:
        payload = read_json(self.case / "n4_artifacts" / "F10_DELIVERY_INTEGRITY.json", {})
        payload["deliveredAt"] = "2026-07-01T10:00:00"
        from forja_n4_common import expected_content_hash

        payload["contentHash"] = expected_content_hash(payload)
        write_artifact(self.case, "F10_DELIVERY_INTEGRITY.json", payload)
        baseline, reasons = resolve_ai_baseline(self.case, received_at="2026-07-02T10:00:00-03:00")
        self.assertIsNotNone(baseline)
        self.assertEqual([], reasons)

    def test_backfill_never_overwrites_f9_or_f10(self) -> None:
        f10 = self.case / "n4_artifacts" / "F10_DELIVERY_INTEGRITY.json"
        before = f10.read_bytes()
        sent = Path(self.temp.name) / "sent.docx"
        write_docx(sent, ["versão enviada"])
        raw = sent.read_bytes()
        import base64

        demand = {"emailsResposta": ["sent-message"]}
        message = {
            "threadId": "thread",
            "payload": {
                "headers": [{"name": "Date", "value": "Wed, 1 Jul 2026 10:00:00 -0300"}],
                "parts": [{
                    "filename": "sent.docx",
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "body": {"attachmentId": "sent-attachment"},
                }],
            },
        }
        result = backfill_baseline_from_gmail(
            self.case,
            demand,
            human_suffix=".docx",
            received_at="2026-07-02T10:00:00-03:00",
            get_message=lambda _id: (message, {"ok": True}),
            get_attachment=lambda _m, _a: (
                {"data": base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")},
                {"ok": True},
            ),
            shadow=False,
        )
        self.assertEqual("created", result["status"])
        self.assertEqual(before, f10.read_bytes())
        self.assertTrue((self.case / "n4_artifacts" / "F10_POST_PROTOCOL_BASELINE_BACKFILL.json").is_file())
        second = backfill_baseline_from_gmail(
            self.case,
            demand,
            human_suffix=".docx",
            received_at="2026-07-02T10:00:00-03:00",
            get_message=lambda _id: (message, {"ok": True}),
            get_attachment=lambda _m, _a: (
                {"data": base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")},
                {"ok": True},
            ),
            shadow=False,
        )
        self.assertEqual("existing", second["status"])

    def test_rebuild_preserves_rejected_review_decisions(self) -> None:
        result = self.ingest()
        candidate_path = self.case / "n4_artifacts" / "F10_LEARNING_CANDIDATE.json"
        candidate = read_json(candidate_path, {})
        candidate["candidates"][0]["status"] = "rejected"
        candidate["candidates"][0]["decision"] = "rejected"
        write_artifact(self.case, "F10_LEARNING_CANDIDATE.json", candidate)
        diff_path = self.case / "n4_artifacts" / "F10_HUMAN_DIFF_CLASSIFICATION.json"
        diff = read_json(diff_path, {})
        diff["changes"][0]["reviewDecision"] = "rejected"
        write_artifact(self.case, "F10_HUMAN_DIFF_CLASSIFICATION.json", diff)
        rebuild_comparison(self.case, result["contentKey"])
        rebuilt_candidate = read_json(candidate_path, {})
        rebuilt_diff = read_json(diff_path, {})
        self.assertEqual("rejected", rebuilt_candidate["candidates"][0]["status"])
        self.assertEqual("rejected", rebuilt_diff["changes"][0]["reviewDecision"])

    def test_second_return_preserves_human_learning_decisions_and_archives_first(self) -> None:
        first = self.ingest(message="return-one")
        candidate_path = self.case / "n4_artifacts" / "F10_LEARNING_CANDIDATE.json"
        diff_path = self.case / "n4_artifacts" / "F10_HUMAN_DIFF_CLASSIFICATION.json"
        candidate = read_json(candidate_path, {})
        target = candidate["candidates"][0]
        target.update({
            "status": "rejected",
            "decision": "rejected",
            "origin": "human_selected",
            "originEvidenceId": "gmail-message:return-one",
        })
        write_artifact(self.case, "F10_LEARNING_CANDIDATE.json", candidate)
        diff = read_json(diff_path, {})
        diff["changes"][0].update({
            "reviewDecision": "rejected",
            "origin": "human_selected",
            "reviewedBy": "office-email:return-one",
        })
        write_artifact(self.case, "F10_HUMAN_DIFF_CLASSIFICATION.json", diff)
        document = Document(self.human)
        document.core_properties.title = "segunda rodada"
        document.save(self.human)
        second = self.ingest(message="return-two")
        self.assertNotEqual(first["contentKey"], second["contentKey"])
        current_candidate = read_json(candidate_path, {})["candidates"][0]
        current_diff = read_json(diff_path, {})["changes"][0]
        self.assertEqual("rejected", current_candidate["status"])
        self.assertEqual("rejected", current_candidate["decision"])
        self.assertEqual("human_selected", current_candidate["origin"])
        self.assertEqual("rejected", current_diff["reviewDecision"])
        archived = (
            self.case
            / "n4_artifacts"
            / "post_protocol_history"
            / first["contentKey"]
        )
        self.assertTrue((archived / "F10_LEARNING_CANDIDATE.json").is_file())
        self.assertTrue((archived / "F10_HUMAN_DIFF_CLASSIFICATION.json").is_file())

    def test_promotion_requires_real_fixture_and_test(self) -> None:
        result = self.ingest()
        candidate_id = result["learningCandidateIds"][0]
        with self.assertRaisesRegex(ForjaN3Error, "fixture não existe"):
            promote_learning(
                self.case,
                candidate_id,
                content_key_value=result["contentKey"],
                approved_by="office-email:test",
                fixture_id="missing.json",
                test_id="test_forja_post_protocol.py",
                evidence_runs=["run-one"],
                evidence_case_ids=[self.case.name],
            )

    def test_promotion_persists_approved_decision_in_both_artifacts(self) -> None:
        result = self.ingest()
        candidate_id = result["learningCandidateIds"][0]
        fixture = Path(self.temp.name) / "promotion-fixture.json"
        fixture.write_text('{"expected": "approved"}', encoding="utf-8")
        test_id = (
            "test_forja_post_protocol.py"
            "::PostProtocolContractTests"
            "::test_idempotency_keys_separate_content_from_evidence"
        )
        with patch(
            "forja_post_protocol.register_promoted_rule",
            return_value={"ruleId": "rule-test"},
        ):
            promoted = promote_learning(
                self.case,
                candidate_id,
                content_key_value=result["contentKey"],
                approved_by="office-email:test",
                fixture_id=str(fixture),
                test_id=test_id,
                evidence_runs=["run-one"],
                evidence_case_ids=[self.case.name],
            )
            repeated = promote_learning(
                self.case,
                candidate_id,
                content_key_value=result["contentKey"],
                approved_by="office-email:test",
                fixture_id=str(fixture),
                test_id=test_id,
                evidence_runs=["run-one"],
                evidence_case_ids=[self.case.name],
            )
        learning = read_json(
            self.case / "n4_artifacts" / "F10_LEARNING_CANDIDATE.json",
            {},
        )
        human_diff = read_json(
            self.case / "n4_artifacts" / "F10_HUMAN_DIFF_CLASSIFICATION.json",
            {},
        )
        stored = next(item for item in learning["candidates"] if item["candidateId"] == candidate_id)
        proposal = next(
            item for item in human_diff["regressionProposals"]
            if item["proposalId"] == candidate_id
        )
        self.assertEqual("approved", promoted["decision"])
        self.assertEqual("approved", repeated["decision"])
        self.assertEqual("approved", stored["decision"])
        self.assertEqual("approved", proposal["decision"])
        self.assertEqual([], validate_learning_candidate(learning))

    def test_two_latest_delivery_hashes_are_ambiguous(self) -> None:
        history = self.case / "n4_artifacts" / "delivery_history"
        history.mkdir(parents=True)
        other = Path(self.temp.name) / "other.docx"
        write_docx(other, ["outra entrega"])
        for index, path in enumerate((self.baseline, other), 1):
            payload = {
                "contentHash": f"{index}" * 64,
                "selectedArtifactId": f"petition-{index}",
                "selectedHash": sha256_file(path),
                "selectedPath": str(path),
                "postDeliveryVerification": {"status": "confirmed"},
                "deliveredAt": "2026-07-01T10:00:00-03:00",
            }
            (history / f"{index}.json").write_text(json.dumps(payload), encoding="utf-8")
        baseline, reasons = resolve_ai_baseline(
            self.case,
            received_at="2026-07-02T10:00:00-03:00",
        )
        self.assertIsNone(baseline)
        self.assertEqual(["PP-BASELINE-AMBIGUOUS"], reasons)


class DocumentExtractionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_material_mutation_classifiers(self) -> None:
        scenarios = {
            "procedural_identity": ("Processo 0000001-11.2026.8.07.0001", "Processo 0000002-22.2026.8.07.0001"),
            "authority_citation": ("Segundo o REsp 1.", "Segundo o REsp 2."),
            "request_relief": ("Requer provimento.", "Requer improvimento."),
            "calculation": ("Valor de R$ 100,00.", "Valor de R$ 200,00."),
        }
        for expected, pair in scenarios.items():
            with self.subTest(expected=expected):
                layer, _cause, impact, _confidence, _reasons = classify_change(*pair)
                self.assertEqual(expected, layer)
                self.assertEqual("material", impact)

    def test_polarity_and_quantifier_changes_are_material(self) -> None:
        for before, after in (
            ("O pedido deve ser acolhido.", "O pedido não deve ser acolhido."),
            ("O provimento é parcial.", "O provimento é integral."),
        ):
            with self.subTest(after=after):
                layer, _cause, impact, _confidence, reasons = classify_change(before, after)
                self.assertEqual("reasoning", layer)
                self.assertEqual("material", impact)
                self.assertIn("deterministic:polarity_or_quantifier", reasons)

    def test_grouped_locator_keeps_subedits_separate(self) -> None:
        before = self.root / "before.txt"
        after = self.root / "after.txt"
        before.write_text("O pedido é parcial e por isso deve ser acolhido.", encoding="utf-8")
        after.write_text("O pedido é integral e por isso não deve ser acolhido.", encoding="utf-8")
        comparison = compare_documents(before, after)
        self.assertGreaterEqual(len(comparison["changes"]), 2)
        self.assertEqual(1, len({item["regionId"] for item in comparison["changes"]}))

    def test_receipt_filename_wins_even_if_it_mentions_petition(self) -> None:
        petition = {"filename": "Petição final.docx", "attachmentId": "p"}
        receipt = {"filename": "Comprovante protocolo Petição Agravo.pdf", "attachmentId": "r"}
        selected, evidence, reason = _select_return_parts([petition, receipt])
        self.assertIsNone(reason)
        self.assertEqual(["p"], [item["attachmentId"] for item in selected])
        self.assertEqual(["r"], [item["attachmentId"] for item in evidence])

    def test_docx_track_changes_uses_accepted_view_and_counts_revision(self) -> None:
        path = self.root / "tracked.docx"
        write_docx(path, ["TEXTO ANTIGO"])
        replacement = self.root / "tracked-new.docx"
        with zipfile.ZipFile(path) as source, zipfile.ZipFile(replacement, "w") as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    text = text.replace(
                        "<w:r><w:t>TEXTO ANTIGO</w:t></w:r>",
                        '<w:del w:id="1"><w:r><w:delText>TEXTO ANTIGO</w:delText></w:r></w:del>'
                        '<w:ins w:id="2"><w:r><w:t>TEXTO FINAL</w:t></w:r></w:ins>',
                    )
                    data = text.encode("utf-8")
                target.writestr(item, data)
        extracted = extract_document(replacement)
        self.assertIn("TEXTO FINAL", extracted.visible_text)
        self.assertNotIn("TEXTO ANTIGO", extracted.visible_text)
        self.assertGreater(extracted.structural["trackChanges"]["insertions"], 0)
        self.assertGreater(extracted.structural["trackChanges"]["deletions"], 0)

    def test_docx_pdf_comparison_suppresses_protocol_noise_and_layout_conclusions(self) -> None:
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF indisponível")
        baseline = self.root / "baseline.docx"
        write_docx(baseline, ["TÍTULO", "A parte requer o provimento."])
        stamped = self.root / "stamped.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "PODER JUDICIÁRIO")
        page.insert_text((72, 100), "Protocolo realizado 02/07/2026")
        page.insert_text((72, 140), "TÍTULO")
        page.insert_text((72, 170), "A parte requer o provimento.")
        document.save(stamped)
        document.close()
        comparison = compare_documents(baseline, stamped)
        self.assertTrue(comparison["comparisonPolicy"]["crossFormat"])
        self.assertTrue(comparison["comparisonPolicy"]["protocolNoiseSuppressed"])
        self.assertFalse(comparison["comparisonPolicy"]["layoutConclusionAllowed"])


class VaultIsolationTests(unittest.TestCase):
    def test_vault_and_full_diff_paths_are_ignored_and_untracked(self) -> None:
        candidates = [
            (
                REPO
                / "_FORJA_HARNESS"
                / "state"
                / "case-test"
                / "PEÇA PROTOCOLADA — TESTE"
                / "COMPARAÇÃO_PRIVADA_IA_VS_HUMANO.json"
            ),
            REPO / "_FORJA_HARNESS" / "reports" / "POST_PROTOCOL_LAST_RUN.json",
        ]
        for candidate in candidates:
            with self.subTest(candidate=candidate):
                ignored = subprocess.run(
                    ["git", "check-ignore", "--no-index", "-q", str(candidate)],
                    cwd=REPO,
                    check=False,
                )
                self.assertEqual(0, ignored.returncode)
        tracked = subprocess.run(
            ["git", "ls-files", "--", "_FORJA_HARNESS/state"],
            cwd=REPO,
            capture_output=True,
            text=True,
            check=True,
        )
        forbidden = [
            line for line in tracked.stdout.splitlines()
            if "/private/post_protocol/" in line
            or "/PEÇA PROTOCOLADA — " in line
            or "/VERSÃO HUMANA FINAL — " in line
        ]
        self.assertEqual([], forbidden)


class GmailBoundaryTests(unittest.TestCase):
    def test_sender_allowlist_is_checked_from_configuration(self) -> None:
        message = {
            "payload": {
                "headers": [{"name": "From", "value": "Advogado <advogado@medinaosorio.adv.br>"}]
            }
        }
        with patch(
            "forja_post_protocol.load_config",
            return_value={
                "postProtocol": {
                    "gmail": {
                        "allowedDomains": ["medinaosorio.adv.br"],
                        "allowedAddresses": [],
                    }
                }
            },
        ):
            self.assertTrue(_sender_allowed(message))
            message["payload"]["headers"][0]["value"] = "Pessoa <fora@example.com>"
            self.assertFalse(_sender_allowed(message))

    def test_duplicate_demand_to_case_mapping_fails_closed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            data = root / "data"
            state = root / "state"
            data.mkdir()
            state.mkdir()
            (data / "forja_case_links.json").write_text(
                json.dumps({"links": {"case-a": "demand-x", "case-b": "demand-x"}}),
                encoding="utf-8",
            )
            (state / "case-a").mkdir()
            (state / "case-b").mkdir()
            with patch("forja_post_protocol.OFFICE_DATA", data), patch("forja_post_protocol.FORJA", root):
                self.assertIsNone(_case_for_demand("demand-x", []))


class LearningRegistryTests(unittest.TestCase):
    def test_devolved_chapters_fixture_is_machine_readable(self) -> None:
        fixture = json.loads(
            (
                ROOT
                / "n4_fixtures"
                / "post_protocol"
                / "devolved_chapters_memoriais_apelacao.json"
            ).read_text(encoding="utf-8")
        )
        self.assertEqual("memoriais de apelação", fixture["productType"])
        self.assertEqual(
            ["all_devolved_chapters", "ancillary_requests"],
            fixture["requiredCoverageKeys"],
        )
        self.assertEqual("prospective_case_suite", fixture["expectedGate"])

    def test_prospective_fixture_carries_promoted_lesson(self) -> None:
        fixture = json.loads(
            (
                ROOT
                / "n4_fixtures"
                / "post_protocol"
                / "prospective_memoriais_apelacao_suite.json"
            ).read_text(encoding="utf-8")
        )
        self.assertTrue(fixture["draftedBeforeFinalText"])
        self.assertEqual("prospective", fixture["executionMode"])
        self.assertEqual(
            "learn-b80a07dd026116a8",
            fixture["tests"][0]["learningCandidateId"],
        )
        self.assertIn("all_devolved_chapters", fixture["tests"][0]["coverageKeys"])
        self.assertIn("ancillary_requests", fixture["tests"][0]["coverageKeys"])

    def test_promoted_product_rule_is_required_in_prospective_suite(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            registry = root / "ACTIVE_RULES.json"
            case = root / "case-future"
            (case / "n4_artifacts").mkdir(parents=True)
            (case / "n4_artifacts" / "F2_N4_CLASSIFICATION.json").write_text(
                json.dumps({"product": "memoriais de apelação"}),
                encoding="utf-8",
            )
            fixture = root / "fixture.json"
            receipt = root / "receipt.json"
            fixture.write_text("{}", encoding="utf-8")
            receipt.write_text(json.dumps({"passed": True, "exitCode": 0}), encoding="utf-8")
            candidate = {
                "candidateId": "learn-test",
                "scope": "product_type",
                "layer": "request_relief",
                "cause": "reasoning",
                "impact": "material",
                "fixtureId": str(fixture),
                "testId": "test-id",
                "testExecutionPath": str(receipt),
                "approvedBy": "office-email:test",
                "promotedAt": "2026-07-29T01:00:00-03:00",
            }
            with patch("forja_learning_registry.REGISTRY_PATH", registry):
                register_promoted_rule(
                    source_case_id="case-source",
                    candidate=candidate,
                    scope_key="memoriais de apelação",
                )
                rules = active_rules(
                    case_id=case.name,
                    product_type="memoriais de apelação",
                )
                self.assertEqual(["learn-test"], [item["candidateId"] for item in rules])
                findings = suite_learning_findings(case, {"tests": []})
                self.assertEqual(["PP-LEARNING-NOT-APPLIED"], [item["code"] for item in findings])
                covered = suite_learning_findings(
                    case,
                    {"tests": [{"learningCandidateId": "learn-test"}]},
                )
                self.assertEqual([], covered)


if __name__ == "__main__":
    unittest.main(verbosity=2)
