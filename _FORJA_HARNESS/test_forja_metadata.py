from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

import fitz
from docx import Document

from forja_metadata import OFFICE, sanitize_final_artifacts


class MetadataTests(unittest.TestCase):
    def test_final_sanitization_preserves_content_and_normalizes_authors(self):
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            docx = root / "piece.docx"
            pdf = root / "piece.pdf"
            document = Document()
            document.add_paragraph("Conteúdo jurídico preservado.")
            document.core_properties.author = "Pessoa A"
            document.core_properties.last_modified_by = "Pessoa B"
            document.save(docx)
            with fitz.open() as output:
                page = output.new_page()
                page.insert_text((72, 72), "Conteudo juridico preservado.")
                output.set_metadata({"author": "Pessoa A", "title": "Peça"})
                output.save(pdf)
            with zipfile.ZipFile(docx) as before:
                document_xml = before.read("word/document.xml")
            sanitize_final_artifacts(docx, pdf)
            reopened = Document(docx)
            self.assertEqual(OFFICE, reopened.core_properties.author)
            self.assertEqual(OFFICE, reopened.core_properties.last_modified_by)
            self.assertIn("Conteúdo jurídico preservado.", "\n".join(p.text for p in reopened.paragraphs))
            with zipfile.ZipFile(docx) as after:
                self.assertEqual(document_xml, after.read("word/document.xml"))
            with fitz.open(pdf) as reopened_pdf:
                self.assertEqual(OFFICE, reopened_pdf.metadata.get("author"))
                self.assertIn("Conteudo juridico preservado.", "".join(page.get_text() for page in reopened_pdf))


if __name__ == "__main__":
    unittest.main()
