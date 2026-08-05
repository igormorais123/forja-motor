"""FORJA N2 - M4: peça piloto retroativa (F6 redação em template + F8 QA estrutural).

Piloto em modo sombra: reconstrói a peça final do caso cobaia SOBRE o template
oficial do escritório e valida o pipeline completo:
  template -> corpo real -> DOCX -> auditoria OOXML/estilo -> memória estrutural.

Artefatos SÓ em _FORJA_HARNESS/state/<caseId>/piloto/. Nada é alterado nas pastas
do caso, no painel ou no template.

Simplificação registrada (sem silent cap): tabelas/figuras da peça original NÃO são
transferidas neste piloto — o objetivo do M4 é validar o pipeline de template e a
QA estrutural; a rota de produção embute SVG nativo no OOXML.
"""

import json
import re
import shutil
import sys
from pathlib import Path

from forja_n3_common import now_iso

FORJA = Path(__file__).resolve().parent
RAIZ = FORJA.parent
sys.path.insert(0, str(RAIZ / "_FERRAMENTAS"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from forja_visual_qa_structural import auditar_documento

TEMPLATE = RAIZ / "_FERRAMENTAS" / "TEMPLATE_MEDINA_OSORIO_PETICAO.docx"
if not TEMPLATE.exists():
    raise SystemExit(f"TEMPLATE do escritório não encontrado em {TEMPLATE} — sem ele nenhuma peça pode ser gerada (timbre é arte do template).")

# Placeholders proibidos no produto final (gate G7.4); colchete de citação de
# dispositivo legal ("[sic]" etc.) não ocorre nas peças da casa.
PLACEHOLDER_RE = re.compile(r"\[[^\]\n]{1,60}\]")


def append_unique_many(existing, values):
    items = list(existing or [])
    for value in values:
        if value not in items:
            items.append(value)
    return items


def limpar_corpo(doc):
    """Remove todos os parágrafos e tabelas do corpo, preservando seção/timbre."""
    body = doc.element.body
    for el in list(body):
        if el.tag.endswith("}p") or el.tag.endswith("}tbl"):
            body.remove(el)


def eh_titulo(texto):
    t = texto.strip()
    if not t or len(t) > 120:
        return False
    if re.match(r"^(I{1,3}V?|IV|V?I{0,3}|\d+)\s*[-–—.]\s+\S", t) and t == t.upper():
        return True
    return t == t.upper() and len(t.split()) <= 12 and not t[0].isdigit()


def montar_piloto(case_key, peca_fonte):
    matches = list((FORJA / "state").glob(f"case-*{case_key}*/FORJA_STATE.json"))
    if not matches:
        raise SystemExit(f"estado nao encontrado para {case_key}")
    state_path = matches[0]
    state = json.loads(state_path.read_text(encoding="utf-8-sig"))
    piloto_dir = state_path.parent / "piloto"
    piloto_dir.mkdir(exist_ok=True)

    destino = piloto_dir / "PILOTO_M4_TEMPLATE.docx"
    shutil.copy2(TEMPLATE, destino)

    fonte = Document(str(peca_fonte))
    paragrafos_fonte = [p.text for p in fonte.paragraphs]

    doc = Document(str(destino))
    limpar_corpo(doc)
    inseridos = 0
    for texto in paragrafos_fonte:
        texto = texto.rstrip()
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        fmt = p.paragraph_format
        fmt.line_spacing = 1.5
        if eh_titulo(texto):
            run.bold = True
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER if inseridos == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Cm(0)
        else:
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Cm(2.0)
        inseridos += 1
    # metadados do escritório (gate G8.2 - nunca "python-docx")
    doc.core_properties.author = "Medina Osório Advogados"
    doc.core_properties.last_modified_by = "Medina Osório Advogados"
    doc.core_properties.title = "Piloto M4 FORJA - reconstrução sobre template"
    doc.save(str(destino))

    # gate G7.4: placeholders no texto final
    texto_final = "\n".join(p.text for p in Document(str(destino)).paragraphs)
    placeholders = sorted({m.group(0) for m in PLACEHOLDER_RE.finditer(texto_final)})

    qa = auditar_documento(destino)
    (piloto_dir / "F8_QA_ESTRUTURAL.json").write_text(
        json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")

    relatorio = piloto_dir / "F8_QA_PILOTO.md"
    linhas = [
        "# M4 — Piloto de redação em template + QA visual (modo sombra)",
        "",
        f"Caso cobaia: `{state['caseId']}` | Peça-fonte: `{Path(peca_fonte).name}` | Gerado: {now_iso()}",
        "",
        f"- Parágrafos transferidos: {inseridos}",
        "- Materialização: DOCX/OOXML auditado sem Word COM, PDF, PNG ou renderização",
        f"- Placeholders no texto final: {len(placeholders)}" + (f" — {placeholders}" if placeholders else " (zero)"),
        "- Metadados: autor = Medina Osório Advogados (gate G8.2)",
        f"- QA estrutural: {'aprovada' if qa['approved'] else 'pendências registradas'} em `F8_QA_ESTRUTURAL.json`",
        "- Simplificação declarada: tabelas e figuras da peça original não transferidas neste piloto.",
        "",
        "Abertura humana do DOCX continua obrigatória para decidir paginação e legibilidade.",
    ]
    relatorio.write_text("\n".join(linhas) + "\n", encoding="utf-8")

    state["updatedAt"] = now_iso()
    state["currentPhase"] = "F6_F8_PILOTO_M4"
    state.setdefault("phaseHistory", []).append(
        {"phase": "F6_F8_PILOTO_M4", "at": now_iso(),
         "status": "ok" if not placeholders else "pendencias"})
    state["artifacts"] = append_unique_many(state.get("artifacts") or [], [str(destino), str(piloto_dir / "F8_QA_ESTRUTURAL.json"), str(relatorio)])
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps({
        "docx": str(destino), "pdf": None, "paginas": None,
        "paragrafos": inseridos, "placeholders": placeholders,
        "qaEstrutural": str(piloto_dir / "F8_QA_ESTRUTURAL.json"),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    montar_piloto(sys.argv[1], sys.argv[2])
