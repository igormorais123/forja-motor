"""Gate estrutural de diagramação OOXML da FORJA.

A rota canônica não renderiza o DOCX. Este módulo lê o OOXML efetivo e falha
fechado quando o texto principal não está:

* justificado;
* em Times New Roman 12 pt;
* tipograficamente uniforme dentro de cada parágrafo.

Títulos, sumário, capa, assinaturas, legendas, tabelas e caixas com função
visual identificável são papéis distintos. Eles podem usar outra composição,
mas continuam sujeitos a legibilidade mínima e não são confundidos com corpo.
Exceções adicionais só valem quando nominadas, justificadas e vinculadas ao
hash do texto do parágrafo; uma lista genérica de índices não passa.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from forja_n3_common import atomic_write_json, now_iso, sha256_file


BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12.0
TABLE_MIN_SIZE_PT = 8.0
# A sans da casa. `medina_visual_kit` a define como `SANS = "Segoe UI"` e a usa
# em rótulo de tabela, legenda, linha de seção e citação recuada desde o padrão
# aprovado pelo dono em 09/07/2026. Medido em 04/08/2026 nas peças reais do
# CASO-04: a V4 (15/07) e a V8 (30/07) misturam Segoe UI e Times New Roman
# dentro das mesmas tabelas, e as duas foram entregues.
#
# Isto está aqui porque eu classifiquei essa mistura como defeito antes de
# conferir, e cheguei a gerar uma "correção" que na verdade desfazia a
# identidade da casa. É o teste-âncora de novo, na direção mais cara: gate que
# reprova o padrão aprovado está errado, não a peça.
HOUSE_SANS = "Segoe UI"
TABLE_FALLBACK_FONT = "Verdana"
TABLE_FALLBACK_SIZE_PT = 8.0
# Calibrado contra o TEMPLATE APROVADO da casa, medido em 04/08/2026: o shape do
# fólio em `_FERRAMENTAS\TEMPLATE_MEDINA_OSORIO_PETICAO.docx` tem 57,3 pt. O
# limiar anterior era 36,0 pt e não veio de medição nenhuma — resultado: o gate
# reprovava por "fólio inseguro" toda peça nascida do template do escritório,
# que é justamente o único jeito autorizado de nascer uma peça. Mesmo teste-âncora
# da síntese executiva: gate que reprova o padrão aprovado está errado, não a peça.
#
# A margem direita do padrão é 3,5 cm (99,2 pt), então 57,3 pt de fólio cabem nela
# com folga. O teto abaixo dá 4 pt de tolerância sobre o valor do template e
# continua reprovando o fólio que avança sobre a mancha de texto — que é o defeito
# que este gate existe para pegar.
FOLIO_SAFE_WIDTH_PT = 61.0
# Alvo da NORMALIZAÇÃO, que é coisa diferente do teto de tolerância. Consertar um
# fólio grande demais significa deixá-lo como o do template aprovado — 57,3 pt,
# medido —, e não como o teto. Normalizar para o teto produziria um documento no
# limite do aceitável em vez de um documento no padrão.
FOLIO_TEMPLATE_WIDTH_PT = 57.3
# Folga da checagem de colisão contra a margem da seção. A largura do shape é
# aproximação da área que ele ocupa — a posição horizontal exata depende de
# âncora e deslocamento —, então cobrar o encaixe ao centésimo transforma meio
# ponto em P0. Dois pontos são 0,07 cm: invisível na página impressa, e acima
# disso o fólio realmente entra na mancha de texto.
FOLIO_MARGIN_SLACK_PT = 2.0
W_NS = qn("w:document").split("}")[0][1:]
BODY_STYLE_HINTS = {
    "normal",
    "body text",
    "bodytext",
    "first paragraph",
    "corpo de texto",
    "texto de corpo",
}
EXCEPTION_ROLES = {
    "cover_metadata",
    "epigraph",
    "section_lede",
    "callout",
    "caption",
    "signature",
}
# Estes papéis podem usar o tamanho deliberado da estrutura aprovada (por
# exemplo, 10,5 pt), mas não escapam da família tipográfica nem da justificação
# exigidas para texto corrido. A exceção é de tamanho, não uma porta dos fundos
# para ignorar o restante da régua.
STRUCTURAL_SIZE_EXCEPTION_ROLES = {
    "executive_summary",
    "qualification",
    "block_quote",
}


def _norm(value: str | None) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().casefold()


def _text_sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def docx_content_signature(path: Path) -> dict:
    """Assina o conteúdo OOXML que uma correção puramente visual deve preservar.

    O hash ignora propriedades de formatação, mas inclui, na ordem original,
    texto visível, texto excluído, instruções de campos e os contadores de
    inserções/exclusões controladas. Assim a IA não pode usar a normalização
    estética para trocar uma palavra, suprimir uma ressalva ou achatar o
    controle de alterações sem derrubar o gate.
    """
    path = Path(path)
    parts: dict[str, dict] = {}
    tracked_insertions = tracked_deletions = deleted_text_nodes = 0
    substantive_insertions = substantive_deletions = 0
    all_inserted_text: list[str] = []
    all_deleted_text: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
        for name in names:
            try:
                root = etree.fromstring(archive.read(name))
            except etree.XMLSyntaxError:
                continue
            # O Word pode dividir ou fundir runs ao salvar, sem mudar uma letra.
            # Por isso o hash respeita limites de parágrafo, não limites de run.
            paragraphs = []
            for paragraph in root.xpath(".//w:p", namespaces={"w": W_NS}):
                value = "".join(
                    node.text or ""
                    for node in paragraph.xpath(
                        ".//w:t | .//w:delText | .//w:instrText",
                        namespaces={"w": W_NS},
                    )
                )
                # Parágrafos vazios pertencem à diagramação e podem ser criados
                # pelo próprio Word ao materializar cabeçalhos. O gate visual os
                # cobre; eles não integram a assinatura do conteúdo jurídico.
                if value:
                    paragraphs.append(value)
                deleted_text_nodes += len(paragraph.xpath(".//w:delText", namespaces={"w": W_NS}))

            inserted_nodes = root.xpath(".//w:ins | .//w:moveTo", namespaces={"w": W_NS})
            deleted_nodes = root.xpath(".//w:del | .//w:moveFrom", namespaces={"w": W_NS})
            inserted_values = [
                "".join(node.text or "" for node in change.xpath(".//w:t | .//w:delText", namespaces={"w": W_NS}))
                for change in inserted_nodes
            ]
            deleted_values = [
                "".join(node.text or "" for node in change.xpath(".//w:t | .//w:delText", namespaces={"w": W_NS}))
                for change in deleted_nodes
            ]
            inserted_text = "".join(inserted_values)
            deleted_text = "".join(deleted_values)
            all_inserted_text.append(inserted_text)
            all_deleted_text.append(deleted_text)
            inserted = len(inserted_nodes)
            deleted = len(deleted_nodes)
            tracked_insertions += inserted
            tracked_deletions += deleted
            substantive_insertions += sum(bool(value) for value in inserted_values)
            substantive_deletions += sum(bool(value) for value in deleted_values)
            if paragraphs or inserted or deleted:
                canonical = json.dumps(paragraphs, ensure_ascii=False, separators=(",", ":"))
                parts[name] = {
                    "paragraphCount": len(paragraphs),
                    "textSha256": _text_sha256(canonical),
                    "trackedInsertions": inserted,
                    "trackedDeletions": deleted,
                    "insertedTextChars": len(inserted_text),
                    "insertedTextSha256": _text_sha256(inserted_text),
                    "deletedTextChars": len(deleted_text),
                    "deletedTextSha256": _text_sha256(deleted_text),
                }
    protected_parts = {
        name: {
            "paragraphCount": item["paragraphCount"],
            "textSha256": item["textSha256"],
            "insertedTextChars": item["insertedTextChars"],
            "insertedTextSha256": item["insertedTextSha256"],
            "deletedTextChars": item["deletedTextChars"],
            "deletedTextSha256": item["deletedTextSha256"],
        }
        for name, item in parts.items()
    }
    canonical_parts = json.dumps(protected_parts, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    inserted_text = "".join(all_inserted_text)
    deleted_text = "".join(all_deleted_text)
    return {
        "schemaVersion": 1,
        "contentSha256": _text_sha256(canonical_parts),
        "trackedInsertions": tracked_insertions,
        "trackedDeletions": tracked_deletions,
        "substantiveInsertions": substantive_insertions,
        "substantiveDeletions": substantive_deletions,
        "insertedTextChars": len(inserted_text),
        "insertedTextSha256": _text_sha256(inserted_text),
        "deletedTextChars": len(deleted_text),
        "deletedTextSha256": _text_sha256(deleted_text),
        "deletedTextNodes": deleted_text_nodes,
        "parts": parts,
    }


def compare_docx_content(source: Path, output: Path) -> dict:
    before = docx_content_signature(source)
    after = docx_content_signature(output)
    fields = (
        "contentSha256",
        "insertedTextChars",
        "insertedTextSha256",
        "deletedTextChars",
        "deletedTextSha256",
    )
    mismatches = [field for field in fields if before.get(field) != after.get(field)]
    return {
        "approved": not mismatches,
        "source": {"path": str(Path(source)), "sha256": sha256_file(Path(source)), **before},
        "output": {"path": str(Path(output)), "sha256": sha256_file(Path(output)), **after},
        "mismatches": mismatches,
    }


def _paragraph_text(paragraph) -> str:
    """Inclui texto em w:ins/w:del, que python-docx não expõe em ``text``."""
    values = paragraph._p.xpath(".//w:t | .//w:delText")
    return "".join(node.text or "" for node in values).strip()


def _unique_header_parts(document):
    """Percorre cabeçalhos reais uma única vez, inclusive os vinculados."""
    seen: set[int] = set()
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            part = header.part
            marker = id(part)
            if marker in seen:
                continue
            seen.add(marker)
            yield part


def _right_margin_pt(section) -> tuple[float | None, bool]:
    """Lê a margem direita sem transformar erro de leitura em passe.

    ``python-docx`` pode levantar ``ValueError`` quando o DOCX traz twips
    fracionários ou inválidos. Valores fracionários válidos ainda são medidas
    legítimas e podem ser lidos do atributo OOXML bruto; o que continuar
    ilegível precisa chegar ao laudo como falha fechada.
    """
    try:
        margem = section.right_margin
        if margem is not None:
            try:
                return float(margem.pt), False
            except (AttributeError, TypeError, ValueError):
                pass
    except (AttributeError, TypeError, ValueError):
        pass

    sect_pr = getattr(section, "_sectPr", None)
    pg_mar = sect_pr.find(qn("w:pgMar")) if sect_pr is not None else None
    raw = pg_mar.get(qn("w:right")) if pg_mar is not None else None
    if raw is not None:
        try:
            # w:right é medido em twips (1/20 de ponto), inclusive quando o
            # valor não é inteiro e o wrapper python-docx não o aceita.
            return float(raw) / 20.0, False
        except (TypeError, ValueError):
            pass
    return None, True


def _folios_com_margem(document):
    """Cada fólio pareado com a margem direita da SEÇÃO que o hospeda.

    O fólio mora dentro da margem direita; mais largo que ela, invade a mancha de
    texto. Essa é a colisão que o nome do achado promete — e não o teto estético
    de ``FOLIO_SAFE_WIDTH_PT``, que o acervo inteiro respeita.

    A margem é a da seção dona do cabeçalho, e não a menor do documento. Comparar
    contra a menor seção reprovaria peça correta cujo fólio está numa seção larga
    enquanto outra seção — capa de imagem, página de foto — tem margem zero. Esse
    é exatamente o falso positivo que já custou uma recalibração hoje. Quando o
    mesmo cabeçalho é compartilhado por seções de margens diferentes, vale a mais
    apertada, porque é onde a colisão de fato aconteceria. Se alguma margem do
    cabeçalho compartilhado for ilegível, o retorno marca a margem como não
    resolvida para que o chamador emita P0, em vez de abrir um bypass.
    """
    por_parte: dict[int, tuple] = {}
    for section in document.sections:
        margem, margem_nao_resolvida = _right_margin_pt(section)
        # Seção de margem direita ZERO não tem banda de margem para o fólio
        # ocupar, e comparar contra ela é comparar contra nada. Medido em
        # 04/08/2026: três documentos do acervo — a capa fotográfica do parecer
        # CASO-17 e dois relatórios autorais — têm a seção de abertura com margem
        # zero e o cabeçalho VINCULADO à seção seguinte, de margem normal. Como a
        # regra do compartilhamento é "vale a mais apertada", o zero da capa
        # arrastava o fólio da peça inteira para reprovação. É o falso positivo
        # que o próprio comentário acima antecipava e que, ainda assim, aconteceu.
        # Margem zero vira ausência de base de comparação, não colisão. Que a
        # seção esteja fora do padrão de 3,5 cm é outro assunto, de outro achado.
        if margem is not None and margem <= 0:
            margem = None
        for header in (section.header, section.first_page_header, section.even_page_header):
            part = header.part
            marca = id(part)
            anterior = por_parte.get(marca)
            if anterior is None:
                por_parte[marca] = (part, margem, margem_nao_resolvida)
            elif margem_nao_resolvida or anterior[2]:
                por_parte[marca] = (part, None, True)
            elif margem is not None and (anterior[1] is None or margem < anterior[1]):
                por_parte[marca] = (part, margem, False)

    for part, margem, margem_nao_resolvida in por_parte.values():
        for rect in part.element.xpath(".//*[local-name()='rect']"):
            instrucoes = " ".join(
                node.text or "" for node in rect.xpath(".//*[local-name()='instrText']")
            ).upper()
            if re.search(r"\bPAGE\b", instrucoes):
                yield rect, margem, margem_nao_resolvida


def _folio_rectangles(document):
    """Localiza o fólio VML pelo campo PAGE, sem confiar no nome da forma."""
    for part in _unique_header_parts(document):
        for rect in part.element.xpath(".//*[local-name()='rect']"):
            instructions = " ".join(
                node.text or "" for node in rect.xpath(".//*[local-name()='instrText']")
            ).upper()
            if re.search(r"\bPAGE\b", instructions):
                yield rect


def _vml_width_pt(rect) -> float | None:
    style = str(rect.get("style") or "")
    match = re.search(r"(?:^|;)\s*width\s*:\s*([0-9]+(?:\.[0-9]+)?)pt(?:;|$)", style, re.I)
    return float(match.group(1)) if match else None


def _set_vml_width_pt(rect, width_pt: float) -> None:
    style = str(rect.get("style") or "")
    replacement = f"width:{width_pt:g}pt"
    if re.search(r"(?:^|;)\s*width\s*:\s*[0-9]+(?:\.[0-9]+)?pt(?:;|$)", style, re.I):
        style = re.sub(
            r"(?:(?<=;)|^)\s*width\s*:\s*[0-9]+(?:\.[0-9]+)?pt(?=;|$)",
            replacement,
            style,
            count=1,
            flags=re.I,
        )
    else:
        style = f"{style.rstrip(';')};{replacement}" if style else replacement
    rect.set("style", style)


def _style_chain(style):
    seen: set[str] = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        yield style
        style = style.base_style


def _effective_alignment(paragraph) -> int:
    if paragraph.alignment is not None:
        return int(paragraph.alignment)
    for style in _style_chain(paragraph.style):
        if style.paragraph_format.alignment is not None:
            return int(style.paragraph_format.alignment)
    # Ausência de w:jc no Word significa alinhamento à esquerda.
    return int(WD_ALIGN_PARAGRAPH.LEFT)


def _font_from_rpr(rpr) -> str | None:
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = rfonts.get(qn(f"w:{key}"))
        if value and not value.startswith("+"):
            return value
    return None


def _size_from_rpr(rpr) -> float | None:
    if rpr is None:
        return None
    node = rpr.find(qn("w:sz"))
    if node is None:
        return None
    try:
        return int(node.get(qn("w:val"))) / 2
    except (TypeError, ValueError):
        return None


def _style_font_size(style) -> tuple[str | None, float | None]:
    font_name = None
    size = None
    for current in _style_chain(style):
        if font_name is None:
            font_name = current.font.name or _font_from_rpr(current.element.rPr)
        if size is None:
            size = current.font.size.pt if current.font.size else _size_from_rpr(current.element.rPr)
        if font_name is not None and size is not None:
            break
    return font_name, size


def _effective_run_font_size(run_element, paragraph) -> tuple[str | None, float | None]:
    rpr = run_element.find(qn("w:rPr"))
    font_name = _font_from_rpr(rpr)
    size = _size_from_rpr(rpr)
    style_font, style_size = _style_font_size(paragraph.style)
    return font_name or style_font, size if size is not None else style_size


def _substantial_run_text(run_element) -> str:
    if run_element.xpath(".//w:footnoteReference | .//w:endnoteReference | .//w:fldChar"):
        return ""
    text = "".join(node.text or "" for node in run_element.xpath(".//w:t | .//w:delText"))
    # Marcadores isolados e pontuação não definem a tipografia do corpo.
    return text if len(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]", text)) >= 2 else ""


def _has_visual_container(paragraph) -> bool:
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return False
    # `w:framePr` é a moldura de texto. O `pull()` do kit aprovado emite o
    # destaque dentro de uma moldura, com borda superior terracota e alinhamento
    # à esquerda deliberado — e sem reconhecê-la o parágrafo caía em citação
    # recuada e era cobrado justificado, contra o desenho da casa.
    return (ppr.find(qn("w:pBdr")) is not None
            or ppr.find(qn("w:shd")) is not None
            or ppr.find(qn("w:framePr")) is not None)


def _is_heading(paragraph, text: str) -> bool:
    style_name = _norm(paragraph.style.name)
    if any(token in style_name for token in ("heading", "titulo", "título", "title", "subtitle")):
        return True
    if style_name.startswith("toc") or "sumário" in style_name or "sumario" in style_name:
        return True
    if paragraph._p.xpath("./w:pPr/w:outlineLvl"):
        return True
    letters = [char for char in text if char.isalpha()]
    upper_ratio = sum(char.isupper() for char in letters) / len(letters) if letters else 0
    heading_prefix = re.match(r"^(?:[IVXLCDM]+(?:[-–—][A-Z])?|\d+(?:\.\d+)+)\s*(?:[-–—]|\.)", text)
    return bool(len(text) <= 220 and (upper_ratio >= 0.82 or heading_prefix))


def _is_caption(text: str) -> bool:
    return bool(len(text) <= 260 and re.match(r"^(?:FIGURA|QUADRO|TABELA|GRÁFICO)\s+\d+\b", text, re.I))


def _is_signature(text: str) -> bool:
    return bool(
        re.match(r"^(?:Brasília|São Paulo|Porto Alegre|Rio de Janeiro|Palmas)(?:[/,]|\b).*(?:19|20)\d{2}\.?$", text, re.I)
        or re.match(r"^(?:OAB[/\s]|Parecerista$|Advogad[oa]s?$)", text, re.I)
        or text in {"Nestes termos, pede deferimento.", "Termos em que, pedem deferimento."}
    )


# Dois elementos ESTRUTURAIS do padrão aprovado da casa que vinham sendo lidos
# como corpo fora de tamanho. Medido em 04/08/2026 contra quatro peças reais e
# entregues: `body_font_size_not_12pt` reprovou o bloco de qualificação das
# partes na V8 do CASO-04 e cinco parágrafos da síntese executiva nos memoriais
# do AI 0011621-15 — todos a 10,5 pt, todos deliberados.
#
# A síntese executiva no estilo do art. 343-A do RISTJ é obrigatória em TODA
# peça desde 07/07/2026, por determinação do Prof. Fábio. Um gate que a reprova
# reprova a regra da casa, e o teste-âncora manda concluir que o errado é o
# gate, não a peça. Se o F8-S tivesse sido ligado antes desta medição, barraria
# o próprio formato do escritório em todas as peças.
#
# A exceção é estreita de propósito: vale para o TAMANHO, e não para a família
# tipográfica nem para a justificação, que continuam cobradas nesses blocos.
_SINTESE_EXECUTIVA = re.compile(
    r"^\s*(?:[IVX]+\s*[.\-–]\s*)?(?:CASO EM EXAME|QUEST(?:ÃO|AO|ÕES|OES) EM DISCUSS(?:ÃO|AO)|"
    r"RAZ(?:ÕES|OES) DE DECIDIR|DISPOSITIVO E TESE|DISPOSITIVO|TESE DE JULGAMENTO|"
    r"S[ÍI]NTESE EXECUTIVA|EMENTA)\b", re.I)
# A abertura de ementa/síntese vem em caixa alta com o ramo do direito.
_ABERTURA_DE_EMENTA = re.compile(
    r"^[A-ZÀ-Ú][A-ZÀ-Ú\s,]{18,}\.\s+[A-ZÀ-Ú]", re.U)
_CITACAO_RECUADA = re.compile(r"^\s*[“\"«]")
_QUALIFICACAO = re.compile(
    r"j[áa]\s+qualificad|qualificad[oa]s?\s+n(?:os|a)\s+(?:autos|ep[íi]grafe)|"
    r"por\s+se(?:u|us)\s+advogad|vem,?\s+respeitosamente|"
    r"na\s+condi[çc][ãa]o\s+de\s+(?:AGRAVAD|RECORRID|EMBARGANT|AGRAVANT|RECORRENT)", re.I)


# Rótulos de capa. Os quatro primeiros vêm do parecer clássico; os demais são os
# que a própria esteira emite no topo dos produtos internos — "Status:",
# "Trilha:", "Fase:", "Classificação:" —, e sem eles a linha de identificação do
# documento era cobrada como texto corrido a 12 pt.
_ROTULO_DE_CAPA = re.compile(
    r"^(?:Consulente|Parecerista|Objeto|Data-base|Natureza do documento|Processo|"
    r"Interessado|Ementa|Status|Trilha|Fase|Classifica[çc][ãa]o|Destinat[áa]rio|"
    r"Origem|Escopo(?: desta entrega)?)\s*:", re.I)


def _e_titulo_centralizado(paragraph, text: str) -> bool:
    """Título de abertura em caixa mista: centralizado, maior que o corpo."""
    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return False
    if re.search(r"[.!?]\s*$", text) or len(text) > 160:
        return False
    tamanhos = {
        float(run.font.size.pt) for run in paragraph.runs if run.font.size is not None
    }
    return bool(tamanhos) and min(tamanhos) > BODY_SIZE_PT


def _role_for(paragraph, text: str, index: int, previous_role: str | None) -> str:
    style_name = _norm(paragraph.style.name)
    # Rótulos isolados como "EMENTA" ou "DISPOSITIVO" são cabeçalhos,
    # ainda que a linha seguinte inaugure a síntese executiva. Não se pode
    # cobrar justificação de um rótulo centralizado só porque ele introduz
    # uma estrutura que, no texto corrido, usa a exceção de tamanho.
    if _is_heading(paragraph, text) and _SINTESE_EXECUTIVA.fullmatch(text.strip()):
        return "heading"
    if _SINTESE_EXECUTIVA.match(text) or (index <= 12 and _ABERTURA_DE_EMENTA.match(text)):
        return "executive_summary"
    if previous_role == "executive_summary" and not _is_heading(paragraph, text):
        return "executive_summary"
    if index <= 8 and _QUALIFICACAO.search(text):
        return "qualification"
    # Citação recuada: menor que o corpo por convenção do texto jurídico e pelo
    # padrão da casa. Apareceu como terceiro falso P0 nos memoriais do AI
    # 0011621-15, num trecho transcrito do próprio acórdão impugnado.
    # A ESTRUTURA vem antes da heurística de pontuação. Um parágrafo dentro de
    # moldura ou caixa é destaque visual por construção; um parágrafo que começa
    # com aspas é citação por palpite. Enquanto o palpite era testado primeiro, o
    # pull quote do kit aprovado — moldura, aspas, texto longo — era lido como
    # citação recuada e cobrado justificado, contra o desenho da casa.
    if _has_visual_container(paragraph):
        return "callout"
    if _CITACAO_RECUADA.match(text) and len(text) >= 120:
        return "block_quote"
    if style_name.startswith("toc") or "sumário" in style_name or "sumario" in style_name:
        return "toc"
    if _is_heading(paragraph, text):
        return "heading"
    if _is_caption(text):
        return "caption"
    if _is_signature(text):
        return "signature"
    if index <= 12 and _ROTULO_DE_CAPA.match(text):
        return "cover_metadata"
    # Título do documento em CAIXA MISTA. O reconhecedor de cabeçalho pega o
    # título em versal — "IMPUGNAÇÃO AO AGRAVO INTERNO" —, e por isso a V8
    # aprovada nunca reprovou. Mas a rota canônica emite o título da peça em
    # caixa mista e centralizado, e ele caía em "corpo": foi reprovado por não
    # estar a 12 pt, sendo que título maior que o corpo é justamente o padrão.
    # O recorte é estreito de propósito — abertura do documento, centralizado,
    # maior que o corpo e sem pontuação final de frase —, para que nenhum
    # parágrafo de texto corrido escape por aqui.
    if index <= 3 and _e_titulo_centralizado(paragraph, text):
        return "heading"
    if _has_visual_container(paragraph):
        return "callout"
    if previous_role == "heading" and len(text) <= 240 and not re.match(r"^\d+[.-]", text):
        return "section_lede"
    if style_name in BODY_STYLE_HINTS or len(text) >= 45 or re.search(r"[.!?;:]\s*$", text):
        return "body"
    return "other"


def _load_exceptions(path: Path | None, docx_sha256: str) -> tuple[dict[int, dict], list[dict]]:
    if path is None:
        return {}, []
    try:
        payload = json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        return {}, [{"severity": "P0", "code": "layout_exceptions_invalid", "detail": str(exc)}]
    findings = []
    if payload.get("docxSha256") != docx_sha256:
        findings.append({"severity": "P0", "code": "layout_exceptions_wrong_docx"})
    entries = payload.get("exceptions") or []
    if not isinstance(entries, list):
        findings.append({"severity": "P0", "code": "layout_exceptions_not_list"})
        return {}, findings
    result: dict[int, dict] = {}
    for entry in entries:
        try:
            idx = int(entry.get("paragraphIndex"))
        except (TypeError, ValueError, AttributeError):
            findings.append({"severity": "P0", "code": "layout_exception_without_index"})
            continue
        role = str(entry.get("role") or "")
        reason = str(entry.get("reason") or "").strip()
        if role not in EXCEPTION_ROLES or len(reason) < 20 or not entry.get("textSha256"):
            findings.append({"severity": "P0", "code": "layout_exception_not_substantiated", "paragraph": idx})
            continue
        if idx in result:
            findings.append({"severity": "P0", "code": "layout_exception_duplicate", "paragraph": idx})
            continue
        result[idx] = entry
    return result, findings


def audit_docx_layout(path: Path, *, exceptions: Path | None = None) -> dict:
    path = Path(path)
    document = Document(str(path))
    docx_sha = sha256_file(path)
    declared, findings = _load_exceptions(exceptions, docx_sha)
    records = []
    issue_buckets: dict[str, list[dict]] = {
        "body_text_not_justified": [],
        "body_font_not_medina": [],
        "body_font_size_not_12pt": [],
        "body_typography_unresolved": [],
        "structural_text_not_justified": [],
        "structural_font_not_medina": [],
        "structural_typography_unresolved": [],
        "table_font_below_minimum": [],
        "table_typography_inconsistent": [],
        "table_typography_unresolved": [],
        "folio_width_unsafe": [],
        "folio_overflows_right_margin": [],
        "folio_width_unresolved": [],
        "folio_margin_unresolved": [],
    }

    previous_role = None
    body_total = body_justified = body_font_ok = body_size_ok = 0
    for index, paragraph in enumerate(document.paragraphs):
        text = _paragraph_text(paragraph)
        if not text:
            continue
        role = _role_for(paragraph, text, index, previous_role)
        exception = declared.get(index)
        if exception:
            if exception.get("textSha256") != _text_sha256(text):
                findings.append({"severity": "P0", "code": "layout_exception_text_mismatch", "paragraph": index})
            else:
                role = str(exception["role"])
        alignment = _effective_alignment(paragraph)
        run_values = []
        for run_element in paragraph._p.xpath(".//w:r"):
            run_text = _substantial_run_text(run_element)
            if not run_text:
                continue
            font_name, size = _effective_run_font_size(run_element, paragraph)
            run_values.append({"font": font_name, "size": size, "chars": len(run_text)})

        record = {
            "paragraph": index,
            "role": role,
            "style": paragraph.style.name,
            "alignment": alignment,
            "textSha256": _text_sha256(text),
            "sample": text[:180],
            "fonts": sorted({str(item["font"]) for item in run_values}),
            "sizesPt": sorted({item["size"] for item in run_values if item["size"] is not None}),
            "declaredException": bool(exception),
        }
        records.append(record)
        if role == "body" or role in STRUCTURAL_SIZE_EXCEPTION_ROLES:
            is_body = role == "body"
            if is_body:
                body_total += 1
            sample = {"paragraph": index, "style": paragraph.style.name, "sample": text[:180]}
            if alignment == int(WD_ALIGN_PARAGRAPH.JUSTIFY):
                if is_body:
                    body_justified += 1
            else:
                bucket = "body_text_not_justified" if is_body else "structural_text_not_justified"
                issue_buckets[bucket].append({**sample, "alignment": alignment, "role": role})
            fonts = {_norm(item["font"]) for item in run_values if item["font"]}
            unresolved_font = any(item["font"] is None for item in run_values) or not run_values
            # O corpo continua sendo Times New Roman e nada mais. Os papéis
            # estruturais — citação recuada, síntese executiva, qualificação —
            # podem usar a sans da casa, que é o que o kit aprovado emite.
            familia_aceita = ({_norm(BODY_FONT)} if is_body
                              else {_norm(BODY_FONT), _norm(HOUSE_SANS)})
            if not unresolved_font and fonts <= familia_aceita and fonts:
                if is_body:
                    body_font_ok += 1
            elif unresolved_font:
                bucket = "body_typography_unresolved" if is_body else "structural_typography_unresolved"
                issue_buckets[bucket].append({**sample, "property": "font", "role": role})
            else:
                bucket = "body_font_not_medina" if is_body else "structural_font_not_medina"
                issue_buckets[bucket].append({**sample, "fonts": sorted(fonts), "role": role})
            sizes = {round(float(item["size"]), 2) for item in run_values if item["size"] is not None}
            unresolved_size = any(item["size"] is None for item in run_values) or not run_values
            if is_body:
                if not unresolved_size and sizes == {BODY_SIZE_PT}:
                    body_size_ok += 1
                elif unresolved_size:
                    issue_buckets["body_typography_unresolved"].append({**sample, "property": "size"})
                else:
                    issue_buckets["body_font_size_not_12pt"].append({**sample, "sizesPt": sorted(sizes)})
            elif unresolved_size:
                # O valor numérico é exceção nesses papéis, mas a ausência de
                # medição ainda é uma falha fechada: não se pode confundir
                # "tamanho deliberadamente diferente" com "tamanho desconhecido".
                issue_buckets["structural_typography_unresolved"].append({
                    **sample, "property": "size", "role": role,
                })
        previous_role = role

    for table_index, table in enumerate(document.tables):
        table_groups = {"header": [], "body": []}
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    text = _paragraph_text(paragraph)
                    if not text:
                        continue
                    for run_element in paragraph._p.xpath(".//w:r"):
                        run_text = _substantial_run_text(run_element)
                        if not run_text:
                            continue
                        font, size = _effective_run_font_size(run_element, paragraph)
                        table_groups["header" if row_index == 0 else "body"].append({
                            "font": font,
                            "size": size,
                            "chars": len(run_text),
                        })
                        if size is not None and size < TABLE_MIN_SIZE_PT:
                            issue_buckets["table_font_below_minimum"].append({
                                "table": table_index,
                                "row": row_index,
                                "cell": cell_index,
                                "paragraph": paragraph_index,
                                "sizePt": round(size, 2),
                                "sample": text[:140],
                            })
                            break
        for group, values in table_groups.items():
            if not values:
                continue
            fonts = {_norm(item["font"]) for item in values if item["font"]}
            sizes = {round(float(item["size"]), 2) for item in values if item["size"] is not None}
            unresolved = any(item["font"] is None or item["size"] is None for item in values)
            if unresolved:
                issue_buckets["table_typography_unresolved"].append({"table": table_index, "group": group})
            # A mistura Times + sans da casa dentro de uma tabela é o padrão
            # aprovado, não descuido: rótulo em sans, conteúdo em serifada.
            # Medido nas V4 e V8 do CASO-04, ambas entregues. O achado existe
            # para pegar tabela colada de OUTRA fonte — Calibri, Arial, Verdana —,
            # e essas continuam reprovando.
            familia_da_casa = {_norm(BODY_FONT), _norm(HOUSE_SANS)}
            fonte_estranha = bool(fonts - familia_da_casa)
            # Uma tabela de família única com tamanhos diferentes continua
            # suspeita: ali não há a hierarquia rótulo/conteúdo que justifica a
            # variação.
            tamanho_solto = len(fonts) == 1 and len(sizes) > 1
            if fonte_estranha or tamanho_solto:
                issue_buckets["table_typography_inconsistent"].append({
                    "table": table_index,
                    "group": group,
                    "fonts": sorted(fonts),
                    "sizesPt": sorted(sizes),
                })

    folios = list(_folios_com_margem(document))
    for folio_index, (rect, margem_direita, margem_nao_resolvida) in enumerate(folios):
        if margem_nao_resolvida:
            issue_buckets["folio_margin_unresolved"].append({"folio": folio_index})
        width = _vml_width_pt(rect)
        if width is None:
            issue_buckets["folio_width_unresolved"].append({"folio": folio_index})
            continue
        if width > FOLIO_SAFE_WIDTH_PT + 0.01:
            issue_buckets["folio_width_unsafe"].append({
                "folio": folio_index,
                "widthPt": round(width, 2),
                "maximumPt": FOLIO_SAFE_WIDTH_PT,
            })
        # Teto absoluto e margem do documento medem coisas diferentes, e só a
        # segunda é colisão de verdade. Medido em 04/08/2026 sobre 287 DOCX com
        # fólio: as larguras do acervo são 28,0 / 36,0 / 57,3 pt e NENHUMA passa
        # do teto — ou seja, sozinho ele nunca dispararia. Num documento de
        # margem estreita, porém, o fólio do template invade o texto, e é esse
        # caso que o nome do achado promete pegar.
        if not margem_nao_resolvida and margem_direita is not None and width > margem_direita + FOLIO_MARGIN_SLACK_PT:
            issue_buckets["folio_overflows_right_margin"].append({
                "folio": folio_index,
                "widthPt": round(width, 2),
                "rightMarginPt": round(margem_direita, 2),
            })

    if body_total == 0:
        findings.append({"severity": "P0", "code": "no_body_paragraphs_detected"})
    if declared and len(declared) > max(12, round(max(body_total, 1) * 0.1)):
        findings.append({"severity": "P0", "code": "excessive_layout_exceptions", "count": len(declared)})
    for code, items in issue_buckets.items():
        if items:
            findings.append({"severity": "P0", "code": code, "count": len(items), "samples": items[:25]})

    return {
        "schemaVersion": 1,
        "generatedAt": now_iso(),
        "docx": {"path": str(path), "sha256": docx_sha},
        "profile": {
            "id": "medina-legal-body-v1",
            "bodyFont": BODY_FONT,
            "bodySizePt": BODY_SIZE_PT,
            "bodyAlignment": "justify",
            "tableMinimumSizePt": TABLE_MIN_SIZE_PT,
            "folioMaximumWidthPt": FOLIO_SAFE_WIDTH_PT,
        },
        "metrics": {
            "bodyParagraphs": body_total,
            "bodyJustified": body_justified,
            "bodyFontCompliant": body_font_ok,
            "bodySizeCompliant": body_size_ok,
            "justificationCoverage": round(body_justified / body_total, 6) if body_total else 0.0,
            "fontCoverage": round(body_font_ok / body_total, 6) if body_total else 0.0,
            "sizeCoverage": round(body_size_ok / body_total, 6) if body_total else 0.0,
            "declaredExceptions": len(declared),
        },
        "paragraphs": records,
        "findings": findings,
        "approved": not findings,
    }


def _set_run_font_size(run_element, font_name: str, size_pt: float) -> None:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font_name)
    half_points = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), half_points)


def normalize_medina_body(source: Path, output: Path) -> dict:
    """Corrige apenas o corpo; preserva texto, tabelas, tracking e papéis visuais."""
    source, output = Path(source), Path(output)
    document = Document(str(source))
    changed = []
    previous_role = None
    for index, paragraph in enumerate(document.paragraphs):
        text = _paragraph_text(paragraph)
        if not text:
            continue
        role = _role_for(paragraph, text, index, previous_role)
        if role == "body":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.5
            for run_element in paragraph._p.xpath(".//w:r"):
                if _substantial_run_text(run_element):
                    _set_run_font_size(run_element, BODY_FONT, BODY_SIZE_PT)
            changed.append(index)
        previous_role = role
    table_changes = []
    for table_index, table in enumerate(document.tables):
        groups = {"header": [], "body": []}
        for row_index, row in enumerate(table.rows):
            group = "header" if row_index == 0 else "body"
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run_element in paragraph._p.xpath(".//w:r"):
                        run_text = _substantial_run_text(run_element)
                        if not run_text:
                            continue
                        font, size = _effective_run_font_size(run_element, paragraph)
                        groups[group].append({
                            "run": run_element,
                            "font": font,
                            "size": size,
                            "chars": len(run_text),
                        })
        for group, values in groups.items():
            if not values:
                continue
            font_counts = Counter()
            size_counts = Counter()
            for item in values:
                font_counts[item["font"] or TABLE_FALLBACK_FONT] += item["chars"]
                size_counts[round(float(item["size"]), 2) if item["size"] is not None else TABLE_FALLBACK_SIZE_PT] += item["chars"]
            selected_font = font_counts.most_common(1)[0][0]
            selected_size = max(TABLE_MIN_SIZE_PT, float(size_counts.most_common(1)[0][0]))
            if len(font_counts) > 1 or len(size_counts) > 1 or any(
                item["font"] is None or item["size"] is None for item in values
            ):
                for item in values:
                    _set_run_font_size(item["run"], selected_font, selected_size)
                table_changes.append({
                    "table": table_index,
                    "group": group,
                    "font": selected_font,
                    "sizePt": selected_size,
                    "runs": len(values),
                })
    folio_changes = []
    for folio_index, rect in enumerate(_folio_rectangles(document)):
        width = _vml_width_pt(rect)
        if width is None or width > FOLIO_SAFE_WIDTH_PT + 0.01:
            _set_vml_width_pt(rect, FOLIO_TEMPLATE_WIDTH_PT)
            folio_changes.append({
                "folio": folio_index,
                "fromWidthPt": width,
                "toWidthPt": FOLIO_TEMPLATE_WIDTH_PT,
            })
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    content_fidelity = compare_docx_content(source, output)
    if not content_fidelity["approved"]:
        raise RuntimeError(
            "normalização alterou conteúdo jurídico ou controle de alterações: "
            + ", ".join(content_fidelity["mismatches"])
        )
    audit = audit_docx_layout(output)
    audit["normalization"] = {
        "source": str(source),
        "sourceSha256": sha256_file(source),
        "changedParagraphs": changed,
        "changedCount": len(changed),
        "tableTypographyChanges": table_changes,
        "folioWidthChanges": folio_changes,
        "contentFidelity": content_fidelity,
    }
    return audit


def main() -> int:
    parser = argparse.ArgumentParser(description="Audita/corrige o corpo Word Medina Osório")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path, help="grava cópia corrigida; nunca sobrescreve a origem")
    parser.add_argument("--json", type=Path)
    parser.add_argument("--exceptions", type=Path)
    args = parser.parse_args()
    if args.output:
        if args.output.resolve() == args.docx.resolve():
            raise SystemExit("a correção deve usar arquivo de saída distinto da origem")
        result = normalize_medina_body(args.docx, args.output)
    else:
        result = audit_docx_layout(args.docx, exceptions=args.exceptions)
    if args.json:
        atomic_write_json(args.json, result)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["approved"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
