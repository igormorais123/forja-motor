"""Gates de fidelidade semântica.

``compare_fidelity`` permanece para artefatos legados que já possuem PDF. A
rota oficial da FORJA não materializa nem renderiza PDF: usa
``compare_docx_fidelity`` para comparar Markdown e o OOXML final diretamente.
"""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from collections import Counter
from pathlib import Path

from forja_context import markdown_blocks
from forja_n3_common import atomic_write_json, now_iso, sha256_file


QUALIFIERS = (
    "não",
    "salvo",
    "exceto",
    "apenas",
    "subsidiariamente",
    "preliminarmente",
    "não verificado",
    "verificar",
)


def _norm(value: str) -> str:
    value = unicodedata.normalize("NFKD", value.casefold())
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = re.sub(r"[`*_#>|]", " ", value)
    return re.sub(r"[^a-z0-9]+", " ", value).strip()


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts += [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(parts)


def _pdf_text(path: Path) -> tuple[str, int]:
    import fitz

    with fitz.open(path) as document:
        return "\n".join(page.get_text("text") for page in document), len(document)


def _segments(block: dict) -> list[str]:
    raw = block["text"]
    lines = []
    for line in raw.splitlines():
        stripped = line.strip()
        stripped = re.sub(r"^\*{0,2}\d{1,3}[.)]\*{0,2}\s+", "", stripped)
        if re.fullmatch(r"\|?[\s:|-]+\|?", stripped):
            continue
        if stripped.startswith("|"):
            lines.extend(cell.strip() for cell in stripped.strip("|").split("|") if cell.strip())
        else:
            lines.append(stripped)
    normalized = [_norm(line) for line in lines]
    return [line for line in normalized if len(line) >= 25]


def _coverage(segments: list[dict], target: str) -> tuple[float, list[dict]]:
    missing = []
    checked = 0
    target_norm = _norm(target)
    for item in segments:
        for segment in item["segments"]:
            checked += 1
            words = segment.split()
            chunks = [" ".join(words[index:index + 8]) for index in range(0, len(words), 8) if len(words[index:index + 8]) >= 4]
            chunk_ratio = 1.0 if segment in target_norm else (
                sum(chunk in target_norm for chunk in chunks) / len(chunks) if chunks else 0.0
            )
            token_ratio = 0.0
            if item.get("unordered"):
                unique_tokens = set(words)
                target_tokens = set(target_norm.split())
                token_ratio = sum(token in target_tokens for token in unique_tokens) / len(unique_tokens) if unique_tokens else 0.0
            if chunk_ratio < 0.75 and token_ratio < 0.9:
                missing.append({
                    "blockId": item["blockId"],
                    "line": item["line"],
                    "sample": segment[:180],
                    "chunkCoverage": round(chunk_ratio, 3),
                    "tokenCoverage": round(token_ratio, 3),
                })
    return (1.0 if not checked else (checked - len(missing)) / checked), missing


def _number_tokens(value: str) -> Counter:
    result = Counter()
    value = re.sub(r"(?<=\d)-\s*\n\s*(?=\d)", "-", value)
    for line in value.splitlines():
        # Numeração estrutural de títulos/listas pode ser recomposta no Word;
        # números jurídicos dentro da frase continuam obrigatoriamente preservados.
        line = re.sub(r"^\s*#{1,6}\s+(?:\d+(?:\.\d+)*[.)]?\s+)?", "", line)
        # O primeiro campo de uma tabela pode ser apenas o rótulo estrutural da
        # linha ("| 2. Tutela |"). No DOCX cada célula vira uma linha e o gate
        # abaixo já ignora esse mesmo prefixo; faça o tratamento simétrico na
        # fonte Markdown para não produzir falso P0.
        line = re.sub(r"^\s*\|\s*\d+[.)]\s+", "| ", line)
        # Parágrafos jurídicos numerados podem chegar como ``**2.** Texto``.
        # O DOCX/PDF não carrega os asteriscos; remova a numeração estrutural
        # de forma simétrica antes de comparar apenas os números jurídicos.
        line = re.sub(r"^\s*\*{0,2}\d+[.)]\*{0,2}\s+", "", line)
        for raw in re.findall(r"(?<!\w)(?:R\$\s*)?\d[\d./,%:-]*", line, re.I):
            token = re.sub(r"\D", "", raw)
            if token:
                result[token] += 1
    return result


def _missing_counter(source: Counter, target: Counter) -> list[dict]:
    return [
        {"token": token, "expected": count, "found": target.get(token, 0)}
        for token, count in source.items()
        if target.get(token, 0) < count
    ]


def _missing_numbers(source: Counter, target_text: str) -> list[dict]:
    target_tokens = _number_tokens(target_text)
    digit_stream = re.sub(r"\D", "", target_text)
    missing = []
    for token in source:
        found = token in digit_stream if len(token) >= 6 else token in target_tokens
        if not found:
            missing.append({"token": token, "expected": "presence", "found": 0})
    return missing


def _qualifier_counts(value: str) -> Counter:
    normalized = " " + _norm(value) + " "
    return Counter({qualifier: normalized.count(" " + _norm(qualifier) + " ") for qualifier in QUALIFIERS})


def compare_fidelity(markdown: Path, docx: Path, pdf: Path) -> dict:
    markdown_text = markdown.read_text(encoding="utf-8")
    docx_text = _docx_text(docx)
    pdf_text, page_count = _pdf_text(pdf)
    blocks = markdown_blocks(markdown_text)
    significant = []
    for block in blocks:
        segments = _segments(block)
        if segments:
            significant.append({
                "blockId": block["blockId"],
                "line": block["startLine"],
                "segments": segments,
                "unordered": block["kind"] == "table",
            })
    docx_coverage, docx_missing = _coverage(significant, docx_text)
    pdf_coverage, pdf_missing = _coverage(significant, pdf_text)
    source_numbers = _number_tokens(markdown_text)
    source_qualifiers = _qualifier_counts(markdown_text)
    findings = []
    if docx_missing:
        findings.append({"severity": "P0", "code": "markdown_block_missing_in_docx", "count": len(docx_missing)})
    if pdf_missing:
        findings.append({"severity": "P0", "code": "markdown_block_missing_in_pdf", "count": len(pdf_missing)})
    docx_numbers_missing = _missing_numbers(source_numbers, docx_text)
    pdf_numbers_missing = _missing_numbers(source_numbers, pdf_text)
    if docx_numbers_missing:
        findings.append({"severity": "P0", "code": "number_missing_in_docx", "count": len(docx_numbers_missing)})
    if pdf_numbers_missing:
        findings.append({"severity": "P0", "code": "number_missing_in_pdf", "count": len(pdf_numbers_missing)})
    docx_qualifiers_missing = _missing_counter(source_qualifiers, _qualifier_counts(docx_text))
    pdf_qualifiers_missing = _missing_counter(source_qualifiers, _qualifier_counts(pdf_text))
    if docx_qualifiers_missing:
        findings.append({"severity": "P0", "code": "qualifier_missing_in_docx", "count": len(docx_qualifiers_missing)})
    if pdf_qualifiers_missing:
        findings.append({"severity": "P0", "code": "qualifier_missing_in_pdf", "count": len(pdf_qualifiers_missing)})
    return {
        "schemaVersion": 1,
        "generatedAt": now_iso(),
        "markdown": {"path": str(markdown), "sha256": sha256_file(markdown)},
        "docx": {"path": str(docx), "sha256": sha256_file(docx)},
        "pdf": {"path": str(pdf), "sha256": sha256_file(pdf), "pageCount": page_count},
        "blocks": {
            "total": len(blocks),
            "significant": len(significant),
            "docxCoverage": round(docx_coverage, 6),
            "pdfCoverage": round(pdf_coverage, 6),
            "docxMissing": docx_missing,
            "pdfMissing": pdf_missing,
        },
        "numbers": {"docxMissing": docx_numbers_missing, "pdfMissing": pdf_numbers_missing},
        "qualifiers": {"docxMissing": docx_qualifiers_missing, "pdfMissing": pdf_qualifiers_missing},
        "findings": findings,
        "approved": not findings,
    }


def compare_docx_fidelity(markdown: Path, docx: Path) -> dict:
    """Compara blocos, números e qualificadores apenas contra o DOCX.

    A inspeção ocorre no texto efetivamente armazenado no pacote OOXML. Isso
    cobre a fidelidade jurídica sem depender de Word COM, PDF, PNG ou de um
    motor de renderização. A preservação de layout é responsabilidade do
    ``forja_docx_layout`` e a presença de SVG do QA estrutural.
    """
    markdown_text = Path(markdown).read_text(encoding="utf-8")
    docx_text = _docx_text(Path(docx))
    blocks = markdown_blocks(markdown_text)
    significant = []
    for block in blocks:
        segments = _segments(block)
        if segments:
            significant.append({
                "blockId": block["blockId"],
                "line": block["startLine"],
                "segments": segments,
                "unordered": block["kind"] == "table",
            })
    coverage, missing = _coverage(significant, docx_text)
    source_numbers = _number_tokens(markdown_text)
    source_qualifiers = _qualifier_counts(markdown_text)
    numbers_missing = _missing_numbers(source_numbers, docx_text)
    qualifiers_missing = _missing_counter(source_qualifiers, _qualifier_counts(docx_text))
    findings = []
    if missing:
        findings.append({"severity": "P0", "code": "markdown_block_missing_in_docx", "count": len(missing)})
    if numbers_missing:
        findings.append({"severity": "P0", "code": "number_missing_in_docx", "count": len(numbers_missing)})
    if qualifiers_missing:
        findings.append({"severity": "P0", "code": "qualifier_missing_in_docx", "count": len(qualifiers_missing)})
    return {
        "schemaVersion": 1,
        "generatedAt": now_iso(),
        "mode": "markdown_docx_ooxml",
        "markdown": {"path": str(markdown), "sha256": sha256_file(Path(markdown))},
        "docx": {"path": str(docx), "sha256": sha256_file(Path(docx))},
        "blocks": {
            "total": len(blocks),
            "significant": len(significant),
            "docxCoverage": round(coverage, 6),
            "docxMissing": missing,
        },
        "numbers": {"docxMissing": numbers_missing},
        "qualifiers": {"docxMissing": qualifiers_missing},
        "findings": findings,
        "approved": not findings,
    }


def write_fidelity(markdown: Path, docx: Path, pdf: Path, output: Path) -> dict:
    result = compare_fidelity(markdown, docx, pdf)
    atomic_write_json(output, result)
    return result


def main() -> None:
    parser = argparse.ArgumentParser(description="Compara fidelidade Markdown, Word e PDF")
    parser.add_argument("markdown", type=Path)
    parser.add_argument("docx", type=Path)
    parser.add_argument("pdf", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    result = write_fidelity(args.markdown, args.docx, args.pdf, args.output)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if result["approved"] else 1)


if __name__ == "__main__":
    main()
