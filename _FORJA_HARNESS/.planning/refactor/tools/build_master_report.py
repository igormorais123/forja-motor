from __future__ import annotations

import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
HARNESS = ROOT.parents[1]
FACTORY = HARNESS.parent
TOOLS = FACTORY / "_FERRAMENTAS"
sys.path.insert(0, str(TOOLS))

from estilo_medina import checar_fontes_svg, largura_recomendada_cm  # noqa: E402
from word_visual_pipeline import (  # noqa: E402
    MMDC,
    docx_para_pdf,
    inserir_emf_word_com,
    svg_para_emf,
)


OUT = ROOT / "deliverables"
DIAGRAM_DIR = OUT / "diagramas_pdf"
DOCX = OUT / "PLANO_MESTRE_REFATORACAO_FORJA_2026-07-15.docx"
PDF = OUT / "PLANO_MESTRE_REFATORACAO_FORJA_2026-07-15.pdf"
LOGO = FACTORY / "Cafelana" / "_revista" / "logo_medina_transp.png"

PETROLEUM = "395C60"
TERRACOTTA = "9C5B38"
GRAPHITE = "49494D"
LIGHT_PETROLEUM = "EFF4F3"
LIGHT_TERRACOTTA = "FBF2EC"
WHITE = "FFFFFF"
LIGHT_GREY = "F5F5F3"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep(paragraph, keep_next=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    p_pr.append(OxmlElement("w:keepLines"))
    p_pr.append(OxmlElement("w:widowControl"))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FORJA R1  •  ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAPHITE)
    for kind in ("PAGE", "NUMPAGES"):
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = kind
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        field_run = paragraph.add_run()
        field_run._r.extend((fld_char_begin, instr_text, fld_char_end))
        if kind == "PAGE":
            paragraph.add_run(" / ")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(GRAPHITE)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 28, PETROLEUM, 0, 10),
        ("Subtitle", 13, GRAPHITE, 0, 10),
        ("Heading 1", 19, PETROLEUM, 16, 8),
        ("Heading 2", 14, TERRACOTTA, 12, 5),
        ("Heading 3", 11.5, PETROLEUM, 9, 3),
        ("Heading 4", 10.5, GRAPHITE, 7, 2),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("MEDINA OSÓRIO  |  PLANO TÉCNICO")
    r.font.name = "Times New Roman"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(PETROLEUM)

    footer = section.footer
    line = footer.paragraphs[0]
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:top")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), PETROLEUM)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_field(line)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO.is_file():
        p.add_run().add_picture(str(LOGO), width=Cm(5.2))

    doc.add_paragraph("PLANO MESTRE", style="Title")
    title = doc.add_paragraph("Refatoração estrutural segura da FORJA", style="Title")
    title.runs[0].font.size = Pt(25)
    doc.add_paragraph(
        "PRD • TDD • ROADMAP • ATLAS MERMAID • TESTES • ROLLBACK • CUTOVER",
        style="Subtitle",
    )

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_fill(cell, PETROLEUM)
    set_cell_margins(cell, 220, 220, 220, 220)
    p = cell.paragraphs[0]
    r = p.add_run(
        "Objetivo: limpar, deduplicar, modularizar e tornar o código navegável por humanos e IAs, "
        "sem alterar comportamento jurídico, trilha de prova ou estado normativo por acidente."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph("")
    metadata = doc.add_table(rows=5, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    metadata.autofit = False
    labels = (
        ("Data-base", "15 de julho de 2026"),
        ("Versão", "FORJA R1.0 — planejamento"),
        ("Base Git", "291fae0d57168e98ae10b0db866629b153ca9d5c"),
        ("Estado", "Planejado; execução não iniciada"),
        ("Escopo", "_FORJA_HARNESS"),
    )
    for row, (label, value) in zip(metadata.rows, labels):
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(11.5)
        set_cell_fill(row.cells[0], LIGHT_PETROLEUM)
        for cell in row.cells:
            set_cell_margins(cell)
        rr = row.cells[0].paragraphs[0].add_run(label)
        rr.bold = True
        rr.font.color.rgb = RGBColor.from_string(PETROLEUM)
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph("")
    warning = doc.add_table(rows=1, cols=1)
    cell = warning.cell(0, 0)
    set_cell_fill(cell, LIGHT_TERRACOTTA)
    set_cell_margins(cell, 160, 180, 160, 180)
    r = cell.paragraphs[0].add_run(
        "LIMITE DA ENTREGA — Este documento autoriza o planejamento, não a execução. "
        "Nenhum arquivo funcional foi movido, removido ou refatorado nesta etapa."
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_page_break()


def add_panel(doc: Document, title: str, text: str, color=LIGHT_PETROLEUM) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, color)
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    r = p.add_run(title.upper() + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(PETROLEUM if color == LIGHT_PETROLEUM else TERRACOTTA)
    p.add_run(text)


def add_inline_runs(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(PETROLEUM)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    parsed = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(parsed) > 1 and is_table_separator(lines[1]):
        parsed.pop(1)
    columns = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(parsed):
        for j in range(columns):
            cell = table.cell(i, j)
            set_cell_margins(cell, 55, 70, 55, 70)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(8.2)
            if i == 0:
                set_cell_fill(cell, PETROLEUM)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            elif i % 2 == 0:
                set_cell_fill(cell, LIGHT_GREY)
    if table.rows:
        set_repeat_table_header(table.rows[0])
    doc.add_paragraph("").paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_GREY)
    set_cell_margins(cell, 90, 110, 90, 110)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code.rstrip())
    r.font.name = "Consolas"
    r.font.size = Pt(7.6)
    r.font.color.rgb = RGBColor.from_string(GRAPHITE)


def render_markdown(doc: Document, path: Path, top_level_offset: int = 0) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```mermaid"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(4, len(heading.group(1)) + top_level_offset)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, heading.group(2))
            set_keep(p, keep_next=True)
        elif re.match(r"^\s*[-*]\s+", line):
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet 2" if indent >= 2 else "List Bullet")
            add_inline_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
            set_keep(p)
        elif re.match(r"^\s*\d+[.)]\s+", line):
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Number 2" if indent >= 2 else "List Number")
            add_inline_runs(p, re.sub(r"^\s*\d+[.)]\s+", "", line))
            set_keep(p)
        elif stripped.startswith(">"):
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_fill(cell, LIGHT_PETROLEUM)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            add_inline_runs(p, stripped.lstrip("> "))
            for run in p.runs:
                run.italic = True
        elif stripped in ("---", "***") or re.fullmatch(r"</?[A-Za-z_][^>]*>", stripped):
            pass
        elif stripped:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped)
            set_keep(p)
        i += 1


def custom_diagrams() -> dict[str, str]:
    return {
        "A01": """flowchart TB
U[\"Pessoa usuária / IA\"] --> F[\"FORJA\"]
C[\"Acervo do caso\"] --> F
F --> G[\"Gates jurídicos e técnicos\"]
G --> O[\"Peça + evidências + relatório\"]
""",
        "A02": """flowchart TB
E[\"Entrypoints e CLI\"] --> A[\"Application: casos de uso\"]
A --> D[\"Domain: regras e contratos\"]
A --> P[\"Ports: fronteiras explícitas\"]
P --> X[\"Adapters: arquivos, Word, pesquisa e gestão\"]
""",
        "A03": """flowchart TB
P[\"Proteção • R0–R2\"] --> M[\"Modularização • R3–R5\"]
M --> D[\"Decomposição e limpeza • R6–R7\"]
D --> V[\"Atlas, replay e decisão • R8–R9\"]
""",
    }


def prepare_diagrams() -> tuple[dict[str, tuple[str, float]], list[dict[str, object]]]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    render_report = json.loads((ROOT / "deliverables" / "diagramas" / "render_report.json").read_text(encoding="utf-8"))
    selected_prefixes = {
        "D01", "D06", "D09", "D10", "D11", "D12", "D13", "D14", "D18", "D19"
    }
    selected_titles = {
        "5. Estados explícitos de fonte",
        "6. Política RED–GREEN–REFACTOR",
        "3. Descoberta da régua",
        "11. Cutover",
    }
    candidates: list[tuple[str, str, Path]] = []
    for item in render_report["rendered"]:
        title = item["title"]
        diagram_id = title.split()[0] if title.startswith("D") else ""
        if diagram_id in selected_prefixes or title in selected_titles:
            candidates.append((diagram_id or f"S{len(candidates)+1:02d}", title, (ROOT / item["svg"]).with_suffix(".mmd")))

    for key, source_text in custom_diagrams().items():
        source = DIAGRAM_DIR / f"{key}.mmd"
        source.write_text(source_text.strip() + "\n", encoding="utf-8")
        candidates.insert(0, (key, {"A01": "Contexto do sistema", "A02": "Camadas-alvo", "A03": "Ondas do programa"}[key], source))

    config = DIAGRAM_DIR / "mermaid_word_config.json"
    config.write_text(
        json.dumps(
            {
                "theme": "base",
                "htmlLabels": False,
                "flowchart": {"htmlLabels": False, "curve": "basis"},
                "themeVariables": {
                    "fontFamily": "Times New Roman",
                    "fontSize": "18px",
                    "primaryColor": "#EFF4F3",
                    "primaryTextColor": "#395C60",
                    "primaryBorderColor": "#395C60",
                    "secondaryColor": "#FBF2EC",
                    "tertiaryColor": "#FFFFFF",
                    "lineColor": "#49494D",
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    compact_overrides = {
        "S02": """flowchart TB
C[\"Caracterizar comportamento\"] --> R[\"RED: falha planejada\"]
R --> G[\"GREEN: mudança mínima\"]
G --> F[\"REFACTOR + suíte ampla\"]
F --> Q{\"Afeta Word, pacote ou integração?\"}
Q -->|sim| X[\"Canário real + telemetria\"]
Q -->|não| Z[\"Concluir\"]
X --> Z
""",
        "D06": """flowchart TB
A[\"prepare attempt\"] --> P[\"produção isolada\"]
P --> R[\"revalidar contexto, contrato e inputs\"]
R --> G{\"gates aprovados?\"}
G -->|não| B[\"preservar attempt bloqueado\"]
G -->|sim| C[\"commit atômico\"]
C --> E[\"evento + visão canônica\"]
""",
        "D10": """flowchart TB
I[\"arquivo inventariado\"] --> S{\"examinado integralmente?\"}
S -->|não| U[\"P0: unscanned + razão\"]
S -->|sim| T{\"ameaça técnica?\"}
T -->|sim| B[\"P0: bloquear\"]
T -->|não| P[\"benigno: prosseguir\"]
""",
        "D13": """flowchart TB
I[\"Markdown + artefatos\"] --> M[\"modelo intermediário\"]
M --> C[\"template + SVG→EMF\"]
C --> W[\"Word COM → PDF\"]
W --> R[\"render de todas as páginas\"]
R --> Q[\"QA visual independente\"]
Q --> G{\"fidelidade aprovada?\"}
""",
    }

    markers: dict[str, tuple[str, float]] = {}
    catalog: list[dict[str, object]] = []
    used: set[str] = set()
    for key, title, source in candidates:
        base = key
        suffix = 2
        while key in used:
            key = f"{base}_{suffix}"
            suffix += 1
        used.add(key)
        if key in compact_overrides:
            source = DIAGRAM_DIR / f"{key}_compact.mmd"
            source.write_text(compact_overrides[key].strip() + "\n", encoding="utf-8")
        svg = DIAGRAM_DIR / f"{key}_word.svg"
        subprocess.run(
            [MMDC, "-i", str(source), "-o", str(svg), "-w", "1200", "-b", "transparent", "-c", str(config)],
            check=True,
            shell=False,
        )
        svg_text = svg.read_text(encoding="utf-8")
        if "<foreignObject" in svg_text or "<text" not in svg_text:
            raise RuntimeError(f"SVG incompatível com EMF/Word em {title}: rótulos não são texto SVG nativo")
        # Mermaid separa palavras em tspans e usa espaço inicial no tspan seguinte.
        # Inkscape/EMF descarta esse espaço sem xml:space="preserve", colando palavras no Word.
        svg_text = re.sub(r"<text(?![^>]*xml:space=)", '<text xml:space="preserve"', svg_text)
        svg_text = re.sub(r"<tspan(?![^>]*xml:space=)", '<tspan xml:space="preserve"', svg_text)
        svg.write_text(svg_text, encoding="utf-8")
        width = largura_recomendada_cm(str(svg), alvo_pt=8.2, maximo_cm=15.5)
        violations = checar_fontes_svg(str(svg), width, min_pt=8.0)
        if violations:
            raise RuntimeError(f"Legibilidade reprovada em {title} a {width} cm: {violations}")
        emf = DIAGRAM_DIR / f"{key}.emf"
        svg_para_emf(str(svg), str(emf), largura_final_cm=width)
        marker = f"{{{{FIG_{key}}}}}"
        markers[marker] = (str(emf), width)
        catalog.append({"key": key, "title": title, "marker": marker, "widthCm": width, "svg": str(svg)})
    return markers, catalog


def add_figure(doc: Document, item: dict[str, object], caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run(str(item["marker"]))
    r.font.size = Pt(1)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run(caption or str(item["title"]))
    rr.italic = True
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor.from_string(GRAPHITE)
    set_keep(cp)


def extract_tag(raw: str, tag: str) -> str:
    match = re.search(fr"<{tag}[^>]*>(.*?)</{tag}>", raw, flags=re.DOTALL)
    if not match:
        return ""
    value = re.sub(r"<[^>]+>", " ", match.group(1))
    return re.sub(r"\s+", " ", value).strip()


def add_plan_catalog(doc: Document) -> None:
    doc.add_heading("8. Catálogo dos 18 planos executáveis", level=1)
    add_panel(
        doc,
        "Formato de execução",
        "Cada plano possui contexto prévio, ameaça, arquivos a ler antes de editar, ação concreta, "
        "critério de aceite, comando de verificação e condição de rollback.",
    )
    for plan in sorted((ROOT / "plans").glob("P*-PLAN.md")):
        raw = plan.read_text(encoding="utf-8")
        title_match = re.search(r"^#\s+(.+)$", raw, flags=re.MULTILINE)
        title = title_match.group(1) if title_match else plan.stem
        requirements_match = re.search(r"^requirements:\s*\[(.*?)\]", raw, flags=re.MULTILINE)
        requirements = requirements_match.group(1) if requirements_match else ""
        objective = extract_tag(raw, "objective")
        threat = extract_tag(raw, "threat_model")
        verification = extract_tag(raw, "verification")

        p = doc.add_paragraph(style="Heading 2")
        add_inline_runs(p, title)
        meta = doc.add_table(rows=1, cols=2)
        meta.style = "Table Grid"
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell_fill(meta.cell(0, 0), LIGHT_PETROLEUM)
        meta.cell(0, 0).paragraphs[0].add_run("Requisitos").bold = True
        meta.cell(0, 1).paragraphs[0].add_run(requirements)
        for cell in meta.rows[0].cells:
            set_cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.3)

        for label, value in (("Objetivo", objective), ("Ameaça controlada", threat), ("Verificação", verification)):
            p = doc.add_paragraph()
            r = p.add_run(label + ": ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(PETROLEUM)
            add_inline_runs(p, value)
            set_keep(p)


def add_executive(doc: Document, diagrams: list[dict[str, object]]) -> None:
    doc.add_heading("1. Síntese executiva", level=1)
    add_panel(
        doc,
        "Decisão central",
        "A refatoração será conduzida como migração observável, por ondas e com fachadas compatíveis. "
        "Nenhuma remoção ocorre antes de equivalência funcional, replay histórico, backup restaurável e ausência de uso do shim.",
    )
    doc.add_paragraph(
        "O programa FORJA R1 organiza a limpeza em dez ondas (R0–R9), 22 requisitos funcionais, "
        "15 requisitos não funcionais, 24 sentinelas de invariantes e 18 planos executáveis. O foco imediato é "
        "construir confiança: descobrir toda a régua de testes, corrigir os quatro gates P0 e congelar um baseline reproduzível."
    )
    bullets = (
        "Preservar N2 vigente, N3 em sombra e N4 em pilot_blocking; refatoração não promove especificação.",
        "Manter F2-A FORJA-F2A-100-v1 obrigatório em casos novos.",
        "Centralizar regras e catálogos sem criar um novo monólito.",
        "Separar domínio, aplicação, portas e adaptadores; entrypoints antigos tornam-se fachadas temporárias.",
        "Tratar eventos, evidências, contratos e estado como patrimônio não descartável.",
        "Produzir atlas Mermaid vivo e manifestos legíveis por IA para reduzir mudanças às cegas.",
    )
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    by_key = {str(item["key"]): item for item in diagrams}
    for key in ("A01", "A02", "A03"):
        add_figure(doc, by_key[key])

    doc.add_heading("Mapa de leitura", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Bloco", "Pergunta respondida", "Fonte canônica")
    for i, header in enumerate(headers):
        set_cell_fill(table.cell(0, i), PETROLEUM)
        rr = table.cell(0, i).paragraphs[0].add_run(header)
        rr.bold = True
        rr.font.color.rgb = RGBColor.from_string(WHITE)
    rows = (
        ("PRD", "O que deve existir e o que não pode mudar?", "01-PRD_REFATORACAO_FORJA.md"),
        ("TDD", "Como separar responsabilidades e testar cada mudança?", "02-TDD_REFATORACAO_FORJA.md"),
        ("Roadmap", "Em que ordem, com quais gates e rollback?", "03-ROADMAP_REFATORACAO_FORJA.md"),
        ("Atlas", "O que chama o quê e onde estão os riscos?", "04-DIAGRAMAS_REFATORACAO_FORJA.md"),
        ("Rastreabilidade", "Qual requisito produz qual teste e evidência?", "05-MATRIZ_RASTREABILIDADE.md"),
        ("Cutover", "Como provar equivalência e voltar atrás?", "06-TESTES_ROLLBACK_E_CUTOVER.md"),
        ("Planos", "Qual é a unidade executável de trabalho?", "18 planos em plans/P00–P16"),
    )
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].paragraphs[0].add_run(value)
            set_cell_margins(cells[i])
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    markers, diagrams = prepare_diagrams()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_executive(doc, diagrams)

    sources = (
        ("2. Contexto e decisões travadas", ROOT / "00-CONTEXT.md"),
        ("3. PRD — requisitos e invariantes", ROOT / "01-PRD_REFATORACAO_FORJA.md"),
        ("4. TDD — desenho técnico e testes", ROOT / "02-TDD_REFATORACAO_FORJA.md"),
        ("5. Roadmap — ondas, dependências e gates", ROOT / "03-ROADMAP_REFATORACAO_FORJA.md"),
        ("6. Rastreabilidade ponta a ponta", ROOT / "05-MATRIZ_RASTREABILIDADE.md"),
        ("7. Testes, rollback e cutover", ROOT / "06-TESTES_ROLLBACK_E_CUTOVER.md"),
    )
    for report_title, source in sources:
        doc.add_page_break()
        doc.add_heading(report_title, level=1)
        p = doc.add_paragraph()
        rr = p.add_run(f"Fonte canônica: {source.name}")
        rr.italic = True
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor.from_string(GRAPHITE)
        render_markdown(doc, source, top_level_offset=1)

    doc.add_page_break()
    add_plan_catalog(doc)

    doc.add_page_break()
    doc.add_heading("9. Atlas visual selecionado", level=1)
    add_panel(
        doc,
        "Cobertura visual",
        "O atlas fonte contém 30 blocos Mermaid validados. Esta edição incorpora os 17 diagramas de maior utilidade operacional; "
        "os demais permanecem no arquivo canônico e no relatório de renderização.",
    )
    for item in diagrams:
        if str(item["key"]).startswith("A"):
            continue
        doc.add_heading(str(item["title"]), level=2)
        add_figure(doc, item, f"Diagrama {item['key']} — {item['title']} • largura final {item['widthCm']} cm")

    final_heading = doc.add_heading("10. Decisão de início", level=1)
    final_heading.paragraph_format.page_break_before = True
    add_panel(
        doc,
        "Gate de autorização",
        "A execução começa somente quando este plano for aprovado como referência e o programa entrar em R0. "
        "A aprovação do documento não autoriza exclusões, cutover, promoção normativa ou comunicação externa.",
        LIGHT_TERRACOTTA,
    )
    doc.add_heading("Primeiro lote recomendado", level=2)
    for item in (
        "P00 — baseline isolado, inventário e contratos observados.",
        "P01 — descoberta canônica de toda a régua de testes.",
        "P02–P04 — correção TDD dos quatro gates P0.",
        "P05 — ambiente reproduzível e esqueleto modular sem mover lógica.",
    ):
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "Conclusão do planejamento: o trabalho está decomposto, rastreável, reversível e pronto para execução por ondas. "
        "A próxima mudança funcional segura é iniciar R0, não mover arquivos nem deduplicar código diretamente."
    )

    props = doc.core_properties
    props.title = "Plano Mestre de Refatoração Segura da FORJA"
    props.subject = "PRD, TDD, roadmap, Mermaid, testes, rollback e cutover"
    props.author = "Medina Osório Advogados — planejamento técnico assistido"
    props.keywords = "FORJA, refatoração, PRD, TDD, roadmap, Mermaid, arquitetura"
    props.comments = "Planejamento somente; execução não autorizada."

    doc.save(DOCX)
    inserir_emf_word_com(str(DOCX), markers)
    docx_para_pdf(str(DOCX), str(PDF))
    import fitz

    pdf_document = fitz.open(PDF)
    pdf_document.set_metadata(
        {
            "title": "Plano Mestre de Refatoração Segura da FORJA",
            "author": "Medina Osório Advogados",
            "subject": "PRD, TDD, roadmap, Mermaid, testes, rollback e cutover",
            "keywords": "FORJA, refatoração, PRD, TDD, roadmap, Mermaid, arquitetura",
            "creator": "Microsoft Word 2021 / pipeline Medina Osório",
            "producer": "Microsoft Word 2021",
        }
    )
    pdf_document.saveIncr()
    pdf_document.close()

    report = {
        "ok": True,
        "docx": str(DOCX),
        "pdf": str(PDF),
        "diagramsEmbedded": len(markers),
        "allSvgPassedMinimumPrintedFontPt": 8.0,
        "executionAuthorized": False,
    }
    (OUT / "build_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    build()
