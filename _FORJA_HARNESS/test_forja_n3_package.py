from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey

from forja_close_cycle import confirm_delivery, create_package, fulfill, register_draft
from forja_f10_contract import REQUIRED_F10_GATES, compute_f10_gates, validate_f10_gates
from forja_fidelity import write_fidelity
from forja_human_review import (
    build_unsigned_visual_receipt,
    canonical_receipt_bytes,
    public_key_id,
)
from forja_n3_common import PHASES, ForjaN3Error, RevisionConflict, atomic_write_json, now_iso, read_json, sha256_file
from forja_package import validate_definition
from forja_state_machine import derive_state, initialize_case, record_event
from forja_visual_qa import inspect_pdf, run_visual_qa
from forja_visual_review import REQUIRED_PAGE_CHECKS, build_pending_review


class ForjaN3PackageTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self._trust_patchers = []
        self.case = Path(self.temp.name) / "case-package"
        self.case.mkdir()
        initialize_case(self.case, demand_id="demand-package")
        revision = 1
        for index, phase in enumerate(PHASES[:9]):
            _, state, _ = record_event(
                self.case, "phase_started", expected_revision=revision,
                idempotency_key=f"{phase}:start", phase=phase,
            )
            revision = state["revision"]
            _, state, _ = record_event(
                self.case, "phase_completed", expected_revision=revision,
                idempotency_key=f"{phase}:done", phase=phase,
            )
            revision = state["revision"]
        _, state, _ = record_event(
            self.case, "phase_started", expected_revision=revision,
            idempotency_key="f9:start", phase=PHASES[9],
        )
        self.revision = state["revision"]
        self.files = self.case / "synthetic"
        self.files.mkdir()
        self.md = self.files / "peca.md"
        self.docx = self.files / "peca.docx"
        self.pdf = self.files / "peca.pdf"
        self.email = self.files / "EMAIL_RESPOSTA.txt"
        self.f7 = self.files / "F7.json"
        self.sources = self.files / "sources.json"
        self.f8 = self.files / "F8.json"
        self.context = self.files / "CONTEXT_VALIDATION.json"
        self.fidelity = self.files / "FORMAT_FIDELITY.json"
        body = (
            "O fundamento jurídico central permanece integralmente preservado nesta peça de teste, "
            "com extensão suficiente para validar o corpo do documento, sua fidelidade material e a "
            "diagramação profissional exigida antes da formação do pacote protocolável. "
            "A conferência é repetida sobre os arquivos reais e não aceita um resultado declarado pelo produtor."
        )
        self.md.write_text("# PEÇA\n\n" + body + "\n", encoding="utf-8")
        document = Document()
        document.add_heading("PEÇA", level=1)
        paragraph = document.add_paragraph(body)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.runs[0].font.name = "Times New Roman"
        paragraph.runs[0].font.size = Pt(12)
        document.save(self.docx)
        import fitz
        pdf_document = fitz.open()
        page = pdf_document.new_page()
        page.insert_textbox(fitz.Rect(72, 72, 520, 500), body, fontsize=11)
        pdf_document.save(self.pdf)
        pdf_document.close()
        self.email.write_text("Segue a peça com fontes conferidas.", encoding="utf-8")
        atomic_write_json(self.f7, {"p0": 0, "p1": 0, "mdSha256": sha256_file(self.md), "citacoesNaoConferidas": [], "verificarRestantes": []})
        atomic_write_json(self.sources, {"entries": []})
        pages_dir = self.files / "pages"
        rendered = inspect_pdf(
            self.pdf,
            pages_dir,
            generator_run_id="visual-generator",
            reviewer_run_id="visual-auditor",
        )["pages"]
        self.visual_review = self.files / "VISUAL_REVIEW.json"
        review = build_pending_review(
            self.visual_review,
            pdf=self.pdf,
            docx=self.docx,
            rendered_pages=rendered,
            generator_run_id="visual-generator",
        )
        review["reviewedAt"] = now_iso()
        review["reviewer"] = {"id": "package-visual-reviewer", "runId": "visual-auditor", "type": "agent_visual"}
        review["reviewMethod"] = "page_by_page_at_100_percent"
        review["approved"] = True
        for page_review in review["pages"]:
            page_review["status"] = "pass"
            page_review["checks"] = {name: True for name in REQUIRED_PAGE_CHECKS}
        atomic_write_json(self.visual_review, review)
        visual_result = run_visual_qa(
            self.pdf,
            self.f8,
            qa_dir=pages_dir,
            generator_run_id="visual-generator",
            reviewer_run_id="visual-auditor",
            docx=self.docx,
            manual_review=self.visual_review,
        )
        self.assertTrue(visual_result["approved"], visual_result["findings"])
        private_key = Ed25519PrivateKey.generate()
        public_raw = private_key.public_key().public_bytes(
            encoding=serialization.Encoding.Raw,
            format=serialization.PublicFormat.Raw,
        )
        key_id = public_key_id(public_raw)
        self.human_trust_store = self.files / "HUMAN_REVIEW_TRUST.json"
        atomic_write_json(self.human_trust_store, {
            "schemaVersion": 1,
            "reviewers": [{
                "reviewerId": "package-human-reviewer",
                "publicKeyId": key_id,
                "algorithm": "Ed25519",
                "publicKeyBase64": base64.b64encode(public_raw).decode("ascii"),
                "enabled": True,
            }],
        })
        self.human_trust_pin = self.files / "HUMAN_REVIEW_TRUST_PIN.json"
        atomic_write_json(self.human_trust_pin, {
            "schemaVersion": 1,
            "trustStorePath": str(self.human_trust_store),
            "trustStoreSha256": sha256_file(self.human_trust_store),
            "status": "configured",
        })
        self._trust_patchers = [
            patch("forja_human_review.DEFAULT_TRUST_STORE", self.human_trust_store),
            patch("forja_human_review.TRUST_STORE_PIN_PATH", self.human_trust_pin),
        ]
        for patcher in self._trust_patchers:
            patcher.start()
        human_receipt = build_unsigned_visual_receipt(
            reviewer_id="package-human-reviewer",
            reviewed_at=now_iso(),
            public_key_id_value=key_id,
            generator_run_id="visual-generator",
            reviewer_run_id="visual-auditor",
            pdf_sha256=sha256_file(self.pdf),
            docx_sha256=sha256_file(self.docx),
            page_count=visual_result["pageCount"],
            page_image_sha256=[page["imageSha256"] for page in visual_result["pages"]],
            required_checks=list(REQUIRED_PAGE_CHECKS),
            visual_attestation_sha256=sha256_file(self.visual_review),
        )
        human_receipt["signatureBase64"] = base64.b64encode(
            private_key.sign(canonical_receipt_bytes(human_receipt))
        ).decode("ascii")
        self.human_visual_receipt = self.files / "HUMAN_VISUAL_REVIEW.json"
        atomic_write_json(self.human_visual_receipt, human_receipt)
        visual_result["humanVisualReceipt"] = {
            "receiptPath": str(self.human_visual_receipt),
            "receiptSha256": sha256_file(self.human_visual_receipt),
        }
        atomic_write_json(self.f8, visual_result)
        atomic_write_json(self.context, {
            "approved": True,
            "p0": 0,
            "p1": 0,
            "markdown": {"sha256": sha256_file(self.md)},
        })
        fidelity_result = write_fidelity(self.md, self.docx, self.pdf, self.fidelity)
        self.assertTrue(fidelity_result["approved"], fidelity_result["findings"])
        for artifact_id, path in {
            "draft_markdown": self.md,
            "docx": self.docx,
            "pdf": self.pdf,
            "email_response": self.email,
            "f7_gate_result": self.f7,
            "verified_source_ledger": self.sources,
            "visual_qa_ledger": self.f8,
            "context_validation": self.context,
            "format_fidelity": self.fidelity,
        }.items():
            _, state, _ = record_event(
                self.case,
                "artifact_promoted",
                expected_revision=self.revision,
                idempotency_key=f"artifact:{artifact_id}",
                phase=PHASES[9],
                payload={"artifactId": artifact_id, "artifact": {"path": str(path), "sha256": sha256_file(path), "size": path.stat().st_size}},
            )
            self.revision = state["revision"]
        self.definition = self.case / "PACKAGE_DEFINITION.json"
        atomic_write_json(self.definition, {
            "caseId": self.case.name,
            "runId": "run-package",
            "emailArtifactId": "email_response",
            "deliverables": [{
                "id": "main",
                "role": "protocolavel",
                "audience": "office_review",
                "releasePolicy": "strict_protocol",
                "mdArtifactId": "draft_markdown",
                "docxArtifactId": "docx",
                "pdfArtifactId": "pdf",
                "f7ArtifactId": "f7_gate_result",
                "sourceLedgerArtifactId": "verified_source_ledger",
                "contextArtifactId": "context_validation",
                "f8ArtifactId": "visual_qa_ledger",
                "fidelityArtifactId": "format_fidelity",
                "attachKinds": ["docx", "pdf"]
            }]
        })

    def tearDown(self) -> None:
        for patcher in reversed(self._trust_patchers):
            patcher.stop()
        self.temp.cleanup()

    def test_full_close_cycle_is_hash_bound(self) -> None:
        package_result = create_package(self.case, self.definition, expected_revision=self.revision)
        manifest = package_result["manifest"]
        state = package_result["state"]
        self.assertEqual("ready_for_review", state["lifecycleStatus"])
        self.assertEqual(0, manifest["emailStyle"]["p0"])
        receipt = self.case / "draft_receipt.json"
        atomic_write_json(receipt, {
            "draftId": "draft-1",
            "threadId": "thread-1",
            "bodySha256": sha256_file(self.email),
            "attachments": [{"artifactId": item["artifactId"], "sha256": item["sha256"], "size": item["size"]} for item in manifest["attachments"]],
        })
        state = register_draft(self.case, receipt, expected_revision=state["revision"])
        self.assertEqual("draft_awaiting_review", state["lifecycleStatus"])
        evidence = self.case / "evidence.json"
        atomic_write_json(evidence, {"type": "email", "externalId": "sent-message-1", "packageHash": manifest["packageHash"]})
        state = confirm_delivery(self.case, evidence, expected_revision=state["revision"])
        metrics = self.case / "FORJA_RUN_METRICS.json"
        retrospective = self.case / "FORJA_RETROSPECTIVE.md"
        atomic_write_json(metrics, {"status": "complete", "caseId": self.case.name})
        retrospective.write_text("# Retrospectiva\n\nNenhuma correção humana registrada neste teste controlado.\n", encoding="utf-8")
        for artifact_id, path in {"run_metrics": metrics, "retrospective": retrospective}.items():
            _, state, _ = record_event(
                self.case,
                "artifact_promoted",
                expected_revision=state["revision"],
                idempotency_key=f"artifact:f10:{artifact_id}",
                phase=PHASES[10],
                payload={"artifactId": artifact_id, "artifact": {"path": str(path), "sha256": sha256_file(path), "size": path.stat().st_size}},
            )
        delivery_revision = state["revision"]
        _, state, _ = record_event(
            self.case, "sync_succeeded", expected_revision=delivery_revision,
            idempotency_key="sync-delivery", payload={"syncedEventSeq": delivery_revision},
        )
        state = fulfill(self.case, expected_revision=state["revision"])
        self.assertEqual("fulfilled_by_forja_f10", state["lifecycleStatus"])
        self.assertEqual(
            {name: "pass" for name in REQUIRED_F10_GATES},
            state["deliveryEvidence"]["gates"],
        )

    def test_f10_recomputes_hash_and_sync_instead_of_accepting_declaration(self) -> None:
        state = {"revision": 8, "sync": {"status": "ok", "lastSyncedEventSeq": 7}}
        gates = compute_f10_gates(
            {"packageHash": "package-real"},
            {"packageHash": "package-forjado", "externalId": "sent-1"},
            state,
            minimum_synced_event_seq=7,
        )
        self.assertEqual("pass", gates["external_identifier_valid"])
        self.assertEqual("fail", gates["package_hash_matches"])
        self.assertEqual("pass", gates["management_synced"])
        result = validate_f10_gates(gates)
        self.assertFalse(result["approved"])
        self.assertTrue(any("package_hash_matches" in item for item in result["findings"]))

    def test_draft_receipt_requires_exact_artifact_ids(self) -> None:
        package_result = create_package(self.case, self.definition, expected_revision=self.revision)
        manifest = package_result["manifest"]
        receipt = self.case / "draft_receipt_missing_ids.json"
        atomic_write_json(receipt, {
            "draftId": "draft-without-ids",
            "threadId": "thread-1",
            "bodySha256": sha256_file(self.email),
            "attachments": [{"sha256": item["sha256"], "size": item["size"]} for item in manifest["attachments"]],
        })
        with self.assertRaises(ForjaN3Error):
            register_draft(self.case, receipt, expected_revision=package_result["state"]["revision"])

    def test_draft_rejects_tampered_package_pointer(self) -> None:
        package_result = create_package(self.case, self.definition, expected_revision=self.revision)
        manifest = package_result["manifest"]
        pointer = self.case / "FORJA_PACKAGE.json"
        changed = read_json(pointer)
        changed["packageHash"] = "hash-adulterado"
        atomic_write_json(pointer, changed)
        receipt = self.case / "draft_receipt_tampered.json"
        atomic_write_json(receipt, {
            "draftId": "draft-tampered",
            "threadId": "thread-1",
            "bodySha256": sha256_file(self.email),
            "attachments": [
                {"artifactId": item["artifactId"], "sha256": item["sha256"], "size": item["size"]}
                for item in manifest["attachments"]
            ],
        })
        with self.assertRaises(ForjaN3Error):
            register_draft(self.case, receipt, expected_revision=package_result["state"]["revision"])

    def test_package_blocks_email_with_ai_writing_vices(self) -> None:
        self.email.write_text(
            "Prezado Fábio,\n\nEspero que este e-mail o encontre bem. Gostaria de informar que a peça está pronta.\n\n"
            "Permaneço à disposição para quaisquer esclarecimentos adicionais.",
            encoding="utf-8",
        )
        state = derive_state(self.case)
        _, state, _ = record_event(
            self.case,
            "artifact_promoted",
            expected_revision=state["revision"],
            idempotency_key="artifact:email:ai-style",
            phase=PHASES[9],
            payload={
                "artifactId": "email_response",
                "artifact": {
                    "path": str(self.email),
                    "sha256": sha256_file(self.email),
                    "size": self.email.stat().st_size,
                },
            },
        )
        result = validate_definition(self.case, atomic_write_and_read(self.definition))
        self.assertFalse(result["approved"])
        self.assertGreater(result["emailStyle"]["p0"], 0)
        self.assertTrue(any("e-mail reprovado" in item for item in result["findings"]))

    def test_draft_requires_hash_of_approved_email_body(self) -> None:
        package_result = create_package(self.case, self.definition, expected_revision=self.revision)
        manifest = package_result["manifest"]
        receipt = self.case / "draft_receipt_without_body_hash.json"
        atomic_write_json(receipt, {
            "draftId": "draft-no-body-hash",
            "threadId": "thread-1",
            "attachments": [
                {"artifactId": item["artifactId"], "sha256": item["sha256"], "size": item["size"]}
                for item in manifest["attachments"]
            ],
        })
        with self.assertRaisesRegex(ForjaN3Error, "hash do corpo"):
            register_draft(self.case, receipt, expected_revision=package_result["state"]["revision"])

    def test_changed_email_after_package_cannot_be_registered(self) -> None:
        package_result = create_package(self.case, self.definition, expected_revision=self.revision)
        manifest = package_result["manifest"]
        self.email.write_text("Corpo alterado depois da aprovação.", encoding="utf-8")
        receipt = self.case / "draft_receipt_changed_body.json"
        atomic_write_json(receipt, {
            "draftId": "draft-changed-body",
            "threadId": "thread-1",
            "bodySha256": sha256_file(self.email),
            "attachments": [
                {"artifactId": item["artifactId"], "sha256": item["sha256"], "size": item["size"]}
                for item in manifest["attachments"]
            ],
        })
        with self.assertRaisesRegex(ForjaN3Error, "diverge do pacote"):
            register_draft(self.case, receipt, expected_revision=package_result["state"]["revision"])

    def test_revision_conflict_does_not_replace_package_pointer(self) -> None:
        pointer = self.case / "FORJA_PACKAGE.json"
        atomic_write_json(pointer, {"packageId": "pkg-anterior"})
        with patch("forja_close_cycle.record_event", side_effect=RevisionConflict("conflito")):
            with self.assertRaises(RevisionConflict):
                create_package(self.case, self.definition, expected_revision=self.revision)
        self.assertEqual("pkg-anterior", read_json(pointer)["packageId"])

    def test_strict_package_blocks_unverified_citation(self) -> None:
        atomic_write_json(self.f7, {"p0": 0, "mdSha256": sha256_file(self.md), "citacoesNaoConferidas": ["TEMA 1365"], "verificarRestantes": []})
        state = derive_state(self.case)
        artifact = state["artifacts"]["f7_gate_result"]
        artifact["sha256"] = sha256_file(self.f7)
        # Record the updated version instead of mutating canonical state.
        _, state, _ = record_event(
            self.case, "artifact_promoted", expected_revision=state["revision"],
            idempotency_key="artifact:f7:unverified", phase=PHASES[9],
            payload={"artifactId": "f7_gate_result", "artifact": artifact},
        )
        result = validate_definition(self.case, atomic_write_and_read(self.definition))
        self.assertFalse(result["approved"])
        self.assertTrue(any("TEMA 1365" in item for item in result["findings"]))

    def test_response_package_requires_adversarial_audit_bundle(self) -> None:
        definition = read_json(self.definition)
        definition["deliverables"][0]["adversarialResponse"] = True
        result = validate_definition(self.case, definition)
        self.assertFalse(result["approved"])
        self.assertTrue(any("sem auditoria" in item for item in result["findings"]))

    def test_updated_markdown_invalidates_previous_f7(self) -> None:
        self.md.write_text("# Peça alterada\n", encoding="utf-8")
        state = derive_state(self.case)
        _, state, _ = record_event(
            self.case,
            "artifact_promoted",
            expected_revision=state["revision"],
            idempotency_key="artifact:md:updated",
            phase=PHASES[9],
            payload={"artifactId": "draft_markdown", "artifact": {"path": str(self.md), "sha256": sha256_file(self.md), "size": self.md.stat().st_size}},
        )
        result = validate_definition(self.case, atomic_write_and_read(self.definition))
        self.assertFalse(result["approved"])
        self.assertTrue(any("F7 foi calculado" in item for item in result["findings"]))

    def test_visual_ledger_requires_complete_run_identity_and_page_sequence(self) -> None:
        atomic_write_json(self.f8, {
            "pdfSha256": sha256_file(self.pdf),
            "pageCount": 1,
            "generatorRunId": "visual-generator",
            "pages": [{"page": 2, "imageSha256": "image-hash", "lint": "pass", "independentReview": {"status": "pass", "runId": "visual-auditor"}}],
            "approved": True,
        })
        state = derive_state(self.case)
        _, state, _ = record_event(
            self.case,
            "artifact_promoted",
            expected_revision=state["revision"],
            idempotency_key="artifact:f8:malformed",
            phase=PHASES[9],
            payload={"artifactId": "visual_qa_ledger", "artifact": {"path": str(self.f8), "sha256": sha256_file(self.f8), "size": self.f8.stat().st_size}},
        )
        result = validate_definition(self.case, atomic_write_and_read(self.definition))
        self.assertFalse(result["approved"])
        self.assertTrue(any("não identifica gerador e revisor" in item for item in result["findings"]))
        self.assertTrue(any("ausentes, duplicadas ou fora de ordem" in item for item in result["findings"]))


def atomic_write_and_read(path: Path) -> dict:
    import json
    return json.loads(path.read_text(encoding="utf-8"))


if __name__ == "__main__":
    unittest.main(verbosity=2)
