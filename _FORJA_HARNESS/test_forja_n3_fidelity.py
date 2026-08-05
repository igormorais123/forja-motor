import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document

from forja_fidelity import compare_fidelity


class FidelityTests(unittest.TestCase):
    def make_files(self, root: Path, *, pdf_text: str):
        source_text = (
            "A pretensão não pode ser acolhida, porque o valor comprovado é R$ 12.500,00 "
            "e apenas subsidiariamente se admite a revisão em 18/07/2026."
        )
        markdown = root / "peça.md"
        docx = root / "peça.docx"
        pdf = root / "peça.pdf"
        markdown.write_text("# TESTE\n\n" + source_text + "\n", encoding="utf-8")
        document = Document()
        document.add_heading("TESTE", level=1)
        document.add_paragraph(source_text)
        document.save(docx)
        pdf_document = fitz.open()
        page = pdf_document.new_page()
        page.insert_textbox(fitz.Rect(72, 72, 520, 300), pdf_text, fontsize=11)
        pdf_document.save(pdf)
        pdf_document.close()
        return markdown, docx, pdf, source_text

    def test_exact_semantic_content_passes(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source_text = (
                "A pretensão não pode ser acolhida, porque o valor comprovado é R$ 12.500,00 "
                "e apenas subsidiariamente se admite a revisão em 18/07/2026."
            )
            markdown, docx, pdf, _ = self.make_files(root, pdf_text=source_text)
            result = compare_fidelity(markdown, docx, pdf)
            self.assertTrue(result["approved"], result["findings"])

    def test_lost_negation_and_number_are_blocked(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            changed = (
                "A pretensão pode ser acolhida, porque o valor comprovado é R$ 15.500,00 "
                "e subsidiariamente se admite a revisão em 18/07/2026."
            )
            markdown, docx, pdf, _ = self.make_files(root, pdf_text=changed)
            result = compare_fidelity(markdown, docx, pdf)
            codes = {item["code"] for item in result["findings"]}
            self.assertFalse(result["approved"])
            self.assertIn("number_missing_in_pdf", codes)
            self.assertIn("qualifier_missing_in_pdf", codes)

    def test_reconstructed_structural_paragraph_number_does_not_break_fidelity(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            sentence = "O pedido deve ser acolhido porque a premissa comprovada conduz diretamente à consequência jurídica indicada."
            markdown = root / "source.md"
            docx = root / "source.docx"
            pdf = root / "source.pdf"
            markdown.write_text("# TESTE\n\n20. " + sentence + "\n", encoding="utf-8")
            document = Document()
            document.add_heading("TESTE", level=1)
            document.add_paragraph("19. " + sentence)
            document.save(docx)
            pdf_document = fitz.open()
            page = pdf_document.new_page()
            page.insert_textbox(fitz.Rect(72, 72, 520, 300), "19. " + sentence, fontsize=11)
            pdf_document.save(pdf)
            pdf_document.close()
            result = compare_fidelity(markdown, docx, pdf)
            self.assertTrue(result["approved"], result["findings"])


if __name__ == "__main__":
    unittest.main()
