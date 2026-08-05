"""Regressão dos 3 defeitos de ferramenta da Lição 41 (RETROSPECTIVAS.md).

(a) forja_render_docx.add_runs_com_negrito: *itálico* virava asterisco literal no DOCX;
(b) extrator de citações: "Tema 1.365" era lido como "Tema 1" (ponto de milhar);
(c) forja_metricas_f7: súmulas nunca eram localizadas no cache (faltava a chave SUMULA).

Rodar após qualquer mudança nesses módulos: python test_licao41.py
"""

import sys
from pathlib import Path

FORJA = Path(__file__).resolve().parent
sys.path.insert(0, str(FORJA))

falhas = []


def caso(nome, cond, detalhe=""):
    status = "OK  " if cond else "FALHOU"
    print(f"[{status}] {nome}" + (f" — {detalhe}" if detalhe and not cond else ""))
    if not cond:
        falhas.append(nome)


# ---------------------------------------------------------------- (a) itálico
from docx import Document
from forja_render_docx import add_runs_com_negrito

doc = Document()
p = doc.add_paragraph()
add_runs_com_negrito(p, "Texto com *ênfase em itálico* e **peso em negrito** no fim.")
texto_final = "".join(r.text for r in p.runs)
caso("(a) nenhum asterisco literal sobra no parágrafo", "*" not in texto_final, texto_final)
caso("(a) trecho em itálico ganha run italic",
     any(r.italic and r.text == "ênfase em itálico" for r in p.runs))
caso("(a) trecho em negrito segue com run bold",
     any(r.bold and r.text == "peso em negrito" for r in p.runs))

p2 = doc.add_paragraph()
add_runs_com_negrito(p2, "Multiplicar 3 * 4 * 5 não é itálico.")
caso("(a) não-trava: asterisco de multiplicação com espaços não vira itálico",
     not any(r.italic for r in p2.runs))

# ------------------------------------------------------- (b) tema com milhar
from forja_citations import extrair_citacoes, procurar_cache_oficial

texto = ("A tese do Tema Repetitivo nº 1.365 do STJ e a do Tema 1368 do STJ convivem; "
         "registre-se também o Tema 1.368 (grafia com ponto).")
cits = extrair_citacoes(texto)
temas = {c["rotulo"]: c for c in cits if c["tipo"] == "TEMA"}
caso("(b) 'Tema 1.365' extraído com o número completo",
     any("1.365" in r or "1365" in r.replace(".", "") for r in temas), str(list(temas)))
caso("(b) nenhum tema extraído como 'Tema 1' isolado",
     not any(r.strip() in ("Tema 1", "Tema 1 STJ") for r in temas), str(list(temas)))
caso("(b) 'Tema 1368' e 'Tema 1.368' deduplicam para uma só citação",
     sum(1 for c in cits if c["tipo"] == "TEMA" and "1368" in c["rotulo"].replace(".", "")) == 1,
     str(list(temas)))
c1368 = next((c for c in cits if c["tipo"] == "TEMA" and "1368" in c["rotulo"].replace(".", "")), None)
caso("(b) Tema 1368 (mesmo grafado com ponto) localiza STJ_TEMA_1368 no cache",
     c1368 is not None and procurar_cache_oficial(c1368, require_live=False) is not None)

# ------------------------------------------------- (c) súmulas no metricas_f7
from forja_metricas_f7 import metricas_f7

md = ("Aplica-se a Súmula 608 do STJ e a Súmula 609 do STJ; ver ainda o Tema 1.368 do STJ "
      "e a Súmula 999 do STJ (sem fonte no cache).")
m = metricas_f7(md, require_live=False)
conferidas = set(m.get("citacoesConferidasRotulos") or [])
nao = set(m.get("citacoesNaoConferidas") or [])
caso("(c) Súmula 608 conferida no cache", any("608" in r for r in conferidas), str(conferidas))
caso("(c) Súmula 609 conferida no cache", any("609" in r for r in conferidas), str(conferidas))
caso("(c) Tema 1.368 conferido no cache (regex de milhar na cópia local)",
     any("1368" in r.replace(".", "") for r in conferidas), str(conferidas))
caso("(c) não-trava: Súmula 999 sem fonte segue como não conferida",
     any("999" in r for r in nao), str(nao))

# ------------------------------------------- (d) kit visual — Lição 43
# quadro_zebrado/caixas do medina_visual_kit inseriam texto cru: **negrito**
# vazava asterisco literal em peça ENTREGUE (achado do teste real 09/07/2026).
import tempfile
sys.path.insert(0, str(FORJA.parent / "_FERRAMENTAS"))
from medina_visual_kit import PecaVisual

with tempfile.TemporaryDirectory() as tmp:
    alvo = str(Path(tmp) / "kit_regressao.docx")
    pv = PecaVisual(alvo)
    pv.quadro_zebrado(["Óbice", "Alcance"],
                      [["**Súmula 182/STJ**", "Responsabilidade **extracontratual** direta"]])
    pv.caixa_acordao("**Tese**", "Conteúdo com **peso** e *ênfase*.")
    pv.doc.save(alvo)
    d = Document(alvo)
    txt_kit = "\n".join(p.text for t in d.tables for row in t.rows
                        for cell in row.cells for p in cell.paragraphs)
    caso("(d) kit visual: zero asterisco literal em quadro/caixa", "*" not in txt_kit, txt_kit[:120])
    runs_kit = [r for t in d.tables for row in t.rows for cell in row.cells
                for p in cell.paragraphs for r in p.runs]
    caso("(d) kit visual: **negrito** de célula vira run bold",
         any(r.bold and r.text == "Súmula 182/STJ" for r in runs_kit))
    caso("(d) kit visual: *itálico* de caixa vira run italic",
         any(r.italic and r.text == "ênfase" for r in runs_kit))

# -------------------------- (e) corte é identidade — review adversarial 09/07
# Achado high do Codex: dedupe por (tipo, numero) fundia Súmula 7/STJ com 7/STF
# e o cache do tribunal errado validava a citação (falso-verde de autoridade).
import forja_metricas_f7 as fm7

texto_cortes = "A Súmula 7 do STJ e a Súmula 7 do STF são autoridades distintas; ver também a Súmula 7."
cits_cortes = fm7.extrair_citacoes_basico(texto_cortes)
sumulas7 = [c for c in cits_cortes if c["tipo"] == "SUMULA" and c["numero"] == "7"]
caso("(e) Súmula 7/STJ e 7/STF NÃO fundem (citações distintas)",
     len([c for c in sumulas7 if c["corte"]]) == 2, str([c["rótulo"] for c in sumulas7]))
caso("(e) menção sem tribunal com DUAS cortes qualificadas fica separada (ambígua)",
     any(c["corte"] is None for c in sumulas7), str([c["rótulo"] for c in sumulas7]))

def lastro_sumula(corte, numero):
    host = "scon.stj.jus.br" if corte == "STJ" else "portal.stf.jus.br"
    return (
        f"Fonte: https://{host}/sumulas/{numero}\n"
        f"Capturado em: 10/07/2026\n{corte} Súmula {numero}.\n"
        + f"Texto integral da Súmula {numero} do {corte} capturado para conferência verbatim material. " * 3
    )

with tempfile.TemporaryDirectory() as tmp:
    cache_fake = Path(tmp) / "cache" / "fontes_oficiais"
    cache_fake.mkdir(parents=True)
    (cache_fake / "STJ_SUMULA_99.txt").write_text(lastro_sumula("STJ", "99"), encoding="utf-8")
    (cache_fake / "STF_SUMULA_99.txt").write_text(lastro_sumula("STF", "99"), encoding="utf-8")
    from forja_official_sources import build_manifest
    build_manifest(cache_fake, cache_fake / "OFFICIAL_SOURCE_MANIFEST.json")
    _orig = fm7.FORJA
    fm7.FORJA = Path(tmp)
    try:
        m_amb = fm7.metricas_f7("Aplica-se a Súmula 99.", require_live=False)
        caso("(e) sem tribunal + cache dos DOIS tribunais = ambíguo, NÃO conferida",
             m_amb["citacoesConferidasEmFonte"] == 0, m_amb["citacoesConferidasRotulos"])
        m_stf = fm7.metricas_f7("Aplica-se a Súmula 99 do STF.", require_live=False)
        caso("(e) com tribunal declarado, confere só no cache daquele tribunal",
             m_stf["citacoesConferidasEmFonte"] == 1 and "STF" in (m_stf["citacoesConferidasRotulos"] or [""])[0],
             str(m_stf["citacoesConferidasRotulos"]))
        (cache_fake / "STF_SUMULA_99.txt").unlink()
        m_um = fm7.metricas_f7("Aplica-se a Súmula 99.", require_live=False)
        caso("(e) sem tribunal mas cache de UM só tribunal = conferida (sem ambiguidade)",
             m_um["citacoesConferidasEmFonte"] == 1, str(m_um["citacoesNaoConferidas"]))
        m_errada = fm7.metricas_f7("Aplica-se a Súmula 99 do STF.", require_live=False)
        caso("(e) tribunal declarado ERRADO para o cache existente = não conferida",
             m_errada["citacoesConferidasEmFonte"] == 0, str(m_errada["citacoesConferidasRotulos"]))
    finally:
        fm7.FORJA = _orig

# ------------------------- (f) F7 fail-closed no render — review adversarial
# Achado high do Codex: exceção do gate era engolida e o render entregava DOCX/PDF
# sem F7. Agora o gate roda ANTES de qualquer artefato e a falha aborta o render.
import forja_render_docx as frd

_verificador_real = sys.modules.get("forja_verificador")
class _VerificadorQuebrado:
    @staticmethod
    def verificar(*a, **k):
        raise RuntimeError("verificador sabotado pelo teste")
sys.modules["forja_verificador"] = _VerificadorQuebrado
try:
    with tempfile.TemporaryDirectory() as tmp:
        md_teste = Path(tmp) / "peca_teste.md"
        md_teste.write_text("# Teste\n\nParágrafo simples.", encoding="utf-8")
        out_teste = Path(tmp) / "build"
        explodiu = False
        try:
            frd.render(md_teste, out_teste, "teste fail-closed")
        except RuntimeError:
            explodiu = True
        caso("(f) falha do verificador ABORTA o render (fail-closed)", explodiu)
        caso("(f) nenhum DOCX nasce quando o gate falha",
             not list(out_teste.glob("*.docx")) if out_teste.exists() else True)
        caso("(f) nenhum F7 falso é persistido quando o gate falha",
             not (out_teste / "F7_VERIFICADOR_FORJA.json").exists() if out_teste.exists() else True)
finally:
    if _verificador_real is not None:
        sys.modules["forja_verificador"] = _verificador_real
    else:
        sys.modules.pop("forja_verificador", None)

print()
if falhas:
    print(f"FALHOU: {len(falhas)} caso(s): {falhas}")
    sys.exit(1)
print("OK: todos os casos da Lição 41 passam")
