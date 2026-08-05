# -*- coding: utf-8 -*-
"""Regressão do gate F8-S — assinatura visual.

Teste de mutação (exigência da auditoria do Igor, 30/07/2026): retirar um
elemento obrigatório de um artefato válido tem de fazer o gate acusar. Um gate
que aprova tudo é indistinguível de gate nenhum — e foi assim que as peças
pobres saíram com sinal verde entre 10 e 30 de julho.
"""
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

import forja_acervo

FORJA = Path(__file__).resolve().parent
sys.path.insert(0, str(FORJA))

from forja_assinatura_visual import _faixa, avaliar  # noqa: E402

REFERENCIA = (FORJA / "state" / forja_acervo.caso("CASO-16")
              / "producao" / "_visual"
              / "MEMORIAIS_LIBRA_SUL_AGINT_ARESP_2578181_VISUAL_LAW.docx")


def _mutar(origem, destino, transformar=None, remover=()):
    """Copia o DOCX aplicando uma mutação no document.xml ou removendo mídia."""
    with zipfile.ZipFile(origem) as z:
        itens = {n: z.read(n) for n in z.namelist()}
    if transformar:
        itens["word/document.xml"] = transformar(
            itens["word/document.xml"].decode("utf-8")).encode("utf-8")
    for padrao in remover:
        for nome in [n for n in list(itens) if re.search(padrao, n)]:
            del itens[nome]
    with zipfile.ZipFile(destino, "w", zipfile.ZIP_DEFLATED) as z:
        for nome, dados in itens.items():
            z.writestr(nome, dados)
    return destino


def test_faixa_por_extensao():
    """Densidade calibrada: peça curta e longa não têm a mesma régua."""
    assert _faixa(4)[0] < _faixa(20)[0]
    assert _faixa(4)[1] < _faixa(30)[1]


def test_referencia_aprovada_e_conforme():
    """A peça que o escritório aprovou em 09/07/2026 tem de passar.

    Âncora contra deriva do gate: se um limiar for endurecido a ponto de
    reprovar o padrão aprovado pelo próprio dono, o gate está errado, não a
    peça.

    RESSALVA (03/08/2026, conferência do Igor): esta peça é âncora de
    ESTRUTURA, não exemplo de perfeição. A Figura 1 dela tem colisão — a caixa
    central cobre o texto das duas caixas inferiores — e a Figura 2 lista
    "2005-2007" antes de "27/05/2005". Ambas foram desenhadas à mão no mapa
    manual de julho. O gate as aprova porque verifica PRESENÇA, não correção do
    desenho: é o limite declarado do instrumento, e a razão de o QA visual
    página a página continuar obrigatório."""
    if not REFERENCIA.exists():
        return
    laudo = avaliar(REFERENCIA, paginas=9)
    assert laudo["conforme"], f"referência aprovada reprovou: {laudo['achados']}"


def test_mutacoes_sao_detectadas():
    """Cada elemento retirado tem de produzir um achado com o código certo."""
    if not REFERENCIA.exists():
        return
    tmp = Path(tempfile.mkdtemp())
    try:
        casos = [
            ("VIS-03", "sem elemento gráfico vetorial",
             dict(remover=(r"word/media/.*\.(emf|wmf)$",))),
            ("VIS-04", "sem destaques de varredura",
             # NB: substituir "<w:framePr" por "<w:framePrX" não remove nada —
             # o primeiro é prefixo do segundo e o contador continua achando.
             dict(transformar=lambda x: re.sub(r"<w:tcBorders>.*?</w:tcBorders>", "", x, flags=re.S)
                  .replace("<w:framePr", "<w:semMoldura"))),
            ("VIS-11", "sem quadros estruturados",
             dict(transformar=lambda x: re.sub(r"<w:tbl>.*?</w:tbl>", "", x, flags=re.S))),
        ]
        for codigo, descricao, kw in casos:
            alvo = _mutar(REFERENCIA, tmp / f"{codigo}.docx", **kw)
            laudo = avaliar(alvo, paginas=9)
            codigos = {a["codigo"] for a in laudo["achados"]}
            assert not laudo["conforme"], f"mutação não detectada: {descricao}"
            assert codigo in codigos, (
                f"{descricao}: esperado {codigo}, obtido {sorted(codigos)}")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_negrito_universal_reprova():
    """Documento inteiro em negrito é ausência de ênfase, não riqueza.

    A revisão de engenharia propôs remover o negrito do padrão alegando
    detecção frágil. Rejeitado: documento com 100% de negrito reprovando é o
    gate funcionando. Este teste fixa a decisão."""
    if not REFERENCIA.exists():
        return
    tmp = Path(tempfile.mkdtemp())
    try:
        alvo = _mutar(REFERENCIA, tmp / "negrito.docx",
                      transformar=lambda x: x.replace("<w:rPr>", "<w:rPr><w:b/>"))
        laudo = avaliar(alvo, paginas=9)
        assert not laudo["conforme"]
        assert laudo["inventario"]["razaoNegrito"] > 0.20
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_docx_fora_do_template_reprova():
    """Peca nascida de Document() vazio — proibida pelo protocolo da casa — tem
    de reprovar por falta de timbre.

    Controle negativo do gate: sem ele, a checagem de timbre passava com
    qualquer imagem no pacote, inclusive um diagrama do corpo (achado da rodada
    2 da revisao cruzada Codex, 03/08/2026)."""
    import tempfile
    from docx import Document
    tmp = Path(tempfile.mkdtemp())
    try:
        alvo = tmp / "sem_template.docx"
        doc = Document()
        doc.add_paragraph("Peca gerada fora do template do escritorio.")
        doc.save(alvo)
        laudo = avaliar(alvo, paginas=2)
        assert not laudo["inventario"]["timbre"], "timbre falso-positivo fora do template"
        assert "VIS-05" in {a["codigo"] for a in laudo["achados"]}
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

if __name__ == "__main__":
    for nome, fn in sorted(globals().items()):
        if nome.startswith("test_") and callable(fn):
            fn()
            print(f"ok {nome}")
    print("assinatura visual: regressão completa")
