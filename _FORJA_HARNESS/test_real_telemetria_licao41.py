"""Bateria REAL com telemetria — Lição 41 (nada de smoke).

Diferença para test_licao41.py (unitário/sintético): aqui TUDO roda sobre
artefatos reais de produção da fábrica, com a pipeline verdadeira:

  B1. Extrator de citações + métricas F7 sobre TODOS os fontes .md reais
      de state/*/producao/ (10 peças/estudos entregues), com invariantes:
      nenhum "Tema 1" espúrio, nenhuma citação duplicada por (tipo, número),
      súmulas conferidas quando o cache oficial tem o verbatim.
  B2. Gate negativo sobre fonte legado viciado + render ponta a ponta REAL
      (template oficial + Word COM + PDF + QA) de peças atuais aprovadas.
      O DIAGNOSTICO_CORSAN preserva o cenário real de 15 spans itálicos, mas
      desde o gate de escrita humana v1 deve ser barrado antes do Word por
      excesso de travessões/dogmatismo; não é mais corpus positivo limpo.
  B3. Varredura de TODOS os DOCX de produção já entregues: nenhum asterisco
      literal remanescente em parágrafo ou célula de tabela.

Saída: telemetria/TELEMETRIA_LICAO41_<data>.json + relatório no console.
Uso: python test_real_telemetria_licao41.py [--sem-render]
"""

import json
import re
import sys
import time
from datetime import datetime
from pathlib import Path

if __name__ != "__main__":
    import unittest

    raise unittest.SkipTest(
        "telemetria real standalone; executar python test_real_telemetria_licao41.py --sem-render"
    )

import forja_acervo  # noqa: E402

FORJA = Path(__file__).resolve().parent
sys.path.insert(0, str(FORJA))

from docx import Document

from forja_citations import extrair_citacoes
from forja_metricas_f7 import metricas_f7

STATE = FORJA / "state"
TELEDIR = FORJA / "telemetria"
TELEDIR.mkdir(exist_ok=True)

SEM_RENDER = "--sem-render" in sys.argv

ITALICO_RE = re.compile(r"(?<!\*)\*(\S(?:[^*\n]*\S)?)\*(?!\*)")
NEGRITO_RE = re.compile(r"\*\*([^*\n]+)\*\*")

tele = {
    "executadoEm": datetime.now().astimezone().isoformat(timespec="seconds"),
    "modo": "REAL (artefatos de produção, pipeline completa)" + (" — render pulado por flag" if SEM_RENDER else ""),
    "baterias": {},
}
falhas = []


def caso(nome, cond, detalhe=""):
    print(f"[{'OK    ' if cond else 'FALHOU'}] {nome}" + (f" — {detalhe}" if detalhe and not cond else ""))
    if not cond:
        falhas.append({"caso": nome, "detalhe": str(detalhe)[:300]})


def texto_docx(path):
    doc = Document(str(path))
    partes = [p.text for p in doc.paragraphs]
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes), doc


# ================================================================ B1
print("=" * 72)
print("B1 — Extração de citações + F7 sobre os fontes REAIS de produção")
print("=" * 72)

fontes = sorted(
    p for p in STATE.glob("*/producao/*.md")
    if p.name not in ("MAPA_IA.md", "F4_BLUEPRINT.md")
)
b1 = {"arquivos": [], "totais": {}}
t0 = time.perf_counter()
tot_cit = tot_conf = tot_nao = 0
for md in fontes:
    texto = md.read_text(encoding="utf-8")
    t1 = time.perf_counter()
    cits = extrair_citacoes(texto)
    # Esta bateria mede o acervo real hash-bound de forma determinística. O
    # replay HTTPS obrigatório de produção é exercitado, com fetcher controlado
    # e caso adversarial, em test_forja_anti_cheat.py; indisponibilidade de rede
    # jamais é convertida em aprovação externa.
    m = metricas_f7(texto, require_live=False)
    dur_ms = round((time.perf_counter() - t1) * 1000, 1)

    rotulos = [c["rotulo"] for c in cits]

    def num_de(c):
        # o extrator guarda o número em dados[0]; fallback: dígitos do rótulo
        bruto = (c.get("dados") or ("",))[0] or c["rotulo"]
        return re.sub(r"[^\d]", "", str(bruto))

    # invariante 1: nenhum tema truncado — número curto que é prefixo de um tema com milhar do fonte
    milhares = {m.replace(".", "") for m in re.findall(r"Tema\s+(?:Repetitivo\s+)?(?:n[oº.]?\s*)?(\d\.\d{3})", texto)}
    truncados = [c["rotulo"] for c in cits if c["tipo"] == "TEMA" and num_de(c)
                 and len(num_de(c)) <= 2 and any(m.startswith(num_de(c)) for m in milhares)]
    # invariante 2: sem duplicata por (tipo, número normalizado)
    chaves = [(c["tipo"], num_de(c)) for c in cits]
    dups = sorted({k for k in chaves if k[1] and chaves.count(k) > 1})

    reg = {
        "arquivo": str(md.relative_to(STATE)),
        "caracteres": len(texto),
        "duracaoMs": dur_ms,
        "citacoes": len(cits),
        "rotulos": rotulos,
        "f7ConferidasEmFonte": m["citacoesConferidasEmFonte"],
        "f7NaoConferidas": m["citacoesNaoConferidas"],
        "temasTruncados": truncados,
        "duplicatasTipoNumero": [f"{t}:{n}" for t, n in dups],
    }
    b1["arquivos"].append(reg)
    tot_cit += len(cits)
    tot_conf += m["citacoesConferidasEmFonte"]
    tot_nao += len(m["citacoesNaoConferidas"])
    print(f"  {md.parent.parent.name[:38]:<38} {md.name[:34]:<34} {len(cits):>2} cit | {m['citacoesConferidasEmFonte']} conferidas | {dur_ms} ms")
    caso(f"B1 {md.name}: nenhum tema truncado", not truncados, truncados)
    caso(f"B1 {md.name}: nenhuma citação duplicada por (tipo, número)", not dups, dups)

b1["totais"] = {
    "arquivos": len(fontes),
    "citacoes": tot_cit,
    "conferidasEmFonte": tot_conf,
    "naoConferidas": tot_nao,
    "duracaoTotalMs": round((time.perf_counter() - t0) * 1000, 1),
}
caso("B1 volume real processado (>= 8 fontes, >= 20 citações)", len(fontes) >= 8 and tot_cit >= 20,
     f"{len(fontes)} fontes, {tot_cit} citações")

# asserções dirigidas nos fontes conhecidos (dados reais, não sintéticos)
estudo_ps = next((r for r in b1["arquivos"] if r["arquivo"].endswith("ESTUDO_ESTRATEGICO_PLANO_SAUDE.md")), None)
if estudo_ps:
    caso("B1 estudo Plano de Saúde: Tema 1365 extraído completo",
         any("1365" in r.replace(".", "") for r in estudo_ps["rotulos"]), estudo_ps["rotulos"])
    caso("B1 estudo Plano de Saúde: Súmulas 608 e 609 no cache real hash-bound",
         estudo_ps["f7ConferidasEmFonte"] >= 2,
         f"conferidas={estudo_ps['f7ConferidasEmFonte']} nao={estudo_ps['f7NaoConferidas']}")
tele["baterias"]["B1_extracao_fontes_reais"] = b1

# ================================================================ B2
print()
print("=" * 72)
print("B2 — Render REAL ponta a ponta (Word COM + PDF + QA) de peças reais")
print("=" * 72)

NEGATIVE_STYLE_TARGETS = [
    STATE / forja_acervo.caso("CASO-07") / "producao" / "DIAGNOSTICO_CORSAN_AGERST.md",
]
RENDER_ALVOS = [
    (
        forja_acervo.caminho("render-alvo-impugnacao-v4"),
        "Impugnação ao agravo interno — CASO-04",
    ),
    (
        STATE / forja_acervo.caso("CASO-16") / "producao"
        / "MEMORIAIS_LIBRA_SUL_AGINT_ARESP_2578181.md",
        "Memoriais — CASO-16",
    ),
]
b2 = {"controlesNegativos": [], "pecas": []}
if SEM_RENDER:
    print("  (pulado por --sem-render)")
else:
    from forja_render_docx import render
    # Controle negativo real: o corpus antigo deve ser recusado antes de nascer DOCX.
    for md in NEGATIVE_STYLE_TARGETS:
        out = TELEDIR / "renders" / (md.stem + "_BLOQUEIO_ESTILO")
        destino_docx = out / f"{md.stem}.docx"
        erro = None
        try:
            render(md, out, "Controle negativo de escrita humana")
        except RuntimeError as exc:
            erro = str(exc)
        gate_json = out / "F7_VERIFICADOR_FORJA.json"
        gate_payload = json.loads(gate_json.read_text(encoding="utf-8")) if gate_json.exists() else {}
        b2["controlesNegativos"].append({
            "peca": md.name,
            "bloqueado": bool(erro),
            "p0": gate_payload.get("p0"),
            "docxGerado": destino_docx.exists(),
        })
        caso(f"B2 negativo {md.stem}: vício real bloqueado antes do Word", bool(erro), erro)
        caso(f"B2 negativo {md.stem}: relatório F7 persistido", gate_json.exists())
        caso(f"B2 negativo {md.stem}: nenhum DOCX gerado", not destino_docx.exists())

    for md, titulo in RENDER_ALVOS:
        nome_md = md.name
        if not md.exists():
            caso(f"B2 fonte existe: {nome_md}", False, str(md))
            continue
        out = TELEDIR / "renders" / md.stem
        texto_md = md.read_text(encoding="utf-8")
        # remover negritos antes de contar itálicos (senão ** casa como dois *)
        texto_sem_bold = NEGRITO_RE.sub(" ", texto_md)
        italicos_fonte = [i for i in ITALICO_RE.findall(texto_sem_bold)
                          if not re.fullmatch(r"[-_*\s]+", i)]
        negritos_fonte = NEGRITO_RE.findall(texto_md)

        print(f"\n  Renderizando {nome_md} ({len(texto_md)} chars, {len(italicos_fonte)} itálicos e {len(negritos_fonte)} negritos no fonte)...")
        t1 = time.perf_counter()
        resumo = render(md, out, titulo)
        dur_s = round(time.perf_counter() - t1, 1)

        txt, doc = texto_docx(resumo["docx"])
        runs_italic = []
        runs_bold = 0
        for p in doc.paragraphs:
            for r in p.runs:
                if r.italic:
                    runs_italic.append(r.text)
                if r.bold:
                    runs_bold += 1
        for tab in doc.tables:
            for row in tab.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.italic:
                                runs_italic.append(r.text)
                            if r.bold:
                                runs_bold += 1

        gates = resumo.get("gatesForjaVerificador", {})
        reg = {
            "peca": nome_md,
            "duracaoRenderSegundos": dur_s,
            "paginasPdf": resumo["paginas"],
            "italicosNoFonte": len(italicos_fonte),
            "runsItalicNoDocx": len(runs_italic),
            "runsBoldNoDocx": runs_bold,
            "asteriscosLiteraisNoDocx": txt.count("*"),
            "placeholdersProibidos": resumo["placeholdersProibidos"],
            "gatesP0": gates.get("p0"), "gatesP1": gates.get("p1"),
            "docx": resumo["docx"], "pdf": resumo["pdf"],
        }
        b2["pecas"].append(reg)
        print(f"    -> {resumo['paginas']} págs em {dur_s}s | itálicos fonte->docx: {len(italicos_fonte)}->{len(runs_italic)} | asteriscos no DOCX: {txt.count('*')} | gates P0={gates.get('p0')} P1={gates.get('p1')}")

        caso(f"B2 {md.stem}: ZERO asterisco literal no DOCX renderizado", "*" not in txt,
             f"{txt.count('*')} asteriscos")
        if italicos_fonte:
            caso(f"B2 {md.stem}: cada *itálico* do fonte virou run italic no DOCX",
                 len(runs_italic) >= len(italicos_fonte),
                 f"fonte={len(italicos_fonte)} docx={len(runs_italic)}")
            faltando = [i for i in italicos_fonte if not any(i in ri or ri in i for ri in runs_italic)]
            caso(f"B2 {md.stem}: conteúdo dos itálicos preservado no DOCX", not faltando, faltando[:3])
        if negritos_fonte:
            caso(f"B2 {md.stem}: negritos do fonte presentes como runs bold", runs_bold > 0, runs_bold)
        caso(f"B2 {md.stem}: PDF real gerado com páginas", resumo["paginas"] >= 1, resumo["paginas"])
        caso(f"B2 {md.stem}: F7 persistido no build",
             (Path(out) / "F7_VERIFICADOR_FORJA.json").exists())
        caso(f"B2 {md.stem}: zero P0 no gate real", gates.get("p0") == 0, gates.get("violacoes"))
tele["baterias"]["B2_render_real"] = b2
if not SEM_RENDER:
    b2["metricasQualidade"] = {
        "artefatosRenderizados": len(b2["pecas"]),
        "artefatosSemRessalvas": sum((item.get("gatesP0") or 0) == 0 and (item.get("gatesP1") or 0) == 0 for item in b2["pecas"]),
        "artefatosComRevisaoP1": sum((item.get("gatesP1") or 0) > 0 for item in b2["pecas"]),
        "totalP0": sum(item.get("gatesP0") or 0 for item in b2["pecas"]),
        "totalP1": sum(item.get("gatesP1") or 0 for item in b2["pecas"]),
    }

# ================================================================ B3
print()
print("=" * 72)
print("B3 — Varredura dos DOCX de produção JÁ ENTREGUES (asterisco literal)")
print("=" * 72)

b3 = {"docx": []}
docx_producao = sorted(
    p for p in STATE.glob("*/producao/**/*.docx")
    if not p.name.startswith("~$") and "_teste" not in str(p) and "telemetria" not in str(p)
)
for d in docx_producao:
    try:
        txt, _ = texto_docx(d)
    except Exception as exc:
        caso(f"B3 legível: {d.name}", False, exc)
        continue
    n_ast = txt.count("*")
    b3["docx"].append({"arquivo": str(d.relative_to(STATE)), "asteriscos": n_ast})
    print(f"  {d.name[:60]:<60} asteriscos: {n_ast}")
    caso(f"B3 {d.name}: sem asterisco literal", n_ast == 0, n_ast)
caso("B3 volume real varrido (>= 12 DOCX de produção)", len(b3["docx"]) >= 12, len(b3["docx"]))
tele["baterias"]["B3_docx_entregues"] = b3

# ================================================================ fecho
quality = (b2.get("metricasQualidade") or {}) if not SEM_RENDER else {}
tele["resultado"] = {
    "falhas": falhas,
    "pipelineAprovado": not falhas,
    "artefatosSemRessalvas": quality.get("artefatosSemRessalvas"),
    "artefatosComRevisaoP1": quality.get("artefatosComRevisaoP1"),
    "totalP0": quality.get("totalP0"),
    "totalP1": quality.get("totalP1"),
    "nota": "Aprovação do pipeline não equivale a liberação jurídica ou editorial dos artefatos.",
}
saida = TELEDIR / f"TELEMETRIA_LICAO41_{datetime.now():%Y-%m-%d_%H%M}.json"
saida.write_text(json.dumps(tele, ensure_ascii=False, indent=2), encoding="utf-8")

print()
print("=" * 72)
if falhas:
    print(f"REPROVADO: {len(falhas)} falha(s). Telemetria: {saida}")
    for f in falhas:
        print(f"  - {f['caso']}: {f['detalhe']}")
    sys.exit(1)
if not SEM_RENDER and quality.get("totalP1"):
    print(f"PIPELINE APROVADO; ARTEFATOS COM RESSALVAS: {quality['artefatosComRevisaoP1']}/{quality['artefatosRenderizados']} e {quality['totalP1']} P1. Telemetria: {saida}")
else:
    print(f"PIPELINE APROVADO em dados reais. Telemetria completa: {saida}")
