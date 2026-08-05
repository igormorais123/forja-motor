"""FORJA N2 - F5 Extração e ledger de citações (modo leitura / sombra).

Para cada caso indicado:
  - extrai citações jurisprudenciais da peça (DOCX ou MD): REsp, AREsp, AgInt, EDcl,
    RE/ARE, HC, Temas, Súmulas, Informativos;
  - para cada citação procura ARQUIVO-FONTE local na pasta do caso (PDF/DOCX/MD com o
    número) — via oficial aceita pela spec N2 quando o portal está inacessível;
  - gera URL de conferência em fonte oficial (SCON/STJ, portal STF);
  - grava F5_CHECKLIST_CITACOES.md em state/<caseId>/ e atualiza o sourceLedger.

Limitação registrada em 2026-07-08: SCON/STJ atrás de WAF Cloudflare (HTTP 403 e
challenge não resolvido em navegador automatizado). Validação web exige Chrome real
com perfil humano (skill testar-navegador) ou conferência manual — o checklist sai
com URL pronta para um clique.
"""

import json
import re
import sys
from pathlib import Path
from urllib.parse import quote

from forja_authorities import extract_authorities, normalize_number, tribunal_from_cnj
from forja_n3_common import now_iso, sha256_file
from forja_official_sources import validate_archived_source, validate_cached_source

FORJA = Path(__file__).resolve().parent
RAIZ = FORJA.parent
STATE_DIR = FORJA / "state"

PADROES = [
    ("CNJ", re.compile(
        r"\b(ADI|A[CÇ][AÃ]O\s+DIRETA\s+DE\s+INCONSTITUCIONALIDADE|APELA[CÇ][AÃ]O|AGRAVO\s+DE\s+INSTRUMENTO)\s*"
        r"(?:n[oº.]?\s*)?(\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4})",
        re.I,
    )),
    ("STJ", re.compile(r"\b(?:AgInt\s+(?:nos?\s+)?|EDcl\s+(?:nos?\s+)?)*(REsp|AREsp|EREsp)\s*(?:n[oº.]?\s*)?([\d.]{5,})\s*[/-]?\s*([A-Z]{2})?", re.I)),
    ("STF", re.compile(r"\b(RE|ARE|ADI|ADPF)\s*(?:n[oº.]?\s*)?([\d.]{5,})\b(?!-\d{2}\.\d{4}\.)", re.I)),
    ("SUMULA_VINCULANTE", re.compile(r"S[úu]mula\s+Vinculante\s+(?:n[oº.]?\s*)?(\d{1,3})", re.I)),
    ("SUMULA", re.compile(r"S[úu]mula\s+(?!Vinculante)(?:n[oº.]?\s*)?(\d{1,4})(?:\s*(?:[/]|\s+d[oe]\s+|[-–—]\s*)\s*(STJ|STF))?", re.I)),
    # Lição 41(b): aceitar ponto de milhar ("Tema 1.365") e o qualificador "Repetitivo"
    ("TEMA", re.compile(r"Tema\s+(?:Repetitivo\s+)?(?:n[oº.]?\s*)?(\d{1,3}(?:\.\d{3})+|\d{1,5})\s*(?:(?:d[oe]\s*)|(?:[-–—]\s*))?(STJ|STF)?", re.I)),
    ("INFORMATIVO", re.compile(r"Informativo\s+(?:n[oº.]?\s*)?(\d{2,4})\s*(?:(?:d[oe]\s*)|(?:[-–—]\s*))?(STJ|STF)?", re.I)),
]

CNJ_TRIBUNAIS = {
    "4.01": "TRF1", "4.02": "TRF2", "4.03": "TRF3", "4.04": "TRF4", "4.05": "TRF5", "4.06": "TRF6",
    "8.07": "TJDFT", "8.26": "TJSP", "8.27": "TJTO", "8.21": "TJRS", "8.19": "TJRJ",
}


def tribunal_numero_cnj(value):
    return tribunal_from_cnj(value)


def texto_da_peca(path):
    path = Path(path)
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        partes = [p.text for p in doc.paragraphs]
        for tabela in doc.tables:
            for linha in tabela.rows:
                partes.extend(c.text for c in linha.cells)
        # notas de rodapé ficam fora do python-docx básico; extrair do XML bruto
        import zipfile
        with zipfile.ZipFile(str(path)) as z:
            for nome in ("word/footnotes.xml", "word/endnotes.xml"):
                if nome in z.namelist():
                    xml = z.read(nome).decode("utf-8", errors="replace")
                    partes.append(re.sub(r"<[^>]+>", " ", xml))
        return "\n".join(partes)
    return path.read_text(encoding="utf-8", errors="replace")


def normalizar_numero(n):
    return normalize_number(n)


def url_oficial(tipo, dados):
    if tipo == "CNJ":
        tribunal = tribunal_numero_cnj(dados[1])
        if tribunal == "TJSP":
            return "https://esaj.tjsp.jus.br/cjsg/consultaCompleta.do"
        return ""
    if tipo == "STJ":
        classe, numero, uf = dados
        q = quote(f'"{classe} {numero}"')
        return f"https://scon.stj.jus.br/SCON/pesquisar.jsp?livre={q}&b=ACOR"
    if tipo == "STF":
        classe, numero = dados[0], dados[1]
        return f"https://jurisprudencia.stf.jus.br/pages/search?base=acordaos&pesquisa_inteiro_teor=false&sinonimo=true&plural=true&queryString={quote(classe + ' ' + numero)}"
    if tipo == "NORMA":
        return "https://www.planalto.gov.br/ccivil_03/"
    if tipo == "SUMULA_VINCULANTE":
        return "https://portal.stf.jus.br/jurisprudencia/sumariosumulas.asp?base=26"
    if tipo == "SUMULA":
        numero, corte = dados
        if (corte or "").upper() == "STF":
            return "https://portal.stf.jus.br/jurisprudencia/sumariosumulas.asp?base=30"
        return f"https://scon.stj.jus.br/SCON/pesquisar.jsp?livre=@NUM={quote(numero)}&b=SUMU"
    if tipo == "TEMA":
        numero, corte = dados
        if (corte or "").upper() == "STF":
            return f"https://portal.stf.jus.br/jurisprudenciaRepercussao/pesquisarProcesso.asp?numeroTema={numero}"
        return f"https://processo.stj.jus.br/repetitivos/temas_repetitivos/pesquisa.jsp?novaConsulta=true&tipo_pesquisa=T&num_processo_classe=&cod_tema_inicial={numero}&cod_tema_final={numero}"
    if tipo == "INFORMATIVO":
        numero, corte = dados
        if (corte or "").upper() == "STF":
            return "https://portal.stf.jus.br/textos/verTexto.asp?servico=informativoSTF"
        return f"https://scon.stj.jus.br/jurisprudencia/externo/informativo/?acao=pesquisarumaedicao&livre=%27{numero}%27.cod."
    return ""


def extrair_citacoes(texto):
    return extract_authorities(texto)


def procurar_cache_oficial(citacao, *, require_live=True):
    """Fonte oficial arquivada em cache/fontes_oficiais (capturada do portal com evidência de URL/data)."""
    cache = FORJA / "cache" / "fontes_oficiais"
    if not cache.exists():
        return None
    tipo = citacao["tipo"]
    numero = None
    for g in citacao["dados"]:
        n = normalizar_numero(g or "")
        if n:
            numero = n
            break
    if not numero:
        return None
    # Corte é identidade (review adversarial 09/07/2026): com tribunal conhecido, só o
    # cache DAQUELE tribunal conta; sem tribunal e com fontes de dois, é ambíguo (None).
    corte = ""
    if tipo in ("SUMULA", "TEMA", "INFORMATIVO") and len(citacao["dados"]) > 1:
        corte = (citacao["dados"][1] or "").upper()
    prefixos = {
        "SUMULA": [f"STJ_SUMULA_{numero}", f"STF_SUMULA_{numero}"],
        "SUMULA_VINCULANTE": [f"STF_SUMULA_VINCULANTE_{numero}"],
        "TEMA": [f"STJ_TEMA_{numero}", f"STF_TEMA_{numero}"],
        "INFORMATIVO": [f"STJ_INFORMATIVO_{numero}", f"STF_INFORMATIVO_{numero}"],
        "STJ": [f"STJ_{citacao['classe'].upper()}_{numero}"],
        "STF": [f"STF_{citacao['classe'].upper()}_{numero}"],
        "NORMA": [
            (
                f"PLANALTO_{citacao['authorityIdentity'].get('code')}_ART_"
                f"{citacao['authorityIdentity'].get('article')}"
            )
            if citacao.get("classe") == "ARTICLE"
            else f"PLANALTO_{citacao['classe']}_{numero}"
        ],
    }.get(tipo, [])
    if corte in ("STJ", "STF"):
        prefixos = [p for p in prefixos if p.startswith(corte + "_")]
    from forja_metricas_f7 import cache_com_lastro
    achados = [cache / (n + ".txt") for n in prefixos
               if (cache / (n + ".txt")).exists()
               and cache_com_lastro(cache / (n + ".txt"), require_live=require_live)]
    if not achados:
        return None
    if not corte and len({a.name.split("_", 1)[0] for a in achados}) > 1:
        return None
    return achados[0]


def normalizar_aspa(texto):
    """Normaliza espaços múltiplos, preserva acentos e caixa."""
    return re.sub(r"\s+", " ", texto).strip()


def conferir_aspas(texto_peca, arquivo_fonte):
    """Extrai e confere aspas na peça contra o verbatim da fonte.

    Procura por padrão "...texto..." na peça e verifica se é substring
    normalizada do arquivo de fonte (espaços colapsados, acentos preservados).
    Retorna lista de tuplas (aspa_encontrada, verbatim_localizado, é_válida).
    """
    if not arquivo_fonte.exists():
        return []

    verbatim = arquivo_fonte.read_text(encoding="utf-8", errors="replace")
    verbatim_norm = normalizar_aspa(verbatim)

    achadas = []
    padrao_aspa = re.compile(r'"([^"]{10,300})"')
    for m in padrao_aspa.finditer(texto_peca):
        aspa_bruta = m.group(1)
        aspa_norm = normalizar_aspa(aspa_bruta)
        eh_valida = aspa_norm in verbatim_norm
        achadas.append((aspa_bruta, aspa_norm, eh_valida))

    return achadas


def procurar_fonte_local(citacao, pasta_caso):
    """Fonte local só conta com sidecar oficial, hash e identidade do precedente."""
    if not pasta_caso or not pasta_caso.exists():
        return None
    numero = None
    for g in citacao["dados"]:
        n = normalizar_numero(g or "")
        if len(n) >= 4:
            numero = n
            break
    if not numero:
        return None
    for arq in pasta_caso.rglob("*"):
        if arq.suffix.lower() not in (".pdf", ".md", ".txt", ".docx"):
            continue
        if numero in normalizar_numero(arq.name) and validate_archived_source(arq)["approved"]:
            return arq
        if arq.suffix.lower() in (".md", ".txt") and arq.stat().st_size < 3_000_000:
            try:
                if (
                    numero in normalizar_numero(arq.read_text(encoding="utf-8", errors="replace"))
                    and validate_archived_source(arq)["approved"]
                ):
                    return arq
            except OSError:
                pass
    return None


def merge_by_id(existing, new_items):
    """Mantém um item por id; itens novos substituem versões antigas do mesmo id."""
    merged = {}
    for item in existing or []:
        item_id = item.get("id")
        if item_id:
            merged[item_id] = item
    for item in new_items or []:
        item_id = item.get("id")
        if item_id:
            merged[item_id] = item
    return list(merged.values())


def append_unique(existing, value):
    items = list(existing or [])
    if value not in items:
        items.append(value)
    return items


def processar(case_key, peca_path):
    matches = list(STATE_DIR.glob(f"case-*{case_key}*/FORJA_STATE.json"))
    if not matches:
        print(f"AVISO: estado não encontrado para {case_key}")
        return None
    state_path = matches[0]
    state = json.loads(state_path.read_text(encoding="utf-8-sig"))
    pasta_caso = Path((state.get("inputs") or {}).get("caseFolder") or "")

    texto = texto_da_peca(peca_path)
    citacoes = extrair_citacoes(texto)

    linhas = [
        "# F5 — Checklist de citações da peça",
        "",
        f"Caso: `{state['caseId']}` | Peça analisada: `{Path(peca_path).name}` | Gerado: {now_iso()}",
        "",
        "Regra N2: citação final exige fonte oficial OU arquivo-fonte arquivado. SCON/STF estão",
        "atrás de anti-bot (registrado em 08/07/2026): use os links abaixo no Chrome normal.",
        "",
        f"Total de citações distintas: **{len(citacoes)}**",
        "",
    ]
    ledger_novos = []
    confirmadas_arquivo = 0
    for c in sorted(citacoes, key=lambda x: x["rotulo"]):
        cache_oficial = procurar_cache_oficial(c)
        fonte_local = procurar_fonte_local(c, pasta_caso)
        url = url_oficial(c["tipo"], c["dados"])
        if cache_oficial:
            provenance = validate_cached_source(cache_oficial)["record"] or {}
            confirmadas_arquivo += 1
            status = f"FONTE OFICIAL arquivada: `{cache_oficial.relative_to(RAIZ)}`"
            classificacao, permitido = "FONTE_OFICIAL", True
            fonte_local = cache_oficial
        elif fonte_local:
            provenance = validate_archived_source(fonte_local)["record"] or {}
            confirmadas_arquivo += 1
            status = f"ARQUIVO-FONTE: `{fonte_local.relative_to(RAIZ)}`"
            classificacao, permitido = "FONTE_ARQUIVO", True
        else:
            provenance = {}
            status = f"PENDENTE conferência oficial: {url}"
            classificacao, permitido = "NAO_VERIFICADO", False
        linhas.append(f"- **{c['rotulo']}** ({c['ocorrencias']}x) — {status}")
        linhas.append(f"  - contexto: ...{c['contexto']}...")
        ledger_novos.append({
            "id": f"cit-{re.sub(r'[^A-Za-z0-9]+', '-', c['rotulo'])[:40]}",
            "claim": f"Citação na peça: {c['rotulo']}",
            "classification": classificacao,
            "sourcePathOrUrl": str(fonte_local) if fonte_local else url,
            "sourceSha256": sha256_file(fonte_local) if fonte_local else None,
            "sourceUrl": provenance.get("sourceUrl") if fonte_local else url,
            "sourceIdentity": provenance.get("identity") if fonte_local else None,
            "pageOrEvent": None,
            "verifiedAt": now_iso() if fonte_local else None,
            "finalUseAllowed": permitido,
        })
    pendentes = len(citacoes) - confirmadas_arquivo
    linhas += [
        "",
        "## Balanço",
        "",
        f"- Confirmadas por arquivo-fonte local: **{confirmadas_arquivo}**",
        f"- Pendentes de conferência em portal oficial (bloqueiam peça FINAL, não a minuta): **{pendentes}**",
    ]

    out = state_path.parent / "F5_CHECKLIST_CITACOES.md"
    out.write_text("\n".join(linhas) + "\n", encoding="utf-8")

    state["updatedAt"] = now_iso()
    state["currentPhase"] = "F5_PESQUISA_OFICIAL"
    state.setdefault("phaseHistory", []).append(
        {"phase": "F5_PESQUISA_OFICIAL", "at": now_iso(),
         "status": "ok" if pendentes == 0 else "pendencias"})
    state["sourceLedger"] = merge_by_id(state.get("sourceLedger") or [], ledger_novos)
    state["artifacts"] = append_unique(state.get("artifacts") or [], str(out))
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    return {"caso": state["caseId"], "peca": Path(peca_path).name, "citacoes": len(citacoes),
            "confirmadasArquivo": confirmadas_arquivo, "pendentesWeb": pendentes, "checklist": str(out)}


# ---------------------------------------------------------------------------
# Política de citação computada — gate `citations_policy_satisfied`
# ---------------------------------------------------------------------------
# Até 04/08/2026 este gate era escrito pelo agente da fase F7: onze execuções,
# onze `pass`, nenhuma reprovação. É o gate de maior volume da esteira e o mais
# caro se for falso — "jurisprudência com atribuição errada" é o erro recorrente
# nº 1 das entregas reais, e o U1 cataloga seis modos de falha de citação.
#
# A política computável é de COBERTURA e LIBERAÇÃO, não de mérito: toda
# autoridade citada no texto final precisa existir no ledger de fontes
# verificadas, e nenhuma pode ser usada com `finalUseAllowed` diferente de true.
# Se a citação é fiel à tese do precedente — ratio ou dictum, superado ou
# vigente — continua sendo trabalho humano do F7, e o gate não finge decidir
# isso. Gate que tenta julgar mérito vira trava, e trava ensina a contornar.
#
# O casamento é determinístico porque os dois lados falam a mesma língua: o
# extrator devolve `authorityIdentity` com corte, classe e número, exatamente o
# formato gravado nas entradas do ledger.

_POLITICA_VERSAO = "FORJA-CITACOES-POLITICA-v1"


def _chave_autoridade(identidade):
    if not isinstance(identidade, dict):
        return None
    corte = str(identidade.get("court") or "").strip().upper()
    classe = str(identidade.get("kind") or "").strip().upper()
    numero = re.sub(r"\D", "", str(identidade.get("number") or ""))
    if not classe or not numero:
        return None
    # Corte ambígua não entra na chave: "Tema 1368" sem tribunal precisa casar
    # com a entrada do ledger que o resolveu. A ambiguidade em si é assunto do
    # gate `citation_identity_and_cnj_tribunal_resolved`, não desta política.
    if corte in {"", "TRIBUNAL_AMBIGUO"}:
        corte = "*"
    return (corte, classe, numero)



def _indice_do_ledger(ledger):
    """Índice identidade -> entrada, lendo os dois esquemas reais do acervo.

    O primeiro esquema traz `entries[]` com `authorityIdentity` estruturada. O
    segundo traz `officialSources[]` com `identifier` em texto ("EREsp
    800.578/MG") e `archivedSha256`. Em vez de escrever um segundo mapeamento à
    mão, o identificador passa pelo MESMO extrator usado no texto da peça: os
    dois lados voltam a falar a mesma língua, e um formato novo de rótulo
    passa a ser entendido nos dois lugares de uma vez.

    Devolve ``None`` quando nenhuma identidade pôde ser derivada — o chamador
    trata isso como "não conferível", nunca como aprovação.
    """
    if not isinstance(ledger, dict):
        return None
    indice = {}

    for entrada in ledger.get("entries") or []:
        if not isinstance(entrada, dict):
            continue
        chave = _chave_autoridade(entrada.get("authorityIdentity"))
        if chave:
            indice[chave] = entrada

    for fonte in ledger.get("officialSources") or []:
        if not isinstance(fonte, dict):
            continue
        rotulo = str(fonte.get("identifier") or "")
        for citacao in extrair_citacoes(rotulo):
            chave = _chave_autoridade(citacao.get("authorityIdentity"))
            if not chave or chave in indice:
                continue
            # Cópia oficial arquivada com hash é a prova de conferência neste
            # esquema; ele não tem campo `finalUseAllowed`, e inventar um
            # significado para a ausência seria fabricar liberação.
            indice[chave] = {
                "authorityIdentity": citacao.get("authorityIdentity"),
                "finalUseAllowed": bool(fonte.get("archivedSha256")),
                "status": fonte.get("status"),
                "blockedReason": (None if fonte.get("archivedSha256")
                                  else "copia oficial arquivada ausente"),
                "origem": "officialSources",
            }

    return indice or None


def validar_politica_citacoes(texto, ledger):
    """Achados e veredito do gate `citations_policy_satisfied`."""
    achados = []
    citacoes = extrair_citacoes(texto or "")

    if not citacoes:
        # Peça sem citação não é peça defeituosa. O gate não pode inventar
        # exigência onde não há afirmação de autoridade.
        return {"versao": _POLITICA_VERSAO, "findings": [],
                "gates": {"citations_policy_satisfied": "pass"}, "citacoes": 0}

    indice = _indice_do_ledger(ledger)
    if indice is None:
        # O acervo tem DOIS esquemas de ledger e nenhum é errado. Quando o
        # ledger não permite derivar identidade de autoridade, o honesto é
        # dizer que não foi possível conferir — nunca "conferi e aprovei", e
        # nunca reprovar em P0 quem escreveu no outro formato. `warn` é
        # exatamente esse terceiro estado, e é o que impede este gate de
        # repetir a MC-15 que ele nasceu para evitar.
        achados.append({
            "gate": "LP1-ledger-nao-conferivel", "sev": "P1",
            "problema": ("texto cita %d autoridade(s) e o ledger de fontes verificadas nao "
                         "expoe identidade de autoridade conferivel automaticamente" % len(citacoes)),
            "acao": "registre as autoridades em entries[] com authorityIdentity, ou confira a mao",
            "versao": _POLITICA_VERSAO})
        return {"versao": _POLITICA_VERSAO, "findings": achados,
                "gates": {"citations_policy_satisfied": "warn"}, "citacoes": len(citacoes)}

    def achar(chave):
        if chave in indice:
            return indice[chave]
        corte, classe, numero = chave
        for (c, k, n), entrada in indice.items():
            if k == classe and n == numero and "*" in (corte, c):
                return entrada
        return None

    for citacao in citacoes:
        chave = _chave_autoridade(citacao.get("authorityIdentity"))
        rotulo = citacao.get("rotulo") or citacao.get("rótulo") or "?"
        if chave is None:
            achados.append({
                "gate": "LP1-identidade", "sev": "P1", "citacao": rotulo,
                "problema": "%s: identidade da autoridade incompleta para conferencia" % rotulo,
                "versao": _POLITICA_VERSAO})
            continue
        entrada = achar(chave)
        if entrada is None:
            achados.append({
                "gate": "LP1-nao-conferida", "sev": "P0", "citacao": rotulo,
                "problema": ("%s citada no texto final sem entrada no ledger de fontes "
                             "verificadas - a autoridade nao foi conferida" % rotulo),
                "acao": "confira a autoridade na fonte oficial e registre a entrada",
                "versao": _POLITICA_VERSAO})
            continue
        if entrada.get("finalUseAllowed") is not True:
            motivo = str(entrada.get("blockedReason") or entrada.get("status")
                         or "sem motivo registrado")
            achados.append({
                "gate": "LP2-uso-bloqueado", "sev": "P0", "citacao": rotulo,
                "problema": ("%s usada no texto final com finalUseAllowed diferente de true "
                             "(%s)" % (rotulo, motivo)),
                "acao": "libere a autoridade no ledger ou remova a citacao da peca",
                "versao": _POLITICA_VERSAO})

    reprovado = any(item["sev"] == "P0" for item in achados)
    return {"versao": _POLITICA_VERSAO, "findings": achados,
            "gates": {"citations_policy_satisfied": "fail" if reprovado else "pass"},
            "citacoes": len(citacoes)}


# ---------------------------------------------------------------------------
# Identidade de autoridade — gate `citation_identity_and_cnj_tribunal_resolved`
# ---------------------------------------------------------------------------
# Âncora real: o P0 mais grave do caso Vale Trading foi afirmar que um agravo
# "envolve as mesmas partes e a mesma liquidação" quando os números CNJ apontam
# para liquidações distintas. O protocolo da casa manda identificar o tribunal
# pelo número CNJ antes de qualquer coisa; até 04/08/2026 esse gate era
# atestado pelo agente da fase.
#
# Duas verificações, ambas determinísticas e ambas com âncora em falha real:
#
#   LI1 — número CNJ citado junto de um tribunal nomeado que o próprio número
#         contradiz. O segmento do CNJ é a fonte da verdade; o rótulo escrito
#         ao lado é a hipótese.
#   LI2 — autoridade cuja corte o extrator não resolveu ("Tema 1368" sem
#         tribunal) e que o ledger também não resolve. Fica em P1: é ambiguidade
#         de redação, não afirmação falsa.
#
# O que NÃO é feito aqui: par súmula x tribunal, que já é P0 do G4 em
# `forja_verificador.py` desde a lição 20. Duplicar gate produz dois vereditos
# para o mesmo defeito e nenhuma autoridade sobre qual vale.

_JANELA_TRIBUNAL = 120
_TRIBUNAL_NOMEADO = re.compile(
    r"\b(TRF\s?-?\s?[1-6]|TJ[A-Z]{2}|TJDFT|STJ|STF|TST|TSE)\b", re.I)


def _normalizar_tribunal(bruto):
    return re.sub(r"[\s-]", "", str(bruto or "")).upper()


def validar_identidade_citacoes(texto, ledger=None):
    """Achados e veredito de `citation_identity_and_cnj_tribunal_resolved`."""
    texto = texto or ""
    achados = []

    # LI1 — CNJ contra tribunal nomeado na vizinhança.
    for match in re.finditer(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}", texto):
        numero = match.group(0)
        esperado = tribunal_numero_cnj(numero)
        if not esperado:
            continue
        inicio = max(0, match.start() - _JANELA_TRIBUNAL)
        janela = texto[inicio:match.end() + _JANELA_TRIBUNAL]
        nomeados = {_normalizar_tribunal(m.group(1)) for m in _TRIBUNAL_NOMEADO.finditer(janela)}
        if not nomeados:
            continue
        # STJ e STF aparecem legitimamente perto de um CNJ de origem — o recurso
        # sobe. Só contradiz quem alega ser o tribunal DE ORIGEM do número.
        origem = {t for t in nomeados if t.startswith(("TRF", "TJ"))}
        if origem and _normalizar_tribunal(esperado) not in origem:
            achados.append({
                "gate": "LI1-cnj-tribunal", "sev": "P0", "citacao": numero,
                "problema": ("%s e do %s pelo segmento CNJ, mas o texto o associa a %s"
                             % (numero, esperado, ", ".join(sorted(origem)))),
                "acao": "confira o numero e o tribunal de origem antes de afirmar identidade",
                "versao": _POLITICA_VERSAO})

    # LI2 — corte não resolvida no texto nem no ledger.
    indice = _indice_do_ledger(ledger) or {}
    resolvidas = {(k, n) for (c, k, n) in indice if c != "*"}
    for citacao in extrair_citacoes(texto):
        identidade = citacao.get("authorityIdentity") or {}
        corte = str(identidade.get("court") or "").upper()
        if corte not in {"", "TRIBUNAL_AMBIGUO"}:
            continue
        classe = str(identidade.get("kind") or "").upper()
        # Norma não tem tribunal. Cobrar corte de "art. 203 do CPC" é ruído, e
        # ruído ensina a ignorar o gate — foi o que a primeira versão fez com
        # três artigos do CPC numa peça correta.
        if str(citacao.get("tipo") or "").upper() == "NORMA" or classe == "ARTICLE":
            continue
        numero = re.sub(r"\D", "", str(identidade.get("number") or ""))
        if (classe, numero) in resolvidas:
            continue
        achados.append({
            "gate": "LI2-corte-ambigua", "sev": "P1",
            "citacao": citacao.get("rotulo") or citacao.get("rótulo") or "?",
            "problema": ("%s citado sem tribunal, e o ledger tambem nao o resolve"
                         % (citacao.get("rotulo") or "?")),
            "acao": "nomeie o tribunal da autoridade no texto ou no ledger",
            "versao": _POLITICA_VERSAO})

    reprovado = any(item["sev"] == "P0" for item in achados)
    return {"versao": _POLITICA_VERSAO, "findings": achados,
            "gates": {"citation_identity_and_cnj_tribunal_resolved":
                      "fail" if reprovado else "pass"}}


def main():
    if len(sys.argv) < 3 or len(sys.argv) % 2 == 0:
        print("uso: python forja_citations.py <chave-do-caso> <caminho-da-peca> [<chave> <peca> ...]")
        return 1
    resultados = []
    pares = list(zip(sys.argv[1::2], sys.argv[2::2]))
    for chave, peca in pares:
        r = processar(chave, peca)
        if r:
            resultados.append(r)
    print(json.dumps({"processados": len(resultados), "resultados": resultados}, ensure_ascii=False, indent=2))
    return 2 if any(int(item.get("pendentesWeb") or 0) for item in resultados) else 0


if __name__ == "__main__":
    sys.exit(main())
