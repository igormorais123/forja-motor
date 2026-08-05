# -*- coding: utf-8 -*-
"""FORJA — camada visual law determinística (padrão padrao-visual-medina).

Motivação (lição 37, 09/07/2026): agentes que "transcrevem" conteúdo para a
edição visual RESUMEM o texto (auditoria: 80-95% dos parágrafos ausentes em
5 de 5 tentativas). Aqui o texto entra no DOCX extraído do próprio markdown
congelado — fidelidade por construção — e a moldura visual (pull quotes,
caixas, figuras, cards) entra por um MAPA declarativo por caso.

Gate embutido: auditoria de cobertura ao final — parágrafo significativo do
md ausente do DOCX é BLOQUEADOR (exceção: linhas listadas em mapa["ignorar"]).

Uso:
    from forja_visual import compor
    compor(md_path, out_docx, mapa)

Mapa (todas as chaves opcionais, exceto nenhuma):
    tipo: "peca" | "estudo"          (default "peca")
    linhas_sintese: {"TÍTULO DA SEÇÃO (como no md, sem numeração)": "linha-síntese sans"}
    pulls: [("âncora: substring EXATA de um parágrafo do md", "texto da pull quote"), ...]
    laterais: [("âncora", "texto da nota lateral"), ...]
    caixas: [("âncora do parágrafo-citação no md", "acordao"|"precedente"|"chave", "TÍTULO DA CAIXA"), ...]
        -> o PARÁGRAFO ancorado vira o corpo da caixa (extraído do md, não redigitado)
    figs: [("âncora do parágrafo APÓS o qual entra a figura", "{{TAG}}", "legenda figcap"), ...]
    cards_apos_titulo: True -> insere {{CARDS}} na capa após o título
    rotulos_sintese: [("RÓTULO", "âncora-início"), ...] -> divide a seção de síntese
        do md em linhas rotuladas; cada trecho é extraído do md por slicing.
    sintese_extratos: [("RÓTULO", "trecho literal do md"), ...] -> cria um
        painel executivo conciso sem consumir nem duplicar a seção textual.
    ignorar: [substrings de linhas do md que NÃO entram (ex.: rótulos internos)]
"""
import json
import re
import math
import sys
import unicodedata
from pathlib import Path

FORJA = Path(__file__).resolve().parent
RAIZ = FORJA.parent
sys.path.insert(0, str(RAIZ / "_FERRAMENTAS"))

from medina_visual_kit import PecaVisual  # noqa: E402


def _norm(s):
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[*_#|>`“”‘’\"']", "", s.lower())
    return re.sub(r"[^a-z0-9]+", " ", s).strip()


def _e_enderecamento(linha):
    t = linha.strip().lstrip("#").strip().strip("*")
    return bool(t) and (t == t.upper() and len(t) > 25 or t.startswith("Excelent"))


def _consome_tabela(linhas, i):
    corpo = []
    while i < len(linhas) and linhas[i].strip().startswith("|"):
        corpo.append([c.strip() for c in linhas[i].strip().strip("|").split("|")])
        i += 1
    return corpo, i


def _larguras_tabela(header, corpo, total=13.1):
    """Distribui largura pelo conteúdo sem deixar uma coluna dominar o quadro."""
    n = len(header)
    if n == 0:
        return ()
    pesos = []
    for coluna in range(n):
        valores = [str(header[coluna])] + [str(row[coluna]) if coluna < len(row) else "" for row in corpo]
        maior = max((len(_norm(value)) for value in valores), default=1)
        pesos.append(min(5.0, max(1.4, math.sqrt(max(1, maior)))))
    soma = sum(pesos)
    larguras = [round(total * peso / soma, 2) for peso in pesos]
    larguras[-1] = round(larguras[-1] + total - sum(larguras), 2)
    return tuple(larguras)


class _Mapa:
    def __init__(self, mapa, texto_md):
        self.tipo = mapa.get("tipo", "peca")
        self.linhas_sintese = {_norm(k): v for k, v in (mapa.get("linhas_sintese") or {}).items()}
        self.ignorar = mapa.get("ignorar") or []
        self.cards = bool(mapa.get("cards_apos_titulo"))
        self.qualification_on_cover = bool(mapa.get("qualification_on_cover"))
        self.allowed_markers = {"{{CARDS}}"} if self.cards else set()
        self.rotulos_sintese = mapa.get("rotulos_sintese") or []
        self.sintese_extratos = mapa.get("sintese_extratos") or []
        self.sintese_titulos = {
            _norm(item) for item in (mapa.get("sintese_titulos") or [])
        }
        for rotulo, trecho in self.sintese_extratos:
            self._valida(trecho, texto_md, f"extrato de síntese ({rotulo})")
        # Opções de fidelidade para fontes canônicas que já trazem rótulos de
        # seção auditados. O padrão histórico continua auto-numerado.
        self.preserve_heading_labels = bool(mapa.get("preserve_heading_labels"))
        self.preserve_paragraph_labels = bool(mapa.get("preserve_paragraph_labels"))
        self.compact_wide_tables = bool(mapa.get("compact_wide_tables"))
        # Variante de capa (parecer Helena, 31/07/2026): sem a quebra de página,
        # a síntese executiva sobe para a capa e ocupa o branco da metade
        # inferior. Opt-in: o layout com capa autônoma foi o aprovado em
        # 09/07/2026 e continua sendo o padrão.
        self.capa_com_sintese = bool(mapa.get("capa_com_sintese"))
        self.force_cover_page_break = bool(mapa.get("force_cover_page_break"))
        self.table_total_cm = float(mapa.get("table_total_cm", 13.1))
        self.split_table_rows = bool(mapa.get("split_table_rows"))
        self.wide_table_chunk_rows = max(0, int(mapa.get("wide_table_chunk_rows", 0)))
        self.page_break_between_wide_table_chunks = bool(
            mapa.get("page_break_between_wide_table_chunks")
        )
        self.wide_table_scan_notes = [
            str(item).strip() for item in (mapa.get("wide_table_scan_notes") or [])
            if str(item).strip()
        ]
        self.study_numbered_items = bool(
            mapa.get("study_numbered_items", self.tipo == "estudo")
        )
        self.tarja_interna = (mapa.get("tarja_interna") or "").strip()
        self.titulo_sintese = mapa.get("titulo_sintese", "SÍNTESE DOS PONTOS ESSENCIAIS")
        # âncoras: validar TODAS contra o md agora (âncora quebrada = erro de mapa)
        self.pulls, self.laterais, self.caixas, self.figs = [], [], [], []
        for anc, txt in mapa.get("pulls") or []:
            self._valida(anc, texto_md, "pull")
            self.pulls.append((_norm(anc), txt))
        for anc, txt in mapa.get("laterais") or []:
            self._valida(anc, texto_md, "lateral")
            self.laterais.append((_norm(anc), txt))
        for anc, estilo, titulo in mapa.get("caixas") or []:
            self._valida(anc, texto_md, "caixa")
            self.caixas.append((_norm(anc), estilo, titulo))
        for anc, tag, legenda in mapa.get("figs") or []:
            self._valida(anc, texto_md, "fig")
            self.figs.append((_norm(anc), tag, legenda))
            self.allowed_markers.add(tag)

    @staticmethod
    def _valida(anc, texto_md, tipo):
        if _norm(anc) not in _norm(texto_md):
            raise SystemExit(f"MAPA INVÁLIDO — âncora de {tipo} não existe no md: {anc[:90]}")


def compor(md_path, out_docx, mapa, *, case_dir=None, ledger_path=None, base_dir=None):
    md_path = Path(md_path)
    texto_md = md_path.read_text(encoding="utf-8")
    tipo = (mapa or {}).get("tipo", "peca")
    from forja_verificador import verificar as verificar_forja
    from forja_lastro import material_economico
    p0 = [item for item in verificar_forja(
        texto_md, tipo, case_dir=case_dir, ledger=(
            json.loads(Path(ledger_path).read_text(encoding="utf-8"))
            if ledger_path else None
        ), base_dir=base_dir,
        exigir_economico=material_economico(texto_md),
    ) if item["sev"] == "P0"]
    if p0:
        resumo = "; ".join(f"{item['gate']}: {item['problema']}" for item in p0[:8])
        raise SystemExit(f"GATE DE ESCRITA HUMANA REPROVADO ({len(p0)} P0): {resumo}")
    m = _Mapa(mapa or {}, texto_md)
    # fólio áureo é identidade de peça protocolável; estudo/parecer interno roda sem
    pv = PecaVisual(
        str(out_docx), folio_aureo=(m.tipo != "estudo"),
        case_dir=case_dir, ledger_path=ledger_path, base_dir=base_dir,
    )
    # A porta única revalida o DOCX final; preserva aqui o tipo já decidido
    # pelo mapa para que uma rota canônica de estudo não seja julgada como peça.
    pv.tipo_produto = tipo

    linhas = texto_md.splitlines()
    i, primeiro_h1, capa = 0, True, True
    wide_table_index = 0
    pos_deferimento = False
    pulls_rest, lats_rest = list(m.pulls), list(m.laterais)
    caixas_rest, figs_rest = list(m.caixas), list(m.figs)

    def pos_paragrafo(par_norm):
        """Moldura ancorada NESTE parágrafo: pulls/laterais antes (frame na margem), caixas/figs depois."""
        for lst, acao in ((pulls_rest, "pull"), (lats_rest, "lat")):
            for item in list(lst):
                if item[0] in par_norm:
                    (pv.pull if acao == "pull" else pv.lateral)(item[1])
                    lst.remove(item)

    def como_caixa(par_norm, par_texto):
        """Se o parágrafo é âncora de caixa, trata-o COMO caixa (substitui o pgf). True se consumiu."""
        for item in list(caixas_rest):
            if item[0] in par_norm:
                estilo, titulo = item[1], item[2]
                if estilo == "acordao":
                    pv.caixa_acordao(titulo, par_texto)
                elif estilo == "precedente":
                    pv.caixa_precedente(titulo, par_texto)
                else:
                    pv.caixa_chave(par_texto)
                caixas_rest.remove(item)
                return True
        return False

    def figs_depois(par_norm):
        for item in list(figs_rest):
            if item[0] in par_norm:
                pv.marcador(item[1])
                pv.figcap(item[2])
                figs_rest.remove(item)

    em_sintese = False
    sintese_buf = []

    def descarrega_sintese():
        nonlocal em_sintese, sintese_buf
        if not em_sintese:
            return
        texto = "\n".join(sintese_buf).strip()
        if m.rotulos_sintese:
            partes = []
            for k, (rot, anc) in enumerate(m.rotulos_sintese):
                ini = texto.find(anc)
                if ini < 0:
                    raise SystemExit(f"MAPA INVÁLIDO — âncora de síntese não achada: {anc[:80]}")
                fim = texto.find(m.rotulos_sintese[k + 1][1]) if k + 1 < len(m.rotulos_sintese) else len(texto)
                partes.append((rot, texto[ini:fim].strip(" \r\n\t*")))
            pv.sintese(partes, titulo=m.titulo_sintese)
        elif texto:
            pv.sintese([("SÍNTESE", texto)], titulo=m.titulo_sintese)
        em_sintese, sintese_buf = False, []

    while i < len(linhas):
        raw = linhas[i]
        linha = raw.rstrip()
        strip = linha.strip()

        if not strip or re.match(r"^\s*([-*_])\1{2,}\s*$", strip):
            i += 1
            continue
        if any(ig in strip for ig in m.ignorar):
            i += 1
            continue

        # tabela markdown -> quadro zebrado
        if strip.startswith("|") and i + 1 < len(linhas) and re.match(r"^\s*\|[\s:|-]+\|\s*$", linhas[i + 1]):
            descarrega_sintese()
            header = [c.strip().replace("**", "") for c in strip.strip("|").split("|")]
            corpo, i = _consome_tabela(linhas, i + 2)
            n = len(header)
            if m.compact_wide_tables and n == 5:
                corpo = [
                    [
                        row[0] if len(row) > 0 else "",
                        f"{header[1]}: {row[1] if len(row) > 1 else ''}\n{header[2]}: {row[2] if len(row) > 2 else ''}",
                        f"{header[3]}: {row[3] if len(row) > 3 else ''}\n{header[4]}: {row[4] if len(row) > 4 else ''}",
                    ]
                    for row in corpo
                ]
                header = [header[0], f"{header[1]} e {header[2]}", f"{header[3]} e {header[4]}"]
                n = 3
            tabela_larga = m.table_total_cm > 13.1
            if tabela_larga:
                pv.abrir_secao_tabela_larga()
            corpo_normalizado = [r[:n] + [""] * (n - len(r[:n])) for r in corpo]
            chunk_rows = m.wide_table_chunk_rows if tabela_larga else 0
            blocos = (
                [corpo_normalizado[k:k + chunk_rows] for k in range(0, len(corpo_normalizado), chunk_rows)]
                if chunk_rows
                else [corpo_normalizado]
            )
            nota_tabela = (
                m.wide_table_scan_notes[min(wide_table_index, len(m.wide_table_scan_notes) - 1)]
                if tabela_larga and m.wide_table_scan_notes
                else ""
            )
            for bloco_n, bloco in enumerate(blocos, start=1):
                if tabela_larga and m.page_break_between_wide_table_chunks:
                    pv.quebra_pagina()
                if nota_tabela:
                    sufixo = f" Bloco {bloco_n} de {len(blocos)}." if len(blocos) > 1 else ""
                    pv.faixa_leitura_tabela(
                        nota_tabela + sufixo,
                        largura_cm=m.table_total_cm,
                    )
                pv.quadro_zebrado(
                    header,
                    bloco,
                    larguras_cm=_larguras_tabela(header, bloco, total=m.table_total_cm),
                    permitir_quebra_linha=m.split_table_rows,
                    alinhar_esquerda=tabela_larga,
                )
                if bloco_n < len(blocos):
                    pv.par("", 4, antes=1, depois=4)
            if tabela_larga:
                wide_table_index += 1
            if tabela_larga:
                pv.fechar_secao_tabela_larga()
            pv.par("", 8, antes=2, depois=6)
            continue

        mh = re.match(r"^(#{1,6})\s+(.*)$", linha)
        if mh:
            nivel, titulo = len(mh.group(1)), mh.group(2).strip()
            t_sem_num = re.sub(r"^[IVXL0-9]+\s*[—\-–.]\s*", "", titulo.replace("**", "")).strip()
            titulo_sintese_alvo = (
                _norm(t_sem_num) in m.sintese_titulos
                if m.sintese_titulos
                else _norm(t_sem_num).startswith(
                    ("sintese executiva", "sintese dos pontos", "sintese")
                )
            )
            if (
                titulo_sintese_alvo
                and nivel >= 2
                and (pv.n_sec == 0 or bool(m.sintese_titulos))
            ):
                if m.sintese_extratos:
                    pv.sintese(m.sintese_extratos, titulo=m.titulo_sintese)
                    capa = False
                    i += 1
                    continue
                em_sintese = True
                capa = False  # a síntese encerra a capa
                i += 1
                continue
            descarrega_sintese()
            if nivel == 1 and primeiro_h1:
                # H1 inicial = título da peça: aguarda endereçamentos (vêm em seguida no md)
                titulo_peca = re.sub(r"\*+", "", titulo)
                j = i + 1
                enderec = []
                while j < len(linhas):
                    s2 = linhas[j].strip()
                    if not s2:
                        j += 1
                        continue
                    if _e_enderecamento(s2):
                        enderec.append(s2.lstrip("#").strip().strip("*"))
                        j += 1
                        continue
                    break
                for e in enderec:
                    pv.enderecamento(e.upper())  # DNA do escritório: endereçamento em caixa alta
                # caixa do processo: primeira linha bold com número processual (aceita pontos/traços)
                while j < len(linhas) and not linhas[j].strip():
                    j += 1
                if j < len(linhas):
                    s2 = linhas[j].strip()
                    if s2.startswith("**") and re.search(r"\d[\d./\-]{3,}", s2):
                        pv.processo_caixa(s2.replace("**", "").strip())
                        j += 1
                # demais dados do processo em bold ficam na capa
                while j < len(linhas):
                    s3 = linhas[j].strip()
                    if not s3:
                        j += 1
                        continue
                    if s3.startswith("**") and len(s3) < 260 and not s3.startswith("#"):
                        pv.par(s3, 11, antes=0, depois=1)
                        j += 1
                        continue
                    break
                pv.titulo(titulo_peca)
                if m.tarja_interna:
                    pv.tarja(m.tarja_interna)
                if m.qualification_on_cover:
                    while j < len(linhas) and not linhas[j].strip():
                        j += 1
                    if j < len(linhas) and not linhas[j].lstrip().startswith(("#", "|", ">")):
                        qualificacao = [linhas[j].strip()]
                        j += 1
                        while j < len(linhas):
                            trecho = linhas[j].strip()
                            if not trecho or trecho.startswith(("#", "|", ">")):
                                break
                            qualificacao.append(trecho)
                            j += 1
                        pv.par(
                            " ".join(qualificacao),
                            10.5,
                            antes=6,
                            depois=5,
                        )
                if m.cards:
                    pv.marcador("{{CARDS}}", antes=10)
                # peça com endereçamento tem capa própria; consultivo sem endereçamento flui
                # na mesma página (capa só com título é "vazio morto" — proibido pela skill)
                if (enderec or m.force_cover_page_break) and not m.capa_com_sintese:
                    pv.quebra_pagina()
                primeiro_h1 = False
                i = j
                continue
            if capa and pv.n_sec == 0 and _e_enderecamento(titulo) and not re.match(r"^[IVXL0-9]+\s*[—\-–.]", titulo):
                pv.enderecamento(titulo.replace("**", ""))
                i += 1
                continue
            capa = False
            if nivel <= 2:
                # t_sem_num: o kit prefixa a numeração romana própria (evita "X — X. PEDIDO")
                if m.preserve_heading_labels:
                    titulo_literal = re.sub(r"\*+", "", titulo)
                    pv.abre(
                        titulo_literal,
                        m.linhas_sintese.get(_norm(titulo_literal), ""),
                        prefixo=False,
                    )
                else:
                    pv.abre(t_sem_num, m.linhas_sintese.get(_norm(t_sem_num), ""))
            else:
                pv.sub(re.sub(r"\*+", "", titulo))
            i += 1
            continue

        # linha de conteúdo
        if em_sintese:
            sintese_buf.append(strip)
            i += 1
            continue

        # lista com marcador
        if re.match(r"^>\s?", strip):
            citacao = []
            while i < len(linhas) and re.match(r"^\s*>\s?", linhas[i]):
                citacao.append(re.sub(r"^\s*>\s?", "", linhas[i]).strip())
                i += 1
            pv.pgf(" ".join(item for item in citacao if item))
            continue

        if re.match(r"^[-•]\s+", strip):
            itens = []
            while i < len(linhas) and re.match(r"^[-•]\s+", linhas[i].strip()):
                itens.append(re.sub(r"^[-•]\s+", "", linhas[i].strip()))
                i += 1
            pv.topicos(itens)
            continue

        # Parágrafo; por padrão a numeração é recomposta. Fontes canônicas já
        # auditadas podem exigir preservação literal, inclusive reinícios 1–N.
        numero_fonte = re.match(r"^\*{0,2}(\d{1,3})\.\*{0,2}\s+", strip)
        alinea_fonte = re.match(r"^([a-z])\)\s+", strip, re.I)
        par = re.sub(r"^\*{0,2}(\d{1,3})\.\*{0,2}\s+", "", strip)
        par_norm = _norm(strip)
        pos_paragrafo(par_norm)

        # capa: dados em bold antes da 1ª seção
        if capa and strip.startswith("**") and len(strip) < 260:
            pv.par(strip, 11, antes=0, depois=1)
            i += 1
            continue

        if re.search(r"pede[m]?\s+deferimento", strip, re.I):
            # O fecho e as assinaturas formam um bloco indivisível. Se forem
            # tratados como parágrafos comuns, a OAB pode ficar órfã na página
            # seguinte e a fórmula de deferimento recebe numeração indevida.
            formula = par.replace("**", "").strip()
            j = i + 1
            while j < len(linhas) and not linhas[j].strip():
                j += 1
            data = ""
            if j < len(linhas):
                candidata = linhas[j].strip().strip("*").strip()
                if re.match(r"^(Bras[íi]lia|Porto Alegre|Rio de Janeiro|S[ãa]o Paulo).{0,60}20\d\d\.?$", candidata):
                    data = candidata
                    j += 1

            assinaturas = []
            while j < len(linhas):
                while j < len(linhas) and not linhas[j].strip():
                    j += 1
                if j >= len(linhas):
                    break
                nome = linhas[j].strip().strip("*").strip()
                if not (nome == nome.upper() and 2 <= len(nome.split()) <= 8):
                    break
                k = j + 1
                while k < len(linhas) and not linhas[k].strip():
                    k += 1
                if k >= len(linhas):
                    break
                oab = linhas[k].strip().strip("*").strip()
                if not re.match(r"^OAB(?:[/ ]|$)", oab, re.I):
                    break
                assinaturas.append((nome, oab))
                j = k + 1

            if data:
                pv.fecho(data, formula=formula)
            else:
                pv.par(formula, antes=8, keep=bool(assinaturas))
            if assinaturas:
                pv.assinaturas(assinaturas)
            pos_deferimento = True
            i = j
            continue
        t_ass = strip.strip("*").strip()
        # fecho local+data liga o modo assinatura mesmo sem "pede deferimento" literal
        eh_fecho_data = re.match(r"^(Bras[íi]lia|Porto Alegre|Rio de Janeiro|S[ãa]o Paulo).{0,40}de\s+20\d\d", t_ass)
        if eh_fecho_data:
            pos_deferimento = True
        eh_assin = pos_deferimento and (
            (t_ass == t_ass.upper() and 2 <= len(t_ass.split()) <= 6)
            or re.match(r"^OAB[/ ]", t_ass, re.I)
            or re.match(r"^(Curador|Curadora|Advogado|Advogada|Representante|Assistente)\b", t_ass, re.I)
            or (strip.startswith("**") and 2 <= len(t_ass.split()) <= 6)
            or eh_fecho_data)
        if eh_assin:
            pv.par(strip, 12, antes=8 if eh_fecho_data else 2,
                   depois=1, align="center")
            i += 1
            continue

        texto_caixa = (
            strip
            if m.preserve_paragraph_labels and numero_fonte
            else par.replace("**", "")
        )
        if not como_caixa(par_norm, texto_caixa):
            if m.preserve_paragraph_labels:
                if numero_fonte:
                    pv.pgf_literal(numero_fonte.group(1), par)
                elif alinea_fonte:
                    pv.item_literal(
                        f"{alinea_fonte.group(1)})",
                        re.sub(r"^[a-z]\)\s+", "", par, flags=re.I),
                    )
                else:
                    pv.par(par)
            elif m.study_numbered_items and numero_fonte:
                pv.item_literal(f"{numero_fonte.group(1)}.", par)
            else:
                pv.pgf(par)
        figs_depois(par_norm)
        i += 1

    descarrega_sintese()
    pv.salvar()

    # ---------- GATE DE FIDELIDADE (bloqueante) ----------
    from docx import Document
    d = Document(str(out_docx))
    texto_docx_bruto = " ".join(p.text for p in d.paragraphs) + " " + " ".join(
        c.text for t in d.tables for r in t.rows for c in r.cells)
    leaks = re.findall(r"(?:^|\s)#{2,6}\s+|(?:^|\s)>\s+", texto_docx_bruto)
    markers = re.findall(r"\{\{[^}]+\}\}", texto_docx_bruto)
    leaks += [marker for marker in markers if marker not in m.allowed_markers]
    marker_counts = {marker: markers.count(marker) for marker in m.allowed_markers}
    invalid_counts = {marker: count for marker, count in marker_counts.items() if count != 1}
    if leaks:
        raise SystemExit(f"GATE DE GRAMÁTICA REPROVADO: marcadores Markdown no DOCX: {leaks[:8]}")
    if invalid_counts:
        raise SystemExit(f"GATE DE MARCADORES REPROVADO: cada marcador declarado deve ocorrer uma vez: {invalid_counts}")
    texto_docx = _norm(texto_docx_bruto)
    faltando = []
    for l in linhas:
        s = l.strip()
        if len(s) > 60 and not s.startswith("|") and not s.startswith("#") \
                and not re.match(r"^\s*([-*_])\1{2,}\s*$", s) \
                and not any(ig in s for ig in m.ignorar):
            s_sem_num = re.sub(r"^\*{0,2}(\d{1,3})\.\*{0,2}\s+", "", s)
            s_sem_num = re.sub(r"^[-•]\s+", "", s_sem_num)
            frag = _norm(s_sem_num)[:150]
            if frag and frag not in texto_docx:
                faltando.append(s[:110])
    if faltando:
        for f in faltando[:12]:
            print("FALTA:", f)
        raise SystemExit(f"GATE DE FIDELIDADE REPROVADO: {len(faltando)} parágrafos do md ausentes do DOCX")
    sobras = [x[1] for x in pulls_rest + lats_rest] + [x[2] for x in caixas_rest] + [x[1] for x in figs_rest]
    if sobras:
        raise SystemExit(f"MAPA NÃO CONSUMIDO — âncoras sem parágrafo correspondente: {sobras}")
    # Lastro persistido do gate (conselho 11/07/2026): sem este JSON com hash batendo,
    # o elo 4-B do forja_delivery não fecha a demanda. Pega DOCX visual de versão
    # errada/desatualizada (modo de falha do caso CASO-19 — Lição 48).
    import hashlib
    import json as _json
    from datetime import datetime as _dt, timezone as _tz
    out_docx = Path(out_docx)
    lastro = {
        "docx": str(out_docx),
        "docxSha256": hashlib.sha256(out_docx.read_bytes()).hexdigest(),
        "mdFonte": str(md_path),
        "mdSha256": hashlib.sha256(texto_md.encode("utf-8")).hexdigest(),
        "paragrafos": pv.n_pgf,
        "gate": "fidelidade 100% (gramática + marcadores + parágrafos + mapa consumido)",
        "geradoEm": _dt.now(_tz.utc).astimezone().isoformat(timespec="seconds"),
    }
    out_docx.with_name("FIDELIDADE_VISUAL.json").write_text(
        _json.dumps(lastro, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"OK fidelidade 100% | pgf: {pv.n_pgf} | seções: {pv.n_sec} | {out_docx}")
    return out_docx
