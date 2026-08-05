#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Diff automático entre duas versões DOCX (nossa vs. protocolada).

Módulo forja_diff_docx: extrai parágrafos de ambos os DOCX (ignorando vazios),
alinha via difflib.SequenceMatcher, classifica mudanças heuristicamente e gera
relatório markdown para realimentação em APRENDIZADOS_FEEDBACK_HUMANO.md.

Uso:
    python forja_diff_docx.py <nossa.docx> <protocolada.docx> [saida.md]

Saída: markdown com blocos de mudança (formato, estilo-voz, conteudo-juridico)
+ cabeçalho com estatísticas por classe + rodapé com sugestão de integração.
"""

import sys
import re
from pathlib import Path
from difflib import SequenceMatcher
from docx import Document


def extrair_paragrafos_docx(caminho_docx):
    """
    Extrai parágrafos de um DOCX, incluindo texto de tabelas.
    Ignora parágrafos vazios (espaços-em-branco apenas).
    Retorna lista de strings (parágrafos).
    """
    doc = Document(caminho_docx)
    paragrafos = []

    for parag in doc.paragraphs:
        texto = parag.text.strip()
        if texto:
            paragrafos.append(texto)

    # Adicionar texto de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for parag in celula.paragraphs:
                    texto = parag.text.strip()
                    if texto:
                        paragrafos.append(texto)

    return paragrafos


def similaridade_ratio(s1, s2):
    """
    Retorna razão de similaridade entre duas strings (0.0 a 1.0).
    """
    matcher = SequenceMatcher(None, s1, s2)
    return matcher.ratio()


def classificar_mudanca(texto_nosso, texto_protocolado):
    """
    Classifica uma mudança em uma das três categorias heurísticas:
    - "formato": mudança superficial (maiúscula, pontuação, espaço, numeração)
    - "estilo-voz": mesmas citações/números, palavras diferentes (ratio 0.5-0.95)
    - "conteudo-juridico": dispositivos/citações/números diferentes (ratio < 0.5)
    """
    ratio = similaridade_ratio(texto_nosso, texto_protocolado)

    # Heurística 1: detectar mudança puramente de formato
    # Remover espaços, caixa, pontuação — se ficar igual, é formato
    def normalizar_superficial(s):
        s = re.sub(r'\s+', '', s)
        s = s.lower()
        s = re.sub(r'[.,;:!?\(\)\[\]\-—"""]', '', s)
        return s

    norm_nosso = normalizar_superficial(texto_nosso)
    norm_protocolado = normalizar_superficial(texto_protocolado)

    if norm_nosso == norm_protocolado:
        return "formato"

    # Heurística 2: estilo-voz (50-95% similaridade)
    # Significa: estrutura parecida, palavras substituídas
    if 0.50 <= ratio < 0.95:
        return "estilo-voz"

    # Heurística 3: conteúdo jurídico (< 50% ou parágrafo inteiro novo)
    return "conteudo-juridico"


def gerar_diff_markdown(paragrafos_nosso, paragrafos_protocolado, saida_path=None):
    """
    Alinha dois conjuntos de parágrafos e gera relatório markdown.
    Usa SequenceMatcher para encontrar correspondências.
    """
    matcher = SequenceMatcher(None, paragrafos_nosso, paragrafos_protocolado)
    opcodes = matcher.get_opcodes()

    mudancas_por_classe = {"formato": 0, "estilo-voz": 0, "conteudo-juridico": 0}
    blocos_markdown = []

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            # Sem mudanças, pular
            continue

        elif tag == "replace":
            # Mudança: i1:i2 nosso -> j1:j2 protocolado
            for idx_nosso, idx_proto in zip(range(i1, i2), range(j1, j2)):
                texto_nosso = paragrafos_nosso[idx_nosso] if idx_nosso < len(paragrafos_nosso) else ""
                texto_proto = paragrafos_protocolado[idx_proto] if idx_proto < len(paragrafos_protocolado) else ""
                classe = classificar_mudanca(texto_nosso, texto_proto)
                mudancas_por_classe[classe] += 1

                bloco = f"""
### Mudança: {classe.upper()}

**Nosso:**
> {texto_nosso[:150]}...

**Protocolado:**
> {texto_proto[:150]}...

---
"""
                blocos_markdown.append(bloco)

            # Se houver desequilíbrio (mais linhas em um lado), registrar
            if i2 - i1 > j2 - j1:
                for idx in range(i1 + (j2 - j1), i2):
                    classe = "conteudo-juridico"
                    mudancas_por_classe[classe] += 1
                    texto_nosso = paragrafos_nosso[idx]
                    blocos_markdown.append(f"""
### Mudança: {classe.upper()} (DELETADO)

**Nosso (removido no protocolado):**
> {texto_nosso[:150]}...

---
""")

            elif j2 - j1 > i2 - i1:
                for idx in range(j1 + (i2 - i1), j2):
                    classe = "conteudo-juridico"
                    mudancas_por_classe[classe] += 1
                    texto_proto = paragrafos_protocolado[idx]
                    blocos_markdown.append(f"""
### Mudança: {classe.upper()} (INSERIDO)

**Protocolado (novo):**
> {texto_proto[:150]}...

---
""")

        elif tag == "delete":
            # Parágrafos removidos (nosso i1:i2, protocolado vazio)
            for idx in range(i1, i2):
                classe = "conteudo-juridico"
                mudancas_por_classe[classe] += 1
                texto_nosso = paragrafos_nosso[idx]
                blocos_markdown.append(f"""
### Mudança: {classe.upper()} (DELETADO)

**Nosso (removido no protocolado):**
> {texto_nosso[:150]}...

---
""")

        elif tag == "insert":
            # Parágrafos inseridos (protocolado j1:j2, nosso vazio)
            for idx in range(j1, j2):
                classe = "conteudo-juridico"
                mudancas_por_classe[classe] += 1
                texto_proto = paragrafos_protocolado[idx]
                blocos_markdown.append(f"""
### Mudança: {classe.upper()} (INSERIDO)

**Protocolado (novo):**
> {texto_proto[:150]}...

---
""")

    # Montar cabeçalho com estatísticas
    total_mudancas = sum(mudancas_por_classe.values())
    cabecalho = f"""# Diff Automático: Nossa Versão vs. Protocolada

## Estatísticas

- Total de mudanças detectadas: {total_mudancas}
- Formato (pontuação, maiúscula, espaço): {mudancas_por_classe['formato']}
- Estilo-voz (mesma estrutura, palavras diferentes): {mudancas_por_classe['estilo-voz']}
- Conteúdo jurídico (dispositivos/citações/estrutura alterada): {mudancas_por_classe['conteudo-juridico']}

---

## Mudanças Detalhadas

"""

    # Montar rodapé com instruções de integração
    rodape = """

---

## Próximos passos

1. Revisar as mudanças de **conteudo-juridico** em detalhes (são as mais relevantes para aprendizado).
2. Classificar manualmente as que a heurística errasse.
3. Colar o bloco de ESTATÍSTICAS revisado em APRENDIZADOS_FEEDBACK_HUMANO.md.
4. Atualizar o protocolo da fábrica (CLAUDE.md) com lições estruturais.

Gerado automaticamente por forja_diff_docx.py.
"""

    markdown_completo = cabecalho + "\n".join(blocos_markdown) + rodape

    # Salvar em arquivo
    if saida_path:
        with open(saida_path, "w", encoding="utf-8") as f:
            f.write(markdown_completo)
        print(f"[OK] Diff salvo em: {saida_path}")
    else:
        print(markdown_completo)

    return markdown_completo, mudancas_por_classe


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    caminho_nosso = Path(sys.argv[1])
    caminho_protocolado = Path(sys.argv[2])
    saida = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    if not caminho_nosso.exists():
        print(f"ERRO: {caminho_nosso} não encontrado", file=sys.stderr)
        sys.exit(1)

    if not caminho_protocolado.exists():
        print(f"ERRO: {caminho_protocolado} não encontrado", file=sys.stderr)
        sys.exit(1)

    print(f"[...] Extraindo parágrafos de {caminho_nosso.name}...")
    paragrafos_nosso = extrair_paragrafos_docx(str(caminho_nosso))
    print(f"      {len(paragrafos_nosso)} parágrafos")

    print(f"[...] Extraindo parágrafos de {caminho_protocolado.name}...")
    paragrafos_protocolado = extrair_paragrafos_docx(str(caminho_protocolado))
    print(f"      {len(paragrafos_protocolado)} parágrafos")

    print("[...] Gerando diff...")
    markdown, stats = gerar_diff_markdown(paragrafos_nosso, paragrafos_protocolado, saida)

    print("\n[OK] Diff completo.")
    print(f"     Formato: {stats['formato']} | Estilo-voz: {stats['estilo-voz']} | Conteúdo: {stats['conteudo-juridico']}")


if __name__ == "__main__":
    main()
