# -*- coding: utf-8 -*-
"""Regressão da materialização SVG nativa da FORJA (sem renderização)."""

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from forja_svg_docx import inserir_svgs
from forja_visual_qa_structural import auditar_documento


SVG_OK = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 60">
<rect x="0" y="0" width="200" height="60" fill="#395C60"/>
<text x="5" y="30" font-size="12" fill="#ffffff">OK</text>
</svg>"""


class SvgDocxTests(unittest.TestCase):
    def _fixture(self, root: Path, marker="{{FIG1}}"):
        document = Document()
        document.add_paragraph("Texto de apoio")
        document.add_paragraph(marker)
        docx = root / "peca.docx"
        document.save(docx)
        svg = root / "fig.svg"
        svg.write_text(SVG_OK, encoding="utf-8")
        return docx, svg

    def test_embute_svg_sem_pdf_e_sem_marcador(self):
        """O pacote leva o vetor, não leva PDF e não deixa marcador para trás.

        A asserção de que o pacote não podia conter PNG foi retirada em
        07/08/2026: ela codificava o desenho quebrado. O suporte a SVG do
        Office é uma extensão do `blip`, e o `blip` precisa apontar para um
        raster; sem ele o Word recusa o documento inteiro. O que a rota
        continua não fazendo é RENDERIZAR para materializar a peça — e isso
        quem prova é `test_qa_registra_explicitamente_ausencia_de_render`.
        """
        with tempfile.TemporaryDirectory() as tmp:
            docx, svg = self._fixture(Path(tmp))
            resultado = inserir_svgs(docx, {"{{FIG1}}": (svg, 10.0)})
            self.assertEqual(resultado["{{FIG1}}"]["contentType"], "image/svg+xml")
            with zipfile.ZipFile(docx) as archive:
                nomes = archive.namelist()
                self.assertTrue(any(nome.endswith(".svg") for nome in nomes))
                self.assertFalse(any(nome.endswith(".pdf") for nome in nomes))
                self.assertNotIn("{{FIG1}}", archive.read("word/document.xml").decode("utf-8"))
                self.assertIn('ContentType="image/svg+xml"', archive.read("[Content_Types].xml").decode("utf-8"))

    def test_estrutura_que_o_word_exige(self):
        """O invariante que faltava: OOXML que o Word aceita, e não só que abre no lxml.

        De 03/08 a 07/08/2026 a rota canônica pendurou `wp:inline` direto em
        `w:p`, sem `w:r/w:drawing`. Toda biblioteca Python lia o pacote sem
        reclamar e o Microsoft Word recusava o arquivo inteiro, com mensagem de
        documento corrompido — o destinatário não veria nem o texto.

        Abrir o Word em teste não é viável aqui, então o teste guarda as três
        marcas estruturais que a correção estabeleceu. Elas não substituem
        abrir o arquivo no programa do leitor; tornam a regressão barata.
        """
        with tempfile.TemporaryDirectory() as tmp:
            docx, svg = self._fixture(Path(tmp))
            inserir_svgs(docx, {"{{FIG1}}": (svg, 10.0)})
            with zipfile.ZipFile(docx) as archive:
                doc = archive.read("word/document.xml").decode("utf-8")
                nomes = archive.namelist()
            self.assertIn("<w:drawing>", doc, "figura fora de w:drawing")
            self.assertNotIn("<w:p><wp:inline", doc, "wp:inline pendurado direto no parágrafo")
            self.assertIn("svgBlip", doc, "vetor não declarado na extensão de SVG do Office")
            self.assertTrue(any(nome.endswith(".png") for nome in nomes),
                            "sem raster de reserva o Word recusa o documento")

    def test_qa_registra_explicitamente_ausencia_de_render(self):
        with tempfile.TemporaryDirectory() as tmp:
            docx, svg = self._fixture(Path(tmp))
            inserir_svgs(docx, {"{{FIG1}}": (svg, 10.0)})
            laudo = auditar_documento(docx, svgs=[svg])
            self.assertFalse(laudo["renderingUsed"])
            self.assertFalse(laudo["pdfCreated"])
            self.assertFalse(laudo["pngCreated"])
            self.assertTrue(laudo["package"]["approved"])

    def test_marcador_multiplo_bloqueia(self):
        with tempfile.TemporaryDirectory() as tmp:
            docx, svg = self._fixture(Path(tmp), marker="{{FIG1}} e {{FIG1}}")
            with self.assertRaises(RuntimeError):
                inserir_svgs(docx, {"{{FIG1}}": (svg, 10.0)})

    def test_svg_invalido_bloqueia_antes_de_alterar_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            docx, svg = self._fixture(Path(tmp))
            svg.write_text("<svg><rect></svg>", encoding="utf-8")
            with self.assertRaises(Exception):
                inserir_svgs(docx, {"{{FIG1}}": (svg, 10.0)})
            with zipfile.ZipFile(docx) as archive:
                self.assertFalse(any(nome.endswith(".svg") for nome in archive.namelist()))


if __name__ == "__main__":
    unittest.main(verbosity=2)
