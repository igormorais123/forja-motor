# -*- coding: utf-8 -*-
"""test_forja_layout_antimoldagem.py — o gate de layout ainda sabe reprovar?

Em 04/08/2026 eu ajustei o gate de diagramação QUATRO vezes, e todas as quatro na
mesma direção: afrouxar, em favor do que a rota canônica e as peças da casa
produzem. Síntese executiva a 10,5 pt, fólio de 57,3 pt, Segoe UI nas tabelas,
título em caixa mista e pull quote em moldura. Cada ajuste tinha justificativa
boa — o gate estava reprovando o padrão aprovado pelo dono, e nesse caso o errado
é o gate. Ao fim do dia a rota passou a marcar 100% nas três dimensões.

Esse é exatamente o formato de um autoengano caro. Quatro afrouxamentos seguidos e
um verde perfeito no fim são indistinguíveis, pelo resultado, de um gate moldado
até aprovar o que quer que o produtor emita.

O que distingue as duas coisas não é a intenção de quem mexeu: é a existência de
uma prova de que o gate CONTINUA reprovando defeito de verdade. Os testes-âncora
provam o lado fácil — que a peça aprovada passa. Este prova o lado difícil.

Método: pegar uma peça aprovada, ESTRAGÁ-LA de propósito no OOXML e exigir que o
gate acuse. Se ele aprovar a peça estragada, o afrouxamento foi longe demais e a
suíte falha — mesmo com todo o resto verde.

Estragos aplicados ao corpo, um de cada vez:
  - `tamanho`  : corpo a 11 pt, o tamanho que não existe no padrão da casa e que
                 foi encontrado de verdade na V9 do nono tópico.
  - `alinhar`  : corpo à esquerda em vez de justificado.
  - `fonte`    : corpo em Calibri, fonte que não pertence à identidade da casa.
  - `tabela`   : tabela em Arial, para provar que a aceitação de Segoe UI não
                 abriu a porta para qualquer fonte.

Uso: python test_forja_layout_antimoldagem.py   (exit 0 = ok; exit 1 = regressão)
"""
from __future__ import annotations

import io
import shutil
import sys
import tempfile
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from forja_docx_layout import audit_docx_layout  # noqa: E402

RAIZ = Path(__file__).resolve().parent
FABRICA = RAIZ.parent

# Peça aprovada e entregue, que hoje passa limpa no gate. É a base do estrago.
from forja_baseline_aprovado import caminho_da_ancora  # noqa: E402

# A peça-base vem do registro de âncoras do acervo, e não de um caminho
# escrito aqui: o motor não guarda nome de pasta de cliente.
BASE = caminho_da_ancora("peca-aprovada-v8") or (FABRICA / "__sem_acervo__.docx")

# Cada estrago tem de produzir PELO MENOS um destes códigos. Não se exige o código
# exato para não amarrar o teste à implementação: exige-se que o gate acuse algo
# da família certa.
ESTRAGOS = {
    "tamanho": {"body_font_size_not_12pt", "body_typography_unresolved"},
    "alinhar": {"body_text_not_justified"},
    "fonte": {"body_font_not_medina", "body_typography_unresolved"},
    "tabela": {"table_typography_inconsistent", "table_font_below_minimum"},
}

# Piso de material: estragar 3 parágrafos num documento de 200 pode não mover
# nenhuma métrica. O estrago tem de alcançar massa suficiente para ser um defeito
# de verdade, e não ruído.
PARAGRAFOS_A_ESTRAGAR = 12


def _e_corpo(paragrafo) -> bool:
    texto = paragrafo.text.strip()
    return len(texto) >= 80 and paragrafo.runs and not texto.isupper()


def _estragar(origem: Path, destino: Path, modo: str) -> int:
    shutil.copy2(origem, destino)
    documento = Document(str(destino))
    tocados = 0

    if modo == "tabela":
        for tabela in documento.tables[:2]:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            if run.text.strip():
                                run.font.name = "Arial"
                                run.font.size = Pt(7)
                                tocados += 1
        documento.save(str(destino))
        return tocados

    for paragrafo in documento.paragraphs:
        if tocados >= PARAGRAFOS_A_ESTRAGAR:
            break
        if not _e_corpo(paragrafo):
            continue
        if modo == "alinhar":
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragrafo.runs:
            if modo == "tamanho":
                run.font.size = Pt(11)
            elif modo == "fonte":
                run.font.name = "Calibri"
        tocados += 1

    documento.save(str(destino))
    return tocados


def main() -> int:
    falhas = 0

    if not BASE.is_file():
        print(f"  FALHOU: a peça-base sumiu do acervo — {BASE.name}")
        print("REGRESSÃO: sem peça aprovada não há como provar que o gate reprova defeito")
        return 1

    limpo = audit_docx_layout(BASE)
    codigos_limpos = {f.get("code") for f in (limpo.get("findings") or [])}
    if codigos_limpos:
        print(f"  FALHOU: a peça-base já vem com achados {sorted(codigos_limpos)} — "
              "o estrago não provaria nada, porque o gate já estava reclamando")
        falhas += 1

    with tempfile.TemporaryDirectory(prefix="forja-antimoldagem-") as tmp:
        for modo, esperados in ESTRAGOS.items():
            destino = Path(tmp) / f"estragado_{modo}.docx"
            tocados = _estragar(BASE, destino, modo)
            if tocados < 3:
                print(f"  FALHOU: o estrago '{modo}' alcançou só {tocados} elemento(s) — "
                      "o teste ficaria verde por não ter estragado nada")
                falhas += 1
                continue
            laudo = audit_docx_layout(destino)
            codigos = {f.get("code") for f in (laudo.get("findings") or [])}
            if not (codigos & esperados):
                print(f"  FALHOU: o gate APROVOU a peça estragada em '{modo}' "
                      f"({tocados} elementos alterados). Esperava algo de {sorted(esperados)}, "
                      f"obteve {sorted(codigos) or 'nada'} — o afrouxamento foi longe demais")
                falhas += 1

    if falhas:
        print(f"REGRESSÃO: {falhas} verificação(ões) anti-moldagem falharam")
        return 1
    print(f"ok: a peça aprovada passa limpa e os {len(ESTRAGOS)} estragos deliberados são "
          "acusados — o gate não foi moldado até aprovar tudo")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
