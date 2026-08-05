# -*- coding: utf-8 -*-
"""test_forja_layout_papeis.py — teste-âncora dos papéis estruturais do padrão da casa.

Em 04/08/2026 os dezesseis gates da F8 foram exercidos pela primeira vez contra
peças REAIS e entregues — quatro delas, aprovadas pelo escritório. O gate de
tipografia reprovou nas quatro, e a maior parte das reprovações era falsa:

  - o bloco de qualificação das partes na V8 do Cafelana (10,5 pt);
  - cinco parágrafos da SÍNTESE EXECUTIVA nos memoriais do AI 0011621-15
    (10,5 pt) — a síntese no estilo do art. 343-A do RISTJ é obrigatória em
    TODA peça desde 07/07/2026, por determinação do Prof. Fábio;
  - uma citação recuada transcrita do próprio acórdão impugnado (10,5 pt).

Todos deliberados, todos parte do padrão aprovado. Vale aqui a regra que o
plano visual já tinha escrito com outras palavras: **gate que reprova o padrão
aprovado pelo dono está errado, não a peça.** Se o F8-S tivesse sido tornado
bloqueante antes desta medição, ele barraria o formato do próprio escritório em
todas as peças que carregam a síntese executiva — ou seja, em todas.

Este teste é o ÂNCORA: ele fixa que os três papéis continuam reconhecidos, e
fixa também o limite do afrouxamento. A exceção vale para o TAMANHO e não para
a família tipográfica nem para a justificação — texto corrido a 10,5 pt fora
desses papéis continua reprovado, senão bastaria começar um parágrafo com aspas
para escapar da régua.

Uso: python test_forja_layout_papeis.py   (exit 0 = ok; exit 1 = regressão)
"""
from __future__ import annotations

import io
import sys
import tempfile
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from forja_docx_layout import _role_for, audit_docx_layout  # noqa: E402


class _Estilo:
    def __init__(self, nome="Normal"):
        self.name = nome


class _XML:
    """Parágrafo sem outlineLvl nem borda: o caso comum do corpo do texto."""

    def xpath(self, _expr):
        return []

    def find(self, _nome):
        return None


class _Par:
    """O mínimo que `_role_for` consulta: estilo, runs e o XML do parágrafo."""

    def __init__(self, nome="Normal"):
        self.style = _Estilo(nome)
        self.runs = []
        self._p = _XML()


def _write_layout_fixture(
    path: Path,
    *,
    wrong_font_role: str | None = None,
    wrong_alignment_role: str | None = None,
    body_size: float = 12.0,
) -> None:
    document = Document()
    entries = [
        ("executive_summary", "I. CASO EM EXAME. Ação fundada em perdas atribuídas a aplicações do IGEPREV/TO."),
        ("heading", "I — QUALIFICAÇÃO"),
        ("qualification", "CAFELANA COMÉRCIO E REPRESENTAÇÕES DO BRASIL LTDA., já qualificadas nos autos em epígrafe, na condição de AGRAVADAS, por seu advogado"),
        ("block_quote", "“Não se exige, nesta fase, exauriente avaliação do elemento subjetivo do tipo, mas tão somente a verificação da presença dos requisitos mínimos exigidos pelo art. 17, § 6º-B do CPC, conforme a prova documental disponível.”"),
        ("body", "O acórdão recorrido, ao manter a decisão agravada, deixou de enfrentar a questão do elemento subjetivo, o que caracteriza omissão qualificada e exige novo exame do fundamento central."),
    ]
    sizes = {"executive_summary": 10.5, "qualification": 10.5, "block_quote": 10.5, "body": body_size}
    for role, text in entries:
        paragraph = document.add_heading(text, level=1) if role == "heading" else document.add_paragraph(text)
        actual_role = role
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if actual_role == wrong_alignment_role
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        run = paragraph.runs[0]
        run.font.name = "Arial" if actual_role == wrong_font_role else "Times New Roman"
        run.font.size = Pt(sizes.get(actual_role, 12.0))
    document.save(path)


def main() -> int:
    falhas = 0
    casos = 0

    def papel(texto, *, index=5, anterior=None, estilo="Normal"):
        return _role_for(_Par(estilo), texto, index, anterior)

    def checar(nome, obtido, esperado):
        nonlocal falhas, casos
        casos += 1
        if obtido != esperado:
            print(f"  FALHOU: {nome} — esperado '{esperado}', obtido '{obtido}'")
            falhas += 1

    # Os três papéis, com o texto REAL das peças que o gate reprovou.
    checar("síntese executiva — I. CASO EM EXAME",
           papel("I. CASO EM EXAME. 1. Ação fundada em perdas atribuídas a aplicações do "
                 "IGEPREV/TO no Fundo Vitória Régia."), "executive_summary")
    checar("síntese executiva — II. QUESTÃO EM DISCUSSÃO",
           papel("II. QUESTÃO EM DISCUSSÃO. 2. Saber se a decisão enfrentou o art. 17."),
           "executive_summary")
    checar("síntese executiva — III. RAZÕES DE DECIDIR",
           papel("III. RAZÕES DE DECIDIR. 3. O acórdão não distinguiu dolo específico."),
           "executive_summary")
    checar("abertura de ementa em caixa alta",
           papel("DIREITO ADMINISTRATIVO SANCIONADOR E PROCESSUAL CIVIL. Improbidade "
                 "administrativa. Agravo de instrumento.", index=4), "executive_summary")
    checar("rótulo EMENTA isolado é cabeçalho",
           papel("EMENTA", index=4), "heading")
    checar("continuação da síntese herda o papel",
           papel("4. Recurso provido em parte para determinar novo exame.",
                 anterior="executive_summary"), "executive_summary")
    checar("bloco de qualificação das partes",
           papel("CAFELANA COMÉRCIO E REPRESENTAÇÕES DO BRASIL LTDA., MACHADENSE CAFÉ LTDA. e "
                 "INTER CONTINENTAL DE CAFÉ S.A., já qualificadas nos autos em epígrafe, na "
                 "condição de AGRAVADAS, por seu advogado", index=3), "qualification")
    checar("citação recuada longa",
           papel("“Não se exige, nesta fase, exauriente avaliação do elemento subjetivo do tipo, "
                 "mas tão somente a verificação da presença dos requisitos mínimos exigidos pelo "
                 "art. 17, § 6º-B do CPC.”"), "block_quote")

    # O LIMITE do afrouxamento — sem isto, a exceção vira porta dos fundos.
    checar("corpo comum continua corpo",
           papel("O acórdão recorrido, ao manter a decisão agravada, deixou de enfrentar a "
                 "questão do elemento subjetivo, o que caracteriza omissão qualificada."),
           "body")
    checar("aspas curtas não viram citação recuada",
           papel("“omissão qualificada”, como se verá."), "body")
    checar("qualificação fora da abertura não é qualificação",
           papel("A parte já qualificada nos autos apresentou nova manifestação sobre o ponto, "
                 "sem inovar quanto à tese central do recurso.", index=40), "body")
    checar("caixa alta longa depois da abertura não é ementa",
           papel("DIREITO ADMINISTRATIVO SANCIONADOR E PROCESSUAL CIVIL. Improbidade "
                 "administrativa novamente mencionada no corpo.", index=60), "body")

    # A exceção de tamanho não pode desligar os outros dois controles do gate.
    # Este fixture reproduz os três papéis em DOCX, e não apenas a classificação
    # textual acima: é a regressão que teria capturado o bypass na primeira leva.
    with tempfile.TemporaryDirectory() as temp:
        root = Path(temp)
        aprovado = root / "aprovado.docx"
        _write_layout_fixture(aprovado)
        laudo = audit_docx_layout(aprovado)
        checar("papéis especiais a 10,5 pt continuam aprovados quando corretos",
               laudo["approved"], True)

        fonte = root / "fonte-estrutural.docx"
        _write_layout_fixture(fonte, wrong_font_role="qualification")
        laudo_fonte = audit_docx_layout(fonte)
        checar("qualificação com fonte errada é reprovada",
               "structural_font_not_medina" in {item["code"] for item in laudo_fonte["findings"]}, True)

        alinhamento = root / "alinhamento-estrutural.docx"
        _write_layout_fixture(alinhamento, wrong_alignment_role="block_quote")
        laudo_alinhamento = audit_docx_layout(alinhamento)
        checar("citação recuada não justificada é reprovada",
               "structural_text_not_justified" in {item["code"] for item in laudo_alinhamento["findings"]}, True)

        corpo = root / "corpo-10-5.docx"
        _write_layout_fixture(corpo, body_size=10.5)
        laudo_corpo = audit_docx_layout(corpo)
        checar("corpo comum a 10,5 pt continua reprovado",
               "body_font_size_not_12pt" in {item["code"] for item in laudo_corpo["findings"]}, True)

    if falhas:
        print(f"REGRESSÃO: {falhas} de {casos} papéis do padrão da casa mudaram")
        return 1
    print(f"ok: {casos} verificações — síntese executiva, qualificação e citação recuada são "
          "reconhecidas como estrutura, e o corpo comum continua cobrado a 12 pt")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
