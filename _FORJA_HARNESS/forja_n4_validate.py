"""Aggregate validator and shadow status materializer for FORJA N4."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Callable

import jsonschema

from forja_case_tests import run_suite, validate_results, validate_suite
from forja_consistency import (
    validate_comparison,
    validate_delivery,
    validate_event_identity,
    validate_global,
    validate_intertemporal,
    validate_quantification,
)
from forja_learning import validate_learning
from forja_learning_registry import suite_learning_findings
from forja_post_protocol_contracts import (
    validate_document_comparison as validate_post_protocol_document_comparison,
    validate_learning_candidate,
    validate_post_protocol_baseline_backfill,
    validate_post_protocol_return,
    validate_protocol_evidence,
)
from forja_metacognition import validate_metacognition
from forja_f8_contract import validate_f8
from forja_n3_common import (
    PHASES,
    ForjaN3Error,
    atomic_write_json,
    canonical_hash,
    load_config,
    read_json,
    resolve_case_dir,
    sha256_file,
)
from forja_n4_common import ARTIFACT_SPECS, SPEC_VERSION, append_trace, issue, validate_envelope
from forja_reasoning import (
    validate_conducts,
    validate_coverage,
    validate_decision_factors,
    validate_graph,
    validate_question_tree,
    validate_recipient_map,
    validate_signature_brief,
    validate_theses,
)
from forja_science import (
    validate_audit,
    validate_claims,
    validate_classification,
    validate_protocol,
    validate_studies,
    validate_synthesis,
)


Validator = Callable[[dict], list[dict]]


def _settlement(payload: dict) -> list[dict]:
    findings = []
    if payload.get("applicability") != "not_applicable":
        if not payload.get("prohibitions"):
            findings.append(issue("N4-SETTLEMENT-PROHIBITIONS", "estratégia sem limites explícitos"))
        for item in payload.get("possibleConcessions") or []:
            if not item.get("trigger") or not item.get("proceduralEffect"):
                findings.append(issue("N4-SETTLEMENT-CONDITION", "concessão sem gatilho ou efeito processual"))
        for field in ("interests", "nonNegotiables"):
            for item in payload.get(field) or []:
                if isinstance(item, dict) and not item.get("sourceId") and not item.get("humanDecisionId"):
                    findings.append(issue("N4-SETTLEMENT-SOURCE", f"{field}: item sem fonte ou decisão humana"))
    return findings


VALIDATORS: dict[str, Validator] = {
    "F2_N4_CLASSIFICATION.json": validate_classification,
    "F2_QUESTION_TREE.json": validate_question_tree,
    "F3_EVENT_IDENTITY.json": validate_event_identity,
    "F3_DOCUMENT_COMPARISON.json": validate_comparison,
    "F3_REASONING_GRAPH.json": validate_graph,
    "F3_CONDUCT_LEDGER.json": validate_conducts,
    "F4_COVERAGE_MATRIX.json": validate_coverage,
    "F4_THESIS_MATURITY.json": validate_theses,
    "F4_CASE_ACCEPTANCE_TESTS.json": validate_suite,
    "F4_DECISION_FACTOR_MAP.json": validate_decision_factors,
    "F4_SETTLEMENT_MAP.json": _settlement,
    "F4_INTERTEMPORAL_MAP.json": validate_intertemporal,
    "F4_QUANTIFICATION_SCENARIOS.json": validate_quantification,
    "F5C_RESEARCH_PROTOCOL.json": validate_protocol,
    "F5C_STUDY_LEDGER.json": validate_studies,
    "F5C_EVIDENCE_SYNTHESIS.json": validate_synthesis,
    "F7_CASE_TEST_RESULTS.json": validate_results,
    "F7_GLOBAL_CONSISTENCY.json": validate_global,
    "F7_METACOGNITIVE_AUDIT.json": validate_metacognition,
    "F7_SCIENCE_AUDIT.json": validate_audit,
    "F9_DELIVERY_SELECTION.json": validate_delivery,
    "F10_DELIVERY_INTEGRITY.json": validate_delivery,
    "F10_HUMAN_DIFF_CLASSIFICATION.json": validate_learning,
    "F10_POST_PROTOCOL_RETURN.json": validate_post_protocol_return,
    "F10_PROTOCOL_EVIDENCE.json": validate_protocol_evidence,
    "F10_POST_PROTOCOL_BASELINE_BACKFILL.json": validate_post_protocol_baseline_backfill,
    "F10_POST_PROTOCOL_DOCUMENT_COMPARISON.json": validate_post_protocol_document_comparison,
    "F10_LEARNING_CANDIDATE.json": validate_learning_candidate,
    "F3_MAPA_DESTINATARIO.json": validate_recipient_map,
    "F4_SIGNATURE_BRIEF.json": validate_signature_brief,
}


# Exigidos apenas quando `forjaAssinaturaLite.mode` não for `off`.
SIGNATURE_LITE_FILES = {"F3_MAPA_DESTINATARIO.json", "F4_SIGNATURE_BRIEF.json"}


def _recipient_map_validator(config: dict) -> Validator:
    """Liga o validador do mapa ao limite de frescor configurado na feature."""
    horas = (config.get("forjaAssinaturaLite") or {}).get("recipientMapFreshnessHours")
    return lambda payload: validate_recipient_map(
        payload, freshness_hours=horas if isinstance(horas, int) else None
    )


FLAG_FILES = {
    "n4QuestionTreeV1": ["F2_QUESTION_TREE.json"],
    "n4CoverageMatrixV1": ["F4_COVERAGE_MATRIX.json"],
    "n4ReasoningGraphV1": ["F3_REASONING_GRAPH.json"],
    "n4CaseTestsV1": ["F4_CASE_ACCEPTANCE_TESTS.json", "F7_CASE_TEST_RESULTS.json"],
    "n4TerminologyV1": ["F3_EVENT_IDENTITY.json"],
    "n4DocumentComparisonV1": ["F3_DOCUMENT_COMPARISON.json"],
    "n4IntertemporalV1": ["F4_INTERTEMPORAL_MAP.json"],
    "n4QuantificationV1": ["F4_QUANTIFICATION_SCENARIOS.json"],
    "n4ScienceEvidenceV1": ["F2_N4_CLASSIFICATION.json", "F5C_RESEARCH_PROTOCOL.json", "F5C_STUDY_LEDGER.json", "F5C_EVIDENCE_SYNTHESIS.json", "F5C_CLAIM_EVIDENCE_MAP.json", "F7_SCIENCE_AUDIT.json"],
    "n4MetacognitiveAuditV1": ["F7_METACOGNITIVE_AUDIT.json"],
    "n4ConditionalStrategyV1": ["F3_CONDUCT_LEDGER.json", "F4_THESIS_MATURITY.json", "F4_DECISION_FACTOR_MAP.json", "F4_SETTLEMENT_MAP.json"],
    "n4LearningV1": ["F10_HUMAN_DIFF_CLASSIFICATION.json"],
    "n4DeliveryIntegrityV1": ["F9_DELIVERY_SELECTION.json", "F10_DELIVERY_INTEGRITY.json"],
}


PILOT_BLOCKING_PREFIXES = (
    "N4-Q-", "N4-COV-", "N4-EVENT-", "N4-TEST-", "N4-SCI-", "N4-GLOBAL-",
    "N4-DELIVERY-", "N4-CROSS-", "N4-ENV-", "N4-SCHEMA", "N4-SOURCE-", "N4-MISSING",
)


def _pilot_blocking_finding(item: dict) -> bool:
    return str(item.get("code") or "").startswith(PILOT_BLOCKING_PREFIXES)


def _target_phase(case_dir: Path) -> str:
    state = read_json(case_dir / "FORJA_N3_STATE.json", {}) or {}
    completed = [phase for phase in state.get("completedPhases") or [] if phase in PHASES]
    cursor = state.get("phaseCursor")
    candidates = completed + ([cursor] if cursor in PHASES else [])
    return max(candidates, key=PHASES.index) if candidates else PHASES[-1]


def _required_files(config: dict) -> set[str]:
    features = config.get("features") or {}
    required: set[str] = set()
    for flag, files in FLAG_FILES.items():
        if features.get(flag):
            required.update(files)
    if required:
        required.update({"F2_N4_CLASSIFICATION.json", "F7_GLOBAL_CONSISTENCY.json"})
    return required


def _source_registry_findings(case_dir: Path, case_manifest: dict, *, require_verifiable: bool = False) -> tuple[set[str], list[dict]]:
    registered: set[str] = set()
    findings: list[dict] = []
    registry = case_manifest.get("n4SourceRegistry") or {}
    for source_id, raw in registry.items():
        if isinstance(raw, str):
            registered.add(raw)
            if require_verifiable:
                findings.append(issue("N4-SOURCE-OPAQUE", f"{source_id}: hash sem caminho verificável não é aceito em modo bloqueante"))
            continue
        if not isinstance(raw, dict):
            findings.append(issue("N4-SOURCE-REGISTRY", f"{source_id}: registro de fonte inválido"))
            continue
        digest = str(raw.get("sha256") or raw.get("hash") or "").strip()
        if digest:
            registered.add(digest)
        status = str(raw.get("status") or "active")
        if status != "active":
            findings.append(issue("N4-SOURCE-REVOKED", f"{source_id}: fonte {status}: {raw.get('reason') or 'sem justificativa'}"))
        source_path = Path(str(raw.get("path") or ""))
        if str(source_path) in {"", "."}:
            if require_verifiable:
                findings.append(issue("N4-SOURCE-NO-PATH", f"{source_id}: fonte ativa sem caminho verificável"))
        else:
            if not source_path.is_absolute():
                source_path = case_dir / source_path
            if not source_path.is_file():
                findings.append(issue("N4-SOURCE-MISSING", f"{source_id}: fonte registrada não localizada: {source_path}"))
            elif digest and sha256_file(source_path) != digest:
                findings.append(issue("N4-SOURCE-DRIFT", f"{source_id}: conteúdo atual diverge do hash registrado"))
        origin_path = Path(str(raw.get("originPath") or ""))
        if str(origin_path) not in {"", "."}:
            if not origin_path.is_absolute():
                origin_path = case_dir / origin_path
            if origin_path.is_file():
                prefix = origin_path.read_text(encoding="utf-8-sig", errors="ignore")[:4096].lower()
                if "esta fonte foi invalidada" in prefix or "source status: revoked" in prefix:
                    findings.append(issue("N4-SOURCE-ORIGIN-REVOKED", f"{source_id}: arquivo de origem declara invalidação: {origin_path}"))
    return registered, findings


def _registered_source_path(case_dir: Path, case_manifest: dict, digest: str) -> Path | None:
    for raw in (case_manifest.get("n4SourceRegistry") or {}).values():
        if not isinstance(raw, dict) or str(raw.get("status") or "active") != "active":
            continue
        registered = str(raw.get("sha256") or raw.get("hash") or "").strip()
        if registered != digest:
            continue
        path = Path(str(raw.get("path") or ""))
        if str(path) in {"", "."}:
            continue
        if not path.is_absolute():
            path = case_dir / path
        if path.is_file() and sha256_file(path) == digest:
            return path
    return None


def _result_core(payload: dict) -> dict:
    return {
        key: payload.get(key)
        for key in ("suiteHash", "draftHash", "results", "approved", "antiFraud")
    }


def _council_ready(theses: dict) -> bool:
    items = theses.get("theses") or []
    accepted = {"adopt", "adopt_with_qualification"}
    return bool(items) and all(
        item.get("helenaDecision") in accepted
        and item.get("ciceroDecision") in accepted
        and item.get("helenaEvidenceId")
        and item.get("ciceroEvidenceId")
        and item.get("helenaDecisionLocator")
        and item.get("ciceroDecisionLocator")
        for item in items
    )


def _global_replay_findings(files: dict[str, dict]) -> list[dict]:
    payload = files.get("F7_GLOBAL_CONSISTENCY.json") or {}
    if not payload or payload.get("applicability") == "not_applicable":
        return []
    findings: list[dict] = []
    checks = {
        check.get("name"): check
        for layer in (payload.get("layerEvidence") or {}).values()
        for check in (layer.get("checks") or [])
        if isinstance(check, dict)
    }

    for source in ((((checks.get("registered_sources_have_current_hashes") or {}).get("evidenceData") or {}).get("sources")) or {}).values():
        path = Path(str(source.get("path") or ""))
        if not path.is_file() or sha256_file(path) != source.get("sha256"):
            findings.append(issue("N4-GLOBAL-REPLAY-C1", f"fonte da medição C1 não reproduzível: {path}"))

    semantic = (checks.get("source_to_final_semantic_fidelity") or {}).get("evidenceData") or {}
    if semantic.get("mode") == "docx_semantic_hash":
        try:
            from docx import Document

            def semantic_text(path: Path) -> list[str]:
                doc = Document(str(path))
                values = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
                values += [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
                return values

            source = Path(str(semantic.get("sourcePath") or ""))
            final = Path(str(semantic.get("docxPath") or ""))
            if not source.is_file() or not final.is_file() or canonical_hash(semantic_text(source)) != canonical_hash(semantic_text(final)):
                findings.append(issue("N4-GLOBAL-REPLAY-C2", "fidelidade semântica DOCX não foi reproduzida"))
        except Exception as exc:
            findings.append(issue("N4-GLOBAL-REPLAY-C2", f"falha ao reproduzir fidelidade semântica: {exc}"))
    elif semantic.get("mode") == "markdown_docx_pdf":
        try:
            from forja_fidelity import compare_fidelity

            fidelity = compare_fidelity(Path(str(semantic.get("sourcePath") or "")), Path(str(semantic.get("docxPath") or "")), Path(str(semantic.get("pdfPath") or "")))
            if fidelity.get("approved") is not True:
                findings.append(issue("N4-GLOBAL-REPLAY-C2", "fidelidade Markdown/DOCX/PDF não foi reproduzida"))
        except Exception as exc:
            findings.append(issue("N4-GLOBAL-REPLAY-C2", f"falha ao reproduzir fidelidade: {exc}"))
    else:
        findings.append(issue("N4-GLOBAL-REPLAY-C2", "medição C2 sem modo reproduzível"))

    test_data = (checks.get("case_tests") or {}).get("evidenceData") or {}
    results = files.get("F7_CASE_TEST_RESULTS.json") or {}
    if test_data.get("suiteHash") != results.get("suiteHash") or test_data.get("draftHash") != results.get("draftHash"):
        findings.append(issue("N4-GLOBAL-REPLAY-C3", "medição C3 diverge do F7 reexecutado"))

    measured_questions = (((checks.get("question_support_integrity") or {}).get("evidenceData") or {}).get("questions")) or []
    current_questions = [
        {"questionId": item.get("questionId"), "supportIds": item.get("supportIds") or []}
        for item in (files.get("F2_QUESTION_TREE.json") or {}).get("questions") or []
    ]
    if measured_questions != current_questions or any(not item.get("supportIds") for item in current_questions):
        findings.append(issue("N4-GLOBAL-REPLAY-C4", "integridade de suporte das questões não foi reproduzida"))

    physical = (checks.get("physical_and_visual_integrity") or {}).get("evidenceData") or {}
    for kind in ("docx", "pdf", "f8"):
        path = Path(str(physical.get(f"{kind}Path") or ""))
        if not path.is_file() or sha256_file(path) != physical.get(f"{kind}Sha256"):
            findings.append(issue("N4-GLOBAL-REPLAY-C5", f"{kind.upper()} da medição física não confere"))
    f8_path = Path(str(physical.get("f8Path") or ""))
    f8 = (read_json(f8_path, {}) or {}) if f8_path.is_file() else {}
    pages = f8.get("pages") or []
    if f8.get("approved") is not True or len(pages) != int(physical.get("pageCount") or 0) or not pages:
        findings.append(issue("N4-GLOBAL-REPLAY-C5", "ledger visual não confirma todas as páginas"))
    for page in pages:
        image = Path(str(page.get("imagePath") or ""))
        review = page.get("independentReview") or {}
        if page.get("lint") != "pass" or review.get("status") != "pass" or not image.is_file() or sha256_file(image) != page.get("imageSha256"):
            findings.append(issue("N4-GLOBAL-REPLAY-C5", f"página {page.get('page')} sem evidência visual reproduzível"))
            break
    # Reutiliza o mesmo gate fail-closed do empacotamento: rerender do PDF,
    # auditoria OOXML e atestado visual página a página. O N4 não confia no
    # campo ``approved`` produzido pela etapa anterior.
    try:
        docx_path = Path(str(physical.get("docxPath") or ""))
        pdf_path = Path(str(physical.get("pdfPath") or ""))
        replay = validate_f8(
            {"path": str(f8_path)},
            files={
                "docx": {"path": str(docx_path), "sha256": sha256_file(docx_path) if docx_path.is_file() else None},
                "pdf": {"path": str(pdf_path), "sha256": sha256_file(pdf_path) if pdf_path.is_file() else None},
            },
        )
        if not replay["approved"]:
            findings.append(issue(
                "N4-GLOBAL-REPLAY-C5",
                "gate visual anti-autocertificação reprovado: " + "; ".join(replay["findings"][:6]),
            ))
    except Exception as exc:
        findings.append(issue("N4-GLOBAL-REPLAY-C5", f"falha ao reexecutar o gate visual integral: {exc}"))
    return findings


MODOS_VALIDOS = {"off", "shadow", "pilot_blocking", "default_on"}


def _effective_named_mode(
    config: dict,
    case_dir: Path,
    namespace: str,
    override: str | None = None,
) -> tuple[str, str]:
    """Resolve o modo efetivo de um namespace de feature.

    Namespace ausente equivale a `off`: uma capacidade que ninguém declarou não
    se ativa por omissão. Em `pilot_blocking`, só bloqueiam os casos nomeados em
    `pilotCases`; os demais ficam em `shadow`, que materializa e reporta sem
    barrar a saída canônica.
    """
    espaco = config.get(namespace) or {}
    configured = str(override or espaco.get("mode") or "off")
    if configured not in MODOS_VALIDOS:
        raise ForjaN3Error(
            f"modo desconhecido em {namespace}: {configured!r}; use um de {sorted(MODOS_VALIDOS)}"
        )
    if configured != "pilot_blocking" or override:
        return configured, configured
    pilots = {str(value) for value in (espaco.get("pilotCases") or [])}
    manifest = read_json(case_dir / "FORJA_CASE_MANIFEST.json", {}) or {}
    identities = {case_dir.name, str(manifest.get("caseId") or "")}
    return configured, "pilot_blocking" if pilots & identities else "shadow"


def _effective_mode(config: dict, case_dir: Path, override: str | None = None) -> tuple[str, str]:
    return _effective_named_mode(config, case_dir, "n4", override)


def effective_signature_lite_mode(
    config: dict, case_dir: Path, override: str | None = None
) -> tuple[str, str]:
    """Modo da feature FORJA-ASSINATURA Lite; `forjaAssinaturaLite.mode` é a única fonte de ativação."""
    return _effective_named_mode(config, case_dir, "forjaAssinaturaLite", override)


def _schema_findings(filename: str, payload: dict) -> list[dict]:
    root = Path(__file__).parent / "n4_schemas"
    catalog = read_json(root / "ARTIFACT_CATALOG.json", {}) or {}
    entry = (catalog.get("artifacts") or {}).get(filename)
    if not entry:
        return [issue("N4-SCHEMA-CATALOG", f"artefato sem esquema registrado: {filename}", artifact=filename)]
    schema = read_json(root / str(entry.get("schema") or ""), None)
    if not isinstance(schema, dict):
        return [issue("N4-SCHEMA-MISSING", f"esquema ausente: {filename}", artifact=filename)]
    errors = sorted(jsonschema.Draft202012Validator(schema).iter_errors(payload), key=lambda error: list(error.path))
    return [
        issue(
            "N4-SCHEMA",
            f"{filename} em {'.'.join(map(str, error.path)) or '<raiz>'}: {error.message}",
            artifact=filename,
        )
        for error in errors
    ]


def _cross_reference_findings(files: dict[str, dict]) -> list[dict]:
    graph = files.get("F3_REASONING_GRAPH.json") or {}
    node_ids = {str(item.get("id")) for item in graph.get("nodes") or [] if item.get("id")}
    if not node_ids:
        return []
    findings: list[dict] = []
    questions = files.get("F2_QUESTION_TREE.json") or {}
    for item in questions.get("questions") or []:
        for support_id in item.get("supportIds") or []:
            if str(support_id) not in node_ids:
                findings.append(issue("N4-CROSS-Q-SUPPORT", f"{item.get('questionId')}: supportId fora do grafo: {support_id}"))
    coverage = files.get("F4_COVERAGE_MATRIX.json") or {}
    for item in coverage.get("items") or []:
        for support_id in item.get("supportIds") or []:
            if str(support_id) not in node_ids:
                findings.append(issue("N4-CROSS-COV-SUPPORT", f"{item.get('coverageId')}: supportId fora do grafo: {support_id}"))
    theses = files.get("F4_THESIS_MATURITY.json") or {}
    for item in theses.get("theses") or []:
        thesis_id = str(item.get("thesisId") or "")
        if thesis_id and thesis_id not in node_ids:
            findings.append(issue("N4-CROSS-THESIS", f"tese fora do grafo: {thesis_id}"))
    return findings


def validate_case(case_dir: Path, *, target_phase: str | None = None, write: bool = True, mode_override: str | None = None) -> dict:
    config = load_config()
    configured_mode, mode = _effective_mode(config, case_dir, mode_override)
    target = target_phase or _target_phase(case_dir)
    target_index = PHASES.index(target)
    required = _required_files(config)
    # A feature tem namespace próprio e não entra por flag booleana concorrente:
    # `forjaAssinaturaLite.mode` é a única fonte de ativação. Em `off`, nada é
    # exigido e nada é materializado.
    _, signature_lite_mode = effective_signature_lite_mode(config, case_dir)
    if signature_lite_mode != "off":
        required.update(SIGNATURE_LITE_FILES)
    artifacts_dir = case_dir / "n4_artifacts"
    case_manifest = read_json(case_dir / "FORJA_CASE_MANIFEST.json", {}) or {}
    registered_hashes, findings = _source_registry_findings(case_dir, case_manifest, require_verifiable=mode in {"pilot_blocking", "default_on"})
    files: dict[str, dict] = {}
    for filename, spec in ARTIFACT_SPECS.items():
        phase_index = PHASES.index(spec["phase"])
        if filename not in required or phase_index > target_index:
            continue
        path = artifacts_dir / filename
        payload = read_json(path, None)
        if not isinstance(payload, dict):
            severity = "p1" if mode == "shadow" else "p0"
            findings.append(issue("N4-MISSING", f"artefato esperado ainda não materializado: {filename}", severity=severity, artifact=filename))
            continue
        files[filename] = payload
        findings.extend(_schema_findings(filename, payload))
        findings.extend(validate_envelope(case_dir, filename, payload))
        source_hashes = payload.get("sourceHashes") or []
        if payload.get("applicability") != "not_applicable" and mode in {"pilot_blocking", "default_on"} and not source_hashes:
            findings.append(issue("N4-SOURCE-EMPTY", "artefato aplicável sem fonte registrada", artifact=filename))
        for source_hash in source_hashes:
            if source_hash not in registered_hashes:
                findings.append(issue("N4-SOURCE-HASH", f"sourceHash não registrado no manifesto do caso: {source_hash}", artifact=filename))
        if payload.get("applicability") != "not_applicable" and payload.get("status") != "approved":
            findings.append(issue("N4-ARTIFACT-STATUS", f"artefato obrigatório não está aprovado: {payload.get('status')}", artifact=filename))
        if payload.get("applicability") != "not_applicable":
            validator = (
                _recipient_map_validator(config)
                if filename == "F3_MAPA_DESTINATARIO.json"
                else VALIDATORS.get(filename)
            )
            if validator:
                findings.extend(validator(payload))
    studies = files.get("F5C_STUDY_LEDGER.json")
    claims = files.get("F5C_CLAIM_EVIDENCE_MAP.json")
    if claims and claims.get("applicability") != "not_applicable":
        findings.extend(validate_claims(claims, studies))
    suite = files.get("F4_CASE_ACCEPTANCE_TESTS.json")
    results = files.get("F7_CASE_TEST_RESULTS.json")
    if suite and suite.get("executionMode") == "prospective":
        findings.extend(suite_learning_findings(case_dir, suite))
    if suite and results:
        draft_path = _registered_source_path(case_dir, case_manifest, str(results.get("draftHash") or ""))
        findings.extend(validate_results(results, suite, draft_path=draft_path))
        if draft_path is None:
            findings.append(issue("N4-TEST-DRAFT-UNRESOLVED", "texto canônico do F7 não foi localizado por hash no registro de fontes"))
        else:
            recomputed = run_suite(
                suite,
                draft_path,
                reviewer_run_id=str(results.get("reviewerRunId") or "forja-n4-validator-replay"),
                producer_run_id=str(results.get("producerRunId") or "") or None,
            )
            if canonical_hash(_result_core(results)) != canonical_hash(_result_core(recomputed)):
                findings.append(issue("N4-TEST-REEXECUTION-DRIFT", "F7 salvo diverge da reexecução determinística sobre o texto canônico"))
    findings.extend(_global_replay_findings(files))
    findings.extend(_cross_reference_findings(files))
    p0 = [item for item in findings if item.get("severity") == "p0"]
    p1 = [item for item in findings if item.get("severity") == "p1"]
    blocking_p0 = p0 if mode == "default_on" else [item for item in p0 if _pilot_blocking_finding(item)] if mode == "pilot_blocking" else []
    material_blocks = len(blocking_p0)
    expected_count = len([f for f in required if PHASES.index(ARTIFACT_SPECS[f]["phase"]) <= target_index])
    evaluated = expected_count > 0
    complete = evaluated and len(files) == expected_count
    effective_approved = complete and not p0
    test_mode = str((suite or {}).get("executionMode") or "legacy")
    anti_fraud = (results or {}).get("antiFraud") or {}
    semantic_score = float(anti_fraud.get("semanticMutationScore") or 0)
    global_report = files.get("F7_GLOBAL_CONSISTENCY.json") or {}
    theses_report = files.get("F4_THESIS_MATURITY.json") or {}
    council_ready = _council_ready(theses_report)
    promotion_requirements = {
        "aggregateApproved": effective_approved,
        "prospectiveSuite": test_mode == "prospective",
        "literalMutationCoverage": float(anti_fraud.get("mutationScore") or 0) >= 0.8,
        "semanticMutationCoverage": semantic_score >= 0.8,
        "globalConsistencyReplayed": global_report.get("approved") is True and global_report.get("measurementContract") == "N4-MEASURED-v1",
        "councilDecisionsVerified": council_ready,
    }
    promotion_eligible = (
        all(promotion_requirements.values())
    )
    shadow_pass = not p0 if mode != "shadow" else True
    # O lint GRAFO-01..06 é diagnóstico estrutural separado da promoção N4. Ele
    # fica visível no laudo, mas nunca entra em `findings` nem altera
    # `blocksCurrentFlow`: o grafo F3 nasce antes da pesquisa oficial e não
    # recebe autoridade jurídica neste PRD.
    graph_lint_report = None
    graph_path = artifacts_dir / "F3_REASONING_GRAPH.json"
    if graph_path.is_file():
        from forja_grafo_lint import lint_file
        graph_lint_report = lint_file(graph_path)
    report = {
        "schemaVersion": 1,
        "specVersion": SPEC_VERSION,
        "caseId": case_dir.name,
        "mode": mode,
        "configuredMode": configured_mode,
        "targetPhase": target,
        "evaluationStatus": "not_evaluated" if not evaluated else "evaluated_approved" if effective_approved else "evaluated_blocked",
        "complete": complete,
        "approved": effective_approved,
        "promotionEligible": promotion_eligible,
        "promotionRequirements": promotion_requirements,
        "metrics": {"caseTestMode": test_mode, "mutationScore": anti_fraud.get("mutationScore"), "semanticMutationScore": anti_fraud.get("semanticMutationScore"), "mutationsKilled": anti_fraud.get("killed"), "mutationsTotal": anti_fraud.get("total")},
        # A fase só pode ser bloqueada por incompletude quando há ao menos um
        # artefato N4 exigível para o alvo. Em F0/F1, `expected_count == 0`:
        # o N4 ainda não foi avaliado e não pode transformar ausência de escopo
        # em um gate impossível sem achado material.
        "blocksCurrentFlow": bool(blocking_p0) or (
            mode in {"pilot_blocking", "default_on"} and evaluated and not complete
        ),
        "shadowAllowsCurrentFlow": shadow_pass,
        "counts": {"expected": expected_count, "present": len(files), "p0": len(p0), "blockingP0": len(blocking_p0), "p1": len(p1)},
        "materialBlocks": material_blocks,
        "findings": findings,
        "artifactIds": sorted(files),
        "graphLint": graph_lint_report,
    }
    report["validationHash"] = canonical_hash({key: value for key, value in report.items() if key != "validationHash"})
    if write:
        artifacts_dir.mkdir(parents=True, exist_ok=True)
        atomic_write_json(artifacts_dir / "N4_VALIDATION.json", report)
        append_trace(case_dir, "validation_completed", run_id="forja-n4-validator", status="pass" if effective_approved else "fail", detail={"targetPhase": target, "validationHash": report["validationHash"], "counts": report["counts"], "blocksCurrentFlow": report["blocksCurrentFlow"]})
    return report


def management_summary(case_dir: Path) -> dict:
    manifest = read_json(case_dir / "FORJA_CASE_MANIFEST.json", {}) or {}
    declared_target = str((manifest.get("n4M6Cycle") or {}).get("targetPhase") or "")
    current_target = _target_phase(case_dir)
    target = max(
        [phase for phase in (current_target, declared_target) if phase in PHASES],
        key=PHASES.index,
    )
    report = validate_case(case_dir, target_phase=target, write=False)
    q = read_json(case_dir / "n4_artifacts" / "F2_QUESTION_TREE.json", {}) or {}
    qcov = q.get("coverage") or {}
    tests = read_json(case_dir / "n4_artifacts" / "F7_CASE_TEST_RESULTS.json", {}) or {}
    results = tests.get("results") or []
    science = read_json(case_dir / "n4_artifacts" / "F2_N4_CLASSIFICATION.json", {}) or {}
    sci_mode = ((science.get("science") or {}).get("mode") or "not_run")
    science_audit = read_json(case_dir / "n4_artifacts" / "F7_SCIENCE_AUDIT.json", {}) or {}
    global_report = read_json(case_dir / "n4_artifacts" / "F7_GLOBAL_CONSISTENCY.json", {}) or {}
    theses = read_json(case_dir / "n4_artifacts" / "F4_THESIS_MATURITY.json", {}) or {}
    material_gaps = sum(bool(item.get("gaps")) for item in theses.get("theses") or [])
    artifacts = sorted(path.name for path in (case_dir / "n4_artifacts").glob("*.json") if path.name != "N4_VALIDATION.json")
    trace = case_dir / "n4_artifacts" / "N4_EXECUTION_TRACE.jsonl"
    if trace.is_file():
        artifacts.append(trace.name)
    validation = case_dir / "n4_artifacts" / "N4_VALIDATION.json"
    if validation.is_file():
        artifacts.append(validation.name)
    first_block = next((item.get("detail") for item in report.get("findings") or [] if item.get("severity") == "p0"), None)
    if first_block:
        next_action = first_block
    elif report.get("approved") and report.get("counts", {}).get("p1", 0):
        next_action = "Baseline estrutural reproduzida; resolver as pendências P1 antes de uso final."
    elif report.get("approved") and report.get("mode") == "pilot_blocking":
        next_action = "Piloto N4 estruturalmente validado; manter revisão jurídica humana."
    elif report.get("evaluationStatus") == "not_evaluated":
        next_action = "N4 ainda não executada para esta fase."
    elif report.get("approved"):
        next_action = "N4 validada em sombra."
    else:
        next_action = "Materializar ou revisar os artefatos N4 pendentes."
    delivery = read_json(case_dir / "n4_artifacts" / "F10_DELIVERY_INTEGRITY.json", {}) or {}
    citation_coverage = manifest.get("n4CitationCoverage") or {}
    material_citations = int(citation_coverage.get("materialTotal") or 0)
    verified_material_citations = int(citation_coverage.get("verifiedMaterial") or 0)
    registry = manifest.get("n4SourceRegistry") or {}
    regimento_registered = any(
        "regimento" in str(source_id).casefold()
        and isinstance(entry, dict)
        and str(entry.get("status") or "active") == "active"
        for source_id, entry in registry.items()
    )
    release_gates = {
        "targetF10": report.get("targetPhase") == "F10_ENTREGA_EVIDENCIA_APRENDIZADO",
        "aggregateApproved": report.get("approved") is True,
        "noMaterialFindings": report.get("counts", {}).get("p0", 0) == 0 and report.get("counts", {}).get("p1", 0) == 0,
        "noThesisGaps": material_gaps == 0,
        "caseTestsApproved": tests.get("approved") is True,
        "globalConsistencyMeasured": global_report.get("approved") is True and global_report.get("measurementContract") == "N4-MEASURED-v1",
        "councilDecisionsVerified": _council_ready(theses),
        "materialCitationsVerified": material_citations > 0 and verified_material_citations == material_citations,
        "regimentoRegistered": regimento_registered,
        "deliveryApplicableAndConfirmed": delivery.get("applicability") != "not_applicable" and delivery.get("status") == "approved",
    }
    legal_release_status = "structurally_clear" if all(release_gates.values()) else "human_review_required"
    return {
        "enabled": bool(_required_files(load_config())),
        "mode": report.get("mode"),
        "configuredMode": report.get("configuredMode"),
        "complete": report.get("complete", False),
        "evaluationStatus": report.get("evaluationStatus", "not_evaluated"),
        "approved": report.get("approved", False),
        "promotionEligible": report.get("promotionEligible", False),
        "promotionRequirements": report.get("promotionRequirements") or {},
        "caseTestMode": (report.get("metrics") or {}).get("caseTestMode"),
        "legalReleaseStatus": legal_release_status,
        "releaseGates": release_gates,
        "citationCoverage": {"verifiedMaterial": verified_material_citations, "materialTotal": material_citations},
        "materialGapCount": material_gaps,
        "questionCoverage": f"{qcov.get('answeredMaterial', 0)}/{qcov.get('material', 0)}",
        "materialBlocks": report.get("materialBlocks", 0),
        "blocksCurrentFlow": report.get("blocksCurrentFlow", False),
        "caseTests": f"{sum(item.get('status') == 'pass' for item in results)}/{len(results)}",
        "caseTestMode": report.get("metrics", {}).get("caseTestMode", "legacy"),
        "mutationScore": ((tests.get("antiFraud") or {}).get("mutationScore")),
        "scienceMode": sci_mode,
        "scienceStatus": "not_applicable" if sci_mode == "not_applicable" else "pass" if science_audit.get("status") == "approved" and report.get("approved") else "in_review",
        "globalConsistency": "pass" if global_report.get("approved") is True else "pending",
        "nextAction": next_action,
        "artifactIds": artifacts,
        "validationHash": report.get("validationHash"),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Validador agregador FORJA N4")
    parser.add_argument("case")
    parser.add_argument("--target-phase", choices=PHASES)
    parser.add_argument("--no-write", action="store_true")
    args = parser.parse_args()
    result = validate_case(resolve_case_dir(args.case), target_phase=args.target_phase, write=not args.no_write)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
