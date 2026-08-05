"""Prepare versioned, non-destructive M6 canary copies for visual review."""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

import forja_acervo

import fitz
from docx import Document
from PIL import Image, ImageDraw

from forja_metadata import sanitize_final_artifacts
from forja_n3_common import FORJA, atomic_write_json, now_iso, sha256_file


ROOT = FORJA.parent

CASES = {
    "patricia": {
        "caseId": forja_acervo.caso("CASO-19"),
        "docx": ROOT / "Memoriais Apelação CASO-19 e Fábio - Proc. 9000003-00.2014.8.19.0000" / "Anexos do email" / "MEMORIAIS - PATRICIA E FABIO - N3 SUPERIOR PARA REVISAO - 10-07-2026.docx",
        "pdf": ROOT / "Memoriais Apelação CASO-19 e Fábio - Proc. 9000003-00.2014.8.19.0000" / "Anexos do email" / "MEMORIAIS - PATRICIA E FABIO - N3 SUPERIOR PARA REVISAO - 10-07-2026.pdf",
        "layoutProfileId": "medina-visual-law-v1",
    },
    "libra": {
        "caseId": forja_acervo.caso("CASO-16"),
        "docx": forja_acervo.caminho("m6-memoriais-docx"),
        "pdf": forja_acervo.caminho("m6-memoriais-pdf"),
        "layoutProfileId": "medina-visual-law-v1",
    },
    "health": {
        "caseId": "case-email-auto-19f3f25cb64df962",
        "docx": FORJA / "telemetria" / "renders" / "MINUTA_INICIAL_TJDFT" / "MINUTA_INICIAL_TJDFT.docx",
        "pdf": FORJA / "telemetria" / "renders" / "MINUTA_INICIAL_TJDFT" / "MINUTA_INICIAL_TJDFT.pdf",
        "layoutProfileId": "medina-word-v1",
    },
}


def _extract_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    return "\n\n".join(lines) + "\n"


def _render(pdf_path: Path, pages_dir: Path) -> list[Path]:
    pages_dir.mkdir(parents=True, exist_ok=True)
    output: list[Path] = []
    with fitz.open(pdf_path) as pdf:
        for index, page in enumerate(pdf, 1):
            target = pages_dir / f"p{index:03d}.png"
            page.get_pixmap(matrix=fitz.Matrix(1.45, 1.45), alpha=False).save(target)
            output.append(target)
    return output


def _contact_sheet(pages: list[Path], target: Path) -> None:
    thumbs: list[Image.Image] = []
    for page in pages:
        image = Image.open(page).convert("RGB")
        image.thumbnail((360, 510))
        canvas = Image.new("RGB", (380, 550), "white")
        canvas.paste(image, ((380 - image.width) // 2, 24))
        draw = ImageDraw.Draw(canvas)
        draw.text((12, 8), page.stem, fill="#333333")
        thumbs.append(canvas)
    columns = 3
    rows = (len(thumbs) + columns - 1) // columns
    sheet = Image.new("RGB", (columns * 380, rows * 550), "#dddddd")
    for index, thumb in enumerate(thumbs):
        sheet.paste(thumb, ((index % columns) * 380, (index // columns) * 550))
    sheet.save(target, quality=90)


def prepare(key: str) -> dict:
    profile = CASES[key]
    case_dir = FORJA / "state" / profile["caseId"]
    cycle_dir = case_dir / "n4_cycle_m6"
    cycle_dir.mkdir(parents=True, exist_ok=True)
    source_docx = Path(profile["docx"])
    source_pdf = Path(profile["pdf"])
    if not source_docx.is_file() or not source_pdf.is_file():
        raise FileNotFoundError(f"produto final ausente para {key}")
    docx = cycle_dir / "FINAL_N4_M6.docx"
    pdf = cycle_dir / "FINAL_N4_M6.pdf"
    shutil.copy2(source_docx, docx)
    shutil.copy2(source_pdf, pdf)
    sanitize_final_artifacts(docx, pdf)
    canonical = cycle_dir / "CANONICAL_TEXT_FROM_FINAL_DOCX.txt"
    canonical.write_text(_extract_text(docx), encoding="utf-8")
    pages = _render(pdf, cycle_dir / "pages")
    contact = cycle_dir / "CONTACT_SHEET.png"
    _contact_sheet(pages, contact)
    f8 = {
        "schemaVersion": 1,
        "caseId": case_dir.name,
        "generatedAt": now_iso(),
        "approved": False,
        "generatorRunId": "n4-m6-render",
        "reviewerRunId": None,
        "layoutProfileId": profile["layoutProfileId"],
        "pageCount": len(pages),
        "pages": [{"page": index, "path": str(path), "status": "pending"} for index, path in enumerate(pages, 1)],
        "contactSheet": str(contact),
        "sourceFiles": {"docx": str(source_docx), "pdf": str(source_pdf)},
        "finalHashes": {"canonical": sha256_file(canonical), "docx": sha256_file(docx), "pdf": sha256_file(pdf)},
    }
    atomic_write_json(cycle_dir / "F8_QA_LEDGER_N4.json", f8)
    return {"case": key, "caseId": case_dir.name, "cycleDir": str(cycle_dir), "contactSheet": str(contact), "pageCount": len(pages), "hashes": f8["finalHashes"]}


def approve(key: str, reviewer: str) -> dict:
    case_dir = FORJA / "state" / CASES[key]["caseId"]
    path = case_dir / "n4_cycle_m6" / "F8_QA_LEDGER_N4.json"
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    if reviewer == payload.get("generatorRunId"):
        raise ValueError("o gerador não pode revisar o próprio PDF")
    from forja_visual_qa import run_visual_qa

    cycle_dir = path.parent
    audited = run_visual_qa(
        cycle_dir / "FINAL_N4_M6.pdf",
        path,
        qa_dir=cycle_dir / "pages_qa_verified",
        generator_run_id=str(payload.get("generatorRunId") or "n4-m6-render"),
        reviewer_run_id=reviewer,
        docx=cycle_dir / "FINAL_N4_M6.docx",
    )
    audited["caseId"] = case_dir.name
    audited["layoutProfileId"] = payload.get("layoutProfileId")
    audited["sourceFiles"] = payload.get("sourceFiles")
    audited["finalHashes"] = payload.get("finalHashes")
    atomic_write_json(path, audited)
    return {"case": key, "approved": audited["approved"], "reviewerRunId": reviewer, "pages": audited.get("pageCount"), "findings": len(audited.get("findings") or [])}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["prepare", "approve"])
    parser.add_argument("case", choices=sorted(CASES))
    parser.add_argument("--reviewer", default="n4-m6-independent-visual-review")
    args = parser.parse_args()
    result = prepare(args.case) if args.command == "prepare" else approve(args.case, args.reviewer)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
