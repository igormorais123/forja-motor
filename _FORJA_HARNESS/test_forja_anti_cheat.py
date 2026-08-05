from __future__ import annotations

import base64
import copy
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey

from forja_docx_layout import (
    FOLIO_TEMPLATE_WIDTH_PT, audit_docx_layout, compare_docx_content, normalize_medina_body)
from forja_citations import extrair_citacoes
from forja_human_review import (
    build_unsigned_claim_receipt,
    build_unsigned_visual_receipt,
    canonical_receipt_bytes,
    public_key_id,
)
from forja_metricas_f7 import extrair_citacoes_basico
from forja_n3_common import atomic_write_json, canonical_hash, now_iso, read_json, sha256_file
from forja_official_sources import (
    OFFICIAL_CACHE,
    OFFICIAL_MANIFEST,
    sidecar_path,
    source_excerpt_sha256,
    validate_archived_source,
    validate_cached_source,
)
from forja_package import validate_f7, validate_f8, validate_fidelity, validate_source_ledger
from forja_visual_qa import inspect_pdf, run_visual_qa
from forja_visual_review import (
    REQUIRED_PAGE_CHECKS,
    build_pending_review,
    validate_visual_review,
)


BODY = (
    "Este parágrafo jurídico desenvolve o fundamento central com extensão suficiente para ser inequivocamente "
    "classificado como texto principal do documento, preservando a leitura contínua, a hierarquia argumentativa "
    "e a apresentação profissional exigida pelo escritório em todas as peças protocoláveis, sem qualquer função "
    "especial de capa, legenda, assinatura, quadro, chamada visual ou título de seção."
)


def make_docx(path: Path, *, alignment, font: str, size: float) -> None:
    document = Document()
    document.add_heading("I — MÉRITO", level=1)
    paragraph = document.add_paragraph(BODY)
    paragraph.alignment = alignment
    run = paragraph.runs[0]
    run.font.name = font
    run.font.size = Pt(size)
    document.save(path)


def make_pdf(path: Path, text: str = BODY) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(fitz.Rect(72, 72, 520, 500), text, fontsize=11)
    document.save(path)
    document.close()


def inject_unsafe_folio(document: Document, *, width_pt: float = 79.3) -> None:
    paragraph = document.sections[0].header.paragraphs[0]
    pict = OxmlElement("w:pict")
    rect = etree.Element("{urn:schemas-microsoft-com:vml}rect")
    rect.set("style", f"position:absolute;width:{width_pt}pt;height:25.95pt")
    textbox = etree.Element("{urn:schemas-microsoft-com:vml}textbox")
    content = OxmlElement("w:txbxContent")
    field_paragraph = OxmlElement("w:p")
    field_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = "PAGE"
    field_run.append(instruction)
    field_paragraph.append(field_run)
    content.append(field_paragraph)
    textbox.append(content)
    rect.append(textbox)
    pict.append(rect)
    paragraph._p.append(pict)


def corrupt_right_margin(path: Path) -> None:
    """Torna ilegível apenas w:pgMar/@w:right no fixture de teste."""
    temp = path.with_suffix(".corrupt.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(temp, "w") as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                data = data.replace(b"w:right=\"1020\"", b"w:right=\"nao-legivel\"")
            target.writestr(info, data)
    temp.replace(path)


class ForjaAntiCheatTests(unittest.TestCase):
    def test_state_court_cnj_adi_cannot_masquerade_as_stf_authority(self) -> None:
        text = (
            "A ADI nº 2080508-98.2020.8.26.0000 foi julgada pelo TJSP. "
            "O Informativo 849 – STJ registra precedente federal distinto."
        )
        metrics = extrair_citacoes_basico(text)
        self.assertTrue(any(item["tipo"] == "TJSP" and item["classe"] == "ADI" for item in metrics))
        self.assertFalse(any(item["tipo"] == "STF" and item["numero"] == "2080508" for item in metrics))
        self.assertTrue(any(item["tipo"] == "INFORMATIVO" and item["corte"] == "STJ" for item in metrics))

        checklist = extrair_citacoes(text)
        self.assertTrue(any(item["tipo"] == "CNJ" and item["rotulo"].startswith("TJSP ADI") for item in checklist))
        self.assertFalse(any(item["tipo"] == "STF" and "2080508" in item["rotulo"] for item in checklist))

    def test_layout_mutations_are_all_killed(self) -> None:
        mutations = {
            "left_aligned": (WD_ALIGN_PARAGRAPH.LEFT, "Times New Roman", 12),
            "wrong_font": (WD_ALIGN_PARAGRAPH.JUSTIFY, "Verdana", 12),
            "wrong_size": (WD_ALIGN_PARAGRAPH.JUSTIFY, "Times New Roman", 10.5),
        }
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            valid = root / "valid.docx"
            make_docx(valid, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font="Times New Roman", size=12)
            self.assertTrue(audit_docx_layout(valid)["approved"])
            survivors = []
            for name, values in mutations.items():
                path = root / f"{name}.docx"
                make_docx(path, alignment=values[0], font=values[1], size=values[2])
                result = audit_docx_layout(path)
                if result["approved"]:
                    survivors.append(name)
            self.assertEqual([], survivors, f"mutações visuais sobreviventes: {survivors}")

    def test_wide_folio_collision_is_blocked_and_normalized(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            unsafe = root / "unsafe-folio.docx"
            fixed = root / "fixed-folio.docx"
            make_docx(unsafe, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font="Times New Roman", size=12)
            document = Document(unsafe)
            inject_unsafe_folio(document)
            document.save(unsafe)
            audit = audit_docx_layout(unsafe)
            self.assertFalse(audit["approved"])
            self.assertIn("folio_width_unsafe", {item["code"] for item in audit["findings"]})

            normalized = normalize_medina_body(unsafe, fixed)
            self.assertTrue(normalized["approved"])
            # O alvo da normalização é a largura MEDIDA no template aprovado, e não
            # o teto de tolerância: consertar um fólio grande demais é deixá-lo como
            # o do escritório, não deixá-lo no limite do aceitável. O número antigo
            # (36 pt) não vinha de medição nenhuma e fazia o gate reprovar o próprio
            # template, cujo fólio mede 57,3 pt.
            self.assertEqual(FOLIO_TEMPLATE_WIDTH_PT,
                             normalized["normalization"]["folioWidthChanges"][0]["toWidthPt"])

    def test_folio_margin_collision_and_unreadable_margin_are_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            narrow = root / "narrow-margin.docx"
            make_docx(narrow, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      font="Times New Roman", size=12)
            document = Document(narrow)
            document.sections[0].right_margin = Pt(51)
            inject_unsafe_folio(document, width_pt=57.3)
            document.save(narrow)
            audit = audit_docx_layout(narrow)
            self.assertFalse(audit["approved"])
            self.assertIn("folio_overflows_right_margin",
                          {item["code"] for item in audit["findings"]})

            corrupt_right_margin(narrow)
            audit_corrupt = audit_docx_layout(narrow)
            self.assertFalse(audit_corrupt["approved"])
            self.assertIn("folio_margin_unresolved",
                          {item["code"] for item in audit_corrupt["findings"]})

    def test_visual_normalizer_cannot_silently_change_legal_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            tampered = root / "tampered.docx"
            make_docx(source, alignment=WD_ALIGN_PARAGRAPH.LEFT, font="Verdana", size=10.5)
            result = normalize_medina_body(source, normalized)
            self.assertTrue(result["normalization"]["contentFidelity"]["approved"])

            document = Document(normalized)
            document.paragraphs[1].runs[0].text = BODY.replace("fundamento central", "fundamento inventado")
            document.save(tampered)
            fidelity = compare_docx_content(source, tampered)
            self.assertFalse(fidelity["approved"])
            self.assertIn("contentSha256", fidelity["mismatches"])

    def test_visual_attestation_mutations_are_all_killed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            pdf = root / "piece.pdf"
            docx = root / "piece.docx"
            make_pdf(pdf)
            make_docx(docx, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font="Times New Roman", size=12)
            rendered = inspect_pdf(
                pdf, root / "pages", generator_run_id="generator-a", reviewer_run_id="reviewer-b"
            )["pages"]
            review_path = root / "review.json"
            review = build_pending_review(
                review_path,
                pdf=pdf,
                docx=docx,
                rendered_pages=rendered,
                generator_run_id="generator-a",
            )
            review["reviewedAt"] = now_iso()
            review["reviewer"] = {"id": "visual-reviewer", "runId": "reviewer-b", "type": "agent_visual"}
            review["reviewMethod"] = "page_by_page_at_100_percent"
            review["approved"] = True
            for page in review["pages"]:
                page["status"] = "pass"
                page["checks"] = {name: True for name in REQUIRED_PAGE_CHECKS}
            atomic_write_json(review_path, review)
            self.assertTrue(validate_visual_review(
                review_path,
                pdf=pdf,
                docx=docx,
                rendered_pages=rendered,
                generator_run_id="generator-a",
                expected_reviewer_run_id="reviewer-b",
            )["approved"])

            def mutate_review(name: str, mutation) -> bool:
                candidate = copy.deepcopy(review)
                mutation(candidate)
                path = root / f"mut-{name}.json"
                atomic_write_json(path, candidate)
                return validate_visual_review(
                    path,
                    pdf=pdf,
                    docx=docx,
                    rendered_pages=rendered,
                    generator_run_id="generator-a",
                    expected_reviewer_run_id="reviewer-b",
                )["approved"]

            mutations = {
                "automated_reviewer": lambda p: p["reviewer"].update(type="automated"),
                "self_review": lambda p: p["reviewer"].update(runId="generator-a"),
                "missing_page": lambda p: p["pages"].clear(),
                "unchecked_alignment": lambda p: p["pages"][0]["checks"].update(bodyAlignment=False),
                "wrong_image": lambda p: p["pages"][0].update(reviewedImageSha256="0" * 64),
                "autofilled": lambda p: p.update(autoFilled=True),
            }
            survivors = [name for name, mutation in mutations.items() if mutate_review(name, mutation)]
            self.assertEqual([], survivors, f"mutações de autocertificação sobreviventes: {survivors}")

    def test_strict_visual_release_requires_signed_human_receipt(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            pdf, docx = root / "piece.pdf", root / "piece.docx"
            make_pdf(pdf)
            make_docx(docx, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font="Times New Roman", size=12)
            rendered = inspect_pdf(
                pdf, root / "pre-pages", generator_run_id="generator-a", reviewer_run_id="reviewer-b"
            )["pages"]
            review_path = root / "review.json"
            review = build_pending_review(
                review_path, pdf=pdf, docx=docx, rendered_pages=rendered, generator_run_id="generator-a"
            )
            review.update({
                "reviewedAt": now_iso(),
                "reviewer": {"id": "visual-agent", "runId": "reviewer-b", "type": "agent_visual"},
                "reviewMethod": "page_by_page_at_100_percent",
                "approved": True,
            })
            for page in review["pages"]:
                page["status"] = "pass"
                page["checks"] = {name: True for name in REQUIRED_PAGE_CHECKS}
            atomic_write_json(review_path, review)

            ledger_path = root / "F8.json"
            ledger = run_visual_qa(
                pdf,
                ledger_path,
                qa_dir=root / "ledger-pages",
                generator_run_id="generator-a",
                reviewer_run_id="reviewer-b",
                docx=docx,
                manual_review=review_path,
            )
            files = {
                "docx": {"path": str(docx), "sha256": sha256_file(docx)},
                "pdf": {"path": str(pdf), "sha256": sha256_file(pdf)},
            }
            self.assertTrue(validate_f8({"path": str(ledger_path)}, files=files)["approved"])
            unsigned = validate_f8(
                {"path": str(ledger_path)}, files=files, release_policy="strict_protocol"
            )
            self.assertFalse(unsigned["approved"])
            self.assertTrue(any("recibo humano" in item for item in unsigned["findings"]))

            private_key = Ed25519PrivateKey.generate()
            public_raw = private_key.public_key().public_bytes(
                encoding=serialization.Encoding.Raw,
                format=serialization.PublicFormat.Raw,
            )
            key_id = public_key_id(public_raw)
            trust_store = root / "human-trust.json"
            atomic_write_json(trust_store, {
                "schemaVersion": 1,
                "reviewers": [{
                    "reviewerId": "advogado-visual",
                    "publicKeyId": key_id,
                    "algorithm": "Ed25519",
                    "publicKeyBase64": base64.b64encode(public_raw).decode("ascii"),
                    "enabled": True,
                }],
            })
            trust_pin = root / "human-trust-pin.json"
            atomic_write_json(trust_pin, {
                "schemaVersion": 1,
                "trustStorePath": str(trust_store),
                "trustStoreSha256": sha256_file(trust_store),
                "status": "configured",
            })
            receipt = build_unsigned_visual_receipt(
                reviewer_id="advogado-visual",
                reviewed_at=now_iso(),
                public_key_id_value=key_id,
                generator_run_id="generator-a",
                reviewer_run_id="reviewer-b",
                pdf_sha256=sha256_file(pdf),
                docx_sha256=sha256_file(docx),
                page_count=ledger["pageCount"],
                page_image_sha256=[page["imageSha256"] for page in ledger["pages"]],
                required_checks=list(REQUIRED_PAGE_CHECKS),
                visual_attestation_sha256=sha256_file(review_path),
            )
            receipt["signatureBase64"] = base64.b64encode(
                private_key.sign(canonical_receipt_bytes(receipt))
            ).decode("ascii")
            receipt_path = root / "human-visual-review.json"
            atomic_write_json(receipt_path, receipt)
            ledger["humanVisualReceipt"] = {
                "receiptPath": str(receipt_path),
                "receiptSha256": sha256_file(receipt_path),
            }
            atomic_write_json(ledger_path, ledger)
            with patch("forja_human_review.DEFAULT_TRUST_STORE", trust_store), patch(
                "forja_human_review.TRUST_STORE_PIN_PATH", trust_pin
            ):
                signed = validate_f8(
                    {"path": str(ledger_path)}, files=files, release_policy="strict_protocol"
                )
                self.assertTrue(signed["approved"], signed["findings"])

                receipt["pdfSha256"] = "0" * 64
                atomic_write_json(receipt_path, receipt)
                ledger["humanVisualReceipt"]["receiptSha256"] = sha256_file(receipt_path)
                atomic_write_json(ledger_path, ledger)
                tampered = validate_f8(
                    {"path": str(ledger_path)}, files=files, release_policy="strict_protocol"
                )
                self.assertFalse(tampered["approved"])

    def test_fake_jurisprudence_and_unregistered_sidecar_do_not_pass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            fake = root / "STJ_TEMA_99999.txt"
            fake.write_text(
                "FONTE OFICIAL STJ TEMA 99999\nURL: https://processo.stj.jus.br/falso\n" + "conteúdo inventado " * 30,
                encoding="utf-8",
            )
            atomic_write_json(sidecar_path(fake), {
                "schemaVersion": 1,
                "sha256": sha256_file(fake),
                "sourceUrl": "https://processo.stj.jus.br/falso",
                "capturedAt": now_iso(),
                "identity": {"court": "STJ", "kind": "TEMA", "number": "99999"},
            })
            result = validate_archived_source(fake)
            self.assertFalse(result["approved"])
            self.assertTrue(any("registr" in item for item in result["findings"]))

    def test_registered_source_copy_is_hash_bound(self) -> None:
        manifest = read_json(OFFICIAL_MANIFEST)
        source_name = "STJ_TEMA_1368.txt"
        record = manifest["entries"][source_name]
        self.assertTrue(validate_cached_source(OFFICIAL_CACHE / source_name)["approved"])
        with tempfile.TemporaryDirectory() as temp:
            target = Path(temp) / source_name
            target.write_bytes((OFFICIAL_CACHE / source_name).read_bytes())
            atomic_write_json(sidecar_path(target), {
                "schemaVersion": 1,
                "registryEntry": source_name,
                "sha256": sha256_file(target),
                "sourceUrl": record["sourceUrl"],
                "capturedAt": record.get("capturedAt") or now_iso(),
                "identity": record["identity"],
            })
            self.assertTrue(validate_archived_source(target)["approved"])
            target.write_text(target.read_text(encoding="utf-8") + "\nALTERAÇÃO", encoding="utf-8")
            self.assertFalse(validate_archived_source(target)["approved"])

    def test_forged_central_manifest_fails_live_official_replay(self) -> None:
        """URL + hash + identidade escritos pela IA não bastam."""
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "STJ_TEMA_99999.txt"
            url = "https://processo.stj.jus.br/repetitivos/tema-99999"
            invented = (
                "FONTE OFICIAL STJ\nURL: " + url + "\n\n"
                "TEMA 99999 STJ\nTESE FIRMADA (verbatim do portal):\n"
                '"A jurisprudência inventada concede automaticamente todo pedido sem necessidade de prova, '
                'contraditório, competência, fato constitutivo ou exame da fonte primária correspondente."\n'
            )
            source.write_text(invented, encoding="utf-8")
            manifest = root / "manifest.json"
            atomic_write_json(manifest, {
                "schemaVersion": 1,
                "entries": {
                    source.name: {
                        "sha256": sha256_file(source),
                        "sourceUrl": url,
                        "capturedAt": now_iso(),
                        "identity": {"court": "STJ", "kind": "TEMA", "number": "99999"},
                    }
                },
            })

            official_live = {
                "ok": True,
                "status": 200,
                "finalUrl": url,
                "bodySha256": "a" * 64,
                "text": (
                    "STJ Tema 99999. A tese oficial trata de questão inteiramente diversa e exige exame "
                    "individualizado dos pressupostos previstos na legislação aplicável ao caso concreto."
                ),
            }
            forged = validate_cached_source(
                source,
                manifest,
                cache_dir=root,
                require_live=True,
                fetcher=lambda _: official_live,
            )
            self.assertFalse(forged["approved"])
            self.assertTrue(any("verbatim" in item for item in forged["findings"]))

            official_live["text"] = invented
            authentic = validate_cached_source(
                source,
                manifest,
                cache_dir=root,
                require_live=True,
                fetcher=lambda _: official_live,
            )
            self.assertTrue(authentic["approved"], authentic["findings"])

    def test_strict_source_ledger_requires_coverage_excerpt_and_human_review(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            ledger = root / "sources.json"
            source = root / "STJ_RESP_2228834.txt"
            excerpt = (
                "A prescrição do fundo de direito nas relações de trato sucessivo exige negativa expressa "
                "formalizada em ato de efeitos concretos e ciência inequívoca da pessoa interessada."
            )
            source.write_text("STJ REsp 2228834\n" + excerpt, encoding="utf-8")
            identity = {"court": "STJ", "kind": "RESP", "number": "2228834"}
            expected = [{"corte": "STJ", "classe": "RESP", "numero": "2228834"}]

            atomic_write_json(ledger, {"entries": []})
            uncovered = validate_source_ledger(
                {"path": str(ledger)}, release_policy="strict_protocol", expected_citations=expected
            )
            self.assertFalse(uncovered["approved"])
            self.assertTrue(any("sem entrada probatória" in item for item in uncovered["blocked"]))

            claim = "O REsp 2.228.834 exige negativa formal e ciência para o fundo de direito."
            markdown = root / "final.md"
            markdown.write_text(claim + "\n", encoding="utf-8")
            paragraph_hash = canonical_hash({"paragraph": claim})
            entry = {
                "id": "resp-2228834",
                "claim": claim,
                "generatorRunId": "legal-writer",
                "finalUseAllowed": True,
                "sourcePathOrUrl": str(source),
                "sourceSha256": sha256_file(source),
                "sourceUrl": "https://processo.stj.jus.br/repetitivos/tema-1410",
                "sourceIdentity": identity,
                "sourceExcerpt": excerpt,
                "sourceExcerptSha256": source_excerpt_sha256(excerpt),
                "documentSha256": sha256_file(markdown),
                "documentProposition": claim,
                "documentPropositionSha256": canonical_hash({"proposition": claim}),
                "documentParagraphIndex": 1,
                "documentParagraphSha256": paragraph_hash,
                "authorityIdentity": identity,
                "authorityIdentitySha256": canonical_hash(identity),
                "claimReview": {
                    "status": "pass",
                    "reviewedAt": now_iso(),
                    "reviewer": {"id": "another-ai", "runId": "legal-review", "type": "agent_legal"},
                    "claimSha256": canonical_hash({"claim": claim}),
                    "sourceExcerptSha256": source_excerpt_sha256(excerpt),
                },
            }
            atomic_write_json(ledger, {"entries": [entry]})
            verified_source = {"approved": True, "findings": [], "record": {
                "sourceUrl": entry["sourceUrl"], "identity": identity,
            }}
            with patch("forja_package.validate_source_path", return_value=verified_source):
                ai_review = validate_source_ledger(
                    {"path": str(ledger)}, release_policy="strict_protocol", expected_citations=expected,
                    markdown={"path": str(markdown), "sha256": sha256_file(markdown)},
                )
                self.assertFalse(ai_review["approved"])
                self.assertTrue(any("recibo humano" in item for item in ai_review["blocked"]))

                # Nem escrever type=human no JSON basta. O PASS só nasce de um
                # recibo Ed25519 assinado por chave confiada fora do workspace.
                private_key = Ed25519PrivateKey.generate()
                public_raw = private_key.public_key().public_bytes(
                    encoding=serialization.Encoding.Raw,
                    format=serialization.PublicFormat.Raw,
                )
                key_id = public_key_id(public_raw)
                trust_store = root / "human-trust.json"
                atomic_write_json(trust_store, {
                    "schemaVersion": 1,
                    "reviewers": [{
                        "reviewerId": "advogado-revisor",
                        "publicKeyId": key_id,
                        "algorithm": "Ed25519",
                        "publicKeyBase64": base64.b64encode(public_raw).decode("ascii"),
                        "enabled": True,
                    }],
                })
                trust_pin = root / "human-trust-pin.json"
                atomic_write_json(trust_pin, {
                    "schemaVersion": 1,
                    "trustStorePath": str(trust_store),
                    "trustStoreSha256": sha256_file(trust_store),
                    "status": "configured",
                })
                receipt = build_unsigned_claim_receipt(
                    reviewer_id="advogado-revisor",
                    reviewed_at=now_iso(),
                    public_key_id_value=key_id,
                    generator_run_id=entry["generatorRunId"],
                    claim=claim,
                    claim_sha256=canonical_hash({"claim": claim}),
                    source_excerpt=excerpt,
                    source_excerpt_sha256=source_excerpt_sha256(excerpt),
                    source_sha256=sha256_file(source),
                    source_url=entry["sourceUrl"],
                    source_identity=identity,
                    source_identity_sha256=canonical_hash(identity),
                    document_sha256=sha256_file(markdown),
                    document_proposition=claim,
                    document_proposition_sha256=canonical_hash({"proposition": claim}),
                    document_paragraph_index=1,
                    document_paragraph_sha256=paragraph_hash,
                    authority_identity=identity,
                    authority_identity_sha256=canonical_hash(identity),
                )
                receipt["signatureBase64"] = base64.b64encode(
                    private_key.sign(canonical_receipt_bytes(receipt))
                ).decode("ascii")
                receipt_path = root / "human-claim-review.json"
                atomic_write_json(receipt_path, receipt)
                entry["claimReview"] = {
                    "status": "pass",
                    "receiptPath": str(receipt_path),
                    "receiptSha256": sha256_file(receipt_path),
                }
                atomic_write_json(ledger, {"entries": [entry]})
                with patch("forja_human_review.DEFAULT_TRUST_STORE", trust_store), patch(
                    "forja_human_review.TRUST_STORE_PIN_PATH", trust_pin
                ):
                    human_review = validate_source_ledger(
                        {"path": str(ledger)}, release_policy="strict_protocol", expected_citations=expected,
                        markdown={"path": str(markdown), "sha256": sha256_file(markdown)},
                    )
                    self.assertTrue(human_review["approved"], human_review["blocked"])

                    entry["claim"] = claim + " Conteúdo adulterado depois da assinatura."
                    atomic_write_json(ledger, {"entries": [entry]})
                    tampered = validate_source_ledger(
                        {"path": str(ledger)}, release_policy="strict_protocol", expected_citations=expected,
                        markdown={"path": str(markdown), "sha256": sha256_file(markdown)},
                    )
                    self.assertFalse(tampered["approved"])
                    self.assertTrue(any("recibo humano reprovado" in item for item in tampered["blocked"]))

    def test_f7_recomputes_pending_jurisprudence_instead_of_trusting_json(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            markdown = root / "piece.md"
            f7 = root / "F7.json"
            markdown.write_text(
                "# PEÇA\n\nO alegado Tema 99.999 STJ resolveria integralmente a controvérsia jurídica examinada.\n",
                encoding="utf-8",
            )
            atomic_write_json(f7, {
                "p0": 0,
                "mdSha256": sha256_file(markdown),
                "citacoesNaoConferidas": [],
                "verificarRestantes": [],
            })
            result = validate_f7(
                {"path": str(f7)},
                document_key=None,
                release_policy="strict_protocol",
                markdown={"path": str(markdown), "sha256": sha256_file(markdown)},
            )
            self.assertFalse(result["approved"])
            self.assertTrue(any("recomputação" in blocker or "diverge" in blocker for blocker in result["blockers"]))

    def test_fidelity_json_cannot_hide_removed_legal_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            markdown = root / "piece.md"
            docx = root / "piece.docx"
            pdf = root / "piece.pdf"
            fake_report = root / "FORMAT_FIDELITY.json"
            markdown.write_text("# PEÇA\n\n" + BODY + "\n", encoding="utf-8")
            document = Document()
            document.add_paragraph("Conteúdo jurídico removido do Word.")
            document.save(docx)
            make_pdf(pdf, "Conteúdo jurídico removido do PDF.")
            atomic_write_json(fake_report, {
                "approved": True,
                "markdown": {"sha256": sha256_file(markdown)},
                "docx": {"sha256": sha256_file(docx)},
                "pdf": {"sha256": sha256_file(pdf)},
                "blocks": {"docxCoverage": 1.0, "pdfCoverage": 1.0},
            })
            files = {
                "md": {"path": str(markdown), "sha256": sha256_file(markdown)},
                "docx": {"path": str(docx), "sha256": sha256_file(docx)},
                "pdf": {"path": str(pdf), "sha256": sha256_file(pdf)},
            }
            result = validate_fidelity({"path": str(fake_report)}, files=files)
            self.assertFalse(result["approved"])
            self.assertTrue(any("recomputação" in item for item in result["findings"]))


if __name__ == "__main__":
    unittest.main(verbosity=2)
