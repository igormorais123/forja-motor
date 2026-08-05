"""Lastro anti-alucinação para fontes normativas e jurisprudenciais.

Um arquivo com nome de julgado e uma URL escrita no corpo não prova nada: a IA
consegue fabricar ambos. A FORJA exige agora um manifesto hash-bound do cache ou
um sidecar de proveniência para fontes arquivadas no caso. A validação é sempre
reexecutada por quem empacota a peça.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import unicodedata
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from forja_n3_common import atomic_write_json, now_iso, read_json, sha256_file


FORJA = Path(__file__).resolve().parent
OFFICIAL_CACHE = FORJA / "cache" / "fontes_oficiais"
OFFICIAL_MANIFEST = OFFICIAL_CACHE / "OFFICIAL_SOURCE_MANIFEST.json"
OFFICIAL_HOST_SUFFIXES = (
    ".stj.jus.br",
    ".stf.jus.br",
    ".tst.jus.br",
    ".tse.jus.br",
    ".cnj.jus.br",
    ".jus.br",
    ".gov.br",
)
SOURCE_FILE_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}


class _OfficialHtmlText(HTMLParser):
    """Extrai texto visível suficiente para conferir identidade e verbatim."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._ignored = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag.casefold() in {"script", "style", "noscript"}:
            self._ignored += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in {"script", "style", "noscript"} and self._ignored:
            self._ignored -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignored and data.strip():
            self.parts.append(data)


def normalize_evidence_text(value: str) -> str:
    """Normalização estável para prova de trecho, sem depender do HTML."""
    value = html.unescape(str(value or ""))
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^0-9A-Za-z]+", " ", value).casefold()
    return re.sub(r"\s+", " ", value).strip()


def source_excerpt_sha256(value: str) -> str:
    return hashlib.sha256(normalize_evidence_text(value).encode("utf-8")).hexdigest()


def _extract_source_text(path: Path) -> str:
    suffix = path.suffix.casefold()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        from docx import Document
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    if suffix == ".pdf":
        import fitz
        document = fitz.open(path)
        try:
            return "\n".join(page.get_text("text") for page in document)
        finally:
            document.close()
    return ""


def _response_text(body: bytes, content_type: str) -> str:
    if body.startswith(b"%PDF") or "application/pdf" in content_type.casefold():
        import fitz
        document = fitz.open(stream=body, filetype="pdf")
        try:
            return "\n".join(page.get_text("text") for page in document)
        finally:
            document.close()
    charset = re.search(r"charset=([A-Za-z0-9._-]+)", content_type, re.I)
    encoding = charset.group(1) if charset else "utf-8"
    decoded = body.decode(encoding, errors="replace")
    if "html" not in content_type.casefold() and "<html" not in decoded[:1000].casefold():
        return decoded
    parser = _OfficialHtmlText()
    parser.feed(decoded)
    return "\n".join(parser.parts)


@lru_cache(maxsize=256)
def _fetch_official(url: str, timeout: int = 25) -> dict:
    """Busca a fonte no host oficial; nunca recebe conteúdo fornecido pelo agente."""
    try:
        request = Request(
            url,
            headers={
                "User-Agent": "FORJA-Official-Source-Verification/1.0",
                "Accept": "text/html,application/pdf,text/plain;q=0.9,*/*;q=0.1",
            },
        )
        with urlopen(request, timeout=timeout) as response:
            body = response.read(20_000_001)
            final_url = response.geturl()
            status = int(getattr(response, "status", 200) or 200)
            content_type = str(response.headers.get("Content-Type") or "")
        if len(body) > 20_000_000:
            return {"ok": False, "error": "resposta oficial excede 20 MB", "status": status, "finalUrl": final_url}
        return {
            "ok": status == 200 and _official_url(final_url) and len(body) >= 150,
            "status": status,
            "finalUrl": final_url,
            "contentType": content_type,
            "bodySha256": hashlib.sha256(body).hexdigest(),
            "text": _response_text(body, content_type),
        }
    except HTTPError as exc:
        return {"ok": False, "error": f"HTTP {exc.code}", "status": exc.code, "finalUrl": url}
    except (URLError, TimeoutError, OSError, ValueError) as exc:
        return {"ok": False, "error": str(exc), "status": None, "finalUrl": url}


def _candidate_anchors(text: str) -> list[str]:
    """Seleciona trechos materiais; URL/cabeçalho jamais contam como prova."""
    candidates = re.findall(r"[\"“]([^\"”]{70,1600})[\"”]", text, re.S)
    candidates += [
        line.strip()
        for line in text.splitlines()
        if len(line.strip()) >= 90
        and not re.match(r"^(?:URL|Fonte|Capturad[oa]|Obs\.?):", line.strip(), re.I)
    ]
    normalized = []
    for value in candidates:
        item = normalize_evidence_text(value)
        words = item.split()
        if 10 <= len(words) <= 240 and item not in normalized:
            normalized.append(item)
    return sorted(normalized, key=len, reverse=True)[:20]


def validate_live_official_source(
    path: Path,
    record: dict,
    *,
    required_excerpt: str | None = None,
    fetcher=None,
) -> dict:
    """Reexecuta a prova no HTTPS oficial e compara identidade e verbatim.

    Isso fecha a brecha em que a própria IA inventava arquivo, URL e manifesto.
    Falha de rede, WAF ou página alterada bloqueia o uso final: não vira PASS.
    """
    findings = []
    url = str(record.get("sourceUrl") or "")
    if not _official_url(url):
        return {"approved": False, "findings": ["URL oficial inválida para verificação viva"], "evidence": None}
    result = (fetcher or _fetch_official)(url)
    if not isinstance(result, dict) or not result.get("ok"):
        detail = (result or {}).get("error") if isinstance(result, dict) else "resposta inválida"
        return {"approved": False, "findings": [f"captura viva da fonte oficial falhou: {detail}"], "evidence": result}
    final_url = str(result.get("finalUrl") or "")
    if not _official_url(final_url):
        findings.append("redirecionamento saiu de domínio oficial")
    live_text = str(result.get("text") or "")
    local_text = _extract_source_text(Path(path)) if Path(path).is_file() else ""
    identity = record.get("identity")
    if identity is not None and not _identity_present(live_text, identity):
        findings.append("identidade jurisprudencial ausente da resposta oficial viva")

    live_normalized = normalize_evidence_text(live_text)
    local_normalized = normalize_evidence_text(local_text)
    matched_excerpt_hash = None
    if required_excerpt is not None:
        excerpt_normalized = normalize_evidence_text(required_excerpt)
        if len(excerpt_normalized.split()) < 10 or len(excerpt_normalized) < 70:
            findings.append("trecho probatório curto ou inespecífico")
        elif excerpt_normalized not in local_normalized:
            findings.append("trecho probatório não consta da captura arquivada")
        elif excerpt_normalized not in live_normalized:
            findings.append("trecho probatório não consta da fonte oficial viva")
        else:
            matched_excerpt_hash = source_excerpt_sha256(required_excerpt)
    else:
        anchors = _candidate_anchors(local_text)
        matched = next((anchor for anchor in anchors if anchor in live_normalized), None)
        if matched is None:
            findings.append("nenhum verbatim material da captura foi reproduzido pela fonte oficial viva")
        else:
            matched_excerpt_hash = hashlib.sha256(matched.encode("utf-8")).hexdigest()
    evidence = {
        "status": result.get("status"),
        "finalUrl": final_url,
        "bodySha256": result.get("bodySha256"),
        "matchedExcerptSha256": matched_excerpt_hash,
    }
    return {"approved": not findings, "findings": findings, "evidence": evidence}


def _official_url(value: str) -> bool:
    try:
        parsed = urlparse(str(value).strip())
    except ValueError:
        return False
    host = (parsed.hostname or "").casefold()
    return parsed.scheme == "https" and any(host == suffix[1:] or host.endswith(suffix) for suffix in OFFICIAL_HOST_SUFFIXES)


def _urls_from_text(text: str) -> list[str]:
    return re.findall(r"https://[^\s<>\]\[\"']+", text, re.I)


def _identity_from_name(path: Path) -> dict | None:
    match = re.match(
        r"^(STJ|STF)_(SUMULA_VINCULANTE|SUMULA|TEMA|INFORMATIVO|RESP|ARESP|ERESP|"
        r"RMS|RHC|HC|MS|CC|RCL|SLS|SS|RE|ARE|ADI|ADC|ADO|ADPF|MI|PET|INQ|AP|EXT)"
        r"_(\d+)\.txt$",
        path.name,
        re.I,
    )
    if match:
        return {"court": match.group(1).upper(), "kind": match.group(2).upper(), "number": match.group(3)}
    article = re.match(r"^PLANALTO_([A-Z0-9_]+)_ART_(\d+[A-Z]?)\.txt$", path.name, re.I)
    if article:
        return {"kind": "ARTICLE", "code": article.group(1).upper(), "article": article.group(2).upper()}
    law = re.match(r"^PLANALTO_(LEI(?:_COMPLEMENTAR)?|DECRETO(?:_LEI)?|EC)_(\d+)(?:_(\d{4}))?\.txt$", path.name, re.I)
    if law:
        return {"kind": law.group(1).upper(), "number": law.group(2), "year": law.group(3) or ""}
    return None


def _identity_present(text: str, identity: dict | None) -> bool:
    if identity is None:
        return True
    digits = re.sub(r"\D", "", text)
    kind = str(identity.get("kind") or "").upper()
    if kind == "ARTICLE":
        article = re.sub(r"\D", "", str(identity.get("article") or ""))
        code = str(identity.get("code") or "").replace("_", " ")
        return bool(article and article in digits and code.casefold() in text.casefold())
    if not identity.get("number"):
        return False
    if not identity.get("court"):
        return str(identity["number"]) in digits
    court = str(identity["court"])
    kind_patterns = {
        "SUMULA": r"s[úu]mulas?",
        "SUMULA_VINCULANTE": r"s[úu]mula\s+vinculante",
        "TEMA": r"temas?",
        "INFORMATIVO": r"informativos?",
        "RESP": r"resp|recurso\s+especial",
        "ARESP": r"aresp|agravo\s+em\s+recurso\s+especial",
        "ERESP": r"eresp|embargos\s+de\s+diverg[eê]ncia",
        "RE": r"recurso\s+extraordin[aá]rio|\bre\b",
        "ARE": r"agravo\s+em\s+recurso\s+extraordin[aá]rio|\bare\b",
        "ADI": r"\badi\b|a[cç][aã]o\s+direta\s+de\s+inconstitucionalidade",
        "ADPF": r"\badpf\b|argui[cç][aã]o\s+de\s+descumprimento",
        "HC": r"\bhc\b|habeas\s+corpus",
        "RHC": r"\brhc\b|recurso\s+em\s+habeas\s+corpus",
        "RMS": r"\brms\b|recurso\s+ordin[aá]rio\s+em\s+mandado",
        "MS": r"\bms\b|mandado\s+de\s+seguran[cç]a",
        "CC": r"\bcc\b|conflito\s+de\s+compet[eê]ncia",
        "RCL": r"\brcl\b|reclama[cç][aã]o",
        "ADC": r"\badc\b|a[cç][aã]o\s+declarat[oó]ria",
        "ADO": r"\bado\b|inconstitucionalidade\s+por\s+omiss[aã]o",
        "MI": r"\bmi\b|mandado\s+de\s+injun[cç][aã]o",
    }
    kind_pattern = kind_patterns.get(kind, re.escape(kind))
    return (
        identity["number"] in digits
        and court.casefold() in text.casefold()
        and re.search(kind_pattern, text, re.I) is not None
    )


def build_manifest(cache_dir: Path = OFFICIAL_CACHE, output: Path = OFFICIAL_MANIFEST) -> dict:
    """Inventaria capturas existentes; não transforma arquivo sem URL em oficial."""
    entries = {}
    rejected = []
    for path in sorted(Path(cache_dir).iterdir()):
        if not path.is_file() or path.name == Path(output).name or path.suffix.lower() not in SOURCE_FILE_SUFFIXES:
            continue
        identity = _identity_from_name(path)
        if path.suffix.lower() not in {".txt", ".md"}:
            # PDFs podem entrar por sidecar próprio; não se inventa URL lendo binário.
            continue
        text = path.read_text(encoding="utf-8", errors="replace")
        urls = [url.rstrip(".,;)") for url in _urls_from_text(text) if _official_url(url.rstrip(".,;)"))]
        if len(text) < 150 or not urls or not _identity_present(text, identity):
            rejected.append(path.name)
            continue
        captured = re.search(r"Capturad[oa]\s+em:\s*([^\r\n]+)", text, re.I)
        entries[path.name] = {
            "sha256": sha256_file(path),
            "sourceUrl": urls[0],
            "capturedAt": captured.group(1).strip() if captured else None,
            "identity": identity,
        }
    payload = {
        "schemaVersion": 1,
        "generatedAt": now_iso(),
        "cacheRoot": str(Path(cache_dir).resolve()),
        "entries": entries,
        "rejected": rejected,
    }
    atomic_write_json(Path(output), payload)
    return payload


def validate_cached_source(
    path: Path,
    manifest_path: Path = OFFICIAL_MANIFEST,
    *,
    cache_dir: Path = OFFICIAL_CACHE,
    require_live: bool = False,
    required_excerpt: str | None = None,
    fetcher=None,
) -> dict:
    path = Path(path)
    findings = []
    try:
        if path.resolve().parent != Path(cache_dir).resolve():
            findings.append("fonte fora do cache oficial")
    except OSError:
        findings.append("caminho de fonte inválido")
    manifest = read_json(Path(manifest_path), None)
    entry = ((manifest or {}).get("entries") or {}).get(path.name) if isinstance(manifest, dict) else None
    if not isinstance(entry, dict):
        findings.append("fonte ausente do manifesto oficial")
        return {"approved": False, "findings": findings, "record": None}
    if not path.is_file() or entry.get("sha256") != sha256_file(path):
        findings.append("hash da fonte oficial divergente")
    if not _official_url(str(entry.get("sourceUrl") or "")):
        findings.append("URL de fonte oficial inválida")
    identity = entry.get("identity")
    if path.suffix.lower() in {".txt", ".md"} and path.is_file():
        text = path.read_text(encoding="utf-8", errors="replace")
        if len(text) < 150 or not _identity_present(text, identity):
            findings.append("identidade jurisprudencial não consta do verbatim")
    live = None
    if require_live and not findings:
        live = validate_live_official_source(
            path,
            entry,
            required_excerpt=required_excerpt,
            fetcher=fetcher,
        )
        findings += live["findings"]
    return {"approved": not findings, "findings": findings, "record": entry, "live": live}


def sidecar_path(source: Path) -> Path:
    return Path(str(source) + ".source.json")


def validate_archived_source(
    path: Path,
    *,
    require_live: bool = False,
    required_excerpt: str | None = None,
    fetcher=None,
) -> dict:
    """Valida cópia local contra entrada já registrada no cache protegido.

    Um sidecar criado pelo mesmo agente não é fonte independente. Por isso ele
    precisa apontar para ``registryEntry`` existente no manifesto central e o
    arquivo local deve ser byte a byte idêntico à captura registrada.
    """
    path = Path(path)
    sidecar = sidecar_path(path)
    payload = read_json(sidecar, None)
    findings = []
    if not path.is_file():
        findings.append("arquivo-fonte ausente")
    if not isinstance(payload, dict) or payload.get("schemaVersion") != 1:
        findings.append("sidecar de proveniência ausente ou inválido")
        return {"approved": False, "findings": findings, "record": None, "sidecar": str(sidecar)}
    if path.is_file() and payload.get("sha256") != sha256_file(path):
        findings.append("hash do arquivo-fonte diverge do sidecar")
    if not _official_url(str(payload.get("sourceUrl") or "")):
        findings.append("sidecar não aponta para fonte oficial HTTPS")
    identity = payload.get("identity")
    identity_ok = isinstance(identity, dict) and (
        (identity.get("court") and identity.get("number"))
        or (identity.get("kind") == "ARTICLE" and identity.get("code") and identity.get("article"))
        or (identity.get("kind") and identity.get("number"))
    )
    if not identity_ok:
        findings.append("sidecar não declara identidade jurídica suficiente")
    if not payload.get("capturedAt"):
        findings.append("sidecar não registra a data da captura")
    registry_name = str(payload.get("registryEntry") or "").strip()
    registry_path = OFFICIAL_CACHE / registry_name
    registry = validate_cached_source(
        registry_path,
        require_live=require_live,
        required_excerpt=required_excerpt,
        fetcher=fetcher,
    ) if registry_name else {
        "approved": False,
        "findings": ["registro central ausente"],
        "record": None,
    }
    if not registry["approved"]:
        findings.append("sidecar não está ancorado em captura oficial registrada")
    else:
        registered = registry.get("record") or {}
        if payload.get("sha256") != registered.get("sha256"):
            findings.append("arquivo local não é cópia exata da captura registrada")
        if payload.get("sourceUrl") != registered.get("sourceUrl"):
            findings.append("URL do sidecar diverge do registro central")
        if payload.get("identity") != registered.get("identity"):
            findings.append("identidade do sidecar diverge do registro central")
    return {
        "approved": not findings,
        "findings": findings,
        "record": payload,
        "sidecar": str(sidecar),
        "live": registry.get("live"),
    }


def validate_source_path(
    path: Path,
    *,
    require_live: bool = False,
    required_excerpt: str | None = None,
    fetcher=None,
) -> dict:
    path = Path(path)
    try:
        in_cache = path.resolve().parent == OFFICIAL_CACHE.resolve()
    except OSError:
        in_cache = False
    return validate_cached_source(
        path,
        require_live=require_live,
        required_excerpt=required_excerpt,
        fetcher=fetcher,
    ) if in_cache else validate_archived_source(
        path,
        require_live=require_live,
        required_excerpt=required_excerpt,
        fetcher=fetcher,
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="Manifesta e valida fontes oficiais da FORJA")
    parser.add_argument("--build-manifest", action="store_true")
    parser.add_argument("--validate", type=Path)
    parser.add_argument("--live", action="store_true", help="reconsulta o HTTPS oficial e confere identidade/verbatim")
    args = parser.parse_args()
    if args.build_manifest:
        result = build_manifest()
    elif args.validate:
        result = validate_source_path(args.validate, require_live=args.live)
    else:
        parser.error("use --build-manifest ou --validate")
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result.get("approved", True) else 1


if __name__ == "__main__":
    raise SystemExit(main())
