"""Independent pre-human visual QA over SVG, DOCX and final PDF."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from forja_n3_common import FORJA, ForjaN3Error, atomic_write_json, now_iso, sha256_file

sys.path.insert(0, str(FORJA.parent / "_FERRAMENTAS"))
from medina_visual_lint import lint_svg  # noqa: E402
from forja_docx_layout import audit_docx_layout
from forja_fidelity import write_fidelity
from forja_visual_review import build_pending_review, validate_visual_review


MARKDOWN_LEAKS = (
    ("markdown_heading", re.compile(r"(?m)^\s*(?:\d+\.\s*)?#{2,6}\s+")),
    ("markdown_blockquote", re.compile(r"(?m)^\s*(?:\d+\.\s*)?>\s+")),
    ("template_marker", re.compile(r"\{\{[^}]+\}\}")),
)


def lint_text(text: str) -> list[dict]:
    findings = []
    for code, pattern in MARKDOWN_LEAKS:
        for match in pattern.finditer(text):
            findings.append({"severity": "P0", "code": code, "sample": match.group(0)[:100]})
    for line in text.splitlines():
        if "FIGURA" not in line.upper():
            continue
        numbers = re.findall(r"FIGURA\s*(\d+)", line, re.I)
        if len(numbers) > 1:
            findings.append({"severity": "P0", "code": "multiple_figure_numbers_in_caption", "sample": line[:160]})
    return findings


def lint_docx(path: Path) -> dict:
    from docx import Document

    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    findings = lint_text(text)
    if not document.paragraphs and not document.tables:
        findings.append({"severity": "P0", "code": "empty_docx"})
    return {"file": str(path), "sha256": sha256_file(path), "approved": not findings, "findings": findings}


def _overlap_ratio(left: tuple[float, float, float, float], right: tuple[float, float, float, float]) -> float:
    width = max(0.0, min(left[2], right[2]) - max(left[0], right[0]))
    height = max(0.0, min(left[3], right[3]) - max(left[1], right[1]))
    intersection = width * height
    area_left = max(1.0, (left[2] - left[0]) * (left[3] - left[1]))
    area_right = max(1.0, (right[2] - right[0]) * (right[3] - right[1]))
    return intersection / min(area_left, area_right)


def inspect_pdf(pdf: Path, qa_dir: Path, *, generator_run_id: str, reviewer_run_id: str, dpi: int = 120) -> dict:
    import fitz

    if generator_run_id == reviewer_run_id:
        raise ForjaN3Error("o gerador não pode aprovar o próprio run")
    qa_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(pdf)
    pages = []
    all_findings = []
    matrix = fitz.Matrix(dpi / 72, dpi / 72)
    for page_number, page in enumerate(document, 1):
        page_findings = []
        text = page.get_text("text")
        page_findings += lint_text(text)
        blocks = [block for block in page.get_text("blocks") if str(block[4]).strip()]
        for index, left in enumerate(blocks):
            left_box = tuple(float(value) for value in left[:4])
            for right in blocks[index + 1:]:
                right_box = tuple(float(value) for value in right[:4])
                if _overlap_ratio(left_box, right_box) > 0.28:
                    page_findings.append({
                        "severity": "P0",
                        "code": "pdf_text_blocks_overlap",
                        "sampleA": " ".join(str(left[4]).split())[:80],
                        "sampleB": " ".join(str(right[4]).split())[:80],
                    })
        if len(text.strip()) < 20 and not page.get_images(full=True):
            page_findings.append({"severity": "P0", "code": "unexpected_blank_page"})
        bounds = page.rect
        for block in blocks:
            if block[0] < -1 or block[1] < -1 or block[2] > bounds.width + 1 or block[3] > bounds.height + 1:
                page_findings.append({"severity": "P0", "code": "pdf_text_clipped", "sample": str(block[4])[:80]})
        image_path = qa_dir / f"p{page_number:03d}.png"
        page.get_pixmap(matrix=matrix, alpha=False).save(image_path)
        status = "pass" if not page_findings else "blocked"
        pages.append({
            "page": page_number,
            "imagePath": str(image_path),
            "imageSha256": sha256_file(image_path),
            "lint": status,
            # Isto é apenas análise automática. Nunca equivale a inspeção
            # visual independente e não pode promover a própria página.
            "automatedReview": {
                "status": status,
                "reviewer": "forja-pdf-static-lint",
                "runId": generator_run_id,
                "reviewedAt": now_iso(),
                "reviewType": "automated",
                "humanReviewed": False,
            },
            "findings": page_findings,
        })
        all_findings += [{**item, "page": page_number} for item in page_findings]
    return {
        "pdfSha256": sha256_file(pdf),
        "pageCount": len(pages),
        "renderDpi": dpi,
        "generatorRunId": generator_run_id,
        "reviewerRunId": reviewer_run_id,
        "pages": pages,
        "findings": all_findings,
        "approved": not all_findings and bool(pages),
    }


def run_visual_qa(
    pdf: Path,
    output: Path,
    *,
    qa_dir: Path,
    generator_run_id: str,
    reviewer_run_id: str,
    docx: Path | None = None,
    markdown: Path | None = None,
    fidelity_output: Path | None = None,
    svgs: list[Path] | None = None,
    manual_review: Path | None = None,
    pending_review_output: Path | None = None,
    layout_exceptions: Path | None = None,
) -> dict:
    svg_results = [lint_svg(path) for path in svgs or []]
    docx_result = lint_docx(docx) if docx else None
    layout_result = audit_docx_layout(docx, exceptions=layout_exceptions) if docx else None
    pdf_result = inspect_pdf(pdf, qa_dir, generator_run_id=generator_run_id, reviewer_run_id=reviewer_run_id)
    fidelity_result = None
    if markdown is not None:
        if docx is None:
            raise ForjaN3Error("fidelidade Markdown exige DOCX")
        fidelity_output = fidelity_output or (output.parent / "FORMAT_FIDELITY.json")
        fidelity_result = write_fidelity(markdown, docx, pdf, fidelity_output)
    findings = []
    for result in svg_results:
        findings += [{**item, "file": result["file"]} for item in result["findings"]]
    if docx_result:
        findings += [{**item, "file": docx_result["file"]} for item in docx_result["findings"]]
    if layout_result:
        findings += [{**item, "file": str(docx)} for item in layout_result["findings"]]
    findings += pdf_result["findings"]
    if fidelity_result and not fidelity_result["approved"]:
        findings += [{**item, "file": str(fidelity_output)} for item in fidelity_result["findings"]]
    pending_review_output = pending_review_output or (output.parent / "VISUAL_REVIEW_PENDING.json")
    if manual_review is None:
        pending = build_pending_review(
            pending_review_output,
            pdf=pdf,
            docx=docx,
            rendered_pages=pdf_result["pages"],
            generator_run_id=generator_run_id,
        )
        review_result = {
            "approved": False,
            "pendingTemplate": str(pending_review_output),
            "pendingTemplateSha256": sha256_file(pending_review_output),
            "pages": pending["pages"],
            "findings": [{"severity": "P0", "code": "manual_visual_review_missing"}],
        }
    else:
        review_result = validate_visual_review(
            manual_review,
            pdf=pdf,
            docx=docx,
            rendered_pages=pdf_result["pages"],
            generator_run_id=generator_run_id,
            expected_reviewer_run_id=reviewer_run_id,
        )
    findings += review_result["findings"]

    reviewed_by_page = {item["page"]: item for item in review_result.get("pages") or []}
    for page in pdf_result["pages"]:
        evidence = reviewed_by_page.get(page["page"])
        page["independentReview"] = {
            "status": evidence.get("status") if evidence else "pending",
            "reviewer": (review_result.get("reviewer") or {}).get("id"),
            "runId": (review_result.get("reviewer") or {}).get("runId"),
            "reviewedAt": review_result.get("reviewedAt"),
            "reviewType": (review_result.get("reviewer") or {}).get("type"),
            "humanReviewed": (review_result.get("reviewer") or {}).get("type") == "human",
            "checks": evidence.get("checks") if evidence else {},
            "reviewedImageSha256": evidence.get("reviewedImageSha256") if evidence else None,
        }

    ledger = {
        "schemaVersion": 2,
        "generatedAt": now_iso(),
        **pdf_result,
        "svg": svg_results,
        "docx": docx_result,
        "layoutAudit": layout_result,
        "manualVisualReview": review_result,
        "fidelity": fidelity_result,
        "findings": findings,
        "approved": not findings and pdf_result["approved"] and review_result["approved"] and (
            layout_result is None or layout_result["approved"]
        ),
    }
    atomic_write_json(output, ledger)
    return ledger


def main() -> None:
    parser = argparse.ArgumentParser(description="QA visual independente FORJA N3")
    parser.add_argument("pdf", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--qa-dir", type=Path, required=True)
    parser.add_argument("--generator-run", required=True)
    parser.add_argument("--reviewer-run", required=True)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--markdown", type=Path)
    parser.add_argument("--fidelity-output", type=Path)
    parser.add_argument("--svg", type=Path, action="append", default=[])
    parser.add_argument("--manual-review", type=Path)
    parser.add_argument("--pending-review-output", type=Path)
    parser.add_argument("--layout-exceptions", type=Path)
    args = parser.parse_args()
    result = run_visual_qa(
        args.pdf, args.output, qa_dir=args.qa_dir,
        generator_run_id=args.generator_run, reviewer_run_id=args.reviewer_run,
        docx=args.docx, markdown=args.markdown, fidelity_output=args.fidelity_output, svgs=args.svg,
        manual_review=args.manual_review, pending_review_output=args.pending_review_output,
        layout_exceptions=args.layout_exceptions,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if result["approved"] else 1)


if __name__ == "__main__":
    main()
