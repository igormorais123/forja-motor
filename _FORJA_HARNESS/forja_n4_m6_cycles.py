"""Materialize the three real M6 canary cycles without changing source petitions."""

from __future__ import annotations

import json
from pathlib import Path

import forja_acervo

from docx import Document

from forja_case_tests import run_suite, suite_hash
from forja_consistency import inspect_physical_document
from forja_fidelity import compare_fidelity
from forja_n3_common import FORJA, atomic_write_json, canonical_hash, now_iso, read_json, sha256_file
from forja_n4_common import build_envelope, write_artifact
from forja_n4_validate import validate_case


PRODUCER = "forja-n4-m6-canary-producer-20260711"
REVIEWER = "forja-n4-m6-independent-review-20260711"

CASES = {
    "patricia": {
        "caseId": forja_acervo.caso("CASO-19"),
        "product": "memoriais de apelação",
        "response": False,
        "layout": "medina-visual-law-v1",
        "science": False,
        "questions": [
            ("Q-001", "fact", "Qual foi o valor do contrato?", "R$ 165.000,00, pagos em três cheques de R$ 55.000,00.", "DOC-FINAL"),
            ("Q-002", "calculation", "Qual parcela não foi entregue?", "R$ 90.000,00, equivalentes a aproximadamente 54,5% do contrato.", "QUANT-001"),
            ("Q-003", "request", "Qual é o pedido principal?", "Majoração para R$ 50.000,00 por autor, com alternativa de outro patamar superior ao fixado.", "REQUEST-001"),
            ("Q-004", "risk", "Qual é a principal fragilidade?", "O patamar de R$ 50.000,00 ainda demanda melhor comparação e individualização do dano extrapatrimonial.", "CICERO"),
            ("Q-005", "visual", "O documento final é legível?", "Sim, nas seis páginas inspecionadas individualmente.", "F8"),
        ],
        "thesis": "O valor indenizatório deve refletir as circunstâncias concretas já reconhecidas, sem tarifamento nem multiplicadores artificiais.",
        "objection": "A restituição material e a falta de repercussões pessoais individualizadas podem justificar a manutenção do quantum.",
        "anchors": [("22 de fevereiro de 2013", "contains"), ("R$ 165.000,00", "contains"), ("três cheques de R$ 55.000,00", "contains"), ("R$ 90.000,00", "contains"), ("aproximadamente 54,5%", "contains"), ("R$ 10.000,00 por autor", "contains"), ("R$ 50.000,00 para cada autor", "contains"), ("método bifásico", "contains"), ("honorários recursais", "not_contains"), ("art. 944, parágrafo único", "not_contains")],
    },
    "CASO-16": {
        "caseId": forja_acervo.caso("CASO-16"),
        "product": "memoriais de impugnação a agravo interno",
        "response": True,
        "layout": "medina-visual-law-v1",
        "science": False,
        "questions": [
            ("Q-001", "procedural_event", "Qual é o recurso atual?", "Agravo interno da CASO-01 contra a decisão que conheceu do agravo para não conhecer do recurso especial.", "DOC-FINAL"),
            ("Q-002", "opponent_response", "Como tratar a Súmula 182/STJ?", "Somente no ponto restrito de impugnação específica; não como pedido de não conhecimento global.", "CMP-001"),
            ("Q-003", "fact", "Todos os fatos precederam todos os administradores?", "Não; a peça separa os fatos de 2008 da transferência de controle de 2011.", "DOC-FINAL"),
            ("Q-004", "request", "Qual é o resultado pretendido?", "Manutenção da decisão agravada, pelos limites de admissibilidade e pelas premissas do acórdão.", "REQUEST-001"),
            ("Q-005", "visual", "A cronologia e os quadros são legíveis?", "Sim, nas sete páginas inspecionadas individualmente.", "F8"),
        ],
        "thesis": "O agravo interno não supera os limites de admissibilidade nem as premissas temporais e fáticas reconhecidas pelo TRF4.",
        "objection": "A CASO-01 pode sustentar que busca apenas requalificação jurídica dos fatos e que impugnou suficientemente a decisão.",
        "anchors": [("Não se pede o não conhecimento global do agravo", "contains"), ("Súmula 182 do STJ", "contains"), ("dois grupos temporais distintos", "contains"), ("ocorreram em 2008", "contains"), ("transferência de controle realizada em 2011", "contains"), ("teoria da asserção", "contains"), ("Súmula 7 do STJ", "contains"), ("art. 1.021, § 1º, do CPC", "contains"), ("art. 259, § 2º, do RISTJ", "contains"), ("[VERIFICAR]", "not_contains")],
    },
    "health": {
        "caseId": "case-email-auto-19f3f25cb64df962",
        "product": "petição inicial com tutela de urgência em saúde",
        "response": False,
        "layout": "medina-word-v1",
        "science": True,
        "questions": [
            ("Q-001", "fact", "Qual foi o ato impugnado?", "Exclusão integral do beneficiário sob acusação de omissão/fraude.", "DOC-FINAL"),
            ("Q-002", "merit", "Qual é a consequência legal máxima da doença preexistente?", "Cobertura parcial temporária por até 24 meses, e não exclusão total.", "THESIS-001"),
            ("Q-003", "science", "A literatura interdisciplinar prova o caso individual?", "Não; apenas contextualiza barreiras de acesso de pessoas autistas.", "SCIENCE"),
            ("Q-004", "request", "Qual é a tutela imediata?", "Reintegração em 48 horas, cirurgia em 15 dias e multa de R$ 5.000,00 por dia.", "REQUEST-001"),
            ("Q-005", "visual", "O documento final é legível?", "Sim, nas doze páginas inspecionadas individualmente.", "F8"),
        ],
        "thesis": "A exclusão total é incompatível com o regime legal da doença preexistente e com o material documental descrito na inicial.",
        "objection": "A operadora pode controverter a ciência prévia, a indicação funcional e a urgência do procedimento.",
        "anchors": [("por, no máximo, 24 meses", "contains"), ("CPT até 19/07/2027", "contains"), ("reintegração em 48 horas", "contains"), ("cirurgia em 15 dias", "contains"), ("R$ 5.000,00/dia", "contains"), ("R$ 60.000,00", "contains"), ("R$ 100.000,00", "contains"), ("Súmula 609 do STJ", "contains"), ("indicação funcional", "contains"), ("[VERIFICAR]", "not_contains")],
    },
}


def _docx_semantic_text(path: Path) -> str:
    doc = Document(path)
    values = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                values.append(" | ".join(cells))
    return "\n".join(values)


def run(key: str) -> dict:
    profile = CASES[key]
    case_dir = FORJA / "state" / profile["caseId"]
    cycle = case_dir / "n4_cycle_m6"
    canonical = cycle / "CANONICAL_TEXT_FROM_FINAL_DOCX.txt"
    docx = cycle / "FINAL_N4_M6.docx"
    pdf = cycle / "FINAL_N4_M6.pdf"
    f8 = cycle / "F8_QA_LEDGER_N4.json"
    helena = case_dir / "F4_PARECER_HELENA.md"
    cicero = case_dir / "F4_PARECER_CICERO.md"
    qa = read_json(f8, {}) or {}
    if qa.get("approved") is not True or qa.get("generatorRunId") == qa.get("reviewerRunId"):
        raise RuntimeError(f"QA visual independente ainda não aprovado: {key}")

    original_docx = Path(str((qa.get("sourceFiles") or {}).get("docx") or ""))
    original_pdf = Path(str((qa.get("sourceFiles") or {}).get("pdf") or ""))
    if not original_docx.is_file() or not original_pdf.is_file():
        raise RuntimeError(f"origem verificável do canário ausente: {key}")
    if key == "health":
        stable_source = case_dir / "producao" / "MINUTA_INICIAL_TJDFT.md"
        if not stable_source.is_file():
            raise RuntimeError(f"fonte Markdown estável ausente: {stable_source}")
        sources = {"m6:source_markdown": stable_source, "m6:canonical": canonical, "m6:docx": docx, "m6:pdf": pdf, "m6:f8": f8, "m6:helena": helena, "m6:cicero": cicero}
    else:
        stable_source = original_docx
        sources = {"m6:source_docx": original_docx, "m6:source_pdf": original_pdf, "m6:canonical": canonical, "m6:docx": docx, "m6:pdf": pdf, "m6:f8": f8, "m6:helena": helena, "m6:cicero": cicero}
    hashes = {name: sha256_file(path) for name, path in sources.items()}
    manifest = read_json(case_dir / "FORJA_CASE_MANIFEST.json", {}) or {}
    registry = manifest.get("n4SourceRegistry") or {}
    if key == "health":
        registry.pop("m6:source_docx", None)
        registry.pop("m6:source_pdf", None)
        science_paths = {
            "science:caseStudy": case_dir / "producao" / "ESTUDO_ESTRATEGICO_PLANO_SAUDE.md",
            "science:discovery": FORJA / "cache" / "science" / "10.3399-bjgpopen20X101030" / "DISCOVERY.json",
            "science:crossref": FORJA / "cache" / "science" / "10.3399-bjgpopen20X101030" / "CROSSREF_DOI.json",
            "science:pmc": FORJA / "cache" / "science" / "10.3399-bjgpopen20X101030" / "PMC7465578.xml",
        }
        for name, path in science_paths.items():
            if not path.is_file():
                raise RuntimeError(f"fonte científica verificável ausente: {path}")
            registry[name] = {"path": str(path), "sha256": sha256_file(path), "status": "active"}
    for name, path in sources.items():
        registry[name] = {"path": str(path), "sha256": hashes[name], "status": "active"}
    manifest["n4SourceRegistry"] = registry
    manifest["specVersionCandidate"] = "N4.0-candidate"
    manifest["n4M6Cycle"] = {"mode": "pilot_blocking", "status": "candidate", "evidenceClass": "retrospective_baseline", "cyclePath": str(cycle), "originalsUntouched": True}
    atomic_write_json(case_dir / "FORJA_CASE_MANIFEST.json", manifest)
    new_hashes = list(hashes.values())

    def save(filename: str, content: dict, *, applicability: str = "required", status: str = "approved", source_hashes: list[str] | None = None) -> dict:
        payload = build_envelope(case_dir, filename, content, source_hashes=source_hashes or new_hashes, producer_run_id=PRODUCER, reviewer_run_id=REVIEWER, applicability=applicability, status=status)
        write_artifact(case_dir, filename, payload)
        return payload

    def na(filename: str, reason: str) -> None:
        save(filename, {"justification": reason}, applicability="not_applicable", status="not_applicable")

    if not profile["science"]:
        save("F2_N4_CLASSIFICATION.json", {"product": profile["product"], "responsePiece": profile["response"], "complexity": "high", "science": {"mode": "not_applicable", "triggerPropositionIds": [], "domains": [], "justification": "A tese material deste canário depende de autos e fontes jurídicas, não de proposição extrajurídica.", "requiredBeforeF6": False}})

    nodes = [
        {"id": "DOC-FINAL", "type": "document", "sourceArtifact": str(canonical)},
        {"id": "THESIS-001", "type": "thesis", "sourceArtifact": "F4_THESIS_MATURITY.json"},
        {"id": "REQUEST-001", "type": "request", "sourceArtifact": "F4_COVERAGE_MATRIX.json"},
        {"id": "F8", "type": "document", "sourceArtifact": str(f8)},
        {"id": "HELENA", "type": "decision", "sourceArtifact": str(helena)},
        {"id": "CICERO", "type": "decision", "sourceArtifact": str(cicero)},
    ]
    if key == "patricia":
        nodes.append({"id": "QUANT-001", "type": "calculation", "sourceArtifact": "F4_QUANTIFICATION_SCENARIOS.json"})
    if key == "CASO-16":
        nodes.append({"id": "CMP-001", "type": "document", "sourceArtifact": "F3_DOCUMENT_COMPARISON.json"})
    if profile["science"]:
        nodes.append({"id": "SCIENCE", "type": "document", "sourceArtifact": "F5C_EVIDENCE_SYNTHESIS.json"})

    questions = []
    for qid, category, text, answer, support in profile["questions"]:
        questions.append({"questionId": qid, "parentId": None, "category": category, "text": text, "origin": "m6_real_canary", "materiality": "decisive", "status": "answered", "answer": answer, "supportIds": [support], "dependsOn": [], "owner": "F4", "reviewStatus": "confirmed"})
    save("F2_QUESTION_TREE.json", {"questions": questions, "coverage": {"total": len(questions), "material": len(questions), "answeredMaterial": len(questions), "blockedMaterial": 0}})
    save("F3_EVENT_IDENTITY.json", {"events": [{"eventId": "EVENT-001", "canonicalLabel": profile["product"], "sourceId": "DOC-FINAL", "locator": "endereçamento e síntese inicial", "allowedParaphrases": [], "forbiddenEquivalents": [], "temporalPosition": 1}], "surfaces": [{"surfaceId": "final", "text": canonical.read_text(encoding="utf-8"), "semanticContrast": False}]})

    if key == "CASO-16":
        save("F3_DOCUMENT_COMPARISON.json", {"comparisonSets": [{"setId": "CMP-CASO16-001", "documents": ["decisão agravada", "agravo interno", "memoriais finais"], "units": [{"unitId": "CMP-001", "priorArgument": "decisão aplicou óbices de admissibilidade e preservou premissas do TRF4", "priorResponse": "agravo interno busca superar os óbices", "currentArgument": "memoriais limitam a Súmula 182 ao ponto efetivamente deficiente", "classification": "repeated_with_new_basis", "novelElements": ["delimitação expressa do alcance da Súmula 182"], "prequestioningAssessment": "considered", "consequence": "rebuttal", "reviewStatus": "confirmed"}]}]})
    else:
        na("F3_DOCUMENT_COMPARISON.json", "O canário não responde a uma peça adversária comparável neste ciclo.")
    save("F3_REASONING_GRAPH.json", {"nodes": nodes, "edges": [
        {"edgeId": "E-001", "from": "DOC-FINAL", "to": "THESIS-001", "relation": "supports", "scope": "partial", "reason": "a versão final registra as premissas usadas pela tese", "reviewStatus": "confirmed"},
        {"edgeId": "E-002", "from": "THESIS-001", "to": "REQUEST-001", "relation": "justifies", "scope": "partial", "reason": "a tese organiza o pedido, sujeito às ressalvas registradas", "reviewStatus": "confirmed"},
        {"edgeId": "E-003", "from": "CICERO", "to": "THESIS-001", "relation": "qualifies", "reason": "a auditoria jurídica preserva ressalvas materiais e não libera a versão para protocolo", "reviewStatus": "confirmed"},
    ]})
    na("F3_CONDUCT_LEDGER.json", "O canário não formula imputação externa de má-fé ou sanção por conduta.")
    save("F4_COVERAGE_MATRIX.json", {"items": [{"coverageId": "COV-001", "kind": "material_issue", "originDocumentId": "m6_real_canary", "originLocator": "síntese e pedidos", "statement": profile["thesis"], "supportIds": ["DOC-FINAL"], "priorResponseIds": [], "currentTreatment": "addressed_with_qualification", "draftParagraphIds": ["final-document"], "requestedConsequence": "REQUEST-001", "materiality": "decisive", "status": "covered"}]})
    save("F4_THESIS_MATURITY.json", {"theses": [{"thesisId": "THESIS-001", "statement": profile["thesis"], "role": "primary", "documentaryStrength": "moderate", "legalStrength": "moderate", "gaps": [profile["objection"]], "bestObjection": profile["objection"], "contaminationRisk": "low", "activationTrigger": "revisão humana das fontes e ressalvas do caso", "properVehicle": profile["product"], "helenaDecision": "review_required", "helenaEvidenceId": "m6:helena", "helenaDecisionLocator": "conclusão executiva e seção específica do caso", "ciceroDecision": "reject_current_version", "ciceroEvidenceId": "m6:cicero", "ciceroDecisionLocator": "veredito ou situação de bloqueio no início do parecer"}]})

    tests = []
    for index, (value, kind) in enumerate(profile["anchors"], 1):
        tests.append({"testId": f"CT-{key.upper()}-{index:03d}", "question": f"Critério literal: {value}", "severity": "blocking", "method": "deterministic", "expected": f"{kind}: {value}", "evidenceRequired": ["canonical_text_from_final_docx"], "immutableFromHash": hashes["m6:canonical"], "status": "pending", "evaluator": {"kind": kind, "value": value, "ignoreCase": True}})
    suite = {"suiteId": f"{key.upper()}-N4-M6-v2", "executionMode": "retrospective_baseline", "draftedBeforeFinalText": False, "retrospectiveReason": "O texto final já existia antes da criação desta suíte; os resultados medem regressão e discriminação, não contam como ciclo prospectivo para promoção.", "tests": tests}
    suite["suiteHash"] = suite_hash(suite)
    save("F4_CASE_ACCEPTANCE_TESTS.json", suite)
    na("F4_DECISION_FACTOR_MAP.json", "Este ciclo valida o produto final; não reconstrói fatores decisórios de julgados integrais.")
    na("F4_SETTLEMENT_MAP.json", "Composição não integra o objetivo do canário.")

    if key == "patricia":
        save("F4_INTERTEMPORAL_MAP.json", {"temporalIssues": [{"issueId": "TEMP-001", "rule": "Lei 14.905/2024 e regime temporal dos arts. 389 e 406 do Código Civil", "triggeringDate": "2024-08-30", "dateSourceId": "DOC-FINAL", "dateStatus": "confirmed", "transitionRuleSourceId": "DOC-FINAL", "conclusion": "explicitar os marcos de cada verba e evitar duplicidade", "status": "confirmed"}]})
        save("F4_QUANTIFICATION_SCENARIOS.json", {"scenarios": [
            {"scenarioId": "QUANT-001", "description": "valor total pago", "formula": "cheques * valor_cheque", "knownInputs": [{"name": "cheques", "value": 3, "sourceId": "DOC-FINAL"}, {"name": "valor_cheque", "value": 55000, "sourceId": "DOC-FINAL"}], "disputedInputs": [], "outputs": {"value": 165000}, "status": "confirmed"},
            {"scenarioId": "QUANT-002", "description": "percentual do contrato correspondente aos closets não entregues", "formula": "closets / contrato * 100", "knownInputs": [{"name": "closets", "value": 90000, "sourceId": "DOC-FINAL"}, {"name": "contrato", "value": 165000, "sourceId": "DOC-FINAL"}], "disputedInputs": [], "outputs": {"value": 54.54545454545455}, "status": "confirmed"},
        ]})
    else:
        na("F4_INTERTEMPORAL_MAP.json", "Nenhum conflito de regimes intertemporais é resolvido como eixo deste canário.")
        na("F4_QUANTIFICATION_SCENARIOS.json", "A quantificação não é o eixo de validação deste canário.")

    if not profile["science"]:
        for filename in ("F5C_RESEARCH_PROTOCOL.json", "F5C_STUDY_LEDGER.json", "F5C_EVIDENCE_SYNTHESIS.json", "F5C_CLAIM_EVIDENCE_MAP.json"):
            na(filename, "Dimensão científica classificada como não aplicável ao eixo material deste canário.")

    results = run_suite(suite, canonical, reviewer_run_id=REVIEWER, producer_run_id=PRODUCER)
    save("F7_CASE_TEST_RESULTS.json", {"suiteHash": results["suiteHash"], "draftHash": results["draftHash"], "results": results["results"], "approved": results["approved"], "findings": results["findings"], "antiFraud": results["antiFraud"]})
    physical = inspect_physical_document(docx_path=docx, pdf_path=pdf, f8_path=f8, layout_profile_id=profile["layout"], expected_docx_hash=hashes["m6:docx"], expected_pdf_hash=hashes["m6:pdf"])
    final_semantic = _docx_semantic_text(docx)
    if key == "health":
        fidelity = compare_fidelity(stable_source, docx, pdf)
        semantic_match = fidelity["approved"]
        semantic_evidence = f"markdown={fidelity['markdown']['sha256']} docxCoverage={fidelity['blocks']['docxCoverage']} pdfCoverage={fidelity['blocks']['pdfCoverage']}"
        semantic_data = {"mode": "markdown_docx_pdf", "sourcePath": str(stable_source), "docxPath": str(docx), "pdfPath": str(pdf)}
    else:
        source_semantic = _docx_semantic_text(stable_source)
        semantic_match = canonical_hash(source_semantic) == canonical_hash(final_semantic)
        semantic_evidence = f"sourceSemantic={canonical_hash(source_semantic)} finalSemantic={canonical_hash(final_semantic)}"
        semantic_data = {"mode": "docx_semantic_hash", "sourcePath": str(stable_source), "docxPath": str(docx)}
    test_pass = results["approved"] and all(item.get("status") == "pass" for item in results["results"])
    mutation_pass = float((results.get("antiFraud") or {}).get("mutationScore") or 0) >= 0.8
    cross_refs_pass = all(question.get("supportIds") for question in questions) and len({question["questionId"] for question in questions}) == len(questions)
    measured_at = now_iso()
    layer_evidence = {
        "C1": {"measuredAt": measured_at, "checks": [{"name": "registered_sources_have_current_hashes", "passed": all(path.is_file() and sha256_file(path) == hashes[name] for name, path in sources.items()), "evidence": f"{len(sources)} fontes com caminho e SHA-256 recalculado", "evidenceData": {"sources": {name: {"path": str(path), "sha256": hashes[name]} for name, path in sources.items()}}}]},
        "C2": {"measuredAt": measured_at, "checks": [{"name": "source_to_final_semantic_fidelity", "passed": semantic_match, "evidence": semantic_evidence, "evidenceData": semantic_data}]},
        "C3": {"measuredAt": measured_at, "checks": [{"name": "case_tests", "passed": test_pass, "evidence": f"{sum(item.get('status') == 'pass' for item in results['results'])}/{len(results['results'])}", "evidenceData": {"suiteHash": results["suiteHash"], "draftHash": results["draftHash"], "passed": sum(item.get("status") == "pass" for item in results["results"]), "total": len(results["results"])}}, {"name": "mutation_testing", "passed": mutation_pass, "evidence": f"score={(results.get('antiFraud') or {}).get('mutationScore')} killed={(results.get('antiFraud') or {}).get('killed')}/{(results.get('antiFraud') or {}).get('total')}", "evidenceData": dict(results.get("antiFraud") or {})}]},
        "C4": {"measuredAt": measured_at, "checks": [{"name": "question_support_integrity", "passed": cross_refs_pass, "evidence": f"{len(questions)} questões únicas com supportIds", "evidenceData": {"questions": [{"questionId": item["questionId"], "supportIds": item.get("supportIds") or []} for item in questions]}}]},
        "C5": {"measuredAt": measured_at, "checks": [{"name": "physical_and_visual_integrity", "passed": physical["approved"], "evidence": f"layout={profile['layout']} pages={qa.get('pageCount')} independentReviewer={qa.get('reviewerRunId')}", "evidenceData": {"f8Path": str(f8), "f8Sha256": sha256_file(f8), "docxPath": str(docx), "docxSha256": sha256_file(docx), "pdfPath": str(pdf), "pdfSha256": sha256_file(pdf), "pageCount": qa.get("pageCount")}}]},
    }
    layers = {layer: "pass" if all(check["passed"] for check in data["checks"]) else "fail" for layer, data in layer_evidence.items()}
    save("F7_GLOBAL_CONSISTENCY.json", {"measurementContract": "N4-MEASURED-v1", "layers": layers, "layerEvidence": layer_evidence, "findings": physical["findings"], "physicalIntegrity": physical, "approved": all(value == "pass" for value in layers.values())})
    save("F7_METACOGNITIVE_AUDIT.json", {"premises": [{"premiseId": "PREM-001", "statement": profile["thesis"], "originType": "final_document_and_audits", "confirmedBySourceIds": ["DOC-FINAL"], "status": "confirmed", "usedInDraft": True}], "consensusChecks": [{"issueId": "META-001", "agentsAgreeing": 0, "independentSourceCount": 1, "verdict": "council_review_required"}], "recommendationChanges": [{"source": "CICERO", "status": "pending_human_resolution", "reason": "o parecer não libera a versão corrente"}], "metricChecks": [], "bestObjection": profile["objection"], "alternativeExplanation": "O produto final sintetiza premissas que ainda exigem fonte primária e resolução das objeções do conselho; não existe consenso independente de aprovação."})
    if not profile["science"]:
        na("F7_SCIENCE_AUDIT.json", "Dimensão científica classificada como não aplicável.")
    na("F9_DELIVERY_SELECTION.json", "O canário M6 é interno e não seleciona anexo para envio.")
    na("F10_DELIVERY_INTEGRITY.json", "Nenhuma entrega externa foi realizada neste canário.")
    na("F10_HUMAN_DIFF_CLASSIFICATION.json", "Não existe versão humana posterior ao canário para classificar.")

    result = validate_case(case_dir, target_phase="F10_ENTREGA_EVIDENCIA_APRENDIZADO")
    manifest = read_json(case_dir / "FORJA_CASE_MANIFEST.json", {}) or {}
    manifest["n4M6Cycle"]["status"] = "validated_baseline" if result["approved"] else "blocked"
    manifest["n4M6Cycle"]["targetPhase"] = "F10_ENTREGA_EVIDENCIA_APRENDIZADO"
    manifest["n4M6Cycle"]["promotionEligible"] = result.get("promotionEligible", False)
    manifest["n4M6Cycle"]["validationHash"] = result["validationHash"]
    atomic_write_json(case_dir / "FORJA_CASE_MANIFEST.json", manifest)
    atomic_write_json(FORJA / "reports" / f"M6_{key.upper()}_ANTIFRAUD_RESULT.json", result)
    return result


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("case", choices=sorted(CASES))
    args = parser.parse_args()
    print(json.dumps(run(args.case), ensure_ascii=False, indent=2))
