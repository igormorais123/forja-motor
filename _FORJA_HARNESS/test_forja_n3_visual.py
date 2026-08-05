from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from forja_n3_common import FORJA, ForjaN3Error, atomic_write_json, now_iso, read_json
from forja_visual_review import REQUIRED_PAGE_CHECKS
from forja_visual import _larguras_tabela, compor
from forja_visual_qa import inspect_pdf, lint_text, run_visual_qa

sys.path.insert(0, str(FORJA.parent / "_FERRAMENTAS"))
from medina_visual_lint import lint_svg  # noqa: E402
from medina_svg_colisao import analisar as analisar_svg_colisao  # noqa: E402


class ForjaN3VisualTests(unittest.TestCase):
    def test_valid_svg_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "valid.svg"
            path.write_text(
                '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 100">'
                '<rect id="box" x="10" y="10" width="180" height="50" fill="#eff4f3"/>'
                '<text x="100" y="40" font-size="10" font-weight="bold" font-style="normal" text-anchor="middle">Texto válido</text>'
                '</svg>', encoding="utf-8"
            )
            result = lint_svg(path)
            self.assertTrue(result["approved"], result["findings"])

    def test_later_opaque_shape_covering_text_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "covered.svg"
            path.write_text(
                '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 100">'
                '<text x="20" y="40" font-size="10" font-weight="normal" font-style="normal" text-anchor="start">Texto coberto</text>'
                '<rect x="15" y="25" width="100" height="30" fill="#000000"/>'
                '</svg>', encoding="utf-8"
            )
            result = lint_svg(path)
            self.assertTrue(any(item["code"] == "later_shape_covers_text" for item in result["findings"]))

    def test_production_diagrams_are_checked_after_correction(self) -> None:
        roots = FORJA / "state"
        paths = [
            roots / "case-email-natura-cabreuva-19f3991ebc75fe03" / "producao" / "_visual" / "fig2_escada_estrategias.svg",
            roots / "case-email-libra-sul-agint-stj-19f3c9350d875062" / "producao" / "_visual" / "fig2_obices_convergentes.svg",
            roots / "case-email-patricia-fabio-memoriais-19f3c68ee6d8fef2" / "producao" / "_visual" / "fig2_metodo_bifasico.svg",
            roots / "case-email-corsan-agerst-19f3dc9ff92081cd" / "producao" / "_visual" / "fig1_vulnerabilidades.svg",
            roots / "case-email-azimut-19f3ed5bdbdcf159" / "producao" / "_visual" / "fig2_selic_vs_juros.svg",
        ]
        existing = [path for path in paths if path.exists()]
        if not existing:
            self.skipTest("corpus jurídico real excluído por política do espelho Hermes")
        self.assertEqual(
            len(paths),
            len(existing),
            "corpus visual real parcialmente ausente; ausência parcial não pode virar falso-verde",
        )
        # Quatro âncoras históricas foram corrigidas no acervo; o Azimut ainda
        # conserva a mutação real de atributos SVG inválidos para provar que o
        # linter continua bloqueante. O teste não congela o acervo numa
        # fotografia em que todo defeito histórico permanece para sempre.
        expected_bad = {"fig2_selic_vs_juros.svg"}
        for path in paths:
            with self.subTest(path=path.name):
                lint = lint_svg(path)
                if path.name in expected_bad:
                    self.assertFalse(lint["approved"], lint["findings"])
                    continue
                self.assertTrue(lint["approved"], lint["findings"])
                collision = analisar_svg_colisao(path)
                self.assertTrue(collision["aprovado"], collision["achados"])

    def test_markdown_and_duplicate_caption_leaks_fail(self) -> None:
        findings = lint_text("46. #### 4.1.1 Cronologia\nFIGURA 1 | FIGURA 2 — teste\n33. > citação")
        codes = {item["code"] for item in findings}
        self.assertIn("markdown_heading", codes)
        self.assertIn("markdown_blockquote", codes)
        self.assertIn("multiple_figure_numbers_in_caption", codes)

    def test_pdf_reviewer_must_be_independent(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            import fitz

            pdf = Path(temp) / "test.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Página de teste com texto suficiente.")
            document.save(pdf)
            document.close()
            with self.assertRaises(ForjaN3Error):
                inspect_pdf(pdf, Path(temp) / "qa", generator_run_id="same", reviewer_run_id="same")

    def test_visual_parser_handles_h4_blockquote_and_weighted_table(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            from docx import Document

            temp = Path(temp)
            markdown = temp / "source.md"
            output = temp / "output.docx"
            markdown.write_text(
                "# MEMORIAL DE TESTE\n\n"
                "## MÉRITO\n\n"
                "#### 4.1 Subtítulo preservado\n\n"
                "> Esta citação possui conteúdo suficientemente longo para ser preservado integralmente sem marcador de blockquote.\n\n"
                "| Código | Descrição extensa |\n|---|---|\n| A | Conteúdo muito mais longo que a primeira coluna |\n",
                encoding="utf-8",
            )
            compor(markdown, output, {})
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("4.1 Subtítulo preservado", text)
            self.assertNotIn("####", text)
            self.assertNotIn("> Esta", text)
            widths = _larguras_tabela(["Código", "Descrição extensa"], [["A", "Conteúdo muito mais longo que a primeira coluna"]])
            self.assertGreater(widths[1], widths[0])
            self.assertAlmostEqual(13.1, sum(widths), places=2)

    def test_visual_qa_runs_semantic_fidelity_gate(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            import fitz
            from docx import Document
            from docx.shared import Pt

            root = Path(temp)
            sentence = (
                "A pretensão não pode ser acolhida, pois o valor comprovado é R$ 12.500,00 "
                "e apenas subsidiariamente se admite revisão em 18/07/2026. "
                "O fundamento jurídico principal permanece integralmente preservado no Word e no PDF, "
                "com redação suficiente para testar a fidelidade material e a diagramação do corpo da peça."
            )
            markdown = root / "source.md"
            docx = root / "source.docx"
            pdf = root / "source.pdf"
            ledger_path = root / "VISUAL_QA.json"
            fidelity_path = root / "FORMAT_FIDELITY.json"
            markdown.write_text("# TESTE\n\n" + sentence + "\n", encoding="utf-8")
            document = Document()
            document.add_heading("TESTE", level=1)
            paragraph = document.add_paragraph()
            paragraph.alignment = 3
            run = paragraph.add_run(sentence)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            document.save(docx)
            pdf_document = fitz.open()
            page = pdf_document.new_page()
            page.insert_textbox(fitz.Rect(72, 72, 520, 300), sentence, fontsize=11)
            pdf_document.save(pdf)
            pdf_document.close()

            pending_path = root / "VISUAL_REVIEW_PENDING.json"
            pending_result = run_visual_qa(
                pdf,
                ledger_path,
                qa_dir=root / "pages",
                generator_run_id="generator-1",
                reviewer_run_id="reviewer-1",
                docx=docx,
                markdown=markdown,
                fidelity_output=fidelity_path,
                pending_review_output=pending_path,
            )
            self.assertFalse(pending_result["approved"])
            self.assertTrue(any(item["code"] == "manual_visual_review_missing" for item in pending_result["findings"]))

            review = read_json(pending_path)
            review["reviewedAt"] = now_iso()
            review["reviewer"] = {"id": "qa-visual-test", "runId": "reviewer-1", "type": "agent_visual"}
            review["reviewMethod"] = "page_by_page_at_100_percent"
            review["approved"] = True
            for page_review in review["pages"]:
                page_review["status"] = "pass"
                page_review["checks"] = {name: True for name in REQUIRED_PAGE_CHECKS}
            atomic_write_json(pending_path, review)

            result = run_visual_qa(
                pdf,
                ledger_path,
                qa_dir=root / "pages",
                generator_run_id="generator-1",
                reviewer_run_id="reviewer-1",
                docx=docx,
                markdown=markdown,
                fidelity_output=fidelity_path,
                manual_review=pending_path,
            )
            self.assertTrue(result["approved"], result["findings"])
            self.assertTrue(result["fidelity"]["approved"])
            self.assertTrue(fidelity_path.is_file())

    def test_composer_accepts_only_declared_visual_markers(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            from docx import Document

            root = Path(temp)
            markdown = root / "source.md"
            output = root / "output.docx"
            markdown.write_text(
                "# MEMORIAL\n\n"
                "EXCELENTÍSSIMOS SENHORES MINISTROS DA TERCEIRA TURMA DO SUPERIOR TRIBUNAL DE JUSTIÇA\n\n"
                "Este parágrafo material possui conteúdo suficiente e receberá uma figura declarada no mapa visual.\n",
                encoding="utf-8",
            )
            compor(
                markdown,
                output,
                {
                    "cards_apos_titulo": True,
                    "figs": [("Este parágrafo material", "{{FIG1}}", "Legenda da figura")],
                },
            )
            text = "\n".join(p.text for p in Document(output).paragraphs)
            self.assertEqual(1, text.count("{{CARDS}}"))
            self.assertEqual(1, text.count("{{FIG1}}"))

    def test_conclusive_summary_after_merit_remains_a_regular_section(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            from docx import Document

            root = Path(temp)
            markdown = root / "source.md"
            output = root / "output.docx"
            markdown.write_text(
                "# MEMORIAL\n\n"
                "EXCELENTÍSSIMOS SENHORES MINISTROS DA TERCEIRA TURMA DO SUPERIOR TRIBUNAL DE JUSTIÇA\n\n"
                "## SÍNTESE EXECUTIVA\n\nSíntese inicial suficientemente clara para orientar o julgamento.\n\n"
                "## I — MÉRITO\n\n1. Fundamento material suficientemente desenvolvido para integrar o documento.\n\n"
                "## II — SÍNTESE CONCLUSIVA E PEDIDOS\n\n2. Pedido conclusivo suficientemente desenvolvido para integrar o documento.\n",
                encoding="utf-8",
            )
            compor(markdown, output, {})
            text = "\n".join(p.text for p in Document(output).paragraphs)
            self.assertIn("II — SÍNTESE CONCLUSIVA E PEDIDOS", text)

    def test_closing_and_signature_are_not_numbered_or_orphaned(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            from docx import Document

            root = Path(temp)
            markdown = root / "source.md"
            output = root / "output.docx"
            markdown.write_text(
                "# IMPUGNAÇÃO AO AGRAVO INTERNO\n\n"
                "EXCELENTÍSSIMA SENHORA MINISTRA RELATORA DO SUPERIOR TRIBUNAL DE JUSTIÇA\n\n"
                "**Agravo em Recurso Especial nº 2.698.443/DF**\n\n"
                "## I — PEDIDOS\n\n"
                "1. Requer-se o desprovimento do recurso pelos fundamentos expostos.\n\n"
                "Nestes termos, pede deferimento.\n\n"
                "Brasília/DF, 10 de julho de 2026.\n\n"
                "**FÁBIO MEDINA OSÓRIO**  \n"
                "OAB/RS 64.975 · OAB/DF 29.786\n",
                encoding="utf-8",
            )
            compor(markdown, output, {})
            doc = Document(output)
            paragraphs = "\n".join(p.text for p in doc.paragraphs)
            tables = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
            self.assertIn("Nestes termos, pede deferimento.", paragraphs)
            self.assertNotIn("2. Nestes termos, pede deferimento.", paragraphs)
            self.assertIn("FÁBIO MEDINA OSÓRIO", tables)
            self.assertIn("OAB/RS 64.975 · OAB/DF 29.786", tables)


if __name__ == "__main__":
    unittest.main(verbosity=2)
