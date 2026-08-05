# -*- coding: utf-8 -*-
"""Extrai o DNA de formatação Word das peças do escritório Medina Osório:
margens, fontes, tamanhos, espaçamento, recuo, cabeçalho/rodapé, numeração."""
import os, sys, json
from collections import Counter
from docx import Document
from docx.shared import Cm

RAIZ = r"C:\Users\IgorPC\.claude\projects\Escritório fabio osório\fabricas de melhoria de petições"

CORPUS = {
    "ORIGINAL escritório — Cafelana CR EDcl (anexo e-mail)":
        r"Cafelana\Anexos do email\CAFELANA_CR_EDCL_1-07-2026 -.docx",
    "ORIGINAL escritório — Memoriais Patrícia e Fábio":
        r"Memoriais Apelação Patrícia e Fábio - Proc. 0014560-09.2014.8.19.0209\Anexos do email\MEMORIAS - PATRICIA E FABIO - VERSÃO FINAL.docx",
    "ORIGINAL escritório — EDcl José Eduardo (anexo e-mail)":
        r"Minuta de Embargos de Declaração — José Eduardo Siqueira Campos\Anexos do email\EMBARGOS DE DECLARAÇÃO - DECISÃO EV 185 - versão final ajuste alertas.docx",
    "ORIGINAL escritório — Quesitos Cabreúva":
        r"Natura Cabreúva - Parecer e Quesitos - prazo 20-07-2026\Anexos do email\Cabreúva - Quesitos Dr. Fábio Osório.docx",
    "FÁBRICA — Cafelana CR EDcl FINAL":
        r"gestao_escritorio\entregas_fabio_osorio\2026-07-06 Re Contrarrazões Cafelana 19f39731\CAFELANA_CR_EDCL_FINAL_02-07-2026.docx",
    "FÁBRICA — Cafelana AgInt AREsp NÍVEL2":
        r"Cafelana\contrarrazões ao AgInt no AREsp nº 2.698.443D\CAFELANA_IMPUGNACAO_AGINT_ARESP_2698443_NIVEL2_06-07-2026.docx",
    "FÁBRICA — Jalusa NÍVEL 4":
        r"Jalusa Prestes Abaide - Proc. 5000447-02.2011.4.04.7102\PETICAO_FINAL_NIVEL_4_JALUSA_EVENTO_183.docx",
    "FÁBRICA — EDcl José Eduardo SUPER AJUSTADA":
        r"Minuta de Embargos de Declaração — José Eduardo Siqueira Campos\EMBARGOS DE DECLARAÇÃO - SUPER VERSÃO FINAL AJUSTADA.docx",
    "FÁBRICA — Memoriais LIBRA SUL":
        r"Memoriais Cautelar Fiscal\MEMORIAIS_LIBRA_SUL.docx",
}

def cm(v):
    return round(v.cm, 2) if v is not None else None

def pt(v):
    return round(v.pt, 1) if v is not None else None

def analisa(path):
    doc = Document(path)
    sec = doc.sections[0]
    info = {
        "página": f"{cm(sec.page_width)}x{cm(sec.page_height)}cm",
        "margens (sup/inf/esq/dir)": [cm(sec.top_margin), cm(sec.bottom_margin),
                                       cm(sec.left_margin), cm(sec.right_margin)],
        "margem cabeçalho/rodapé": [cm(sec.header_distance), cm(sec.footer_distance)],
    }
    st = doc.styles["Normal"]
    pf = st.paragraph_format
    info["estilo Normal"] = {
        "fonte": st.font.name, "tamanho_pt": pt(st.font.size),
        "entrelinhas": pf.line_spacing,
        "espaço antes/depois_pt": [pt(pf.space_before), pt(pf.space_after)],
        "recuo 1ª linha_cm": cm(pf.first_line_indent),
    }
    fontes, tamanhos, alinhas, entrelinhas = Counter(), Counter(), Counter(), Counter()
    recuos = Counter()
    n_par = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        n_par += 1
        alinhas[str(p.alignment)] += 1
        ls = p.paragraph_format.line_spacing
        if ls is not None:
            entrelinhas[str(ls)] += 1
        fli = p.paragraph_format.first_line_indent
        if fli is not None:
            recuos[cm(fli)] += 1
        for r in p.runs:
            if r.font.name:
                fontes[r.font.name] += 1
            if r.font.size:
                tamanhos[pt(r.font.size)] += 1
    info["parágrafos com texto"] = n_par
    info["fontes nos runs (top)"] = fontes.most_common(4)
    info["tamanhos nos runs (top)"] = tamanhos.most_common(5)
    info["alinhamentos (top)"] = alinhas.most_common(3)
    info["entrelinhas explícitas"] = entrelinhas.most_common(3)
    info["recuos 1ª linha explícitos"] = recuos.most_common(3)
    # cabeçalho e rodapé
    hdr = sec.header
    ftr = sec.footer
    def resumo_parte(parte):
        txt = " | ".join(p.text.strip() for p in parte.paragraphs if p.text.strip())[:200]
        n_img = str(parte.part.blob).count("Relationship") if False else None
        xml = parte.paragraphs[0].part.element.xml if parte.paragraphs else ""
        return {"texto": txt, "tem_imagem": "blip" in (parte.part.element.xml if hasattr(parte.part, "element") else "")}
    try:
        info["cabeçalho"] = resumo_parte(hdr)
        info["rodapé"] = resumo_parte(ftr)
    except Exception as e:
        info["cabeçalho/rodapé_erro"] = str(e)[:80]
    # primeiros 3 parágrafos (endereçamento)
    info["abertura"] = [p.text.strip()[:110] for p in doc.paragraphs if p.text.strip()][:3]
    return info

saida = {}
for nome, rel in CORPUS.items():
    path = os.path.join(RAIZ, rel)
    if not os.path.exists(path):
        saida[nome] = {"ERRO": "arquivo não encontrado"}
        continue
    try:
        saida[nome] = analisa(path)
    except Exception as e:
        saida[nome] = {"ERRO": str(e)[:200]}

print(json.dumps(saida, ensure_ascii=False, indent=1))
with open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "padrao_word_extraido.json"),
          "w", encoding="utf-8") as f:
    json.dump(saida, f, ensure_ascii=False, indent=1)
