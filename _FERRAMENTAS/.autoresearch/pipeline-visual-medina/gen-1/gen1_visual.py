# -*- coding: utf-8 -*-
"""Geração 1 — variantes visuais de um diagrama real (linha do tempo processual)
avaliadas por métrica objetiva + inspeção visual, e teste E2E Word completo:
SVG -> EMF (gate) -> docx com marcadores -> Word COM -> PDF -> PNGs de QA."""
import os, sys, json, re

BASE = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".."))
sys.path.insert(0, BASE)
import estilo_medina as em
import word_visual_pipeline as wvp

OUT = os.path.dirname(os.path.abspath(__file__))
C = em.CORES
F = em.FONTE_DIAGRAMA

EVENTOS = [("1997", "Ajuizamento da execução"),
           ("2011", "Ação rescisória"),
           ("2024", "Acórdão embargado"),
           ("2026", "2º embargos de declaração")]

def timeline(nome, fs_rotulo, fs_ano, destaque_ultimo=False, ancoras=False,
             grossura_linha=2):
    """Gera uma variante de linha do tempo no estilo Medina."""
    w, h = 600, 150 if not ancoras else 190
    xs = [70, 230, 390, 540]
    y_linha = 78
    s = [f'<rect x="0" y="0" width="{w}" height="{h}" fill="white"/>']
    s.append(f'<line x1="30" y1="{y_linha}" x2="580" y2="{y_linha}" '
             f'stroke="{C["petroleo"]}" stroke-width="{grossura_linha}"/>')
    for i, (x, (ano, txt)) in enumerate(zip(xs, EVENTOS)):
        ultimo = (i == len(EVENTOS) - 1)
        cor_no = C["terracota_escura"] if (destaque_ultimo and ultimo) else C["petroleo"]
        raio = 7 if (destaque_ultimo and ultimo) else 5
        s.append(f'<circle cx="{x}" cy="{y_linha}" r="{raio}" fill="{cor_no}"/>')
        peso_ano = "bold"
        s.append(f'<text x="{x}" y="{y_linha-16}" font-family="{F}" font-size="{fs_ano}" '
                 f'font-weight="{peso_ano}" fill="{cor_no}" text-anchor="middle">{ano}</text>')
        palavras = txt.split()
        meio = (len(palavras) + 1) // 2
        linhas = [" ".join(palavras[:meio]), " ".join(palavras[meio:])]
        for j, ln in enumerate(l for l in linhas if l):
            s.append(f'<text x="{x}" y="{y_linha+22+j*(fs_rotulo+2)}" font-family="{F}" '
                     f'font-size="{fs_rotulo}" fill="{C["grafite"]}" text-anchor="middle">{ln}</text>')
    if ancoras:
        s.append(f'<rect x="30" y="{h-38}" width="550" height="30" fill="{C["painel_verde"]}" rx="4"/>')
        s.append(f'<text x="305" y="{h-18}" font-family="{F}" font-size="12.5" '
                 f'fill="{C["petroleo_escuro"]}" text-anchor="middle" font-weight="bold">'
                 f'29 anos de processo · 0 vícios apontados · 2º embargos sucessivos</text>')
    svg = (f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {w} {h}">' +
           "".join(s) + "</svg>")
    path = os.path.join(OUT, nome)
    with open(path, "w", encoding="utf-8") as f:
        f.write(svg)
    return path

PALETA_OK = {v.upper() for v in C.values()} | {"WHITE", "#FFFFFF", "NONE"}

def pontuar(svg_path):
    """Métrica objetiva: legibilidade (gate) + fidelidade de paleta."""
    viol = em.checar_fontes_svg(svg_path, 15)
    with open(svg_path, encoding="utf-8") as f:
        svg = f.read()
    cores_usadas = {c.upper() for c in re.findall(r'(?:fill|stroke)="(#[0-9A-Fa-f]{6})"', svg)}
    fora = cores_usadas - PALETA_OK
    fontes = {m for m in re.findall(r'font-family="([^"]+)"', svg)}
    fontes_fora = fontes - {F, em.FONTE_TEXTO, em.FONTE_DADOS}
    score = 10.0 - 4*len(viol) - 2*len(fora) - 2*len(fontes_fora)
    return {"svg": os.path.basename(svg_path), "viol_legibilidade": viol,
            "cores_fora_da_paleta": sorted(fora), "fontes_fora": sorted(fontes_fora),
            "score_objetivo": score}

variantes = {
    "variant-a": timeline("variant-a.svg", fs_rotulo=12, fs_ano=13),
    "variant-b": timeline("variant-b.svg", fs_rotulo=13, fs_ano=15, grossura_linha=2.5),
    "variant-c": timeline("variant-c.svg", fs_rotulo=13, fs_ano=15, destaque_ultimo=True),
    "variant-d": timeline("variant-d.svg", fs_rotulo=13, fs_ano=15, destaque_ultimo=True,
                          ancoras=True),
}

scores = {k: pontuar(v) for k, v in variantes.items()}
print(json.dumps(scores, ensure_ascii=False, indent=1))
with open(os.path.join(OUT, "scores.json"), "w", encoding="utf-8") as f:
    json.dump(scores, f, ensure_ascii=False, indent=1)

# ---------- E2E Word: as 4 variantes + Graphviz + matplotlib no mesmo docx ----------
import matplotlib
matplotlib.use("Agg")
em.aplicar_estilo_matplotlib()
import matplotlib.pyplot as plt
fig, ax = plt.subplots(figsize=(6, 2.6))
ax.bar(["1997", "2011", "2024", "2026"], [1, 3, 2, 4], width=0.5)
ax.set_title("Atos processuais por marco (exemplo)")
mpl_svg = os.path.join(OUT, "e2e_matplotlib.svg")
fig.savefig(mpl_svg); plt.close(fig)

dot = os.path.join(OUT, "e2e_fluxo.dot")
with open(dot, "w", encoding="utf-8") as f:
    f.write('digraph G { rankdir=LR; ' + em.ESTILO_GRAPHVIZ +
            ' a[label="Acórdão"]; b[label="1º EDcl\\n(rejeitados)"]; '
            'c[label="2º EDcl\\n(pendentes)"]; a->b->c; }')
gv_svg = os.path.join(OUT, "e2e_fluxo.svg")
wvp.dot_para_svg(dot, gv_svg)

figuras = list(variantes.values()) + [gv_svg, mpl_svg]
emfs = {}
problemas = []
for i, svg in enumerate(figuras, 1):
    emf = svg.replace(".svg", ".emf")
    try:
        wvp.svg_para_emf(svg, emf, largura_final_cm=15)
        emfs[f"{{{{FIG{i}}}}}"] = emf
    except ValueError as e:
        problemas.append(f"FIG{i} reprovada no gate: {str(e)[:160]}")
print("\nEMFs aprovados:", len(emfs), "| reprovados:", len(problemas))
for p in problemas:
    print(" ", p)

# docx com timbre simplificado Medina
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("TESTE E2E — PIPELINE VISUAL MEDINA OSÓRIO (geração 1)")
r.bold = True; r.font.color.rgb = RGBColor(0x39, 0x5C, 0x60)
legendas = ["Variante A — linha do tempo base (12px)",
            "Variante B — contraste reforçado (13/15px)",
            "Variante C — destaque Von Restorff no evento decisivo",
            "Variante D — destaque + faixa de âncoras numéricas",
            "Fluxo Graphviz com ESTILO_GRAPHVIZ",
            "Gráfico matplotlib com aplicar_estilo_matplotlib()"]
for i, leg in enumerate(legendas, 1):
    if f"{{{{FIG{i}}}}}" not in emfs:
        continue
    doc.add_paragraph(f"{{{{FIG{i}}}}}")
    pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = pl.add_run(f"Figura {i} — {leg}")
    rl.italic = True; rl.font.size = Pt(9)
    rl.font.color.rgb = RGBColor(0x49, 0x49, 0x4D)
docx_path = os.path.join(OUT, "e2e_teste.docx")
doc.save(docx_path)

wvp.inserir_emf_word_com(docx_path, emfs, largura_cm=15)
pdf_path = os.path.join(OUT, "e2e_teste.pdf")
wvp.docx_para_pdf(docx_path, pdf_path)
paginas = wvp.render_paginas(pdf_path, os.path.join(OUT, "qa_paginas"), dpi=130)
print("\nPDF gerado. Páginas para QA visual:")
for p in paginas:
    print(" ", p)
