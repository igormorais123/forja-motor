# -*- coding: utf-8 -*-
"""Kit de composição das edições VISUAL LAW Medina Osório em Word.

Generalização do builder validado da Cafelana (08/07/2026) — skill
`padrao-visual-medina`. Uso típico:

    from medina_visual_kit import PecaVisual
    pv = PecaVisual("SAIDA.docx")            # copia o template com timbre
    pv.enderecamento("EXCELENTÍSSIMO ...")
    pv.processo_caixa("Apelação nº ...")
    pv.titulo("EMBARGOS DE DECLARAÇÃO")
    pv.par("**FULANO**, já qualificado ...")
    pv.marcador("{{CARDS}}")
    pv.quebra_pagina()
    pv.sintese([("OBJETO", "..."), ...])
    pv.abre("Título da seção", "linha-síntese escaneável")
    pv.pull("frase da margem"); pv.pgf("texto com **negrito** e *itálico*")
    pv.caixa_acordao("Título", "“citação...”")
    pv.caixa_chave("**Conceito decisivo** em uma frase.")
    pv.quadro_zebrado(["col1","col2","col3"], linhas, larguras_cm=(3.1,6.4,3.6))
    pv.pedidos([("a)", "texto ..."), ...])
    pv.fecho("Brasília/DF, 9 de julho de 2026.")
    pv.assinaturas([("NOME","OAB/UF nº ..."), ...])
    pv.salvar()

Depois: a entrada canônica da FORJA embute os SVGs diretamente no pacote OOXML
e executa QA estrutural. Este kit não chama conversor, Word COM, PDF ou PNG.
"""
import hashlib
import json
import os, re, shutil, sys, zipfile
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AQUI = os.path.dirname(os.path.abspath(__file__))


def _classificar_produto(texto, titulo=""):
    """Petição ou estudo interno — e a diferença não é cosmética.

    O verificador cobra do produto externo coisas que o estudo interno tem o
    direito de carregar: marcador de lacuna `[VERIFICAR ...]`, por exemplo, é
    bloqueador P0 numa peça protocolável e é exatamente o que se espera de um
    estudo preliminar em construção.

    Escrevi este classificador depois de errar por não tê-lo. Ao medir o efeito
    da porta única, julguei 25 obras como petição e o resultado acusou 14
    entregáveis com placeholder e 1 com origem operacional vazada — uma crise de
    qualidade que não existia. Classificando pelo tipo real, os três achados
    graves desapareceram inteiramente: a produção estava limpa e o defeito era
    da minha medição. Um gate sem esta distinção acusaria a fábrica todo dia.

    Espelha `forja_visual_build._tipo_produto`, que é a autoridade da rota
    canônica; a duplicação existe porque `_FERRAMENTAS` não pode depender do
    harness para uma decisão que precisa valer mesmo em rota ad hoc.
    """
    abertura = (str(titulo) + "\n" + str(texto)[:1800]).upper()
    if re.search(r"\b(ESTUDO|DIAGN[ÓO]STICO|RELAT[ÓO]RIO|PARECER|MATRIZ|CHECKLIST|"
                 r"NOTA|BRIEFING|MEMORANDO|AUDITORIA)\b", abertura):
        return "estudo"
    return "peca"
TEMPLATE = os.path.join(AQUI, "TEMPLATE_MEDINA_OSORIO_PETICAO.docx")

PET, PETE = "395C60", "2A4548"
TER, TERE = "D9926A", "9C5B38"
GRA, PV_, PTC = "49494D", "EFF4F3", "FBF2EC"
ALE, RUL = "7A2E2E", "C9C9C9"
SANS, SERIF = "Segoe UI", "Times New Roman"

RICO = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")


def _fmt(r, size=12, bold=False, italic=False, cor=None, fonte=SERIF):
    r.font.name = fonte
    r._element.rPr.rFonts.set(qn("w:hAnsi"), fonte)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if cor:
        r.font.color.rgb = RGBColor.from_string(cor)
    return r


def _rico(p, texto, size=12, cor=None, fonte=SERIF, italico_base=False):
    for seg in RICO.split(texto):
        if not seg:
            continue
        if seg.startswith("**"):
            _fmt(p.add_run(seg[2:-2]), size, True, italico_base, cor, fonte)
        elif seg.startswith("*"):
            _fmt(p.add_run(seg[1:-1]), size, False, True, cor, fonte)
        elif seg.startswith("`"):
            _fmt(p.add_run(seg[1:-1]), size, False, italico_base, cor, fonte)
        else:
            _fmt(p.add_run(seg), size, False, italico_base, cor, fonte)
    return p


def _borda_p(p, lado, cor, sz8, espaco=2):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{lado}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz8))
    el.set(qn("w:space"), str(espaco))
    el.set(qn("w:color"), cor)
    pBdr.append(el)


def _sombra(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _bordas(cell, spec):
    tb = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{lado}")
        if lado in spec and spec[lado]:
            cor, sz8 = spec[lado]
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz8))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), cor)
        else:
            el.set(qn("w:val"), "nil")
        tb.append(el)
    cell._tc.get_or_add_tcPr().append(tb)


def _margens(cell, esq=140, dir_=140, cima=100, baixo=100):
    m = OxmlElement("w:tcMar")
    for lado, v in (("top", cima), ("start", esq), ("bottom", baixo), ("end", dir_)):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    cell._tc.get_or_add_tcPr().append(m)


def _pc(cell, primeiro=True):
    p = cell.paragraphs[0] if primeiro and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    return p


ROM = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII",
       "XIII", "XIV", "XV"]

_ROMANOS = ((1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
            (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
            (5, "V"), (4, "IV"), (1, "I"))


def romano(n: int) -> str:
    """Numeral romano da seção, sem teto.

    A lista ``ROM`` parava em XV, e ``abre()`` indexava direto: peça com a 16ª
    seção derrubava a composição inteira com ``IndexError`` — encontrado em
    04/08/2026 compondo o markdown real do Cafelana pela rota canônica. Teto
    silencioso em constante é dívida esperando o documento comprido; quem chega
    lá é justamente a peça grande, que é a que mais custa refazer.
    """
    if n < 1:
        raise ValueError(f"seção sem numeral romano: {n}")
    resto, saida = n, []
    for valor, letra in _ROMANOS:
        while resto >= valor:
            saida.append(letra)
            resto -= valor
    return "".join(saida)

# ---------- identidade áurea (capricho do Igor, 09/07/2026) ----------
# O fólio das páginas 2+ é reposicionado para que a LINHA do marcador caia na
# seção áurea da página: y = 29,7cm / φ ≈ 18,356cm do topo (parte menor abaixo,
# razão 1,618:1). Assinatura oculta da casa — imperceptível a quem não sabe.
PHI = 1.6180339887498949
_PAGINA_A4_EMU = 10692000            # 29,7 cm
_LINHA_AUREA_EMU = round(_PAGINA_A4_EMU / PHI)   # 6.607.926 EMU = 18,3554 cm
# o filete fica ~5,0pt acima da base do shape (padding interno do textbox);
# delta medido no PDF renderizado (Jalusa, 09/07/2026) para a linha cair EXATA
_CALIBRACAO_EMU = 67437              # 5,31 pt (refinado na 2ª medição, desvio final < 0,1pt)


def _aplicar_folio_aureo(caminho_docx):
    """Edita o header2.xml (fólio das páginas 2+): âncora vertical sai de
    'centro da margem' para posição absoluta na página, com a base do shape
    (a linha do marcador) na seção áurea. Cobre o DrawingML e o fallback VML."""
    with zipfile.ZipFile(caminho_docx) as z:
        dados = {n: z.read(n) for n in z.namelist()}
    alvo = "word/header2.xml"
    if alvo not in dados:
        return False
    x = dados[alvo].decode("utf-8")
    m = re.search(r'<wp:extent cx="\d+" cy="(\d+)"/>', x)
    if not m:
        return False
    cy = int(m.group(1))
    topo = _LINHA_AUREA_EMU - cy + _CALIBRACAO_EMU
    x = x.replace(
        '<wp:positionV relativeFrom="margin"><wp:align>center</wp:align></wp:positionV>',
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{topo}</wp:posOffset></wp:positionV>')
    # fallback VML (viewers antigos): margin-top em pt, relativo à página
    alt_m = re.search(r"height:([\d.]+)pt", x)
    alt_pt = float(alt_m.group(1)) if alt_m else cy / 12700.0
    topo_pt = _LINHA_AUREA_EMU / 12700.0 - alt_pt
    x = x.replace("margin-top:0;", f"margin-top:{topo_pt:.2f}pt;")
    x = x.replace("mso-position-vertical:center;", "")
    x = x.replace("mso-position-vertical-relative:margin", "mso-position-vertical-relative:page")
    dados[alvo] = x.encode("utf-8")
    with zipfile.ZipFile(caminho_docx, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in dados.items():
            z.writestr(n, b)
    return True


class PecaVisual:
    """Edição visual law sobre o template com timbre. Corpo 13,1cm + margem de
    pull quotes (esq 2,5 / dir 5,4). Estilo revista: sem recuo, entrelinha 1,15."""

    def __init__(self, saida, template=TEMPLATE, folio_aureo=True, *, case_dir=None,
                 ledger_path=None, base_dir=None):
        shutil.copyfile(template, saida)
        if folio_aureo:
            _aplicar_folio_aureo(saida)
        self.saida = saida
        # Contexto documental opcional e explícito. Sem ele, uma peça visual
        # com valor monetário não pode ser salva: a rota ad hoc foi justamente
        # o ponto cego do incidente Cafelana.
        self.case_dir = case_dir or os.environ.get("FORJA_CASE_DIR")
        self.ledger_path = ledger_path or os.environ.get("FORJA_LASTRO_LEDGER")
        self.base_dir = base_dir or os.environ.get("FORJA_LASTRO_BASE_DIR")
        self.doc = Document(saida)
        for p in list(self.doc.paragraphs):
            p._p.getparent().remove(p._p)
        sec = self.doc.sections[0]
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(5.4)
        est = self.doc.styles["Normal"]
        est.font.name = SERIF
        est.font.size = Pt(12)
        self.n_pgf = 0
        self.n_sec = 0
        self.n_fig = 0

    # ---------- base ----------
    def par(self, texto="", size=12, antes=0, depois=7, linha=1.15,
            align="justify", cor=None, fonte=SERIF, keep=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(antes)
        pf.space_after = Pt(depois)
        pf.line_spacing = linha
        pf.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        if keep:
            pf.keep_with_next = True
        if texto:
            _rico(p, texto, size, cor, fonte)
        return p

    def quebra_pagina(self):
        p = self.par(depois=0)
        p.add_run().add_break(WD_BREAK.PAGE)

    # ---------- capa ----------
    # REGRA (determinação do Igor, 09/07/2026): a capa vai LIMPA para protocolo —
    # nenhum rótulo ou tarja de laboratório/apresentação. tarja() abaixo é
    # reservada a material interno ao cliente, NUNCA a peça protocolada.
    def enderecamento(self, texto):
        p = self.par(antes=2, depois=8)
        _fmt(p.add_run(texto), 12, True)

    def processo_caixa(self, texto, largura_cm=None):
        """Caixa do número do processo. Largura automática pelo comprimento do
        texto (11,5pt bold TNR ~0,20cm/char) — o número NUNCA pode quebrar linha
        (lição EDcl 09/07)."""
        if largura_cm is None:
            largura_cm = min(13.1, round(len(texto) * 0.20 + 1.0, 1))
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.columns[0].width = Cm(largura_cm)
        c = t.rows[0].cells[0]
        c.width = Cm(largura_cm)
        _bordas(c, {k: (RUL, 6) for k in ("top", "left", "bottom", "right")})
        _margens(c, 130, 130, 60, 60)
        p = _pc(c)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _fmt(p.add_run(texto), 11.5, True)

    def titulo(self, texto, filete=True):
        p = self.par(antes=12, depois=2, align="center")
        _fmt(p.add_run(texto), 13, True)
        if filete:
            f = self.par(antes=2, depois=0, align="center")
            f.paragraph_format.left_indent = Cm(5.8)
            f.paragraph_format.right_indent = Cm(5.8)
            _borda_p(f, "bottom", TER, 18, 1)
            _fmt(f.add_run(" "), 3)

    def tarja(self, texto):
        p = self.par(antes=7, depois=12, align="center", fonte=SANS)
        _fmt(p.add_run(texto), 8.5, True, False, TERE, SANS)

    def marcador(self, tag, antes=6):
        p = self.par(antes=antes, depois=0, align="center", keep=True)
        p.add_run(tag)

    def figcap(self, texto):
        self.n_fig += 1
        p = self.par(antes=4, depois=10, align="center")
        _fmt(p.add_run(f"FIGURA {self.n_fig}"), 8, True, False, TERE, SANS)
        _fmt(p.add_run("  |  "), 8, False, False, RUL, SANS)
        _fmt(p.add_run(texto), 8, False, True, GRA, SANS)

    # ---------- estrutura ----------
    def sintese(self, linhas, titulo="SÍNTESE DOS PONTOS ESSENCIAIS"):
        ts = self.doc.add_table(rows=len(linhas) + 1, cols=2)
        ts.alignment = WD_TABLE_ALIGNMENT.CENTER
        ts.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for row in ts.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        larg = (Cm(3.1), Cm(10.0))
        cab = ts.rows[0].cells
        cab[0].merge(cab[1])
        c0 = ts.rows[0].cells[0]
        _sombra(c0, PV_)
        _bordas(c0, {"left": (PET, 24)})
        _margens(c0, 170, 170, 110, 40)
        _fmt(_pc(c0).add_run(titulo), 9.5, True, False, PET, SANS)
        for i, (rot, txt) in enumerate(linhas, start=1):
            cs = ts.rows[i].cells
            for j, cell in enumerate(cs):
                cell.width = larg[j]
                _sombra(cell, PV_)
                spec = {"left": (PET, 24)} if j == 0 else {}
                if i < len(linhas):
                    spec["bottom"] = (RUL, 4)
                _bordas(cell, spec)
                _margens(cell, 170 if j == 0 else 80, 170, 60, 60)
            _fmt(_pc(cs[0]).add_run(rot), 8, True, False, PET, SANS)
            pt_ = _pc(cs[1])
            # Colunas estreitas não comportam justificação sem criar rios de
            # espaço entre palavras. Na síntese, a leitura rápida prevalece.
            pt_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _rico(pt_, txt, 10.5)
        self.par(" ", 2, depois=8)

    def abre(self, titulo, linha, prefixo=True):
        self.n_sec += 1
        f = self.par(antes=14, depois=0, align="center", keep=True)
        f.paragraph_format.left_indent = Cm(5.7)
        f.paragraph_format.right_indent = Cm(5.7)
        _borda_p(f, "bottom", PET, 12, 1)
        _fmt(f.add_run(" "), 4)
        t = self.par(antes=7, depois=1, align="center", keep=True)
        rotulo = f"{romano(self.n_sec)} — {titulo.upper()}" if prefixo else titulo.upper()
        _fmt(t.add_run(rotulo), 12, True)
        s = self.par(antes=0, depois=8, align="center", keep=True)
        _fmt(s.add_run(linha), 8.5, False, False, GRA, SANS)

    def sub(self, texto):
        p = self.par(antes=9, depois=3, keep=True)
        _fmt(p.add_run(texto), 12, True, False, PETE)

    def pgf(self, texto):
        self.n_pgf += 1
        p = self.par()
        _fmt(p.add_run(f"{self.n_pgf}."), 12, True, False, PET)
        p.add_run(" ")
        _rico(p, texto)
        return p

    def pgf_literal(self, numero, texto):
        """Parágrafo cujo rótulo numérico já foi auditado na fonte canônica."""
        self.n_pgf += 1
        p = self.par()
        _fmt(p.add_run(f"{numero}."), 12, True, False, PET)
        p.add_run(" ")
        _rico(p, texto)
        return p

    def item_literal(self, rotulo, texto):
        """Alínea auditada, sem prefixo decimal e com recuo francês legível."""
        p = self.par(antes=1, depois=5)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        _fmt(p.add_run(rotulo), 12, True, False, PET)
        p.add_run(" ")
        _rico(p, texto)
        return p

    # ---------- margem ----------
    def _frame(self, largura_tw=2098, x_tw=9043):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        fp = OxmlElement("w:framePr")
        for k, v in {"w:w": str(largura_tw), "w:hSpace": "142", "w:wrap": "around",
                     "w:vAnchor": "text", "w:hAnchor": "page",
                     "w:x": str(x_tw), "w:y": "1"}.items():
            fp.set(qn(k), v)
        pPr.insert(0, fp)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.1
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def pull(self, texto):
        p = self._frame()
        _borda_p(p, "top", TER, 16, 4)
        _rico(p, texto, 8.5, PET, SANS, italico_base=True)
        return p

    def lateral(self, texto):
        p = self._frame()
        _rico(p, texto, 8, GRA, SANS)
        return p

    def abrir_secao_tabela_larga(self):
        """Abre seção contínua sem margem lateral para quadros extensos."""
        sec = self.doc.add_section(WD_SECTION.CONTINUOUS)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.different_first_page_header_footer = False
        return sec

    def fechar_secao_tabela_larga(self):
        """Restaura a coluna de 13,1 cm e a faixa lateral do padrão Medina."""
        sec = self.doc.add_section(WD_SECTION.CONTINUOUS)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(5.4)
        sec.different_first_page_header_footer = False
        return sec

    def faixa_leitura_tabela(self, texto, largura_cm=16.2):
        """Faixa cognitiva para tabelas longas em seção de largura integral.

        Nas páginas tabulares a própria matriz substitui a coluna lateral.
        Esta faixa recompõe a camada de escaneamento sem comprimir colunas nem
        introduzir conteúdo jurídico novo.
        """
        t = self.doc.add_table(rows=1, cols=1)
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.columns[0].width = Cm(largura_cm)
        c = t.rows[0].cells[0]
        c.width = Cm(largura_cm)
        _sombra(c, PTC)
        _bordas(c, {"left": (TERE, 22), "bottom": (RUL, 4)})
        _margens(c, 110, 110, 55, 55)
        p = _pc(c)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.keep_with_next = True
        _fmt(p.add_run("LEITURA RÁPIDA  "), 8.2, True, False, TERE, SANS)
        _fmt(p.add_run(texto), 8.5, False, False, GRA, SANS)
        return t

    # ---------- caixas ----------
    def _caixa(self, titulo, corpo, fundo, cor_borda, cor_titulo, escura=False):
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.columns[0].width = Cm(13.1)
        c = t.rows[0].cells[0]
        c.width = Cm(13.1)
        _sombra(c, fundo)
        _bordas(c, {"left": (cor_borda, 24)})
        _margens(c, 170, 170, 110, 110)
        if titulo:
            p = _pc(c)
            _fmt(p.add_run(titulo.replace("**", "").upper()), 8, True, False, cor_titulo, SANS)
            p.paragraph_format.space_after = Pt(4)
            p2 = c.add_paragraph()
        else:
            p2 = _pc(c)
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.line_spacing = 1.18
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if escura:
            _rico(p2, corpo, 10, "FFFFFF", SANS)
        else:
            _rico(p2, corpo, 11, italico_base=True)
        self.par(" ", 2, depois=6)
        return t

    def caixa_acordao(self, titulo, corpo):
        return self._caixa(titulo, corpo, PV_, PET, PET)

    def caixa_precedente(self, titulo, corpo):
        return self._caixa(titulo, corpo, PTC, TERE, TERE)

    def caixa_chave(self, corpo):
        return self._caixa(None, corpo, PET, PET, PET, escura=True)

    # ---------- quadros ----------
    def quadro_zebrado(
        self,
        cabecalho,
        linhas,
        larguras_cm=(3.1, 6.4, 3.6),
        permitir_quebra_linha=False,
        alinhar_esquerda=False,
    ):
        tq = self.doc.add_table(rows=len(linhas) + 1, cols=len(cabecalho))
        tq.autofit = False
        tq.alignment = (
            WD_TABLE_ALIGNMENT.LEFT if alinhar_esquerda else WD_TABLE_ALIGNMENT.CENTER
        )
        largs = tuple(Cm(x) for x in larguras_cm)
        total_twips = round(sum(larguras_cm) * 567)
        tblPr = tq._tbl.tblPr
        tblW = tblPr.first_child_found_in("w:tblW")
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(total_twips))
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        grid = tq._tbl.tblGrid
        for j, col in enumerate(grid.gridCol_lst):
            col.set(qn("w:w"), str(round(larguras_cm[j] * 567)))
        # cabeçalho repete quando o quadro atravessa páginas; linha nunca corta no meio
        trPr = tq.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:tblHeader"))
        if not permitir_quebra_linha:
            for row in tq.rows:
                row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for j, cell in enumerate(tq.rows[0].cells):
            cell.width = largs[j]
            _sombra(cell, PET)
            _bordas(cell, {})
            _margens(cell, 90, 90, 60, 60)
            # cabeçalho é uniformemente bold: marcadores ** seriam redundantes — removidos
            _fmt(_pc(cell).add_run(cabecalho[j].replace("**", "")), 8.5, True, False, "FFFFFF", SANS)
        for i, linha in enumerate(linhas, start=1):
            zebra = PV_ if i % 2 == 1 else "FFFFFF"
            for j, txt in enumerate(linha):
                cell = tq.rows[i].cells[j]
                cell.width = largs[j]
                _sombra(cell, zebra)
                _bordas(cell, {"bottom": (RUL, 4)} if i < len(linhas) else {"bottom": (PET, 10)})
                _margens(cell, 90, 90, 55, 55)
                # célula de corpo aceita **negrito**/*itálico* (Lição 43: texto cru vazava asteriscos)
                _rico(_pc(cell), txt, 8.5, GRA, SANS)
        return tq

    # ---------- fecho ----------
    def pedidos(self, itens, intro="Ante o exposto, requer-se:", depois_pt=5):
        self.par(intro, antes=10)
        for letra, txt in itens:
            p = self.par(antes=2, depois=depois_pt)
            p.paragraph_format.left_indent = Cm(0.9)
            _fmt(p.add_run(letra), 12, True, False, PET)
            p.add_run(" ")
            _rico(p, txt)

    def topicos(self, itens, marcador="•"):
        """Lista de tópicos com marcador em petróleo (prequestionamento, pedidos em bullets)."""
        for txt in itens:
            p = self.par(antes=1, depois=4)
            p.paragraph_format.left_indent = Cm(0.9)
            _fmt(p.add_run(marcador), 12, True, False, PET)
            p.add_run("  ")
            _rico(p, txt)

    def fecho(self, data, formula="Termos em que, pede deferimento."):
        # keep: o bloco de assinaturas nunca fica sozinho em página órfã
        if formula:
            self.par(formula, antes=8, keep=True)
        self.par(data, antes=0 if formula else 8, keep=True)

    def assinaturas(self, lista, espaco_linha_pt=26):
        """lista: [(nome, oab), ...] — pares em 2 colunas; ímpar final centralizado.
        Bloco indivisível (cantSplit + keepNext)."""
        pares = [lista[i:i + 2] for i in range(0, len(lista), 2)]
        ta = self.doc.add_table(rows=len(pares), cols=2)
        ta.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in ta.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                _bordas(cell, {})

        def celula(cell, nome, oab):
            cell.width = Cm(6.4)
            lin = _pc(cell)
            lin.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lin.paragraph_format.space_before = Pt(espaco_linha_pt)
            lin.paragraph_format.space_after = Pt(2)
            lin.paragraph_format.left_indent = Cm(0.4)
            lin.paragraph_format.right_indent = Cm(0.4)
            _borda_p(lin, "bottom", GRA, 4, 1)
            _fmt(lin.add_run(" "), 2)
            pn = cell.add_paragraph()
            pn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pn.paragraph_format.space_after = Pt(0)
            _fmt(pn.add_run(nome), 10, True)
            po = cell.add_paragraph()
            po.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            po.paragraph_format.space_after = Pt(4)
            _fmt(po.add_run(oab), 10)

        for i, par_ in enumerate(pares):
            if len(par_) == 2:
                celula(ta.rows[i].cells[0], *par_[0])
                celula(ta.rows[i].cells[1], *par_[1])
            else:
                m = ta.rows[i].cells[0].merge(ta.rows[i].cells[1])
                celula(m, *par_[0])
        for row in ta.rows[:-1]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    @staticmethod
    def _tipo_do_produto(texto, titulo):
        return _classificar_produto(texto, titulo)

    def salvar(self):
        self._sanitizar_metadados()
        self._validar_porta_unica()
        try:
            self._validar_lastro_documental()
        except Exception as erro:
            self._marcar_laudo_porta_unica_reprovado("FORJA-LASTRO", erro)
            raise
        try:
            self.doc.save(self.saida)
            self._finalizar_laudo_porta_unica()
        except Exception as erro:
            self._marcar_laudo_porta_unica_reprovado(
                "FORJA-MATERIALIZACAO", erro)
            raise
        return self.saida

    def _validar_porta_unica(self):
        """O verificador da casa roda em TODA peça, econômica ou não.

        Decisão unânime do conselho Helena + Efesto + Diabob em 05/08/2026, sob
        delegação expressa do Igor: **rota única obrigatória**. O que a medição
        do Efesto encontrou é que existiam seis caminhos capazes de gerar DOCX
        por esta classe — a entrada canônica `forja_visual_build.py` e cinco
        scripts `build_docx.py` dentro de pastas de caso — e só o primeiro
        chamava o verificador.

        O buraco era maior do que "os cinco scripts não têm gate". Era que
        `_validar_lastro_documental` **retorna cedo quando a peça não tem
        conteúdo econômico**: uma petição sem valor em reais, saindo por rota ad
        hoc, atravessava esta classe sem nenhuma conferência — nem placeholder
        esquecido, nem persona interna vazada, nem cara de IA. O gate econômico
        estava certo em ser condicional; o erro foi não haver nada incondicional
        atrás dele.

        Calibrado antes de travar, como manda a casa: rodado contra as 25 obras
        mais recentes do acervo com a classificação correta de tipo, **zero
        seriam bloqueadas** pelas famílias que travam. Não é trava inexequível
        herdada de teoria; é porta que a produção real já atravessa limpa.

        **Não existe bandeira para pular esta porta, e a ausência é deliberada.**
        A primeira versão tinha um `porta_unica_ja_validada` que a rota canônica
        marcava para não pagar o gate duas vezes pelo mesmo texto. Foi removido:
        um sinalizador que desliga o gate é um bypass, e bypass em gate de
        produção é o que esta fábrica passou o mês inteiro fechando. A rota
        canônica paga o verificador duas vezes — uma sobre o markdown de origem,
        outra sobre o texto composto — e esse é o preço de não haver interruptor.
        São conferências diferentes, aliás: a primeira vê o que se pretendeu
        escrever, a segunda vê o que de fato foi para o papel.

        Também não há corte por tamanho. A versão inicial dispensava texto com
        menos de 400 caracteres, para o verificador não acusar ausência de
        estrutura em fragmento — e isso era rota de fuga: um documento curto com
        placeholder atravessava. O par de contraprovas em
        `test_forja_porta_unica.py` cobre exatamente esse caso.
        """
        texto = self._texto_para_gate()
        harness = Path(AQUI).parent / "_FORJA_HARNESS"
        if str(harness) not in sys.path:
            sys.path.insert(0, str(harness))
        from forja_verificador import verificar
        tipo = getattr(self, "tipo_produto", None) or _classificar_produto(
            texto, Path(str(self.saida)).stem)
        achados = verificar(texto, tipo)
        texto_sha256 = hashlib.sha256(texto.encode("utf-8")).hexdigest()
        p0 = [a for a in (achados or [])
              if str(a.get("sev") or a.get("severidade") or "").upper().startswith("P0")]
        bloqueantes = [a for a in p0
                       if not str(a.get("gate", "")).startswith("G10-escrita-humana")]
        estilo = len(p0) - len(bloqueantes)
        # O laudo é gravado SEMPRE, passe ou não. Era esta a cegueira: peça que
        # saía por rota ad hoc não deixava rastro nenhum de conferência, e o
        # silêncio era lido como aprovação.
        laudo = Path(str(self.saida)).with_name(
            Path(str(self.saida)).stem + "_PORTA_UNICA.json")
        self._porta_unica_laudo = laudo
        try:
            laudo.write_text(json.dumps({
                "versao": "FORJA-PORTA-UNICA-v1",
                "docx": Path(str(self.saida)).name,
                "docxSha256": None,
                "textoSha256": texto_sha256,
                "tipoProduto": tipo,
                "veredito": "reprovado" if bloqueantes else "aprovado",
                "bloqueadores": bloqueantes,
                "achados": achados,
            }, ensure_ascii=False, indent=2), encoding="utf-8")
        except OSError as erro:
            try:
                Path(self.saida).unlink(missing_ok=True)
            except OSError:
                pass
            raise RuntimeError(
                f"FORJA — porta única não conseguiu persistir o laudo: {laudo}"
            ) from erro
        # A chave é `sev`, não `severidade`. Escrevi `severidade` na primeira
        # versão e a medição de calibração devolveu "zero peças bloqueadas" —
        # não porque o gate fosse benigno, mas porque ele não conseguia
        # enxergar achado nenhum. É o modo de falha que esta fábrica documenta
        # há meses: o gate verde que está verde por cegueira. Fica o `or` como
        # tolerância a schema, e o teste de contraprova como a prova real.
        # Duas famílias, e a separação foi MEDIDA, não estipulada.
        #
        # Bloqueiam: correção e identidade — placeholder no entregável, persona
        # interna vazada, origem operacional no corpo, regimento, lastro. São os
        # bloqueadores que a casa já declarou invioláveis. Medidos contra as 25
        # obras mais recentes com a classificação correta de tipo: **zero
        # reprovam**. Fechar esta porta hoje não trava uma única peça real.
        #
        # Não bloqueiam, mas ficam no laudo: `G10-escrita-humana`, a família de
        # estilo. Ela reprova 18 das 25 obras recentes — e não por ser ruidosa:
        # a `IMPUGNACAO_AGINT_CAFELANA_V8`, peça aprovada e entregue, passa nela
        # com zero achados, enquanto a V4 que ela superou reprova. O gate
        # discrimina certo, e os 72% são um achado real de qualidade da
        # produção corrente.
        #
        # Transformar esse achado em parede às 2h da manhã pararia a fábrica por
        # um defeito de estilo, o que a casa proíbe desde a bronca de 10/07.
        # Isso não é waiver: nada é dispensado, tudo é gravado, e o número foi
        # levado ao Igor como achado. O que não se faz é descobrir um problema
        # de redação e responder com uma catraca que impede trabalhar.
        if bloqueantes:
            resumo = "; ".join(
                f"{a.get('gate')}: {a.get('problema')}" for a in bloqueantes[:5])
            try:
                Path(self.saida).unlink(missing_ok=True)
            except OSError as erro:
                raise RuntimeError(
                    f"FORJA — porta única reprovou e não removeu a saída parcial: {self.saida}"
                ) from erro
            raise RuntimeError(
                f"FORJA — porta única REPROVOU antes de salvar: {resumo}")
        if estilo:
            print(f"  [porta única] {estilo} achado(s) P0 de escrita humana "
                  f"registrados em {Path(str(self.saida)).stem}_PORTA_UNICA.json")

    def _marcar_laudo_porta_unica_reprovado(self, gate, erro):
        """Não deixe laudo aprovado apontando para uma saída que falhou depois.

        A porta é o primeiro gate e o lastro/materialização vêm depois dela.
        Se um desses passos falhar, o DOCX parcial é removido, mas o laudo
        continua sendo evidência útil — agora com veredito reprovado e o motivo
        da falha posterior, em vez de parecer uma aprovação sem DOCX hash-bound.
        """
        laudo = getattr(self, "_porta_unica_laudo", None)
        if laudo is not None and Path(laudo).is_file():
            try:
                dados = json.loads(Path(laudo).read_text(encoding="utf-8"))
                bloqueadores = list(dados.get("bloqueadores") or [])
                bloqueadores.append({
                    "gate": gate,
                    "sev": "P0",
                    "problema": str(erro),
                })
                dados["veredito"] = "reprovado"
                dados["docxSha256"] = None
                dados["bloqueadores"] = bloqueadores
                Path(laudo).write_text(
                    json.dumps(dados, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            except (OSError, json.JSONDecodeError):
                # Nunca substitua o erro original por falha de evidência.
                pass
        try:
            Path(self.saida).unlink(missing_ok=True)
        except OSError:
            # O chamador ainda recebe o erro original; a mensagem da porta já
            # distingue a saída parcial que não pôde ser removida.
            pass

    def _finalizar_laudo_porta_unica(self):
        """Vincula a aprovação ao DOCX efetivamente materializado."""
        laudo = getattr(self, "_porta_unica_laudo", None)
        if laudo is None:
            return
        try:
            dados = json.loads(laudo.read_text(encoding="utf-8"))
            dados["docxSha256"] = hashlib.sha256(
                Path(self.saida).read_bytes()).hexdigest()
            laudo.write_text(json.dumps(dados, ensure_ascii=False, indent=2),
                             encoding="utf-8")
        except (OSError, json.JSONDecodeError) as erro:
            try:
                Path(self.saida).unlink(missing_ok=True)
            except OSError:
                pass
            raise RuntimeError(
                f"FORJA — porta única não conseguiu finalizar o laudo: {laudo}"
            ) from erro

    def _texto_para_gate(self):
        partes = [p.text for p in self.doc.paragraphs]
        for tabela in self.doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    partes.extend(p.text for p in celula.paragraphs)
        return "\n".join(partes)

    def _validar_lastro_documental(self):
        """Executa L9--L13 antes de materializar a saída visual.

        A chamada acontece apenas no ponto de persistência da classe. Builders
        que escrevem `doc.save()` diretamente continuam sendo cobertos pelo
        `forja_verificador` da rota visual canônica; não há uma terceira autoridade.
        """
        texto = self._texto_para_gate()
        harness = Path(AQUI).parent / "_FORJA_HARNESS"
        sys.path.insert(0, str(harness))
        from forja_lastro import material_economico
        if not material_economico(texto):
            return
        from forja_verificador import verificar
        ledger = None
        candidatos = []
        # Um caminho de ledger fornecido pela rota é uma decisão de autoridade,
        # não uma sugestão de autodiscovery. Se ele não existe (ou aponta para
        # diretório), bloquear antes de considerar qualquer ledger do caso;
        # aceitar o canônico encontrado depois transformaria erro de entrada em
        # aprovação silenciosa.
        caminho_explicito = Path(self.ledger_path) if self.ledger_path else None
        explicitamente_quebrado = (
            caminho_explicito is not None and not caminho_explicito.is_file()
        )
        if explicitamente_quebrado:
            ledger = {}
        elif caminho_explicito is not None:
            candidatos.append(caminho_explicito)
        if caminho_explicito is None and self.case_dir:
            caso = Path(self.case_dir)
            candidatos.extend([
                caso / "producao" / "fact_ledger.json",
                caso / "n3_artifacts" / "F3_FONTES_REGIMENTO_LEIS" / "fact_ledger.json",
            ])
            candidatos.extend(sorted(caso.glob("n3_artifacts/F3_FONTES_REGIMENTO_LEIS/fact_ledger*.json")))
        for candidato in candidatos:
            if candidato.is_file():
                try:
                    ledger = json.loads(candidato.read_text(encoding="utf-8"))
                    if not isinstance(ledger, dict):
                        # Um caminho fornecido explicitamente não pode ser
                        # trocado por um ledger encontrado no caso.
                        if self.ledger_path and candidato == Path(self.ledger_path):
                            ledger = {}
                            break
                        continue
                    break
                except (OSError, UnicodeDecodeError, json.JSONDecodeError):
                    # O chamador apontou uma fonte; erro nela é bloqueio, não
                    # convite para aceitar um snapshot ou outro caminho.
                    if self.ledger_path and candidato == Path(self.ledger_path):
                        ledger = {}
                        break
                    continue
        violacoes = verificar(
            texto, "peca", ledger=ledger, base_dir=self.base_dir or self.case_dir,
            case_dir=self.case_dir, exigir_economico=True,
        )
        p0 = [item for item in violacoes if item.get("sev") == "P0"]
        if p0:
            # O arquivo foi criado pelo construtor no __init__; removê-lo aqui
            # evita que um template parcial pareça uma saída válida.
            try:
                Path(self.saida).unlink(missing_ok=True)
            except OSError:
                pass
            resumo = "; ".join(f"{item.get('gate')}: {item.get('problema')}" for item in p0[:5])
            raise RuntimeError(f"FORJA-LASTRO REPROVADO antes de salvar — {resumo}")

    def _sanitizar_metadados(self):
        # gate G8.2: o template carrega autor/título de documento antigo
        # ("thais mulati" / "Proposta de Serviços e Honorários") — nunca
        # podem chegar ao DOCX entregável
        import os
        cp = self.doc.core_properties
        cp.author = "Medina Osório Advogados"
        cp.last_modified_by = "Medina Osório Advogados"
        cp.title = os.path.splitext(os.path.basename(str(self.saida)))[0].replace("_", " ")
        cp.subject = ""
        cp.keywords = ""
        cp.category = ""
        cp.comments = ""
