# -*- coding: utf-8 -*-
"""
Pipeline de diagramação de excelência para petições em Word.
Uso pelas IAs da fábrica de melhoria de petições.

Funções:
  svg_para_emf(svg, emf)        — converte SVG em EMF vetorial (Inkscape) para embutir no Word sem perda
  mermaid_para_svg(mmd, svg)    — renderiza diagrama Mermaid em SVG (mmdc)
  dot_para_svg(dot, svg)        — renderiza Graphviz em SVG
  inserir_imagem_docx(...)      — insere EMF/PNG em docx via python-docx com legenda estilo forense
  docx_para_pdf(docx, pdf)      — conversão de alta fidelidade via Word COM (requer MS Word)
  render_paginas(pdf, pasta)    — renderiza páginas em PNG (pymupdf) para o gate de QA visual

Regra da casa: diagramas SEMPRE vetoriais (EMF) dentro do Word; PNG só como fallback.
Paleta sóbria forense: navy #1B2A4A, grafite #3A3F47, bronze #8C6A2F, painéis #F4F2ED / #EEF1F6.
"""
import gc
import subprocess, shutil, os, sys, re, tempfile, time

INKSCAPE = shutil.which("inkscape") or r"C:\Program Files\Inkscape\bin\inkscape.exe"
MMDC = shutil.which("mmdc") or os.path.expandvars(r"%APPDATA%\npm\mmdc.cmd")
DOT = shutil.which("dot") or r"C:\Program Files\Graphviz\bin\dot.exe"


def svg_para_emf(svg_path: str, emf_path: str, largura_final_cm: float = None) -> str:
    """SVG -> EMF vetorial. O Word renderiza EMF sem serrilhado em qualquer zoom/impressão.

    Se largura_final_cm for informada, roda o gate de legibilidade ANTES da conversão
    e levanta erro se algum texto ficar abaixo do mínimo impresso (causa raiz das
    'fontes pequenas' nas petições)."""
    if largura_final_cm:
        from estilo_medina import checar_fontes_svg, MIN_PT_ROTULO, fonte_px_minima
        import re as _re
        viol = checar_fontes_svg(svg_path, largura_final_cm)
        if viol:
            with open(svg_path, encoding="utf-8") as f:
                vb = _re.search(r'viewBox\s*=\s*"([^"]+)"', f.read()).group(1)
            w = float([p for p in _re.split(r"[\s,]+", vb.strip()) if p][2])
            raise ValueError(
                f"LEGIBILIDADE REPROVADA em {svg_path} @ {largura_final_cm}cm: "
                f"font-sizes {viol} (px, pt_impresso) abaixo de {MIN_PT_ROTULO}pt. "
                f"Aumente para >= {fonte_px_minima(w, largura_final_cm):.0f}px ou "
                f"reduza o viewBox.")
    # Gate de desenho (03/08/2026). Roda SEMPRE, e não só quando a largura é
    # informada: colisão e oclusão não dependem da largura de inserção. Este é o
    # ponto por onde todo SVG passa antes de entrar no Word, inclusive o
    # desenhado à mão, que é onde a falha vive — o gate de `medina_svg_kit`
    # sozinho só cobriria os diagramas gerados por código.
    from medina_svg_colisao import checar as _checar_desenho
    _checar_desenho(svg_path)
    subprocess.run([INKSCAPE, svg_path, "--export-type=emf",
                    "--export-filename=" + emf_path], check=True)
    return emf_path


def gerar_imagem_ia(prompt: str, png_saida: str, modelo: str = "infsh/z-image-turbo") -> str:
    """Gera imagem por IA via inference.sh (infsh) e baixa como PNG.

    Uso em contexto jurídico: capas de relatórios/pareceres ao cliente, ícones e
    texturas institucionais discretas. NUNCA para retratar fatos, pessoas ou provas
    em peça protocolada (risco de indução a erro do juízo). Ver política na skill
    fabrica-visual-peticoes."""
    import json, urllib.request
    infsh = shutil.which("infsh") or os.path.expanduser(r"~\.local\bin\infsh.exe")
    r = subprocess.run([infsh, "app", "run", modelo, "--input",
                        json.dumps({"prompt": prompt})],
                       capture_output=True, text=True, check=True)
    m = re.search(r'https://\S+\.(?:png|jpg|jpeg|webp)', r.stdout)
    if not m:
        raise RuntimeError("infsh não retornou URL de imagem:\n" + r.stdout[-2000:])
    urllib.request.urlretrieve(m.group(0), png_saida)
    return png_saida


def svg_para_png(svg_path: str, png_path: str, dpi: int = 300) -> str:
    subprocess.run([INKSCAPE, svg_path, "--export-type=png", f"--export-dpi={dpi}",
                    "--export-filename=" + png_path], check=True)
    return png_path


def mermaid_para_svg(mmd_path: str, svg_path: str, width: int = 1400) -> str:
    subprocess.run([MMDC, "-i", mmd_path, "-o", svg_path, "-w", str(width),
                    "-b", "transparent"], check=True, shell=False)
    return svg_path


def dot_para_svg(dot_path: str, svg_path: str) -> str:
    subprocess.run([DOT, "-Tsvg", dot_path, "-o", svg_path], check=True)
    return svg_path


def inserir_emf_word_com(docx_path: str, marcadores: dict, largura_cm: float = 15.0):
    """Substitui placeholders de texto (ex.: '{{FIG1}}') por imagens EMF vetoriais.

    SEMPRE via Word COM: python-docx não reconhece o cabeçalho EMF
    (UnrecognizedImageError — regressão flagrada em teste real 09/07/2026).
    Valores do dict: caminho do EMF, ou (caminho, largura_cm) por figura."""
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        word.AutomationSecurity = 3  # sem macros
    except Exception:
        pass
    doc = None
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        for marcador, valor in marcadores.items():
            emf_path, larg = valor if isinstance(valor, (tuple, list)) else (valor, largura_cm)
            find = word.Selection.Find
            word.Selection.HomeKey(6)  # wdStory
            find.ClearFormatting()
            if find.Execute(marcador):
                word.Selection.Text = ""
                shape = word.Selection.InlineShapes.AddPicture(os.path.abspath(emf_path), False, True)
                escala = (larg * 28.3465) / shape.Width
                shape.Width = int(shape.Width * escala)
                shape.Height = int(shape.Height * escala)
                word.Selection.ParagraphFormat.Alignment = 1  # centralizado
            else:
                raise RuntimeError(f"MARCADOR NÃO ENCONTRADO no DOCX: {marcador}")
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()
    return docx_path


def inserir_imagem_docx(doc, caminho_img: str, largura_cm: float = 15.0,
                        legenda: str = None):
    """Insere imagem centralizada com legenda em itálico 9pt cinza (padrão forense).
    Aceita PNG/JPG/WMF. Para EMF vetorial use inserir_emf_word_com (python-docx
    não reconhece o cabeçalho EMF)."""
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(caminho_img, width=Cm(largura_cm))
    if legenda:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pl.add_run(legenda)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x3A, 0x3F, 0x47)
    return doc


def _docx_para_pdf_once(docx_path: str, temp_pdf: str):
    worker = os.path.join(os.path.dirname(os.path.abspath(__file__)), "word_pdf_worker.py")
    pid_file = temp_pdf + ".word-pid"
    status_file = temp_pdf + ".word-status"
    timeout = max(15, int(os.environ.get("FORJA_WORD_PDF_TIMEOUT_SECONDS", "75")))
    try:
        proc = subprocess.run(
            [sys.executable, worker, os.path.abspath(docx_path), os.path.abspath(temp_pdf), pid_file, status_file],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
            shell=False,
        )
        if proc.returncode != 0:
            detail = (proc.stderr or proc.stdout or "falha sem detalhe")[-1500:]
            raise RuntimeError(f"worker Word terminou com código {proc.returncode}: {detail}")
    except subprocess.TimeoutExpired as exc:
        # O PID pertence à instância privada criada por DispatchEx neste worker.
        # Não toca em qualquer Word que o usuário já estivesse usando.
        try:
            pid = int(open(pid_file, encoding="ascii").read().strip())
            os.kill(pid, 15)
        except (FileNotFoundError, OSError, ValueError):
            pass
        try:
            stage = open(status_file, encoding="ascii").read().strip()
        except (FileNotFoundError, OSError):
            stage = "etapa desconhecida"
        raise TimeoutError(f"Word excedeu {timeout}s na conversão isolada ({stage})") from exc
    finally:
        gc.collect()
        try:
            os.unlink(pid_file)
        except FileNotFoundError:
            pass
        try:
            os.unlink(status_file)
        except FileNotFoundError:
            pass


def docx_para_pdf(docx_path: str, pdf_path: str):
    """Conversão atômica via Word COM, com uma retomada após falha transitória."""
    source = os.path.abspath(docx_path)
    destination = os.path.abspath(pdf_path)
    os.makedirs(os.path.dirname(destination), exist_ok=True)
    last_error = None
    for attempt in range(2):
        fd, temp_name = tempfile.mkstemp(
            prefix=f".{os.path.basename(destination)}.", suffix=".pdf", dir=os.path.dirname(destination))
        os.close(fd)
        os.unlink(temp_name)
        try:
            _docx_para_pdf_once(source, temp_name)
            if not os.path.isfile(temp_name) or os.path.getsize(temp_name) == 0:
                raise RuntimeError("Word não produziu um PDF válido.")
            os.replace(temp_name, destination)
            return pdf_path
        except Exception as exc:
            last_error = exc
            if attempt == 0:
                time.sleep(1.0)
        finally:
            try:
                os.unlink(temp_name)
            except FileNotFoundError:
                pass
    raise RuntimeError(f"Falha na conversão Word após uma retomada: {last_error}") from last_error


def render_paginas(pdf_path: str, pasta_saida: str, dpi: int = 100):
    """Gate de QA visual: renderiza cada página em PNG para inspeção obrigatória."""
    import fitz
    os.makedirs(pasta_saida, exist_ok=True)
    # Um PDF novo pode ter menos páginas que a versão anterior. Sem limpar
    # somente os PNGs canônicos produzidos por esta função, p10-p13 antigos
    # sobrevivem ao render de uma peça agora com 9 páginas e falsificam o QA.
    for nome in os.listdir(pasta_saida):
        if re.fullmatch(r"p\d+\.png", nome, flags=re.IGNORECASE):
            caminho = os.path.join(pasta_saida, nome)
            if os.path.isfile(caminho):
                os.unlink(caminho)
    d = fitz.open(pdf_path)
    saidas = []
    for i, pg in enumerate(d):
        out = os.path.join(pasta_saida, f"p{i+1:02d}.png")
        pg.get_pixmap(dpi=dpi).save(out)
        saidas.append(out)
    return saidas


if __name__ == "__main__":
    print("Ferramentas detectadas:")
    for nome, caminho in [("Inkscape", INKSCAPE), ("mermaid-cli", MMDC), ("Graphviz dot", DOT)]:
        existe = caminho and (os.path.exists(caminho) or shutil.which(str(caminho)))
        print(f"  {'OK ' if existe else '-- '}{nome}: {caminho}")
